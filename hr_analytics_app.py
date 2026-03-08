# ===================================================
# منصة تحليلات الموارد البشرية الذكية v5.0
# رسال الود لتقنية المعلومات
# المجموعة أ: تحليل الرواتب + Headcount + حاسبة المستحقات + الأداء
# + المرحلة 2: ميزانية التدريب + ROI + الاحتياجات التدريبية
# ===================================================

import streamlit as st
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
import io, math, json, sqlite3, os, smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
import openpyxl
from datetime import datetime, date
from dateutil.relativedelta import relativedelta
import hashlib, urllib.request

# =====================================================================
# ENTERPRISE ARCHITECTURE: 3 Core Modules
# =====================================================================

# ==================== 1. MODEL ORCHESTRATOR ====================
class ModelOrchestrator:
    """Independent orchestration layer for AI model management.
    Handles: prompt building, model routing, context assembly, caching, error management."""

    MODELS = {
        'claude': {'url':'https://api.anthropic.com/v1/messages','model':'claude-sonnet-4-20250514','max_tokens':4000},
        'groq': {'url':'https://api.groq.com/openai/v1/chat/completions','model':'llama-3.3-70b-versatile','max_tokens':4000},
        'gemini': {'url':'https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent','model':'gemini-2.0-flash','max_tokens':2000},
        'openrouter': {'url':'https://openrouter.ai/api/v1/chat/completions','model':'meta-llama/llama-3.3-70b-instruct:free','max_tokens':2000},
        'huggingface': {'url':'https://api-inference.huggingface.co/v1/chat/completions','model':'mistralai/Mistral-7B-Instruct-v0.3','max_tokens':1500},
    }

    # Fallback models for OpenRouter
    OR_MODELS = ['meta-llama/llama-3.3-70b-instruct:free','google/gemma-3-27b-it:free','mistralai/mistral-7b-instruct:free']

    # Prompt Templates (editable registry)
    PROMPT_TEMPLATES = {
        'labor_law': 'labor_system_prompt',
        'hr_expert': 'hr_system_prompt',
        'general': 'You are a helpful HR assistant. Answer in the same language as the question.',
    }

    # Context size limits per provider
    CONTEXT_LIMITS = {'claude': 15000, 'groq': 12000, 'gemini': 8000, 'openrouter': 4000, 'huggingface': 3000}

    def __init__(self):
        self._cache = {}
        self._call_count = 0
        self._token_estimate = 0
        # Pre-cached instant responses for common questions
        self._instant = {
            "ما هي حقوقي عند الفصل وفق المادة 77؟": "**المادة 77 من نظام العمل السعودي:**\n\nإذا أُنهي العقد لسبب غير مشروع، يحق للطرف المتضرر تعويض:\n\n**عقد غير محدد المدة:** أجر 15 يوماً عن كل سنة خدمة\n**عقد محدد المدة:** أجر المدة الباقية من العقد\n**الحد الأدنى:** لا يقل التعويض عن أجر شهرين في كلا الحالتين\n\n**بالإضافة إلى:**\n- مكافأة نهاية الخدمة (المادة 84)\n- بدل إجازة غير مستخدمة\n- أجر الشهر الأخير كاملاً\n- شهادة خبرة",
            "متى يتحول العقد المحدد لغير محدد المدة؟": "**المادة 55 من نظام العمل:**\n\nيتحول العقد محدد المدة إلى غير محدد في الحالات التالية:\n\n1. **التجديد 3 مرات متتالية** أو بلوغ 4 سنوات أيهما أقل\n2. **استمرار العمل** بعد انتهاء العقد دون تجديد\n3. **نص العقد** على التحول التلقائي\n\nيحق للعامل بعد التحول الاستفادة من مزايا العقد غير المحدد في الإشعار والتعويض.",
            "كيف تُحسب مكافأة نهاية الخدمة؟": "**المادة 84 من نظام العمل:**\n\n**الحساب:**\n- **أول 5 سنوات:** نصف راتب شهري عن كل سنة\n- **بعد 5 سنوات:** راتب شهري كامل عن كل سنة\n\n**عند الاستقالة (المادة 85):**\n- أقل من سنتين: لا يستحق\n- 2-5 سنوات: ثلث المكافأة\n- 5-10 سنوات: ثلثا المكافأة\n- أكثر من 10 سنوات: المكافأة كاملة\n\n**المدة:** يُلزم صاحب العمل بدفعها خلال أسبوع (المادة 88).",
            "ما هي نسبة اشتراكات التأمينات الاجتماعية؟": "**نسب الاشتراك في التأمينات الاجتماعية (GOSI):**\n\n**السعوديون:**\n- المعاشات: 9.75% (الموظف) + 9.75% (صاحب العمل)\n- الأخطار المهنية: 2% (صاحب العمل)\n- ساند (التعطل): 0.75% + 0.75%\n- **إجمالي خصم الموظف: 10.5%**\n- **إجمالي على الشركة: 12.5%**\n\n**غير السعوديين:**\n- الأخطار المهنية فقط: 2% (صاحب العمل)\n- لا يوجد خصم على الموظف",
            "ما هي فترة التجربة وشروطها؟": "**المادة 53 من نظام العمل (المعدلة بالمرسوم م/44 - نافذة 19/2/2025):**\n\n- **الحد الأقصى:** 180 يوماً من البداية مباشرة في العقد\n- **التمديد:** محظور بعد التعديل. لو كُتبت 90 يوماً لا يمكن تمديدها\n- **لا تدخل في الحساب:** إجازات عيدي الفطر والأضحى والإجازات المرضية\n- **حق الإنهاء:** لكلا الطرفين (لم يعد يجوز قصره على طرف واحد)\n- **لا تعويض:** لا يستحق أي طرف تعويضاً عند الإنهاء خلالها\n- **لا مكافأة:** لا يستحق العامل مكافأة نهاية الخدمة\n- **عدم التكرار:** لا يجوز وضع العامل تحت التجربة أكثر من مرة إلا في مهنة مختلفة أو بعد 6 أشهر",
            "ما هي حقوق المرأة العاملة في نظام العمل؟": "**حقوق المرأة في نظام العمل السعودي (المعدل 2025):**\n\n- **إجازة وضع:** 12 أسبوعاً بأجر كامل (6 إلزامية بعد الولادة + 6 توزعها حسب رغبتها) - المادة 151 المعدلة\n- **ساعة رضاعة:** ساعة يومياً لمدة 24 شهراً\n- **حماية من الفصل:** أثناء الحمل وإجازة الوضع\n- **إجازة وفاة زوج:** 4 أشهر و10 أيام (عدة)\n- **المساواة:** أجر متساوٍ للعمل المتساوي\n- **مكافأة كاملة:** إذا أنهت العقد خلال 6 أشهر من الزواج أو 3 أشهر من الوضع\n- **طفل مريض:** إجازة شهر إضافي بأجر كامل",
            "كيف أبني خطة استقطاب فعالة؟": "**خطة استقطاب فعالة (Talent Acquisition Strategy):**\n\n**1. التحليل:**\n- تحديد الاحتياج الفعلي (Workforce Planning)\n- تحليل سوق العمل والرواتب\n\n**2. التصميم:**\n- وصف وظيفي واضح ومحدد\n- Employee Value Proposition (EVP)\n- قنوات الاستقطاب (LinkedIn, مواقع توظيف, تزكيات)\n\n**3. التنفيذ:**\n- ATS لتتبع المتقدمين\n- مقابلات منظمة (Structured Interviews)\n- تقييم الكفاءات (Competency Assessment)\n\n**4. القياس:**\n- Time-to-Hire\n- Cost-per-Hire\n- Quality of Hire\n- Source Effectiveness",
            "ما الفرق بين OKRs و KPIs؟": "**OKRs vs KPIs:**\n\n**KPIs (مؤشرات الأداء الرئيسية):**\n- تقيس الأداء المستمر\n- أرقام محددة (مثل: معدل دوران 15%)\n- تتبع الوضع الحالي\n- ثابتة نسبياً\n\n**OKRs (الأهداف والنتائج الرئيسية):**\n- تحدد أهداف طموحة للمستقبل\n- هدف + 3-5 نتائج قابلة للقياس\n- تتغير كل ربع سنة\n- طموحة (70% إنجاز = ممتاز)\n\n**مثال HR:**\n- **KPI:** معدل الدوران = 12%\n- **OKR:** هدف: تحسين الاحتفاظ بالموظفين\n  - NR1: خفض الدوران من 15% إلى 10%\n  - NR2: رفع رضا الموظفين إلى 85%\n  - NR3: تنفيذ برنامج تطوير لـ 50 موظف",
            "ما هو نموذج Phillips ROI للتدريب؟": "**نموذج Phillips ROI (5 مستويات):**\n\n**المستوى 1: Reaction (رد الفعل)**\n- رضا المتدربين عن البرنامج\n\n**المستوى 2: Learning (التعلم)**\n- المعرفة والمهارات المكتسبة\n\n**المستوى 3: Application (التطبيق)**\n- مدى تطبيق ما تعلموه في العمل\n\n**المستوى 4: Impact (الأثر)**\n- التأثير على مؤشرات الأعمال\n\n**المستوى 5: ROI (العائد على الاستثمار)**\n- ROI % = (الفوائد - التكاليف) / التكاليف × 100\n\n**مثال:** تدريب بتكلفة 50,000 ريال أدى لزيادة إنتاجية بقيمة 150,000 ريال\nROI = (150,000 - 50,000) / 50,000 × 100 = **200%**",
        }

    def select_provider(self, question_type='general'):
        provider = st.session_state.get('ai_provider', 'auto')
        claude_key = st.session_state.get('claude_api_key', '')
        groq_key = st.session_state.get('groq_api_key', '')
        or_key = st.session_state.get('openrouter_api_key', '')
        hf_key = st.session_state.get('huggingface_api_key', '')
        gemini_key = st.session_state.get('gemini_api_key', '')
        claude_ok = claude_key and not st.session_state.get('_claude_no_credit')

        if provider == 'auto':
            legal_kw = ['مادة','قانون','نظام العمل','تعويض','فصل','إنهاء','article','labor law','termination']
            is_legal = any(kw in question_type.lower() for kw in legal_kw)
            if is_legal and claude_ok: return 'claude'
            if gemini_key: return 'gemini'
            if groq_key: return 'groq'
            if or_key: return 'openrouter'
            if hf_key: return 'huggingface'
            if claude_ok: return 'claude'
            return None
        if provider == 'gemini' and gemini_key: return 'gemini'
        if provider == 'claude' and claude_ok: return 'claude'
        if provider == 'groq' and groq_key: return 'groq'
        if provider == 'huggingface' and hf_key: return 'huggingface'
        if provider == 'openrouter' and or_key: return 'openrouter'
        for p, k in [('gemini',gemini_key),('groq',groq_key),('openrouter',or_key),('huggingface',hf_key),('claude',claude_ok)]:
            if k: return p
        return None

    def build_context(self, system_prompt, user_message, model_type='general', provider='claude'):
        """Assemble context with STRICT advisor_type separation."""
        max_ctx = self.CONTEXT_LIMITS.get(provider, 8000)
        enhanced = system_prompt[:max_ctx // 2]

        # Determine advisor_type from system prompt
        advisor_type = "legal" if 'المستشار القانوني' in system_prompt else "hr"

        # Layer 1: RAG Knowledge Base (FILTERED by advisor_type)
        if hasattr(st.session_state, '_knowledge_engine'):
            rag_ctx = st.session_state._knowledge_engine.search(user_message, advisor_type=advisor_type)
            if rag_ctx:
                enhanced += f"\n\n**RETRIEVED KNOWLEDGE:**\n{rag_ctx[:6000]}"

        # Layer 2: Learned from past (FILTERED by model_type)
        if hasattr(st.session_state, '_learning_system'):
            learned = st.session_state._learning_system.get_relevant_history(user_message, model_type)
            if learned:
                enhanced += f"\n\n**LEARNED FROM PAST:**\n{learned[:2000]}"

        # Layer 3: Legal documents (ONLY for legal advisor)
        if advisor_type == "legal":
            legal_ctx = st.session_state.get('legal_docs_context', '')
            if legal_ctx:
                enhanced += f"\n\n**LEGAL REFERENCES:**\n{legal_ctx[:6000]}"

        if len(enhanced) > max_ctx:
            enhanced = enhanced[:max_ctx]

        return enhanced

    def get_cache_key(self, message, model_type):
        """Generate cache key for response caching."""
        return hashlib.md5(f"{message[:100]}_{model_type}".encode()).hexdigest()

    def call(self, system_prompt, user_message, chat_history=None, model_type='general'):
        """Main orchestrated call with routing, context, caching, fallback."""
        # Check instant cached responses first (< 1 second)
        for q, a in self._instant.items():
            if user_message.strip() == q.strip() or q in user_message:
                return a, None

        # Check response cache
        cache_key = self.get_cache_key(user_message, model_type)
        if cache_key in self._cache:
            return self._cache[cache_key], None

        provider = self.select_provider(user_message)
        if not provider:
            return None, "يرجى إدخال API Key (Groq مجاني أو Claude)"

        # Build enhanced context (size depends on provider)
        enhanced_prompt = self.build_context(system_prompt, user_message, model_type, provider)

        # Build messages
        messages = []
        if chat_history:
            # Limit history for speed (last 4 messages only)
            for msg in chat_history[-4:]:
                messages.append({"role": msg['role'], "content": msg['content'][:500]})
        messages.append({"role": "user", "content": user_message})

        # Estimate tokens
        self._token_estimate += len(enhanced_prompt.split()) + len(user_message.split())

        # Try primary provider, fallback to others
        providers_to_try = [provider]
        all_providers = ['gemini','groq','openrouter','huggingface','claude']
        for p in all_providers:
            if p != provider and st.session_state.get(f'{p}_api_key',''):
                providers_to_try.append(p)

        last_error = None
        for prov in providers_to_try:
            result, error = self._call_provider(prov, enhanced_prompt, messages)
            if result:
                self._call_count += 1
                self._cache[cache_key] = result
                if len(self._cache) > 50: self._cache = dict(list(self._cache.items())[-50:])
                if hasattr(st.session_state, '_learning_system'):
                    st.session_state._learning_system.save_interaction(user_message, result, model_type)
                return result, None
            if error != "fallback":
                last_error = error
        return None, last_error or "⚠️ لم يتمكن أي مزود من الإجابة. تحقق من مفاتيح API."

    def _call_provider(self, provider, system_prompt, messages):
        """Execute API call to specific provider."""
        api_key = st.session_state.get(f'{provider}_api_key', '')
        if not api_key: return None, "fallback"

        config = self.MODELS.get(provider)
        if not config: return None, "fallback"

        # Skip Claude if previously failed with no credit
        if provider == 'claude' and st.session_state.get('_claude_no_credit'):
            return None, "fallback"

        if provider == 'gemini':
            try:
                import requests as req_lib
                url = f"{config['url']}?key={api_key}"
                user_text = messages[-1]['content'] if messages else ""
                gemini_contents = []
                for m in messages[-4:]:
                    role = "user" if m['role'] == 'user' else "model"
                    gemini_contents.append({"role":role,"parts":[{"text":m['content'][:500]}]})
                resp = req_lib.post(url, json={
                    "contents": gemini_contents,
                    "systemInstruction": {"parts":[{"text":system_prompt[:4000]}]},
                    "generationConfig": {"maxOutputTokens": config['max_tokens'], "temperature": 0.3}
                }, timeout=30)
                if resp.status_code == 200:
                    text = resp.json().get('candidates',[{}])[0].get('content',{}).get('parts',[{}])[0].get('text','')
                    if text: return text, None
            except: pass
            return None, "fallback"

        elif provider in ('groq', 'openrouter', 'huggingface'):
            try:
                import requests as req_lib
            except:
                import urllib.request as _ur
                # Fallback to urllib if requests not available
                models_to_try = self.OR_MODELS if provider == 'openrouter' else [config['model']]
                for model_name in models_to_try:
                    try:
                        payload = json.dumps({"model":model_name,"max_tokens":config['max_tokens'],"messages":[{"role":"system","content":system_prompt}]+messages,"temperature":0.3})
                        headers = {'Content-Type':'application/json','Authorization':f'Bearer {api_key}'}
                        req = _ur.Request(config['url'], data=payload.encode('utf-8'), headers=headers, method='POST')
                        with _ur.urlopen(req, timeout=30) as resp:
                            text = json.loads(resp.read().decode()).get('choices',[{}])[0].get('message',{}).get('content','')
                            if text: return text, None
                    except: continue
                return None, "fallback"

            models_to_try = self.OR_MODELS if provider == 'openrouter' else [config['model']]
            for model_name in models_to_try:
                try:
                    headers = {'Authorization':f'Bearer {api_key}'}
                    if provider == 'openrouter':
                        headers['HTTP-Referer'] = 'https://hr-analytics-risal.streamlit.app'
                    resp = req_lib.post(config['url'],
                        json={"model":model_name,"max_tokens":config['max_tokens'],
                            "messages":[{"role":"system","content":system_prompt}]+messages,
                            "temperature":0.3},
                        headers=headers, timeout=30)
                    if resp.status_code == 200:
                        text = resp.json().get('choices',[{}])[0].get('message',{}).get('content','')
                        if text: return text, None
                except: continue
            return None, "fallback"

        else:  # claude
            try:
                import requests as req_lib
            except:
                return None, "fallback"
            for use_web in [True, False]:
                payload_dict = {
                    "model": config['model'], "max_tokens": config['max_tokens'],
                    "system": system_prompt, "messages": messages,
                }
                if use_web:
                    payload_dict["tools"] = [{"type":"web_search_20250305","name":"web_search"}]
                try:
                    resp = req_lib.post(config['url'], json=payload_dict,
                        headers={'x-api-key':api_key,'anthropic-version':'2023-06-01'},
                        timeout=90)
                    if resp.status_code == 200:
                        text = "\n".join([b.get('text','') for b in resp.json().get('content',[]) if b.get('type')=='text'])
                        if text: return text, None
                    elif resp.status_code == 400 and use_web:
                        body = resp.text[:200]
                        if 'credit' in body.lower() or 'balance' in body.lower():
                            st.session_state['_claude_no_credit'] = True
                            return None, "fallback"
                        continue
                    else:
                        if 'credit' in resp.text.lower():
                            st.session_state['_claude_no_credit'] = True
                        return None, "fallback"
                except:
                    if use_web: continue
                    return None, "fallback"
            return None, "fallback"

    def get_stats(self):
        """Return orchestrator statistics."""
        return {"calls": self._call_count, "cached": len(self._cache), "est_tokens": self._token_estimate}


# ==================== 2. KNOWLEDGE ENGINE (Semantic RAG) ====================
class KnowledgeEngine:
    """Enterprise RAG with per-advisor vector indices and full semantic search."""

    def __init__(self):
        self._vectorizer = None
        self._vectors = None
        self._cosine = None
        self._chunks = []
        self._loaded = False
        self._legal_vect = None
        self._legal_vectors = None
        self._legal_chunks = []
        self._hr_vect = None
        self._hr_vectors = None
        self._hr_chunks = []

    def _ensure_vectorizer(self):
        if self._vectorizer is None:
            try:
                from sklearn.feature_extraction.text import TfidfVectorizer
                from sklearn.metrics.pairwise import cosine_similarity
                self._vectorizer = TfidfVectorizer(max_features=5000, ngram_range=(1,2))
                self._cosine = cosine_similarity
            except ImportError:
                self._vectorizer = "unavailable"

    def load_from_db(self):
        if self._loaded: return
        self._loaded = True
        try:
            conn = get_conn(); c = conn.cursor()
            c.execute(f"SELECT value FROM app_config WHERE key = {_ph()}", ("rag_chunks",))
            row = c.fetchone(); conn.close()
            if row:
                self._chunks = json.loads(row[0])
                self._legal_chunks = [ch for ch in self._chunks if ch.get('advisor_type','') == 'legal' or ch.get('type','') == 'legal']
                self._hr_chunks = [ch for ch in self._chunks if ch.get('advisor_type','') == 'hr' or ch.get('type','') == 'hr']
                self._build_all_indices()
        except: pass

    def _build_all_indices(self):
        """Build separate persistent vector indices for legal and HR."""
        self._ensure_vectorizer()
        if self._vectorizer == "unavailable": return
        try:
            from sklearn.feature_extraction.text import TfidfVectorizer
            from sklearn.metrics.pairwise import cosine_similarity
            self._cosine = cosine_similarity
            legal_texts = [ch.get('text','') for ch in self._legal_chunks if ch.get('text','').strip()]
            if legal_texts:
                self._legal_vect = TfidfVectorizer(max_features=5000, ngram_range=(1,2))
                self._legal_vectors = self._legal_vect.fit_transform(legal_texts)
            hr_texts = [ch.get('text','') for ch in self._hr_chunks if ch.get('text','').strip()]
            if hr_texts:
                self._hr_vect = TfidfVectorizer(max_features=5000, ngram_range=(1,2))
                self._hr_vectors = self._hr_vect.fit_transform(hr_texts)
            all_texts = [ch.get('text','') for ch in self._chunks if ch.get('text','').strip()]
            if all_texts:
                self._vectors = self._vectorizer.fit_transform(all_texts)
        except: pass

    def _build_index(self):
        self._build_all_indices()

    def search(self, query, top_k=5, advisor_type=None):
        """Full vector search using pre-built per-advisor index."""
        self.load_from_db()
        if not self._chunks: return ""

        # Select pre-built index per advisor
        if advisor_type == "legal" and self._legal_vectors is not None:
            chunks = self._legal_chunks; vectors = self._legal_vectors; vect = self._legal_vect
        elif advisor_type == "hr" and self._hr_vectors is not None:
            chunks = self._hr_chunks; vectors = self._hr_vectors; vect = self._hr_vect
        elif self._vectors is not None:
            chunks = self._chunks; vectors = self._vectors; vect = self._vectorizer
        else:
            chunks = self._chunks; vectors = None; vect = None

        if not chunks: return ""

        # Vector search (primary)
        if vectors is not None and vect is not None and self._cosine is not None:
            try:
                valid_chunks = [ch for ch in chunks if ch.get('text','').strip()]
                query_vec = vect.transform([query])
                scores = self._cosine(query_vec, vectors).flatten()
                top_indices = scores.argsort()[-top_k:][::-1]
                results = []
                for idx in top_indices:
                    if idx < len(valid_chunks) and scores[idx] > 0.02:
                        results.append(f"[{valid_chunks[idx].get('source','')}]\n{valid_chunks[idx]['text']}")
                if results:
                    return "\n\n---\n".join(results)
            except: pass

        # Keyword fallback (only if vector completely unavailable)
        query_words = set(query.lower().split())
        scored = []
        for ch in chunks:
            chunk_words = set(ch.get('text','').lower().split())
            match = len(query_words & chunk_words)
            if match > 0:
                scored.append((match / max(len(query_words),1), ch))
        scored.sort(key=lambda x: x[0], reverse=True)
        return "\n\n---\n".join([f"[{c[1].get('source','')}]\n{c[1]['text']}" for c in scored[:top_k]])

    def ingest(self, text, source, doc_type="legal", version=None, advisor_type=None):
        """Ingest document with chunking, versioning, and advisor_type tagging."""
        if not advisor_type:
            advisor_type = doc_type
        chunks = self._chunk_text(text)
        ver = version or datetime.now().strftime("%Y%m%d_%H%M")
        try:
            conn = get_conn(); c = conn.cursor()
            c.execute(f"SELECT value FROM app_config WHERE key = {_ph()}", ("rag_chunks",))
            row = c.fetchone()
            existing = json.loads(row[0]) if row else []
            self._save_version(c, source, ver, len(chunks))
            existing = [ch for ch in existing if ch.get('source') != source]
            for i, chunk in enumerate(chunks):
                existing.append({
                    "text": chunk, "source": source, "type": doc_type,
                    "advisor_type": advisor_type,
                    "version": ver, "chunk_id": i,
                    "added": datetime.now().strftime("%Y-%m-%d"),
                    "hash": hashlib.md5(chunk.encode()).hexdigest()[:12]
                })
            _upsert_config(c, "rag_chunks", json.dumps(existing, ensure_ascii=False))
            conn.commit(); conn.close()
            self._chunks = existing
            self._build_index()
            return len(chunks)
        except: return 0

    def _chunk_text(self, text, chunk_size=500, overlap=50):
        """Split text into overlapping chunks."""
        chunks = []
        words = text.split()
        for i in range(0, len(words), chunk_size - overlap):
            chunk = " ".join(words[i:i + chunk_size])
            if chunk.strip(): chunks.append(chunk)
        return chunks

    def _save_version(self, cursor, source, version, n_chunks):
        """Track document versions for governance."""
        try:
            cursor.execute(f"SELECT value FROM app_config WHERE key = {_ph()}", ("rag_versions",))
            row = cursor.fetchone()
            versions = json.loads(row[0]) if row else []
            versions.append({"source":source,"version":version,"chunks":n_chunks,
                "date":datetime.now().strftime("%Y-%m-%d %H:%M"),
                "user":st.session_state.get('user_name','النظام')})
            if len(versions) > 200: versions = versions[-200:]
            _upsert_config(cursor, "rag_versions", json.dumps(versions, ensure_ascii=False))
        except: pass

    def get_versions(self, source=None):
        """Get version history for governance."""
        try:
            conn = get_conn(); c = conn.cursor()
            c.execute(f"SELECT value FROM app_config WHERE key = {_ph()}", ("rag_versions",))
            row = c.fetchone(); conn.close()
            versions = json.loads(row[0]) if row else []
            if source: versions = [v for v in versions if v.get('source') == source]
            return versions
        except: return []

    def get_stats(self):
        """Knowledge base statistics."""
        self.load_from_db()
        sources = {}
        for ch in self._chunks:
            src = ch.get('source','unknown')
            sources[src] = sources.get(src, 0) + 1
        return {
            "total_chunks": len(self._chunks),
            "sources": len(sources), "source_detail": sources,
            "total_words": sum(len(ch.get('text','').split()) for ch in self._chunks),
            "vector_index": self._vectors is not None,
            "versions": len(self.get_versions()),
        }


# ==================== 3. LEARNING SYSTEM ====================
class LearningSystem:
    """Continuous learning from user interactions with dataset building capability.
    Tracks Q&A, feedback, builds fine-tuning datasets, analyzes patterns."""

    def __init__(self):
        self._history = None
        self._loaded = False

    def _load(self):
        """Load interaction history from DB."""
        if self._loaded: return
        self._loaded = True
        try:
            conn = get_conn(); c = conn.cursor()
            c.execute(f"SELECT value FROM app_config WHERE key = {_ph()}", ("rag_qa_history",))
            row = c.fetchone(); conn.close()
            self._history = json.loads(row[0]) if row else []
        except:
            self._history = []

    def save_interaction(self, question, answer, model_type, feedback=None):
        """Save Q&A pair for learning."""
        self._load()
        self._history.append({
            "q": question[:500], "a": answer[:1500], "model": model_type,
            "feedback": feedback, "date": datetime.now().strftime("%Y-%m-%d %H:%M"),
            "user": st.session_state.get('user_name',''),
            "provider": st.session_state.get('ai_provider','auto'),
            "q_hash": hashlib.md5(question.encode()).hexdigest()[:10]
        })
        if len(self._history) > 500: self._history = self._history[-500:]
        self._save_to_db()

    def set_feedback(self, index, feedback):
        """Update feedback for a specific interaction."""
        self._load()
        if 0 <= index < len(self._history):
            self._history[index]['feedback'] = feedback
            self._save_to_db()

    def _save_to_db(self):
        """Persist history to database."""
        try:
            conn = get_conn(); c = conn.cursor()
            _upsert_config(c, "rag_qa_history", json.dumps(self._history, ensure_ascii=False))
            conn.commit(); conn.close()
        except: pass

    def get_relevant_history(self, query, model_type, top_k=3):
        """Get relevant past Q&A for context injection (learning from past)."""
        self._load()
        relevant = [h for h in self._history if h.get('model') == model_type and h.get('feedback') != 'bad']
        if not relevant: return ""

        query_words = set(query.lower().split())
        scored = []
        for h in relevant:
            q_words = set(h.get('q','').lower().split())
            match = len(query_words & q_words)
            if match > 1: scored.append((match, h))
        scored.sort(key=lambda x: x[0], reverse=True)
        top = scored[:top_k]
        if not top: return ""
        return "\n\n".join([f"Q: {h[1]['q']}\nA: {h[1]['a'][:300]}" for h in top])

    def build_finetune_dataset(self, min_feedback='good'):
        """Build fine-tuning dataset from positively-rated interactions."""
        self._load()
        dataset = []
        for h in self._history:
            if h.get('feedback') == min_feedback:
                dataset.append({
                    "messages": [
                        {"role": "user", "content": h['q']},
                        {"role": "assistant", "content": h['a']}
                    ],
                    "model_type": h.get('model','general')
                })
        return dataset

    def get_analytics(self):
        """Comprehensive learning analytics."""
        self._load()
        if not self._history: return {}

        total = len(self._history)
        by_model = {}; by_feedback = {}; by_date = {}; by_user = {}
        for h in self._history:
            m = h.get('model','general'); by_model[m] = by_model.get(m,0)+1
            f = h.get('feedback','none'); by_feedback[f] = by_feedback.get(f,0)+1
            d = h.get('date','')[:10]; by_date[d] = by_date.get(d,0)+1
            u = h.get('user',''); by_user[u] = by_user.get(u,0)+1

        good = by_feedback.get('good', 0)
        bad = by_feedback.get('bad', 0)
        satisfaction = round(good / max(good+bad, 1) * 100, 1)

        # Topic analysis
        all_text = " ".join([h.get('q','') for h in self._history]).lower()
        topics = {}
        topic_kw = {
            "نهاية الخدمة":["مكافأة","نهاية","خدمة","eos"],
            "الفصل":["فصل","تعويض","77","إنهاء"],
            "الإجازات":["إجازة","سنوية","مرضية"],
            "الرواتب":["راتب","رواتب","بدل"],
            "التأمينات":["تأمين","gosi"],
            "التوظيف":["توظيف","استقطاب","recruitment"],
            "التدريب":["تدريب","تطوير","roi"],
            "الأداء":["أداء","تقييم","kpi"],
        }
        for topic, kws in topic_kw.items():
            count = sum(1 for kw in kws if kw in all_text)
            if count > 0: topics[topic] = count

        finetune_ready = len(self.build_finetune_dataset())

        return {
            "total": total, "by_model": by_model, "by_feedback": by_feedback,
            "by_date": by_date, "by_user": by_user, "satisfaction": satisfaction,
            "topics": topics, "finetune_ready": finetune_ready,
            "maturity": min(100, (total * 5 + good * 10 + finetune_ready * 20) // 10)
        }

    def get_history(self, limit=50):
        """Get recent history."""
        self._load()
        return list(reversed(self._history[-limit:]))


# =====================================================================
# Initialize Architecture Singletons
# =====================================================================
@st.cache_resource
def _init_orchestrator():
    return ModelOrchestrator()

@st.cache_resource
def _init_knowledge():
    return KnowledgeEngine()

@st.cache_resource
def _init_learning():
    return LearningSystem()

st.set_page_config(page_title="تحليلات HR | رسال الود", page_icon="📊", layout="wide", initial_sidebar_state="expanded")

# ===== DATABASE LAYER (Cloud + Local) =====
DB_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "hr_personality.db")

# Detect cloud database availability
def _is_cloud_db():
    """Check if cloud database (Supabase PostgreSQL) is configured"""
    try:
        if hasattr(st, 'secrets') and 'database' in st.secrets:
            return bool(st.secrets["database"].get("url") or st.secrets["database"].get("host"))
    except: pass
    return bool(os.environ.get("DATABASE_URL"))

@st.cache_resource(ttl=300)
def _get_cached_conn_str():
    """Cache DB connection string"""
    if _is_cloud_db():
        try:
            if hasattr(st, 'secrets') and 'database' in st.secrets:
                db = st.secrets["database"]
                return db.get("url", ""), "cloud"
        except: pass
    return DB_PATH, "local"

def get_conn():
    """Get database connection - PostgreSQL (cloud) or SQLite (local)"""
    conn_str, db_type = _get_cached_conn_str()
    if db_type == "cloud" and conn_str:
        import psycopg2
        try:
            return psycopg2.connect(conn_str, sslmode="require")
        except Exception as e:
            st.warning(f"⚠️ Cloud DB error, falling back to local: {e}")
    return sqlite3.connect(DB_PATH)

def _ph():
    """Get placeholder symbol: %s for PostgreSQL, ? for SQLite"""
    return "%s" if _is_cloud_db() else "?"

def _serial():
    """Get auto-increment syntax"""
    return "SERIAL PRIMARY KEY" if _is_cloud_db() else "INTEGER PRIMARY KEY AUTOINCREMENT"

def _upsert_config(c, key, value):
    """Insert or update config - handles both engines"""
    if _is_cloud_db():
        p = "%s"
        c.execute(f"INSERT INTO app_config (key, value) VALUES ({p}, {p}) ON CONFLICT (key) DO UPDATE SET value = {p}", (key, value, value))
    else:
        c.execute("INSERT OR REPLACE INTO app_config (key, value) VALUES (?, ?)", (key, value))

def init_db():
    """Initialize database tables - works with both PostgreSQL and SQLite"""
    conn = get_conn()
    c = conn.cursor()
    serial = _serial()
    c.execute(f'''CREATE TABLE IF NOT EXISTS test_results (
        id {serial},
        emp_name TEXT NOT NULL,
        emp_dept TEXT,
        test_type TEXT NOT NULL,
        test_date TEXT NOT NULL,
        scores_json TEXT NOT NULL,
        mbti_type TEXT,
        dominant TEXT,
        secondary TEXT,
        is_mandatory INTEGER DEFAULT 0,
        assigned_by TEXT,
        created_at TEXT NOT NULL,
        created_by TEXT
    )''')
    c.execute(f'''CREATE TABLE IF NOT EXISTS test_assignments (
        id {serial},
        emp_name TEXT NOT NULL,
        emp_dept TEXT,
        test_name TEXT NOT NULL,
        deadline TEXT,
        status TEXT DEFAULT 'pending',
        assigned_by TEXT,
        is_mandatory INTEGER DEFAULT 1,
        created_at TEXT NOT NULL
    )''')
    c.execute(f'''CREATE TABLE IF NOT EXISTS email_log (
        id {serial},
        to_email TEXT, emp_name TEXT, tests TEXT,
        status TEXT, sent_at TEXT, sent_by TEXT
    )''')
    c.execute(f'''CREATE TABLE IF NOT EXISTS app_config (
        key TEXT PRIMARY KEY, value TEXT
    )''')
    c.execute(f'''CREATE TABLE IF NOT EXISTS users_store (
        username TEXT PRIMARY KEY,
        password_hash TEXT, role TEXT, name TEXT,
        email TEXT, dept TEXT, sections TEXT
    )''')
    conn.commit()
    conn.close()

def db_save_result(result, created_by=""):
    """Save a test result to database"""
    conn = get_conn()
    c = conn.cursor()
    p = _ph()
    c.execute(f'''INSERT INTO test_results
        (emp_name, emp_dept, test_type, test_date, scores_json, mbti_type, dominant, secondary, is_mandatory, assigned_by, created_at, created_by)
        VALUES ({p},{p},{p},{p},{p},{p},{p},{p},{p},{p},{p},{p})''',
        (result.get("الاسم",""), result.get("القسم",""), result.get("type",""),
         result.get("التاريخ",""), json.dumps(result.get("scores",{}), ensure_ascii=False),
         result.get("mbti_type",""), result.get("dominant",""), result.get("secondary",""),
         1 if result.get("إجباري") else 0, result.get("معيّن_بواسطة",""),
         datetime.now().strftime("%Y-%m-%d %H:%M:%S"), created_by))
    conn.commit()
    # Get last inserted id
    if _is_cloud_db():
        c.execute("SELECT lastval()")
        rid = c.fetchone()[0]
    else:
        rid = c.lastrowid
    conn.close()
    return rid

def db_load_results(test_type=None, emp_name=None):
    """Load test results from database"""
    conn = get_conn()
    c = conn.cursor()
    p = _ph()
    query = "SELECT * FROM test_results WHERE 1=1"
    params = []
    if test_type and test_type != "الكل":
        query += f" AND test_type = {p}"
        params.append(test_type)
    if emp_name:
        query += f" AND emp_name = {p}"
        params.append(emp_name)
    query += " ORDER BY created_at DESC"
    c.execute(query, params)
    rows = c.fetchall()
    cols = [d[0] for d in c.description]
    conn.close()
    results = []
    for row in rows:
        r = dict(zip(cols, row))
        r["scores"] = json.loads(r["scores_json"]) if r["scores_json"] else {}
        r["الاسم"] = r["emp_name"]
        r["القسم"] = r["emp_dept"]
        r["type"] = r["test_type"]
        r["التاريخ"] = r["test_date"]
        r["إجباري"] = bool(r["is_mandatory"])
        r["معيّن_بواسطة"] = r["assigned_by"] or ""
        results.append(r)
    return results

def db_delete_result(result_id):
    """Delete a test result (admin only)"""
    conn = get_conn()
    c = conn.cursor()
    c.execute(f"DELETE FROM test_results WHERE id = {_ph()}", (result_id,))
    conn.commit()
    conn.close()

def db_delete_all_results():
    """Delete all test results (admin only)"""
    conn = get_conn()
    c = conn.cursor()
    c.execute("DELETE FROM test_results")
    conn.commit()
    conn.close()

def db_save_assignment(assignment):
    """Save test assignment to database"""
    conn = get_conn()
    c = conn.cursor()
    p = _ph()
    c.execute(f'''INSERT INTO test_assignments (emp_name, emp_dept, test_name, deadline, status, assigned_by, is_mandatory, created_at)
        VALUES ({p},{p},{p},{p},{p},{p},{p},{p})''',
        (assignment.get("الموظف",""), assignment.get("القسم",""), assignment.get("الاختبار",""),
         assignment.get("الموعد النهائي",""), assignment.get("الحالة","لم يبدأ"),
         assignment.get("معيّن_بواسطة",""), 1 if assignment.get("إجباري") else 0,
         datetime.now().strftime("%Y-%m-%d %H:%M:%S")))
    conn.commit()
    conn.close()

def db_load_assignments(emp_name=None):
    """Load test assignments from database"""
    conn = get_conn()
    c = conn.cursor()
    p = _ph()
    if emp_name:
        c.execute(f"SELECT * FROM test_assignments WHERE emp_name = {p} ORDER BY created_at DESC", (emp_name,))
    else:
        c.execute("SELECT * FROM test_assignments ORDER BY created_at DESC")
    rows = c.fetchall()
    cols = [d[0] for d in c.description]
    conn.close()
    assignments = []
    for row in rows:
        r = dict(zip(cols, row))
        assignments.append({
            "id": r["id"], "الموظف": r["emp_name"], "القسم": r["emp_dept"],
            "الاختبار": r["test_name"], "الموعد النهائي": r["deadline"],
            "الحالة": r["status"], "معيّن_بواسطة": r["assigned_by"],
            "إجباري": bool(r["is_mandatory"])
        })
    return assignments

def db_delete_assignments():
    """Delete all assignments (admin only)"""
    conn = get_conn()
    c = conn.cursor()
    c.execute("DELETE FROM test_assignments")
    conn.commit()
    conn.close()

def db_count_results():
    """Get total count of saved results"""
    conn = get_conn()
    c = conn.cursor()
    c.execute("SELECT COUNT(*) FROM test_results")
    count = c.fetchone()[0]
    conn.close()
    return count

def db_save_users(users_db):
    """Save users to cloud database for persistence"""
    try:
        conn = get_conn()
        c = conn.cursor()
        p = _ph()
        for uname, udata in users_db.items():
            if _is_cloud_db():
                c.execute(f"""INSERT INTO users_store (username, password_hash, role, name, email, dept, sections)
                    VALUES ({p},{p},{p},{p},{p},{p},{p})
                    ON CONFLICT (username) DO UPDATE SET
                    password_hash=EXCLUDED.password_hash, role=EXCLUDED.role, name=EXCLUDED.name,
                    email=EXCLUDED.email, dept=EXCLUDED.dept, sections=EXCLUDED.sections""",
                    (uname, udata.get("password",""), udata.get("role",""), udata.get("name",""),
                     udata.get("email",""), udata.get("dept",""), udata.get("sections","")))
            else:
                c.execute(f"""INSERT OR REPLACE INTO users_store (username, password_hash, role, name, email, dept, sections)
                    VALUES ({p},{p},{p},{p},{p},{p},{p})""",
                    (uname, udata.get("password",""), udata.get("role",""), udata.get("name",""),
                     udata.get("email",""), udata.get("dept",""), udata.get("sections","")))
        conn.commit()
        conn.close()
    except: pass

def db_load_users():
    """Load users from database"""
    try:
        conn = get_conn()
        c = conn.cursor()
        c.execute("SELECT * FROM users_store")
        rows = c.fetchall()
        cols = [d[0] for d in c.description]
        conn.close()
        if rows:
            users = {}
            for row in rows:
                r = dict(zip(cols, row))
                users[r["username"]] = {
                    "password": r["password_hash"], "role": r["role"], "name": r["name"],
                    "email": r.get("email",""), "dept": r.get("dept",""), "sections": r.get("sections","all")
                }
            return users
    except: pass
    return None

def _get_arabic_font_path():
    """Find an Arabic-supporting TTF font on the system"""
    candidates = [
        "/usr/share/fonts/truetype/freefont/FreeSerif.ttf",
        "/usr/share/fonts/truetype/freefont/FreeSans.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/noto/NotoSansArabic-Regular.ttf",
    ]
    for p in candidates:
        if os.path.exists(p): return p
    return None

def _reshape_arabic(text):
    """Reshape Arabic text for proper PDF rendering"""
    try:
        import arabic_reshaper
        from bidi.algorithm import get_display
        reshaped = arabic_reshaper.reshape(str(text))
        return get_display(reshaped)
    except:
        return str(text)

def _create_pdf_with_arabic():
    """Create FPDF instance with Arabic font support"""
    from fpdf import FPDF
    pdf = FPDF()
    font_path = _get_arabic_font_path()
    if font_path:
        pdf.add_font("ArFont", "", font_path, uni=True)
        pdf.add_font("ArFont", "B", font_path.replace(".ttf", ".ttf"), uni=True)
        # Also try bold variant
        bold_path = font_path.replace("FreeSerif", "FreeSerifBold").replace("FreeSans", "FreeSansBold").replace("DejaVuSans", "DejaVuSans-Bold")
        if os.path.exists(bold_path):
            pdf.add_font("ArFont", "B", bold_path, uni=True)
    return pdf, "ArFont" if font_path else "Helvetica"

def _pdf_cell(pdf, font_name, w, h, text, **kwargs):
    """Write text to PDF cell with Arabic reshaping if needed"""
    txt = str(text)
    has_arabic = any('\u0600' <= c <= '\u06FF' or '\uFB50' <= c <= '\uFEFF' for c in txt)
    if has_arabic and font_name != "Helvetica":
        txt = _reshape_arabic(txt)
    pdf.cell(w, h, txt, **kwargs)

def _pdf_mcell(pdf, font_name, w, h, text, **kwargs):
    """Write multi_cell with Arabic reshaping"""
    txt = str(text)
    has_arabic = any('\u0600' <= c <= '\u06FF' or '\uFB50' <= c <= '\uFEFF' for c in txt)
    if has_arabic and font_name != "Helvetica":
        txt = _reshape_arabic(txt)
    pdf.multi_cell(w, h, txt, **kwargs)

def generate_employee_pdf(result):
    """Generate PDF report for a single employee test - with Arabic support"""
    try:
        pdf, fn = _create_pdf_with_arabic()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=20)

        # Header
        pdf.set_fill_color(15, 76, 92)
        pdf.rect(0, 0, 210, 28, 'F')
        pdf.set_text_color(255, 255, 255)
        pdf.set_font(fn, 'B', 14)
        _pdf_cell(pdf, fn, 0, 14, 'Personality Assessment Report', align='C', ln=True)
        pdf.set_font(fn, '', 10)
        _pdf_cell(pdf, fn, 0, 14, 'Risal Al-Wud IT', align='C', ln=True)

        pdf.ln(15)
        pdf.set_text_color(0, 0, 0)

        # Employee info
        pdf.set_fill_color(230, 240, 250)
        pdf.set_font(fn, 'B', 12)
        _pdf_cell(pdf, fn, 0, 10, _reshape_arabic('  معلومات الموظف / Employee Information') if fn != 'Helvetica' else '  Employee Information', fill=True, ln=True)
        pdf.set_font(fn, '', 10)
        pdf.ln(3)
        emp_name = result.get('الاسم','')
        _pdf_cell(pdf, fn, 0, 7, f"{_reshape_arabic('الاسم') if fn != 'Helvetica' else 'Name'}: {emp_name}", ln=True)
        _pdf_cell(pdf, fn, 0, 7, f"{_reshape_arabic('القسم') if fn != 'Helvetica' else 'Department'}: {result.get('القسم','')}", ln=True)
        _pdf_cell(pdf, fn, 0, 7, f"{_reshape_arabic('نوع الاختبار') if fn != 'Helvetica' else 'Test Type'}: {result.get('type','')}", ln=True)
        _pdf_cell(pdf, fn, 0, 7, f"{_reshape_arabic('التاريخ') if fn != 'Helvetica' else 'Date'}: {result.get('التاريخ','')}", ln=True)
        mandatory = _reshape_arabic("إجباري") if result.get('إجباري') and fn != 'Helvetica' else ("Mandatory" if result.get('إجباري') else "Voluntary")
        _pdf_cell(pdf, fn, 0, 7, f"{_reshape_arabic('الحالة') if fn != 'Helvetica' else 'Status'}: {mandatory}", ln=True)
        if result.get('معيّن_بواسطة'):
            _pdf_cell(pdf, fn, 0, 7, f"Assigned by: {result['معيّن_بواسطة']}", ln=True)
        pdf.ln(5)

        # Scores
        pdf.set_fill_color(46, 117, 182)
        pdf.set_text_color(255, 255, 255)
        pdf.set_font(fn, 'B', 12)
        _pdf_cell(pdf, fn, 0, 10, _reshape_arabic('  نتائج التقييم / Assessment Results') if fn != 'Helvetica' else '  Assessment Results', fill=True, ln=True)
        pdf.set_text_color(0, 0, 0)
        pdf.ln(3)

        # Scores table
        pdf.set_font(fn, 'B', 9)
        pdf.set_fill_color(230, 240, 250)
        _pdf_cell(pdf, fn, 60, 8, 'Scale', border=1, fill=True, align='C')
        _pdf_cell(pdf, fn, 30, 8, 'Score', border=1, fill=True, align='C')
        _pdf_cell(pdf, fn, 30, 8, 'Level', border=1, fill=True, align='C')
        _pdf_cell(pdf, fn, 70, 8, 'Description', border=1, fill=True, align='C')
        pdf.ln()
        pdf.set_font(fn, '', 8)

        scores = result.get("scores", {})
        test_type = result.get("type", "")

        for key, pct in scores.items():
            level = "High" if pct >= 70 else ("Medium" if pct >= 40 else "Low")
            name = key; desc = ""
            if test_type == "Big Five" and key in BIG5_TRAITS:
                info = BIG5_TRAITS[key]; name = info['en']; desc = info.get('high','')[:40] if pct >= 60 else info.get('low','')[:40]
            elif test_type == "Thomas PPA" and key in THOMAS_SCALES:
                info = THOMAS_SCALES[key]; name = info['en']; desc = info.get('high','')[:40] if pct >= 60 else info.get('low','')[:40]
            elif test_type == "Hogan HPI" and key in HOGAN_SCALES:
                info = HOGAN_SCALES[key]; name = info['en']; desc = info.get('impact','')[:40]
            elif test_type == "DISC" and key in DISC_STYLES:
                info = DISC_STYLES[key]; name = info['en']; desc = info.get('high','')[:40] if pct >= 60 else info.get('low','')[:40]

            _pdf_cell(pdf, fn, 60, 7, name[:30], border=1)
            _pdf_cell(pdf, fn, 30, 7, f"{pct}%", border=1, align='C')
            _pdf_cell(pdf, fn, 30, 7, level, border=1, align='C')
            _pdf_cell(pdf, fn, 70, 7, desc, border=1)
            pdf.ln()

        # MBTI type
        if result.get("mbti_type"):
            pdf.ln(5)
            t = result["mbti_type"]; ti = MBTI_TYPES.get(t, {})
            pdf.set_font(fn, 'B', 11)
            _pdf_cell(pdf, fn, 0, 8, f"MBTI Type: {t} - {ti.get('name','')}", ln=True)
            pdf.set_font(fn, '', 9)
            _pdf_cell(pdf, fn, 0, 7, f"Strengths: {ti.get('strengths','')}", ln=True)
            _pdf_cell(pdf, fn, 0, 7, f"Careers: {ti.get('careers','')}", ln=True)

        if result.get("dominant"):
            pdf.ln(5)
            pdf.set_font(fn, 'B', 11)
            _pdf_cell(pdf, fn, 0, 8, f"Dominant: {result['dominant']}  |  Secondary: {result.get('secondary','')}", ln=True)

        # Footer
        pdf.ln(10)
        pdf.set_draw_color(200, 200, 200)
        pdf.line(10, pdf.get_y(), 200, pdf.get_y())
        pdf.ln(3)
        pdf.set_font(fn, '', 7)
        pdf.set_text_color(128, 128, 128)
        _pdf_cell(pdf, fn, 0, 5, f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')} | HR Analytics - Risal Al-Wud IT", align='C', ln=True)
        _pdf_cell(pdf, fn, 0, 5, "CONFIDENTIAL - For authorized use only", align='C')

        return bytes(pdf.output())
    except Exception as e:
        return None

# Initialize database on startup
init_db()

# ===== STYLES =====
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Noto+Sans+Arabic:wght@300;400;500;600;700;800&display=swap');
*{font-family:'Noto Sans Arabic',sans-serif}
.main .block-container{padding-top:.8rem;max-width:1400px}
[data-testid="stSidebar"]{background:linear-gradient(180deg,#0F4C5C 0%,#1A1A2E 100%)}
[data-testid="stSidebar"] *{color:white !important}
[data-testid="stMetric"]{background:white;border-radius:12px;padding:14px 18px;box-shadow:0 1px 3px rgba(0,0,0,.06);border:1px solid #E2E8F0}
[data-testid="stMetric"] label{font-size:12px !important;color:#64748B !important}
[data-testid="stMetric"] [data-testid="stMetricValue"]{font-size:20px !important;font-weight:700 !important}
h1{color:#0F4C5C !important;font-weight:800 !important}
.hdr{background:linear-gradient(135deg,#0F4C5C,#1A1A2E);padding:20px 28px;border-radius:14px;margin-bottom:20px;color:white}
.hdr h1{color:white !important;margin:0;font-size:24px}
.hdr p{color:rgba(255,255,255,.7);margin:4px 0 0;font-size:13px}
.ibox{background:#EFF6FF;border-radius:10px;padding:12px 16px;border-right:4px solid #3B82F6;margin-bottom:8px;font-size:13px;line-height:1.7}
.ibox.warn{background:#FFF7ED;border-right-color:#F97316}
.ibox.ok{background:#F0FDF4;border-right-color:#22C55E}
.ibox.bad{background:#FEF2F2;border-right-color:#EF4444}
.kpi{background:linear-gradient(135deg,#0F4C5C,#1B4D5C);color:white;border-radius:12px;padding:16px;text-align:center;margin-bottom:10px}
.kpi h3{font-size:24px;margin:6px 0 2px;font-weight:800}
.kpi p{font-size:11px;opacity:.7;margin:0}
#MainMenu,footer{visibility:hidden}
</style>
""", unsafe_allow_html=True)

CL = {'p':'#0F4C5C','a':'#E36414','s':'#2D6A4F','d':'#9A031E','dept':px.colors.qualitative.Set2,'sal':px.colors.qualitative.Pastel}

def hdr(t,s=""): st.markdown(f'<div class="hdr"><h1>{t}</h1><p>{s}</p></div>',unsafe_allow_html=True)
def ibox(t,tp="info"):
    c={"info":"ibox","warning":"ibox warn","success":"ibox ok","danger":"ibox bad"}
    ic={"info":"💡","warning":"⚠️","success":"✅","danger":"🚨"}
    st.markdown(f'<div class="{c.get(tp,"ibox")}">{ic.get(tp,"💡")} {t}</div>',unsafe_allow_html=True)
def kpi(l,v): st.markdown(f'<div class="kpi"><p>{l}</p><h3>{v}</h3></div>',unsafe_allow_html=True)

# ===== LOCAL KNOWLEDGE ENGINE (No API needed) =====
LABOR_KB = {
    "مستحقات|تسوية|صرف|متى يجب|دفع|سداد|حقوقي المالية": "**تسوية المستحقات (المادة 88):**\n\nيلتزم صاحب العمل بتصفية جميع حقوق العامل خلال **أسبوع** من تاريخ انتهاء العلاقة. وإذا كان العامل هو من أنهى العقد فخلال **أسبوعين**.\n\n**المستحقات تشمل:**\n- الراتب حتى آخر يوم عمل\n- مكافأة نهاية الخدمة (المادة 84)\n- بدل الإجازات غير المستخدمة\n- أي بدلات مستحقة\n- شهادة الخبرة\n\n**عند التأخير:** يحق للعامل تقديم شكوى لمكتب العمل.",
    "تقاعد|معاش|pension|retirement|احتساب التقاعد|راتب تقاعدي|تقاعد مبكر": "**نظام التقاعد (التأمينات الاجتماعية):**\n\n**شروط استحقاق المعاش:**\n- بلوغ سن 60 سنة مع اشتراك لا يقل عن 120 شهراً (10 سنوات)\n- أو إكمال 300 شهر اشتراك (25 سنة) بغض النظر عن العمر (تقاعد مبكر)\n\n**حساب المعاش التقاعدي:**\n- معاش شهري = (متوسط الأجر في آخر سنتين × عدد أشهر الاشتراك) / 480\n- **مثال:** راتب 10,000 × 240 شهر / 480 = **5,000 ريال شهرياً**\n\n**الحد الأقصى:** 100% من متوسط الأجر\n**الحد الأدنى:** 1,984 ريال\n\n**تقاعد مبكر:** يمكن التقاعد قبل 60 سنة إذا أكمل 300 شهر اشتراك مع خصم 5% عن كل سنة قبل سن الـ 60.\n\n**معاش الأخطار المهنية:** 100% من الأجر في حالة العجز الكلي.",
    "ساند|استحقاق ساند|تعطل|بطالة|unemployment|sand": "**نظام ساند (التأمين ضد التعطل عن العمل):**\n\n**شروط الاستحقاق:**\n1. أن يكون سعودي الجنسية\n2. أن يكون عمره أقل من 60 سنة\n3. ألا يكون مفصولاً بسبب راجع إليه (المادة 80)\n4. أن يكون قادراً على العمل\n5. أن يكون مسجلاً في وزارة الموارد البشرية باحثاً عن عمل\n6. ألا يكون له دخل من عمل خاص\n7. مدة اشتراك لا تقل عن 12 شهراً في 36 شهراً الأخيرة\n8. أن يتقدم للتأمينات خلال 90 يوماً من ترك العمل\n\n**قيمة التعويض:**\n- أول 3 أشهر: **60%** من متوسط الأجر\n- الأشهر 4 إلى 12: **50%** من متوسط الأجر\n- **الحد الأقصى:** 9,000 ريال شهرياً\n- **الحد الأدنى:** 2,000 ريال شهرياً\n- **المدة القصوى:** 12 شهراً\n\n**التقديم:** عبر موقع التأمينات الاجتماعية gosi.gov.sa أو تطبيق تأميناتي.",
    "تأمينات|تأمين اجتماعي|gosi|اشتراك|خصم|نسبة الاشتراك|أخطار مهنية": "**نظام التأمينات الاجتماعية (GOSI):**\n\n**فرع المعاشات (سعوديون فقط):**\n- الموظف: 9.75%\n- صاحب العمل: 9.75%\n\n**فرع الأخطار المهنية (الجميع):**\n- صاحب العمل: 2%\n\n**ساند - التعطل عن العمل (سعوديون):**\n- الموظف: 0.75%\n- صاحب العمل: 0.75%\n\n**الإجمالي:**\n- **خصم الموظف السعودي: 10.5%**\n- **تحمّل الشركة للسعودي: 12.5%**\n- **تحمّل الشركة لغير السعودي: 2% فقط**\n\n**الأجر الخاضع:** الراتب الأساسي + بدل السكن\n**الحد الأقصى للأجر الخاضع:** 45,000 ريال\n\n**ساند (التعطل):** يصرف 60% من الأجر لمدة 3 أشهر ثم 50% لـ 9 أشهر (بحد أقصى 9,000 ريال).",
    "تأمين طبي|تأمين صحي|ضمان صحي|medical|insurance|علاج|مستشفى": "**التأمين الطبي (نظام الضمان الصحي التعاوني CCHI):**\n\n**الإلزامية:** على جميع منشآت القطاع الخاص\n**التغطية تشمل:** الموظف + من يعولهم (الزوجة والأبناء)\n**التكلفة:** يتحملها صاحب العمل كاملة\n\n**التغطية الأساسية:**\n- الكشف الطبي والعلاج\n- التحاليل والأشعة\n- الأدوية\n- العمليات الجراحية\n- الولادة والأمومة\n- طوارئ\n- الأسنان (حسب الوثيقة)\n- النظارات (حسب الوثيقة)\n\n**لا يجوز:** خصم تكلفة التأمين من راتب الموظف.",
    "إجازة|إجازات|سنوية|مرضية|annual|leave|عطلة|راحة": "**الإجازات في نظام العمل:**\n\n**السنوية (المادة 109):**\n- أول 5 سنوات: 21 يوماً\n- بعد 5 سنوات: 30 يوماً\n- لا يجوز التنازل عنها\n\n**المرضية (المادة 113):**\n- 30 يوماً بأجر كامل\n- 60 يوماً بثلاثة أرباع الأجر\n- 30 يوماً بدون أجر\n\n**إجازات أخرى:**\n- وفاة: 5 أيام | زواج: 5 أيام | مولود: 3 أيام\n- حج: 10-15 يوماً (لمرة واحدة)\n- عيد الفطر: 4 أيام | عيد الأضحى: 4 أيام\n- اليوم الوطني + يوم التأسيس",
    "ساعات|دوام|عمل اضافي|أوفرتايم|overtime|وقت العمل": "**ساعات العمل (المادة 98):**\n\n- العادية: 8 ساعات/يوم أو 48 ساعة/أسبوع\n- رمضان: 6 ساعات/يوم أو 36 ساعة/أسبوع\n- العمل الإضافي (المادة 107): أجر الساعة + 50%\n- لا يزيد الإضافي عن 720 ساعة سنوياً\n- راحة: 30 دقيقة بعد كل 5 ساعات متواصلة\n- يوم راحة أسبوعي: الجمعة عادة",
    "استقالة|أستقيل|resignation|ترك العمل|إشعار": "**الاستقالة:**\n\n**الإشعار (المادة 75):** 60 يوماً (بأجر شهري) أو 30 يوماً (بغيره)\n\n**المكافأة عند الاستقالة (المادة 85):**\n- أقل من سنتين: لا يستحق\n- 2-5 سنوات: ثلث المكافأة\n- 5-10 سنوات: ثلثا المكافأة\n- أكثر من 10: كاملة\n\n**ترك العمل بدون إشعار (المادة 81):** عند عدم دفع الأجر أو الاعتداء أو الغش في العقد.",
    "فصل|إنهاء|طرد|termination|إقالة|فسخ|عزل": "**إنهاء العقد:**\n\n**فصل مشروع بدون مكافأة (المادة 80):** اعتداء، إخلال جوهري، غياب 20 يوم متقطعة أو 10 متصلة، تزوير، إفشاء أسرار\n\n**فصل غير مشروع (المادة 77):**\n- غير محدد: 15 يوم/سنة\n- محدد: أجر المدة الباقية\n- الحد الأدنى: شهرين\n+ مكافأة نهاية الخدمة + بدل إجازات",
    "مكافأة|نهاية الخدمة|end of service|eos|مكافاة": "**مكافأة نهاية الخدمة (المادة 84):**\n\n- أول 5 سنوات: نصف راتب عن كل سنة\n- بعد 5 سنوات: راتب كامل عن كل سنة\n- أجزاء السنة تُحسب نسبياً\n\n**عند الاستقالة (المادة 85):**\n- أقل من سنتين: صفر\n- 2-5 سنوات: ثلث\n- 5-10: ثلثان\n- أكثر من 10: كاملة\n\n**التسوية:** خلال أسبوع (المادة 88).",
    "نطاقات|سعودة|توطين|nitaqat|نسبة السعودة": "**نطاقات:** أحمر (عقوبات) → أصفر → أخضر (منخفض/متوسط/مرتفع) → بلاتيني (أفضل المزايا)\n\nالنسب تختلف حسب حجم المنشأة ونشاطها والمنطقة.\n\n**مزايا البلاتيني:** تأشيرات فورية + نقل خدمات بدون موافقة + تغيير مهن.",
    "عقد|عقود|contract|أنواع العقود|محدد|غير محدد": "**أنواع العقود:** محدد المدة | غير محدد (سعوديين) | مهمة | موسمي | تدريب | دوام جزئي\n\nالمادة 55: يتحول المحدد لغير محدد بعد 3 تجديدات أو 4 سنوات.\nالمادة 50: يجب أن يكون مكتوباً من نسختين.",
    "فترة التجربة|تجربة|probation": "**فترة التجربة (المادة 53):**\n- 90 يوماً كحد أقصى\n- تمديد إلى 180 يوماً بموافقة مكتوبة\n- لا تشمل الأعياد والإجازات المرضية\n- لأي طرف الإنهاء بدون تعويض\n- لا تتكرر لدى نفس صاحب العمل",
    "حقوق المرأة|المرأة|امرأة|حامل|وضع|رضاعة|أمومة|maternity": "**حقوق المرأة العاملة:**\n- إجازة وضع: 10 أسابيع (المادة 151)\n- رضاعة: ساعة يومياً لـ 24 شهراً\n- حماية من الفصل أثناء الحمل والوضع\n- إجازة وفاة الزوج: 4 أشهر و10 أيام\n- المساواة في الأجر\n- بيئة عمل آمنة",
    "راتب|رواتب|أجر|أجور|بدل|salary|wage|حد أدنى": "**الأجور:**\n- يُدفع بالريال السعودي (المادة 89)\n- يُدفع خلال 7 أيام من استحقاقه\n- بدل السكن: عادة 25% من الأساسي\n- بدل النقل: عادة 10% من الأساسي\n- لا يجوز الخفض بدون موافقة مكتوبة\n- الحد الأدنى للسعودي في القطاع الخاص: 4,000 ريال (لاحتساب نطاقات)",
    "حقوق العامل|حقوق الموظف|حق العامل|حقوقي": "**حقوق العامل في نظام العمل السعودي:**\n\n1. الأجر العادل والمنتظم في موعده\n2. بيئة عمل آمنة وصحية\n3. إجازة سنوية مدفوعة (21 أو 30 يوماً)\n4. تأمين طبي شامل له ولأسرته\n5. مكافأة نهاية الخدمة\n6. شهادة خبرة عند المغادرة\n7. عدم التمييز والمساواة\n8. ساعات عمل لا تتجاوز 8 يومياً\n9. راحة أسبوعية\n10. السلامة المهنية ومعدات الحماية",
    "واجبات صاحب العمل|التزامات صاحب العمل|مسؤوليات الشركة|واجبات المنشأة|واجبات الشركة": "**واجبات صاحب العمل في نظام العمل السعودي:**\n\n1. **دفع الأجور** في موعدها بدون تأخير (المادة 90)\n2. **عقد مكتوب** من نسختين (المادة 51)\n3. **التأمين الطبي** للموظف وأسرته\n4. **التأمينات الاجتماعية** تسجيل وسداد الاشتراكات\n5. **بيئة آمنة** وصحية (المادة 121)\n6. **منح الإجازات** المستحقة\n7. **شهادة الخبرة** عند انتهاء العلاقة (المادة 64)\n8. **مكافأة نهاية الخدمة** خلال أسبوع (المادة 88)\n9. **السلامة المهنية** توفير معدات الوقاية\n10. **عدم التمييز** بين العاملين\n11. **التدريب** والتأهيل للسعوديين\n12. **السكن أو بدله** توفير سكن أو بدل",
    "واجبات العامل|التزامات الموظف|واجبات الموظف": "**واجبات العامل في نظام العمل السعودي:**\n\n1. أداء العمل بإتقان وعناية\n2. اتباع أوامر صاحب العمل المشروعة\n3. التعامل بأدب واحترام\n4. عدم إفشاء أسرار العمل\n5. المحافظة على أدوات ومعدات العمل\n6. الالتزام بمواعيد العمل\n7. تقديم المساعدة في حالات الكوارث\n8. عدم منافسة صاحب العمل خلال العقد\n9. إبلاغ أي تغيير في البيانات الشخصية\n10. الخضوع للفحوصات الطبية المطلوبة",
    "شكوى|نزاع|خلاف|محكمة عمالية|مكتب العمل|تظلم": "**تسوية النزاعات العمالية:**\n\n1. **الود:** محاولة حل ودي بين الطرفين\n2. **مكتب العمل:** تقديم شكوى خلال 12 شهراً من المخالفة\n3. **المحكمة العمالية:** خلال 12 شهراً من رفض التسوية الودية\n\n**المدة:** تنظر الدعوى خلال أسابيع\n**مجاناً:** لا رسوم على الدعاوى العمالية\n\n**منصة ودي:** تسوية إلكترونية عبر وزارة الموارد البشرية.",
    "نقل كفالة|نقل خدمات|تحويل|كفيل|sponsor|كفالتي|أنقل|تنقل وظيفي|خروج نهائي": "**نقل الخدمات (مبادرة تحسين العلاقة التعاقدية):**\n\nمبادرة تحسين العلاقة التعاقدية تتيح:\n- **التنقل الوظيفي:** الانتقال لصاحب عمل آخر بعد إكمال السنة الأولى أو انتهاء العقد\n- **الخروج والعودة:** بدون موافقة صاحب العمل\n- **تأشيرة الخروج النهائي:** بدون موافقة صاحب العمل\n\n**الشرط:** إشعار صاحب العمل الحالي قبل 90 يوماً.\n\n**الخطوات:**\n1. الدخول على منصة قوى (qiwa.sa)\n2. اختيار خدمة التنقل الوظيفي\n3. تقديم طلب الانتقال\n4. موافقة صاحب العمل الجديد\n5. إشعار صاحب العمل الحالي",
    "أداء|تقييم|performance|kpi|أهداف|okr": "**تقييم الأداء في نظام العمل السعودي:**\n\n**المادة 80 (الفصل بسبب الأداء):** يحق لصاحب العمل فسخ العقد إذا لم يؤدِّ العامل التزاماته الجوهرية المترتبة على عقد العمل.\n\n**المادة 75 (الإشعار):** إذا أراد صاحب العمل إنهاء العقد لضعف الأداء في العقد غير المحدد: إشعار مسبق لا يقل عن 60 يوماً.\n\n**المادة 77 (التعويض):** إذا تم الفسخ بدون سبب مشروع يستحق العامل تعويضاً.\n\n**اللائحة التنفيذية:** يجب توثيق ضعف الأداء كتابياً وإعطاء العامل فرصة للتحسين قبل اتخاذ إجراء.",
    "تدريب|تطوير|training|أطور|تأهيل|دورة|دورات": "**التدريب والتأهيل في نظام العمل السعودي (الباب الرابع: المواد 42-48):**\n\n**المادة 42:** يجب على كل صاحب عمل إعداد عماله السعوديين وتأهيلهم مهنياً لإحلالهم في الوظائف.\n\n**المادة 43:** على صاحب العمل تدريب ما لا يقل عن 12% من مجموع عماله السعوديين سنوياً.\n\n**المادة 44:** تحدد الوزارة المهن والأعمال التي يتم التدريب عليها والبرامج اللازمة.\n\n**المادة 45:** يجب أن يتضمن عقد التأهيل والتدريب: المهنة المتدرب عليها، مدة التدريب، مراحله المتتابعة، والمكافأة المستحقة في كل مرحلة.\n\n**المادة 46:** لصاحب العمل إلزام المتدرب بالعمل لديه مدة مماثلة لمدة التدريب بعد انتهائه، وإذا رفض المتدرب يحق لصاحب العمل مطالبته بتكاليف التدريب.\n\n**المادة 47:** لصاحب العمل إنهاء عقد التأهيل والتدريب إذا ثبت عدم قدرة المتدرب على إتمام التدريب بصورة مفيدة.\n\n**المادة 48:** يتحمل صاحب العمل نفقات التدريب والتأهيل.",
    "استقطاب|توظيف|recruitment|hiring|مقابلة|وظيفة شاغرة|أوظف|تعيين|موظف جديد": "**التوظيف في نظام العمل السعودي:**\n\n**المادة 26:** يجب على صاحب العمل الاستعانة بالعمال السعوديين وألا تقل نسبتهم عن 75% (قابلة للتعديل بقرار وزاري).\n\n**المادة 28:** يحظر على صاحب العمل توظيف عامل بدون رخصة عمل سارية.\n\n**المادة 33:** يجب تسجيل العامل في التأمينات الاجتماعية خلال 15 يوماً من مباشرته.\n\n**المادة 50:** عقد العمل يجب أن يكون مكتوباً من نسختين.\n\n**المادة 51:** يجب أن يتضمن العقد: اسم العامل وجنسيته، صاحب العمل، الأجر، نوع العمل، مكانه، تاريخ الالتحاق، المدة.\n\n**المادة 53:** فترة التجربة 90 يوماً كحد أقصى.\n\n**نطاقات:** الالتزام بنسب السعودة حسب نشاط وحجم المنشأة.",
    "خدمة|سنوات الخدمة|أقدمية|tenure": "**احتساب سنوات الخدمة:**\n- تبدأ من تاريخ المباشرة الفعلي\n- تشمل فترة التجربة\n- تشمل الإجازات المدفوعة\n- لا تشمل الإجازات بدون أجر (إلا بالاتفاق)\n\n**تؤثر على:**\n- مكافأة نهاية الخدمة\n- الإجازة السنوية (21 يوم أو 30 يوم)\n- المعاش التقاعدي",
    "بحار|بحاره|بحري|سفن|ملاحة|maritime|عمال البحر|مكافآت البحاره": "**أحكام العمل البحري (الباب العاشر من نظام العمل):**\n\n**عقد العمل البحري (المادة 164):**\n- يجب أن يكون مكتوباً ومصدقاً من الجهة المختصة\n- يتضمن: اسم السفينة، تاريخ الالتحاق، الأجر، المدة\n\n**الأجور والمكافآت:**\n- يستحق البحار أجره كاملاً طوال مدة العقد\n- مكافأة نهاية الخدمة تُحسب وفق المادة 84 كبقية العمال\n- إذا غرقت السفينة: يستحق أجر شهرين إضافيين (المادة 177)\n- أجر إضافي عن العمل وقت الأخطار والإنقاذ\n\n**الإجازات:**\n- إجازة سنوية بنسبة يوم عن كل 15 يوم عمل بحري\n- إجازة شاطئية بعد كل رحلة طويلة\n\n**التعويضات:**\n- إصابة العمل: علاج كامل + أجر\n- الوفاة: تعويض الورثة\n- فقدان الأمتعة بسبب غرق السفينة: تعويض بحد أقصى أجر شهر\n\n**إنهاء العقد:**\n- لا يجوز فصل البحار أثناء الرحلة\n- يُعاد البحار إلى ميناء الشحن على نفقة صاحب العمل",
    "عمالة منزلية|خادم|سائق|عامل منزلي|domestic": "**نظام العمالة المنزلية:**\n\n- ينظمه نظام خاص (لائحة العمالة المنزلية) وليس نظام العمل\n- فترة التجربة: 90 يوماً\n- ساعات العمل: لا تتجاوز 15 ساعة يومياً مع فترات راحة\n- يوم راحة أسبوعي\n- إجازة شهرية مدفوعة بعد سنتين\n- التأمين الطبي إلزامي\n- مكافأة نهاية خدمة: راتب شهر عن كل 4 سنوات\n\n**منصة مساند:** لاستقدام وإدارة العمالة المنزلية",
    "إصابة عمل|إصابة مهنية|حادث عمل|سلامة|occupational": "**إصابات العمل (المادة 133-141):**\n\n- صاحب العمل يتحمل علاج المصاب بالكامل\n- أجر كامل خلال فترة العلاج (حتى سنة)\n- **عجز كلي:** 100% من الأجر كمعاش\n- **عجز جزئي:** نسبة من الأجر حسب نسبة العجز\n- **وفاة:** تعويض الورثة بأجر 3 سنوات (بحد أدنى 54,000 ريال)\n\n**واجبات صاحب العمل:** توفير وسائل السلامة + التدريب + التأمين",
}

# ===== HR KNOWLEDGE BASE (Frameworks & Best Practices ONLY - NO legal articles) =====
HR_KB = {
    "تدريب|تطوير|training|أطور|تأهيل|دورة|دورات": "**التدريب والتطوير** *(APTD - Instructional Design | PHRi - L&D)*\n\n**نموذج ADDIE** *(APTD):*\n1. Analysis: تحليل الاحتياجات التدريبية\n2. Design: تصميم البرنامج والأهداف\n3. Development: تطوير المحتوى والمواد\n4. Implementation: تنفيذ البرنامج\n5. Evaluation: تقييم الفعالية\n\n**تقييم التدريب Kirkpatrick** *(APTD/PHRi):*\n- المستوى 1: رد الفعل (رضا المتدربين)\n- المستوى 2: التعلم (المعرفة المكتسبة)\n- المستوى 3: السلوك (التطبيق في العمل)\n- المستوى 4: النتائج (الأثر على الأعمال)\n\n**70-20-10** *(CIPD L5):* 70% خبرة عملية + 20% تعلم اجتماعي + 10% تدريب رسمي",
    "استقطاب|توظيف|recruitment|hiring|مقابلة|أوظف|تعيين": "**الاستقطاب والتوظيف** *(PHRi - Talent Acquisition | SHRM - People)*\n\n**العملية** *(PHRi):*\n1. تحليل الاحتياج (Workforce Planning)\n2. الوصف الوظيفي (Job Description)\n3. EVP (عرض قيمة صاحب العمل)\n4. المصادر: LinkedIn, Referrals, Job Boards\n5. فرز السير الذاتية (ATS)\n6. مقابلات منظمة (Structured Interviews)\n7. تقييم الكفاءات (Competency Assessment)\n8. العرض والتهيئة (Offer & Onboarding)\n\n**مؤشرات** *(SPHRi):* Time-to-Hire | Cost-per-Hire | Quality of Hire",
    "أداء|تقييم|performance|kpi|أهداف|okr": "**إدارة الأداء** *(SPHRi - Strategic HR | CIPD L7 - Performance)*\n\n**أهداف SMART** *(PHRi):* محددة + قابلة للقياس + قابلة للتحقيق + ذات صلة + محددة زمنياً\n\n**9-Box Grid** *(SPHRi):* مصفوفة الأداء × الإمكانات لتصنيف المواهب\n\n**OKRs vs KPIs** *(SHRM):*\n- KPIs: مؤشرات ثابتة تقيس الأداء الحالي\n- OKRs: أهداف طموحة ربع سنوية (70% إنجاز = ممتاز)\n\n**التغذية الراجعة** *(CIPD L5):* مستمرة + بنّاءة + 360 درجة",
    "رواتب|تعويضات|compensation|أجور|هيكل رواتب|بدلات": "**التعويضات والمزايا** *(PHRi - Compensation | SPHRi - Total Rewards)*\n\n**Total Rewards** *(SHRM):*\n1. الراتب الأساسي\n2. المزايا والبدلات\n3. التوازن بين العمل والحياة\n4. التقدير والاعتراف\n5. التطوير المهني\n\n**هيكل الرواتب** *(CIPD L7 - Reward):*\n- Job Evaluation → Job Grading → Salary Bands (Min-Mid-Max)\n- Compa-Ratio = الراتب الفعلي / وسط النطاق\n- Salary Survey: P25, P50, P75",
    "دوران|turnover|استبقاء|retention|احتفاظ": "**إدارة الدوران والاحتفاظ** *(PHRi - HR Analytics | SPHRi - Workforce Planning)*\n\n**المعادلة:** Turnover = (المغادرين / متوسط العدد) × 100\n\n**المعايير** *(SHRM):* أقل من 10% ممتاز | 10-15% جيد | أكثر من 25% مشكلة\n\n**استراتيجيات الاحتفاظ** *(CIPD L7):*\n1. Total Rewards تنافسية\n2. مسار وظيفي واضح\n3. بيئة عمل محفزة\n4. تقدير ومكافآت\n5. تطوير القيادات\n6. استبيانات رضا + خطط تحسين",
    "تجربة الموظف|employee experience|ex|onboarding|تهيئة": "**تجربة الموظف** *(SHRM - EX | CIPD L7)*\n\n1. **الاستقطاب** *(PHRi):* عملية سلسة وشفافة\n2. **التهيئة** *(aPHRi):* خطة 30/60/90 يوم + Buddy System\n3. **التطوير** *(APTD):* تدريب مستمر + 70-20-10\n4. **الاحتفاظ** *(SPHRi):* Total Rewards + مرونة\n5. **الانتقال** *(PHRi):* مقابلة خروج + Alumni Network",
    "قيادة|leader|إدارة|management|تطوير قيادي": "**تطوير القيادات** *(CIPD L7 - Leadership | SPHRi)*\n\n**نموذج Ulrich** *(CIPD L7):*\n- شريك استراتيجي\n- وكيل تغيير\n- خبير إداري\n- بطل الموظفين\n\n**Kotter 8 Steps** *(CIPD L7):* إدارة التغيير المؤسسي\n\n**Coaching vs Mentoring** *(APTD):*\n- Coaching: قصير المدى، مهارة محددة\n- Mentoring: طويل المدى، تطوير شامل",
    "تحليلات|analytics|بيانات|data|مؤشرات": "**تحليلات الموارد البشرية** *(PHRi - HR Analytics | SPHRi)*\n\n**4 مستويات:**\n1. **وصفي:** ماذا حدث؟ (تقارير، dashboards)\n2. **تشخيصي:** لماذا حدث؟ (تحليل الأسباب)\n3. **تنبؤي:** ماذا سيحدث؟ (نماذج تنبؤية)\n4. **توصيفي:** ماذا نفعل؟ (توصيات)\n\n**مؤشرات أساسية:** Turnover Rate | Time-to-Hire | Cost-per-Hire | Employee NPS | Absence Rate",
    "تغيير|change|تطوير مؤسسي|od|تنظيمي": "**إدارة التغيير والتطوير المؤسسي** *(CIPD L7 - OD)*\n\n**Burke-Litwin Model:** تحليل أداء المنظمة وعوامل التغيير\n\n**Kotter 8 Steps:**\n1. خلق الإلحاح\n2. بناء تحالف قيادي\n3. تطوير الرؤية\n4. التواصل\n5. تمكين العمل\n6. تحقيق مكاسب سريعة\n7. البناء على المكاسب\n8. ترسيخ التغيير\n\n**ADKAR** *(SHRM):* Awareness → Desire → Knowledge → Ability → Reinforcement",
    "كفاءات|competency|مهارات|skills|جدارات": "**إدارة الكفاءات** *(SHRM - Competency | PHRi)*\n\n**Competency Framework:** تحديد الكفاءات المطلوبة لكل وظيفة\n\n**أنواع الكفاءات:**\n- Core: قيم المنظمة (للجميع)\n- Functional: مهارات الوظيفة\n- Leadership: كفاءات قيادية\n\n**التطبيق:** التوظيف + التطوير + التقييم + الترقية + التعاقب الوظيفي",
}


# =====================================================
# ADVISORY REASONING ENGINE v3
# Party-aware, source-grounded, hallucination-resistant
# For BOTH Legal Advisor and HR Advisor
# =====================================================

# --- Asking Party Patterns ---
PARTY_PATTERNS = {
    "employee": ["أنا موظف","أنا عامل","تم فصلي","فصلوني","أستقيل","حقوقي","مستحقاتي","راتبي","إجازتي","عقدي","كموظف","بصفتي عامل","حقي"],
    "employer": ["أنا صاحب عمل","كصاحب عمل","شركتي","منشأتي","موظف عندي","أريد فصل","كيف أفصل","هل يحق لي فصل","التزاماتي كصاحب"],
    "hr_officer": ["كمسؤول موارد","أنا HR","إدارة الموارد","قسم الموارد","سياسة الشركة","كيف أعد","كيف أصمم","نحتاج سياسة"],
    "manager": ["كمدير","فريقي","موظفيني","أحد موظفيني","تحت إدارتي","أريد تقييم"],
    "neutral": ["ما هو","ما هي","عرّف","اشرح","ما الفرق","ما تعريف","كيف يتم","ما حكم"],
}

# --- Intent Detection ---
INTENT_PATTERNS = {
    "explanation":     ["ما هو","ماهو","ما هي","ماهي","تعريف","مفهوم","اشرح","وضح","what is","define"],
    "case_application":["حالتي","تم فصلي","فصلوني","عندي مشكلة","في حالة","إذا كان","لو","أنا"],
    "calculation":     ["حساب","احسب","كم","مبلغ","نسبة","كيف أحسب","مقدار","قيمة"],
    "comparison":      ["الفرق","مقارنة","أفضل","أيهما","مقابل","vs","difference"],
    "process":         ["خطوات","مراحل","إجراءات","كيف أقدم","كيف أرفع","طريقة تقديم"],
    "methodology":     ["كيف أبني","كيف أصمم","كيف أطور","كيف أنشئ","كيف أعد"],
    "recommendation":  ["أنصح","أفضل طريقة","ماذا تنصح","ما رأيك","أقترح"],
    "diagnosis":       ["مشكلة","تحدي","ضعف","انخفاض","ارتفاع","تراجع","سبب","لماذا"],
    "policy":          ["سياسة","لائحة","نظام داخلي","ضوابط","إعداد سياسة"],
    "rights":          ["حقوق","حقوقي","يحق لي","هل يحق","واجبات","التزامات"],
    "kpi_design":      ["مؤشرات","KPI","قياس","كيف أقيس","تقييم فعالية"],
    "pros_cons":       ["مزايا","عيوب","إيجابيات","سلبيات"],
}

# --- Legal Topics with VERIFIED article facts ---
LEGAL_TOPICS = {
    "termination":     {"kw": ["فصل","إنهاء","طرد","إقالة","فسخ","فصلي","فُصلت","فصل تعسفي","إنهاء خدمة","فصل بدون سبب","فصل بسبب"], "ref": "م74-82", "label": "إنهاء العقد",
        "facts": "م74: حالات انتهاء العقد (اتفاق، انتهاء مدة، إرادة أحد الطرفين، قوة قاهرة، تقاعد). م75: إشعار 60 يوم (أجر شهري) أو 30 يوم. م77: تعويض فسخ غير مشروع: 15 يوم/سنة (غير محدد) أو المدة الباقية (محدد)، حد أدنى شهرين. م80: فسخ مشروع بدون تعويض (اعتداء، إخلال، غياب 30 متفرقة أو 15 متتالية). م81: حق العامل بالترك بدون إشعار (إخلال صاحب العمل، غش، اعتداء)."},
    "end_of_service":  {"kw": ["مكافأة","نهاية الخدمة","مستحقات","مستحقاتي","تسوية","EOS"], "ref": "م84-88", "label": "مكافأة نهاية الخدمة",
        "facts": "م84: نصف شهر/سنة (أول 5) + شهر/سنة (بعدها). تُحسب على آخر أجر. م85: عند الاستقالة: ثلث (2-5 سنوات)، ثلثان (5-10)، كاملة (10+). أقل من سنتين: لا شيء. م88: التسوية خلال أسبوع."},
    "resignation":     {"kw": ["استقالة","أستقيل","ترك العمل","تقديم استقالة"], "ref": "م75,م81,م84,م85,م88", "label": "الاستقالة",
        "facts": "م75: إشعار مسبق 60 يوم. م81: ترك بدون إشعار إذا أخل صاحب العمل. م84: حساب المكافأة الأساسية. م85: نسب الاستقالة. م88: تسوية خلال أسبوع."},
    "contracts":       {"kw": ["عقد","عقود","محدد المدة","غير محدد","تجديد","عقد عمل","تحويل العقد"], "ref": "م49-60", "label": "العقود",
        "facts": "م50: عقد مكتوب من نسختين. م51: يتضمن الاسم والجنسية والأجر ونوع العمل والمدة. م53: فترة تجربة 90 يوم قابلة للتمديد لـ 180. م55: يتحول لغير محدد بعد 3 تجديدات أو 4 سنوات."},
    "wages":           {"kw": ["راتب","أجر","رواتب","بدل","بدلات","خصم","تأخير رواتب","أجور"], "ref": "م89-97", "label": "الأجور",
        "facts": "م89: يُدفع بالريال. م90: التزام بالدفع في الموعد. م92: لا يجوز خصم أكثر من نصف الأجر. م94: حماية الأجر."},
    "working_hours":   {"kw": ["ساعات","ساعات العمل","دوام","عمل إضافي","أوفرتايم","overtime"], "ref": "م98-107", "label": "ساعات العمل",
        "facts": "م98: 8 ساعات/يوم أو 48/أسبوع. م99: رمضان 6 ساعات/36 أسبوع. م101: راحة نصف ساعة كل 5 ساعات. م107: إضافي 150%."},
    "leave":           {"kw": ["إجازة","إجازات","سنوية","مرضية","وفاة","زواج","حج","مولود","إجازة سنوية"], "ref": "م109-116", "label": "الإجازات",
        "facts": "م109: سنوية 21 يوم (أول 5 سنوات) ثم 30 يوم. م110: يجوز تأجيلها بموافقة. م112: وفاة 5 أيام، زواج 5، مولود 3. م113: مرضية 30 يوم كامل + 60 يوم 75% + 30 بدون. م115: حج 10-15 يوم لمرة واحدة."},
    "training":        {"kw": ["تدريب","تأهيل","تدريب مهني","دورة تدريبية","عقد تأهيل","التدريب","برنامج تدريب","نسبة التدريب","تكاليف التدريب"], "ref": "م42-48", "label": "التدريب والتأهيل",
        "facts": "م42: على صاحب العمل تأهيل عماله السعوديين. م43: تدريب 12% سنوياً. م44: الوزارة تحدد المهن. م45: عقد التأهيل يتضمن المهنة والمدة والمراحل والمكافأة. م46: إلزام المتدرب بالعمل أو رد التكاليف. م47: إنهاء التدريب إذا ثبت عدم القدرة. م48: النفقات على صاحب العمل."},
    "probation":       {"kw": ["تجربة","فترة التجربة","فترة الاختبار"], "ref": "م53-54", "label": "فترة التجربة",
        "facts": "م53: 90 يوم كحد أقصى، تمديد لـ 180 بموافقة كتابية. لا تدخل فيها إجازة العيدين والمرضية. لكلا الطرفين الإنهاء بدون تعويض أو مكافأة."},
    "women":           {"kw": ["المرأة","حامل","وضع","رضاعة","أمومة","إجازة وضع","حقوق المرأة"], "ref": "م121-130,م151", "label": "المرأة العاملة",
        "facts": "م151: إجازة وضع 10 أسابيع (4 قبل كحد أقصى). ساعة رضاعة يومياً لـ 24 شهر. حماية من الفصل أثناء الحمل والوضع وبعده 180 يوم. عدة 4 أشهر و10 أيام."},
    "insurance":       {"kw": ["تأمين","تأمينات","تأمينات اجتماعية","gosi","ساند","تقاعد","اشتراك","معاش"], "ref": "نظام التأمينات", "label": "التأمينات الاجتماعية",
        "facts": "سعودي: 10.5% خصم + 12.5% شركة. غير سعودي: 2% أخطار فقط. ساند: 0.75%+0.75%. التقاعد: سن 60 + 120 شهر أو 300 شهر مبكر. ساند: 60% لـ 3 أشهر ثم 50% لـ 9 أشهر أقصى 9000."},
    "medical":         {"kw": ["تأمين طبي","تأمين صحي","ضمان صحي","علاج","CCHI","طبي","صحي"], "ref": "نظام CCHI", "label": "التأمين الطبي",
        "facts": "التأمين الصحي إلزامي على صاحب العمل لجميع العاملين ومعاليهم. الحد الأدنى للتغطية حسب وثيقة الضمان الصحي الموحدة."},
    "disciplinary":    {"kw": ["جزاء","عقوبة","إنذار","خصم","لفت نظر","مخالفة","تأديب","جزاءات","اعتداء","اعتداء لفظي","تحرش","سب","شتم","إهانة","ضرب","عنف","تهديد","إساءة"], "ref": "م66-72,م80,م81", "label": "الجزاءات والاعتداء",
        "facts": "م66: لائحة جزاءات معتمدة من الوزارة. م67: الجزاءات: إنذار، غرامة، تأجيل ترقية، إيقاف بدون أجر، فصل. م69: لا يجوز جزاء بعد 30 يوم من اكتشاف المخالفة. م71: التحقيق كتابياً. م80: يحق لصاحب العمل فسخ العقد بدون مكافأة إذا اعتدى العامل على صاحب العمل أو المدير أو أحد رؤسائه أثناء العمل بالقول أو الفعل. م81: يحق للعامل ترك العمل بدون إشعار إذا وقع عليه اعتداء من صاحب العمل أو من يمثله."},
    "safety":          {"kw": ["سلامة","إصابة","حادث عمل","إصابة مهنية","خطر","وفاة عمل"], "ref": "م131-155", "label": "السلامة وإصابات العمل",
        "facts": "م133-141: صاحب العمل يتحمل العلاج. أجر كامل خلال العلاج (سنة). عجز كلي: معاش كامل. وفاة: تعويض أجر 3 سنوات (أدنى 54,000)."},
    "disputes":        {"kw": ["شكوى","نزاع","محكمة عمالية","مكتب العمل","خلاف","تظلم","ودي"], "ref": "م200-231", "label": "تسوية النزاعات",
        "facts": "منصة ودي: تسوية ودية 21 يوم. المحكمة العمالية: بعد فشل التسوية. مدة التقادم: 12 شهر."},
    "saudization":     {"kw": ["نطاقات","سعودة","توطين","نسبة السعودة"], "ref": "م26+قرارات", "label": "نطاقات والتوطين",
        "facts": "نطاقات: بلاتيني > أخضر عالي > أخضر منخفض > أصفر > أحمر. النسب حسب النشاط والحجم."},
    "transfer":        {"kw": ["نقل كفالة","نقل خدمات","كفالتي","تنقل وظيفي"], "ref": "مبادرة تعاقدية", "label": "نقل الخدمات",
        "facts": "نظام تعاقدي جديد عبر منصة قوى. إشعار 90 يوم. التنقل حق للعامل بعد انتهاء العقد."},
    "absence":         {"kw": ["غياب","تغيب","انقطاع","عدم حضور"], "ref": "م80", "label": "الغياب",
        "facts": "م80: فسخ بدون تعويض إذا تغيب 30 يوم متفرقة أو 15 متتالية في السنة بعد إنذار كتابي."},
    "certificate":     {"kw": ["شهادة خبرة","شهادة عمل","خطاب تعريف"], "ref": "م64", "label": "الشهادات",
        "facts": "م64: يلتزم صاحب العمل بإعطاء شهادة خبرة مجاناً عند انتهاء العقد."},
}

# --- HR Topics with framework references ---
HR_TOPICS = {
    "learning_development": {"kw": ["تدريب","تطوير","L&D","دورة","تعليم","برنامج تدريبي","تصميم تعليمي","التدريب"], "fw": "ADDIE,Kirkpatrick,Phillips ROI,70-20-10,Bloom,SAM", "certs": "APTD,PHRi,CIPD L5", "label": "التعلم والتطوير",
        "hr_facts": "ADDIE (APTD): Analysis→Design→Development→Implementation→Evaluation. Kirkpatrick (APTD/PHRi): Level1 Reaction→Level2 Learning→Level3 Behavior→Level4 Results. Phillips ROI (SPHRi): Level5=(Benefits-Costs)/Costs×100. 70-20-10 (CIPD): 70% خبرة + 20% اجتماعي + 10% رسمي. Bloom's Taxonomy (APTD): Knowledge→Comprehension→Application→Analysis→Synthesis→Evaluation."},
    "performance":          {"kw": ["أداء","تقييم","تقييم الأداء","KPI","OKR","أهداف","إدارة الأداء","الأداء"], "fw": "9-Box,MBO,BARS,360-degree,BSC,SMART", "certs": "SPHRi,CIPD L7,SHRM", "label": "إدارة الأداء",
        "hr_facts": "9-Box Grid (SPHRi/PHRi): مصفوفة الأداء×الإمكانات. SMART Goals (PHRi): Specific+Measurable+Achievable+Relevant+Time-bound. BSC (SPHRi): Financial+Customer+Internal+Learning perspectives. MBO (SHRM): Management By Objectives. 360-degree (CIPD L7): تقييم من الرئيس+الزملاء+المرؤوسين+النفس."},
    "recruitment":          {"kw": ["استقطاب","توظيف","مقابلة","ATS","اختيار","تعيين","وصف وظيفي","التوظيف"], "fw": "Competency-Based,Structured Interview,EVP,Employer Brand", "certs": "PHRi,SHRM,CIPD L5", "label": "الاستقطاب والتوظيف",
        "hr_facts": "Competency-Based Selection (PHRi): اختيار على أساس الكفاءات. Structured Interview (SHRM): أسئلة موحدة ومعايير تقييم ثابتة. EVP (CIPD L5): Employee Value Proposition. KPIs: Time-to-Hire, Cost-per-Hire, Quality of Hire, Source Effectiveness."},
    "compensation":         {"kw": ["رواتب","تعويضات","هيكل رواتب","بدلات","مكافآت","حوافز","درجات وظيفية","التعويضات"], "fw": "Total Rewards,Job Evaluation,Compa-Ratio,Salary Survey,Hay", "certs": "PHRi,SPHRi,CIPD L7", "label": "التعويضات والمزايا",
        "hr_facts": "Total Rewards (SHRM/SPHRi): Compensation+Benefits+Work-Life+Recognition+Development. Job Evaluation (PHRi): Point Factor, Ranking, Classification. Compa-Ratio (PHRi): الراتب الفعلي/وسط النطاق. Salary Survey: P25, P50, P75 benchmarks."},
    "engagement":           {"kw": ["رضا","تجربة الموظف","احتفاظ","دوران","engagement","تحفيز","ولاء"], "fw": "Gallup Q12,eNPS,Pulse Survey,EX Journey Map", "certs": "SHRM,CIPD L5,PHRi", "label": "تجربة الموظف"},
    "org_development":      {"kw": ["تطوير مؤسسي","تغيير","OD","هيكلة","إعادة هيكلة","تحول"], "fw": "Burke-Litwin,Kotter 8-Step,ADKAR,Lewin,McKinsey 7S", "certs": "CIPD L7,SPHRi", "label": "التطوير المؤسسي"},
    "leadership":           {"kw": ["قيادة","تطوير قيادي","إدارة","كوتشنج","coaching","توجيه","mentoring"], "fw": "Ulrich,Situational Leadership,GROW,Coaching Models", "certs": "CIPD L7,SPHRi,APTD", "label": "القيادة"},
    "talent":               {"kw": ["مواهب","talent","تعاقب","succession","خلافة","مسار وظيفي","career"], "fw": "9-Box,Talent Pipeline,HIPO,Career Lattice", "certs": "SPHRi,CIPD L7,PHRi", "label": "إدارة المواهب"},
    "analytics":            {"kw": ["تحليلات","analytics","بيانات","مؤشرات","تقارير","لوحة معلومات"], "fw": "Descriptive→Diagnostic→Predictive→Prescriptive", "certs": "PHRi,SPHRi,CIPD L7", "label": "تحليلات HR"},
    "competency":           {"kw": ["كفاءات","competency","مهارات","جدارات","إطار كفاءات"], "fw": "Competency Framework,Skills Matrix,Gap Analysis,TNA", "certs": "SHRM,PHRi,CIPD L5", "label": "إدارة الكفاءات"},
    "onboarding":           {"kw": ["تهيئة","onboarding","تعريف","استقبال","موظف جديد"], "fw": "30-60-90 Plan,Buddy System,Orientation", "certs": "aPHRi,PHRi,SHRM", "label": "التهيئة"},
    "culture":              {"kw": ["ثقافة","بيئة عمل","قيم","diversity","DEI","شمولية"], "fw": "Schein,Denison,Competing Values,Great Place to Work", "certs": "CIPD L7,SHRM,SPHRi", "label": "الثقافة المؤسسية"},
    "workforce_planning":   {"kw": ["تخطيط القوى","workforce planning","احتياج وظيفي"], "fw": "Supply-Demand,Scenario Planning,FTE", "certs": "SPHRi,CIPD L7", "label": "تخطيط القوى العاملة"},
    "policy_design":        {"kw": ["سياسة","لائحة داخلية","نظام داخلي","ضوابط","إعداد سياسة"], "fw": "Policy Framework,Governance,Compliance", "certs": "SHRM,SPHRi", "label": "تصميم السياسات"},
    "job_design":           {"kw": ["تصميم وظيفي","وصف وظيفي","تحليل وظيفي","job design"], "fw": "Job Analysis,Job Description,Job Evaluation", "certs": "PHRi,SHRM", "label": "التصميم الوظيفي"},
    "employee_relations":   {"kw": ["علاقات موظفين","شكاوى داخلية","تظلم","صراع","نزاع داخلي"], "fw": "Conflict Resolution,Mediation,Grievance Process", "certs": "PHRi,SHRM", "label": "علاقات الموظفين"},
    "wellbeing":            {"kw": ["رفاهية","صحة نفسية","wellbeing","توازن","إرهاق"], "fw": "Wellbeing Framework,EAP,Burnout Prevention", "certs": "CIPD L5,SHRM", "label": "رفاهية الموظفين"},
}

def identify_asking_party(question):
    """Identify who is asking: employee, employer, hr_officer, manager, neutral."""
    q = question.lower()
    scores = {}
    for party, patterns in PARTY_PATTERNS.items():
        score = sum(1 for p in patterns if p in q)
        if score > 0: scores[party] = score
    if not scores: return "neutral"
    return max(scores, key=scores.get)

def analyze_question(question, advisor_type):
    """Rich question analysis: topic, intent, party, entities, confidence."""
    q = question.lower().strip()
    analysis = {
        "topic": "general", "subtopic": "", "intent": "explanation",
        "label": "", "reference": "", "certs": "", "facts": "",
        "asking_party": identify_asking_party(question),
        "entities": {}, "needs_calculation": False, "needs_comparison": False,
        "needs_case_application": False, "missing_info": [], "confidence": 0.0,
    }
    topics = LEGAL_TOPICS if advisor_type == "legal" else HR_TOPICS
    best_score = 0
    for key, tdef in topics.items():
        score = sum(3 if kw in q else (2 if kw.replace('ال','') in q.replace('ال','') else 0) for kw in tdef['kw'])
        if score > best_score:
            best_score = score
            analysis['topic'] = key
            analysis['label'] = tdef['label']
            analysis['reference'] = tdef.get('ref', tdef.get('fw', ''))
            analysis['certs'] = tdef.get('certs', '')
            analysis['facts'] = tdef.get('facts', '')
            analysis['hr_facts'] = tdef.get('hr_facts', '')
            analysis['confidence'] = min(score / 9.0, 1.0)
    for intent, patterns in INTENT_PATTERNS.items():
        if any(p in q for p in patterns):
            analysis['intent'] = intent; break
    analysis['needs_calculation'] = analysis['intent'] == 'calculation'
    analysis['needs_comparison'] = analysis['intent'] in ('comparison', 'pros_cons')
    analysis['needs_case_application'] = analysis['intent'] == 'case_application' or analysis['asking_party'] != 'neutral'
    import re
    years = re.findall(r'(\d+)\s*(?:سنة|سنوات|عام|أعوام|شهر|أشهر)', q)
    if years: analysis['entities']['duration'] = years[0]
    amounts = re.findall(r'(\d[\d,]+)\s*(?:ريال|راتب|أجر)', q)
    if amounts: analysis['entities']['amount'] = amounts[0].replace(',','')
    if analysis['needs_calculation']:
        if 'duration' not in analysis['entities']: analysis['missing_info'].append('المدة/سنوات الخدمة')
        if 'amount' not in analysis['entities']: analysis['missing_info'].append('المبلغ/الراتب')
    return analysis

def apply_reasoning_rules(analysis, advisor_type):
    """Infer reasoning mode from topic + intent + party."""
    rules = {"reasoning_mode": "general", "emphasis": [], "depth": "standard"}
    intent = analysis.get('intent', 'explanation')
    party = analysis.get('asking_party', 'neutral')
    if advisor_type == "legal":
        if intent == "case_application" or party in ("employee","employer"):
            rules["reasoning_mode"] = "legal_case_analysis"; rules["depth"] = "deep"
        elif intent == "calculation": rules["reasoning_mode"] = "legal_calculation"
        elif intent == "comparison": rules["reasoning_mode"] = "legal_comparison"
        elif intent == "process": rules["reasoning_mode"] = "legal_procedure"
        elif intent == "rights": rules["reasoning_mode"] = "rights_obligations"
        elif intent == "diagnosis": rules["reasoning_mode"] = "legal_diagnosis"
        else: rules["reasoning_mode"] = "legal_explanation"
    else:
        if intent == "methodology": rules["reasoning_mode"] = "hr_methodology"
        elif intent == "diagnosis": rules["reasoning_mode"] = "hr_diagnosis"; rules["depth"] = "deep"
        elif intent == "comparison": rules["reasoning_mode"] = "hr_comparison"
        elif intent == "kpi_design": rules["reasoning_mode"] = "hr_kpi_design"
        elif intent == "policy": rules["reasoning_mode"] = "hr_policy_design"
        elif intent == "process": rules["reasoning_mode"] = "hr_process_design"
        elif intent == "recommendation": rules["reasoning_mode"] = "hr_recommendation"
        elif intent == "case_application" or party in ("hr_officer","manager"):
            rules["reasoning_mode"] = "hr_applied_guidance"; rules["depth"] = "deep"
        else: rules["reasoning_mode"] = "hr_explanation"
    return rules

def build_reasoning_context(question, advisor_type, analysis, retrieved_context=""):
    """Build dynamic, party-aware reasoning prompt with verified facts."""
    rules = apply_reasoning_rules(analysis, advisor_type)
    mode = rules["reasoning_mode"]
    label = analysis.get("label", "")
    ref = analysis.get("reference", "")
    facts = analysis.get("facts", "")
    certs = analysis.get("certs", "")
    party = analysis.get("asking_party", "neutral")
    confidence = analysis.get("confidence", 0)

    # Party context
    party_labels = {"employee":"عامل/موظف","employer":"صاحب عمل","hr_officer":"مسؤول موارد بشرية","manager":"مدير","neutral":"مستفسر عام"}
    party_text = party_labels.get(party, "مستفسر")

    if advisor_type == "legal":
        header = f"أنت مستشار قانوني متخصص في نظام العمل السعودي. أجب كما يجيب محامٍ خبير يتحدث مع موكّله.\n"
        header += f"المسألة: {label} | المواد: {ref} | السائل: {party_text} | نوع التحليل: {mode}\n"
        header += "❌ لا تذكر أي إطار منهجي HR. فقط مواد قانونية.\n"
        if facts:
            header += f"\n📋 **المواد الصحيحة لهذا الموضوع:**\n{facts}\n"
        else:
            header += "\n📋 **فهرس نظام العمل السعودي (245 مادة) والأنظمة ذات العلاقة:**\n"
            header += "**الباب 1 - التعريفات (م1-7):** تعريفات: عامل، صاحب عمل، أجر، عقد عمل، عمل مؤقت، عمل عرضي، عمل جزئي، مهنة حرة\n"
            header += "**الباب 2 - التوظيف (م8-24):** م12 مكاتب التوظيف الأهلية، م16 تسجيل العمال، م19 وحدات التوظيف في المنشآت\n"
            header += "**الباب 3 - توظيف غير السعوديين (م25-41):** م26 نسبة السعودة 75%، م28 حظر العمل بدون رخصة، م33 شروط رخصة العمل، م38 كفالة العامل، م39 حظر التنازل عن الكفالة\n"
            header += "**الباب 4 - التدريب والتأهيل (م42-48):** م42 تأهيل السعوديين، م43 تدريب 12% سنوياً، م44 تحديد المهن، م45 عقد التأهيل (المهنة/المدة/المراحل/المكافأة)، م46 إلزام المتدرب بالعمل أو رد التكاليف، م47 إنهاء التدريب لعدم القدرة، م48 النفقات على صاحب العمل\n"
            header += "**الباب 5 - علاقات العمل/العقود (م49-60):** م49 تنظيم العلاقة التعاقدية، م50 عقد مكتوب نسختين، م51 محتويات العقد (الاسم/الجنسية/الأجر/نوع العمل/المدة)، م52 خصوصية العمل، م53 فترة التجربة 90 يوم (تمديد 180 بموافقة كتابية)، م54 لا تجربة مكررة، م55 تحول العقد لغير محدد بعد 3 تجديدات أو 4 سنوات، م56 عقد المشروع، م57 البيانات الصحيحة، م58 نقل العامل، م59 شرط عدم المنافسة، م60 وفاة العامل\n"
            header += "**الباب 6 - شروط العمل وظروفه (م61-73):** م61 واجبات العامل، م62 واجبات صاحب العمل، م63 تعديل الشروط، م64 شهادة الخبرة مجاناً، م65 ملف العامل، م66 لائحة تنظيم العمل معتمدة، م67 لائحة الجزاءات (إنذار/غرامة/تأجيل ترقية/إيقاف/فصل)، م68 شروط الجزاء، م69 لا جزاء بعد 30 يوم من الاكتشاف، م70 لا يُوقع جزاءان، م71 التحقيق كتابياً، م72 حق التظلم\n"
            header += "**الباب 7 - إنهاء عقد العمل (م74-82):** م74 حالات الانتهاء (اتفاق/انتهاء مدة/إرادة أحد الطرفين/قوة قاهرة/إغلاق/تقاعد/وفاة)، م75 الإشعار 60 يوم (أجر شهري) أو 30 يوم (غير شهري)، م76 تعويض الإشعار، م77 تعويض الفسخ غير المشروع: 15 يوم/سنة (غير محدد) أو المدة الباقية (محدد) حد أدنى شهرين، م78 يوم الفصل، م79 إنهاء لأسباب غير مشروعة، م80 فسخ بدون مكافأة: اعتداء على صاحب العمل أو المدير/إخلال جسيم/غياب 30 يوم متفرقة أو 15 متتالية/تزوير/في فترة التجربة، م81 حق العامل بالترك بدون إشعار: إخلال صاحب العمل/غش في الشروط/عمل خطر/اعتداء من صاحب العمل، م82 لا يؤثر الفسخ على الحقوق\n"
            header += "**الباب 8 - مكافأة نهاية الخدمة (م83-88):** م84 المكافأة: نصف شهر/سنة أول 5 سنوات + شهر/سنة بعدها (على آخر أجر)، م85 مكافأة الاستقالة: ثلث المكافأة (2-5 سنوات)/ثلثان (5-10)/كاملة (10+)/لا شيء (أقل من سنتين)، م86 المرأة تستحق الكاملة إذا أنهت للزواج أو الإنجاب، م87 استثناء القوة القاهرة، م88 تسوية المستحقات خلال أسبوع (من أنهى صاحب العمل) أو أسبوعين (من أنهى العامل)\n"
            header += "**الباب 9 - الأجور (م89-97):** م89 الأجر بالريال السعودي، م90 الالتزام بالدفع في الموعد المحدد، م91 الأجر مقابل العمل، م92 لا يجوز خصم أكثر من نصف الأجر، م93 حالات الخصم، م94 حماية الأجر/حساب بنكي، م95 الأجر عن أيام الراحة، م96 الاستقطاعات النظامية، م97 أولوية دين العامل\n"
            header += "**الباب 10 - ساعات العمل والراحة (م98-107):** م98 ثمان ساعات يومياً/48 أسبوعياً، م99 رمضان 6 ساعات/36 أسبوعياً، م100 استثناء الأعمال المتقطعة والحراسة (12 ساعة)، م101 راحة نصف ساعة كل 5 ساعات متصلة، م102 لا يبقى في مكان العمل أكثر من 12 ساعة، م104 يوم الجمعة راحة أسبوعية مدفوعة، م105 استبدال يوم الراحة، م106 العمل في الأعياد والمناسبات، م107 أجر العمل الإضافي: الأجر + 50% (إجمالي 150%)\n"
            header += "**الباب 11 - الإجازات (م108-116):** م109 سنوية 21 يوم عادي (أول 5 سنوات) ثم 30 يوم (بعدها)، م110 تأجيل الإجازة بموافقة العامل، م111 أجر الإجازة مقدماً، م112 مناسبات: وفاة 5 أيام/زواج 5/مولود 3، م113 مرضية: 30 يوم كامل الأجر + 60 يوم ثلاثة أرباع + 30 يوم بدون أجر (في السنة)، م114 مرضية لم تنته: فسخ بعد 120 يوم، م115 حج 10-15 يوم لمرة واحدة (بعد سنتين)، م116 إجازة بدون أجر بالاتفاق\n"
            header += "**الباب 12 - تشغيل الأحداث (م117-120):** م117 سن العمل 15 سنة كحد أدنى، م118 ست ساعات يومياً، م119 حظر العمل الخطر والليلي، م120 إجازة الامتحانات\n"
            header += "**الباب 13 - تشغيل المرأة (م121-130):** م121 المساواة في الأحكام، م122 حظر العمل الليلي (استثناءات بقرار وزاري)، م126 حق العودة بعد الوضع، م128 توفير أماكن مخصصة\n"
            header += "**الباب 14 - الوقاية من المخاطر المهنية (م131-145):** م131 توفير وسائل السلامة والوقاية، م135 تدريب العمال على السلامة، م140 الفحص الطبي عند التعيين ودورياً\n"
            header += "**الباب 15 - إصابات العمل والتعويضات (م146-155):** م146 تعريف إصابة العمل، م149 العلاج على نفقة صاحب العمل، م150 أجر كامل خلال العلاج حتى سنة، م151 إجازة الوضع 10 أسابيع + ساعة رضاعة يومياً لـ24 شهر + حماية من الفصل أثناء الحمل والوضع و180 يوم بعده + عدة 4 أشهر و10 أيام، م152 حق الجمع بين التعويضات\n"
            header += "**الباب 16 - العمل البحري (م156-178):** م164 عقد العمل البحري مكتوب ومصدق، م177 غرق السفينة: أجر شهرين إضافيين\n"
            header += "**الباب 17 - العمل في المناجم والمحاجر (م179-199)**\n"
            header += "**الباب 18 - تسوية الخلافات العمالية (م200-231):** م215 هيئات تسوية الخلافات، منصة ودي: تسوية ودية خلال 21 يوم، المحكمة العمالية: بعد فشل التسوية، مدة التقادم: 12 شهر من انتهاء العلاقة\n"
            header += "**التأمينات الاجتماعية (GOSI):** سعودي: معاشات 9.75%+9.75% + أخطار 2% + ساند 0.75%+0.75% = خصم العامل 10.5% + تحمل الشركة 12.5%. غير سعودي: أخطار 2% فقط على الشركة. ساند: 60% لأول 3 أشهر ثم 50% لـ9 أشهر (أقصى 9000). التقاعد: سن 60 + 120 شهر اشتراك أو 300 شهر مبكر\n"
            header += "**الضمان الصحي (CCHI):** التأمين الصحي إلزامي على صاحب العمل لجميع العاملين وأسرهم. الحد الأدنى حسب الوثيقة الموحدة\n"
            header += "**نطاقات (Nitaqat):** بلاتيني > أخضر عالي > أخضر منخفض > أصفر > أحمر. النسب حسب النشاط والحجم. المزايا تتدرج حسب النطاق\n"
            header += "**العمالة المنزلية:** نظام خاص (لائحة العمالة المنزلية) وليس نظام العمل. تجربة 90 يوم، 15 ساعة يومياً، إجازة شهرية بعد سنتين، مكافأة شهر/4 سنوات\n\n"
        header += "⚠️ إذا لم تجد المادة الدقيقة، قل 'لم أتمكن من تحديد المادة بدقة' ولا تخترع رقماً.\n\n"
        # Party-aware instructions
        party_instructions = {
            "employee": "السائل عامل/موظف. ركّز على حقوقه ومستحقاته. اشرح له وضعه القانوني بوضوح وقدم له خطوات عملية يتخذها.\n",
            "employer": "السائل صاحب عمل. ركّز على التزاماته القانونية والإجراءات الصحيحة. وضّح المخاطر والعقوبات عند المخالفة.\n",
            "hr_officer": "السائل مسؤول موارد بشرية. قدم إرشادات تنفيذية مع النماذج والمدد الزمنية والإجراءات التفصيلية.\n",
            "manager": "السائل مدير. ركّز على صلاحياته ومسؤولياته والتبعات القانونية لقراراته.\n",
        }
        if party in party_instructions:
            header += party_instructions[party]
        # Smart reasoning style
        header += "\n**أسلوب الإجابة الإلزامي:**\n"
        header += "- أجب كمحامٍ متخصص يتحدث مع موكّله. حلّل بعمق ولا تكتفِ بسرد المواد.\n"
        header += "- اشرح لماذا وكيف تنطبق كل مادة على حالة السائل تحديداً.\n"
        header += "- قدم السيناريوهات المحتملة والنتائج المختلفة لكل سيناريو.\n"
        header += "- إذا كانت الحالة تحتمل أكثر من تفسير، اذكر جميع الاحتمالات بوضوح.\n"
        header += "- قدم أمثلة حسابية واقعية عند الحاجة (مثلاً: راتب 10,000 × 5 سنوات = ...).\n"
        header += "- اذكر المدد الزمنية والمواعيد النهائية (خلال أسبوع/60 يوم/12 شهر...).\n"
        header += "- وجّه للجهة المختصة عند الحاجة (مكتب العمل/منصة ودي/المحكمة العمالية/التأمينات).\n"
        header += "- لا تقدم ملخصاً عاماً. قدم استشارة متخصصة مفصّلة.\n\n"
        # Dynamic template
        templates = {
            "legal_case_analysis": f"حلّل الحالة القانونية بعمق:\n🔍 **المسألة القانونية:** حدد الواقعة الجوهرية وكيّفها قانونياً\n📋 **المواد المنطبقة:** اسرد كل مادة مع شرح مختصر لكل واحدة\n⚖️ **التحليل والتطبيق:** طبّق كل مادة على وقائع الحالة. اشرح لماذا تنطبق وماذا يترتب عليها\n👤 **موقف {party_text} القانوني:** حقوقه والتزاماته تحديداً في هذه الحالة\n🔄 **السيناريوهات:** قدم 2-3 سيناريوهات محتملة مع نتائجها\n💡 **التوصية العملية:** خطوات مرقّمة يتخذها {party_text} فوراً\n⏰ **المدد والمواعيد:** المهل القانونية\n🏛️ **الجهة المختصة:** أين يتوجه\n⚠️ **تحذيرات:** مخاطر يجب تجنبها",
            "legal_calculation": f"احسب المستحقات بدقة:\n📋 **الأساس القانوني:** المواد التي تحكم الحساب مع نصوصها\n🔢 **المعادلة:** اكتبها بوضوح مع تعريف كل متغير\n📊 **الحساب التفصيلي:** خطوة بخطوة مع أرقام واضحة\n💰 **النتيجة:** المبلغ النهائي المستحق لـ{party_text}\n🔄 **سيناريوهات مختلفة:** ماذا لو تغيرت المدة أو المبلغ؟\n⚠️ **استثناءات وملاحظات:** حالات يختلف فيها الحساب",
            "legal_comparison": f"قارن بعمق وتحليل:\n📋 **الحالة/الخيار الأول:** الوصف + المواد + النتائج القانونية\n📋 **الحالة/الخيار الثاني:** الوصف + المواد + النتائج القانونية\n⚖️ **المقارنة التفصيلية:** الفروقات في الحقوق والالتزامات والتبعات\n💡 **التوصية لـ{party_text}:** أيهما أفضل لوضعه ولماذا",
            "legal_procedure": f"اشرح الإجراءات خطوة بخطوة:\n📋 **الأساس النظامي:** المواد ذات العلاقة\n📝 **الخطوات:** رقّم كل خطوة مع شرح ما يجب فعله بالضبط\n⏰ **المدد:** المهل الزمنية لكل خطوة\n🏛️ **الجهات:** الجهة المسؤولة لكل إجراء\n📄 **المستندات:** قائمة الأوراق المطلوبة\n💡 **نصائح عملية:** ما ينجح وما يجب تجنبه",
            "rights_obligations": f"اشرح الحقوق والالتزامات بشمولية:\n👤 **حقوق {party_text}:** مرقّمة مع المواد لكل حق\n📋 **التزامات {party_text}:** ما يجب عليه فعله\n🏢 **الطرف الآخر:** حقوقه والتزاماته (لتوضيح الصورة)\n⚠️ **تبعات الإخلال:** العقوبات لكل طرف\n💡 **نصيحة:** كيف يحمي {party_text} حقوقه عملياً",
            "legal_diagnosis": f"شخّص المشكلة وقدم الحلول:\n🔍 **تحليل المشكلة:** الإشكالية القانونية الأساسية والفرعية\n📋 **المخالفة:** هل هناك مخالفة؟ من أي طرف؟ أي مادة؟\n⚖️ **الخيارات القانونية:** رتّبها من الأفضل للأضعف\n📝 **خطة العمل:** إجراءات مع تسلسل زمني\n🏛️ **الجهة:** أين يتوجه {party_text}\n⚠️ **المخاطر:** ما يجب الحذر منه",
            "legal_explanation": f"اشرح بعمق وشمولية:\n🔍 **المسألة:** ما الموضوع القانوني بالتحديد\n📋 **الإطار النظامي:** الباب والمواد مع شرح كل مادة\n⚖️ **الشرح التفصيلي:** القاعدة القانونية وفلسفتها ونطاقها\n🔄 **التطبيق العملي:** كيف تُطبّق مع أمثلة واقعية\n💡 **الأثر على {party_text}:** كيف يؤثر عليه\n⚠️ **استثناءات:** متى لا ينطبق الحكم",
        }
        body = templates.get(mode, templates["legal_explanation"])

    else:  # HR
        header = f"أنت مستشار موارد بشرية محترف ومعتمد دولياً. أجب كخبير استشاري يقدم حلولاً عملية ومنهجية.\n"
        header += f"الموضوع: {label} | المناهج: {certs} | الأطر: {ref} | السائل: {party_text} | نوع التحليل: {mode}\n"
        header += "❌ لا تذكر أي مادة قانونية. فقط مفاهيم ومناهج مهنية.\n"
        header += "في كل نقطة اذكر المنهج مثل **(PHRi - المجال)** أو **(CIPD L7 - المجال)**\n"
        header += "\n⚠️ إذا لم تتأكد من إطار منهجي محدد، قل 'بحسب أفضل الممارسات المهنية' ولا تخترع مرجعاً.\n"
        # Inject verified HR facts
        hr_facts = analysis.get('hr_facts', '')
        if not hr_facts:
            t = analysis.get('topic','')
            if t in HR_TOPICS and 'hr_facts' in HR_TOPICS[t]:
                hr_facts = HR_TOPICS[t]['hr_facts']
        if hr_facts:
            header += f"\n📚 **الأطر الموثّقة لهذا الموضوع:**\n{hr_facts}\n\n"
        else:
            header += "\n📚 **المناهج المعتمدة (7 شهادات دولية):**\n"
            header += "**PHRi (HRCI):** Talent Acquisition, HR Admin, Talent Management, Compensation, Employee Relations, HR Analytics\n"
            header += "**aPHRi (HRCI):** HR Operations, Recruitment, Payroll, Training Coordination\n"
            header += "**SPHRi (HRCI):** Strategic HR, Workforce Planning, Total Rewards, Risk Management\n"
            header += "**SHRM-SCP:** Leadership, Ethical Practice, Business Acumen, People/Organization/Workplace\n"
            header += "**CIPD Level 5:** People Management, Performance, Evidence-Based Practice, L&D, Reward\n"
            header += "**CIPD Level 7:** Strategic People Management, OD, Strategic Reward, High-Performance Work Systems\n"
            header += "**APTD (ATD):** Instructional Design (ADDIE/SAM), Training Delivery, Learning Tech, HPT, Kirkpatrick, Phillips ROI\n"
            header += "**الأطر والنماذج الرئيسية:** ADDIE | Kirkpatrick 4 Levels | Phillips ROI L5 | 70-20-10 | Bloom's Taxonomy | 9-Box Grid | Ulrich HR Model | Burke-Litwin | Kotter 8-Step | ADKAR | BSC | Total Rewards | Competency Framework | EVP | GROW Coaching\n\n"
        # Party instructions
        party_instructions = {
            "hr_officer": "السائل مسؤول موارد بشرية. قدم إرشادات تنفيذية مفصّلة مع أدوات ونماذج وخطوات عملية وجداول زمنية.\n",
            "manager": "السائل مدير. ركّز على التطبيق الإداري والقيادي وكيف ينفّذ مع فريقه.\n",
            "employee": "السائل موظف. اشرح له حقوقه المهنية وكيف يطوّر مساره ويستفيد من الأنظمة المتاحة.\n",
            "employer": "السائل صاحب عمل. ركّز على الاستراتيجيات وأفضل الممارسات والعائد على الاستثمار.\n",
        }
        if party in party_instructions:
            header += party_instructions[party]
        # Smart reasoning style
        header += "\n**أسلوب الإجابة الإلزامي:**\n"
        header += "- أجب كاستشاري محترف يقدم حلاً عملياً لا ملخصاً أكاديمياً.\n"
        header += "- اربط كل توصية بمنهج معتمد مع ذكر المصدر.\n"
        header += "- قدم خطوات تنفيذية واضحة ومرقّمة يمكن تطبيقها مباشرة.\n"
        header += "- اذكر KPIs قابلة للقياس لتقييم النجاح.\n"
        header += "- قدم أمثلة عملية من شركات أو ممارسات معروفة عند الإمكان.\n"
        header += "- وضّح الأدوات والتقنيات المطلوبة (مثل: Excel, LMS, ATS, HRIS).\n"
        header += "- اقترح جدولاً زمنياً واقعياً للتنفيذ.\n"
        header += "- لا تقدم تعريفاً عاماً فقط. قدم استشارة متخصصة قابلة للتطبيق.\n\n"
        templates = {
            "hr_methodology": f"قدم المنهجية كاستشاري محترف:\n🎯 **تحديد المجال والهدف:** ماذا نريد تحقيقه\n📚 **الإطار المنهجي:** النموذج الأمثل ({ref}) مع شرح كل مرحلة\n🔬 **التطبيق التفصيلي:** خطوات مرقّمة لـ{party_text}\n🛠️ **الأدوات:** تقنيات وبرامج مطلوبة\n📊 **مؤشرات النجاح:** 5 KPIs مع طريقة القياس\n⏰ **الجدول الزمني:** مراحل التنفيذ\n🌍 **أفضل الممارسات:** أمثلة عالمية ناجحة",
            "hr_diagnosis": f"شخّص المشكلة كخبير:\n🔍 **تحليل الوضع:** ما المشكلة بالضبط وأعراضها\n🔬 **الأسباب الجذرية:** تحليل (5 Whys / Ishikawa / Root Cause)\n💡 **الحل المقترح:** الإطار المنهجي المناسب ({ref}) ولماذا\n📝 **خطة التنفيذ لـ{party_text}:** مراحل مرقّمة + جدول زمني + مسؤوليات\n📊 **قياس التحسن:** مؤشرات قبل/بعد\n⚠️ **مخاطر التنفيذ:** وكيفية تفاديها",
            "hr_comparison": f"قارن كخبير يساعد {party_text} في اتخاذ القرار:\n📚 **النهج الأول:** التعريف + المصدر + المراحل + نقاط القوة والضعف\n📚 **النهج الثاني:** التعريف + المصدر + المراحل + نقاط القوة والضعف\n⚖️ **جدول المقارنة:** معايير واضحة\n💡 **التوصية لـ{party_text}:** أيهما يناسب وضعه ولماذا\n📊 **KPIs المقترحة:** لقياس نجاح النهج المختار",
            "hr_kpi_design": f"صمم منظومة مؤشرات لـ{party_text}:\n🎯 **الهدف الاستراتيجي:** ماذا نريد قياسه ولماذا\n📊 **المؤشرات:** 5-7 KPIs لكل واحد: الاسم + المعادلة + المستهدف + التكرار + المسؤول\n🎯 **المعايير المرجعية:** benchmarks من السوق\n🛠️ **آلية جمع البيانات:** الأدوات والأنظمة\n📈 **لوحة المتابعة:** كيف يتابع {party_text} النتائج",
            "hr_policy_design": f"صمم السياسة باحترافية:\n🎯 **الهدف والمبررات:** لماذا نحتاج هذه السياسة\n📋 **النطاق:** من يشملهم ومن يُستثنى\n📝 **البنود الرئيسية:** القواعد والضوابط بالتفصيل\n🔄 **الإجراءات التنفيذية:** خطوات التطبيق مرقّمة\n👥 **المسؤوليات:** RACI لكل مرحلة\n📊 **القياس:** كيف نقيّم فعالية السياسة\n🔄 **المراجعة:** دورة التحديث والتطوير",
            "hr_process_design": f"صمم العملية بمنهجية:\n📋 **نظرة عامة:** الأهداف والغاية\n🔄 **المراحل:** مرقّمة مع وصف كل مرحلة\n📥 **المدخلات والمخرجات:** لكل مرحلة بوضوح\n👥 **مصفوفة RACI:** المسؤوليات\n🛠️ **الأدوات:** تقنيات وبرامج مطلوبة\n⏰ **الجدول الزمني:** SLA لكل مرحلة\n📊 **مؤشرات كفاءة العملية:** كيف نعرف أنها تعمل بشكل صحيح",
            "hr_recommendation": f"قدم توصية استشارية لـ{party_text}:\n🔍 **تحليل الوضع:** فهم السياق والتحديات\n📋 **الخيارات المتاحة:** 2-3 بدائل مع مزايا وعيوب كل واحد\n💡 **التوصية:** البديل الأفضل + المبررات المنهجية والعملية\n📝 **خطة التنفيذ:** كيف يبدأ {party_text} خطوة بخطوة\n📊 **قياس النجاح:** كيف يعرف أن التوصية نجحت\n⚠️ **مخاطر:** وخطة بديلة",
            "hr_applied_guidance": f"إرشاد تطبيقي متخصص لـ{party_text}:\n🎯 **الموضوع:** تحديد دقيق للاحتياج\n📚 **الإطار المنهجي:** النموذج الأمثل ({ref}) ولماذا هو الأنسب\n🔬 **خطوات التطبيق:** مفصّلة ومرقّمة لـ{party_text}\n🛠️ **الأدوات والموارد:** ما يحتاجه للبدء\n📊 **كيف يقيس النجاح:** مؤشرات واضحة\n⚠️ **تحديات شائعة:** وكيف يتعامل معها\n💡 **نصائح من الخبرة:** ما ينجح وما يجب تجنبه",
            "hr_explanation": f"اشرح المفهوم كخبير لـ{party_text}:\n🎯 **تحديد المجال:** أين يقع هذا المفهوم في HR\n📚 **التعريف المهني:** من المناهج المعتمدة ({certs})\n🔬 **لماذا يهم:** القيمة المضافة والأثر على المنظمة\n🔄 **كيف يُطبّق:** خطوات عملية مع أمثلة\n📊 **كيف يُقاس:** مؤشرات النجاح\n🌍 **أفضل الممارسات:** تجارب شركات رائدة\n💡 **نصيحة لـ{party_text}:** كيف يستفيد من هذا المفهوم",
        }
        body = templates.get(mode, templates["hr_explanation"])

    # Confidence warning
    if confidence < 0.3:
        body += f"\n\n⚠️ مستوى الثقة في تصنيف السؤال منخفض ({confidence:.0%}). أجب بحذر واذكر أنك غير متأكد من التصنيف الدقيق."

    # Missing info
    missing = analysis.get('missing_info', [])
    if missing:
        body += f"\n\n📝 **معلومات ناقصة:** {', '.join(missing)}. اذكر أنك تحتاج هذه المعلومات لإجابة أدق."

    prompt = header + body
    if retrieved_context:
        prompt += f"\n\n**سياق من قاعدة المعرفة:**\n{retrieved_context[:1500]}"
    return prompt

def verify_answer(answer, analysis, advisor_type):
    """Verify answer: no hallucinated citations, domain compliance, party awareness."""
    if not answer or len(answer) < 30: return False, ["too_short"], 0.0
    issues = []
    confidence = 0.7

    if advisor_type == "legal":
        import re
        cited_articles = re.findall(r'(?:المادة|مادة)\s*(\d+)', answer)
        facts = analysis.get('facts', '')
        if cited_articles:
            # Validate against specific facts if available
            if facts:
                facts_articles = set(re.findall(r'م(\d+)', facts))
                wrong = [a for a in cited_articles if a not in facts_articles]
                if len(wrong) > len(cited_articles) // 2:
                    issues.append(f"hallucinated_articles:{','.join(wrong)}")
                    confidence -= 0.3
            # Always check for impossible article numbers
            for a in cited_articles:
                if int(a) > 245:
                    issues.append(f"invalid_article:{a}")
                    confidence -= 0.4
        # Check HR contamination
        if any(m in answer for m in ['PHRi','SHRM','CIPD','APTD','ADDIE','Kirkpatrick']):
            issues.append("hr_contamination")
            confidence -= 0.2
        # Check has legal basis
        if not any(m in answer for m in ['المادة','مادة','م4','م5','م7','م8','م9','م1','الباب','نظام']):
            if analysis.get('intent') not in ('explanation',):
                issues.append("no_legal_basis")
                confidence -= 0.1
    else:
        # Check legal contamination
        if any(m in answer for m in ['المادة ','م42','م50','م53','م74','م77','م80','م84','م88','م98','م109','م151']):
            issues.append("legal_contamination")
            confidence -= 0.2
        # Check has framework reference
        if not any(m in answer for m in ['PHRi','SHRM','CIPD','APTD','SPHRi','aPHRi','framework','model','نموذج','إطار','منهج']):
            issues.append("no_hr_framework")
            confidence -= 0.1

    # Check party awareness
    party = analysis.get('asking_party', 'neutral')
    if party != 'neutral':
        party_labels = {"employee":"عامل","employer":"صاحب العمل","hr_officer":"موارد بشرية","manager":"مدير"}
        if party in party_labels and party_labels[party] not in answer.lower():
            confidence -= 0.05

    return len(issues) == 0, issues, max(confidence, 0.1)

def filter_response(text, consultant_type):
    """Remove domain-inappropriate content from responses."""
    if not text: return text
    lines = text.split('\n')
    filtered = []
    for line in lines:
        if consultant_type == "hr":
            skip = False
            for marker in ['المادة ','مادة ','نظام العمل','اللائحة التنفيذية',
                'المحكمة العمالية','مكتب العمل','الباب ال',
                'م42','م43','م44','م45','م46','م47','م48',
                'م50','م51','م53','م55','م74','م75','م77','م80','م81',
                'م84','م85','م88','م89','م90','م98','م99','م107',
                'م109','م112','م113','م115','م151','م164']:
                if marker in line and not any(h in line for h in ['PHRi','SHRM','CIPD','APTD','SPHRi']):
                    skip = True; break
            if not skip: filtered.append(line)
        elif consultant_type == "legal":
            skip = False
            for marker in ['PHRi','SHRM','CIPD','APTD','SPHRi','aPHRi','ADDIE',
                'Kirkpatrick','Phillips ROI','Instructional Design','Ulrich',
                '9-Box','Balanced Scorecard','70-20-10','HPT','Bloom']:
                if marker in line: skip = True; break
            if not skip: filtered.append(line)
        else:
            filtered.append(line)
    result = '\n'.join(filtered).strip()
    if len(result) < 30 and len(text) > 100:
        if consultant_type == "hr":
            return "هذا السؤال يتعلق بالجانب القانوني. يرجى استخدام **المستشار القانوني ⚖️**."
        else:
            return "هذا السؤال يتعلق بالموارد البشرية. يرجى استخدام **مستشار الموارد البشرية 📚**."
    return result if len(result) > 20 else text

def generate_advisor_answer(question, advisor_type, system_prompt=None):
    """Full advisory pipeline: analyze → rules → retrieve → reason → generate → verify."""
    import requests as req_lib

    # Step 1: Analyze
    analysis = analyze_question(question, advisor_type)

    # Step 2: Rules
    rules = apply_reasoning_rules(analysis, advisor_type)

    # Step 3: Retrieve from RAG (advisor-specific vector search)
    retrieved = ""
    try:
        if '_knowledge_engine' in st.session_state:
            retrieved = st.session_state._knowledge_engine.search(question, advisor_type=advisor_type)
    except: pass

    # Step 4: Build reasoning context
    reasoning_prompt = build_reasoning_context(question, advisor_type, analysis, retrieved)

    # Step 5: Generate
    answer = _call_llm(question, reasoning_prompt, req_lib)

    # Step 6: Verify
    if answer:
        is_valid, issues, conf = verify_answer(answer, analysis, advisor_type)
        if not is_valid:
            # Filter contamination
            answer = filter_response(answer, advisor_type)
            if answer and len(answer) > 30 and "hallucinated" not in str(issues):
                return answer
            # Retry with stronger instructions
            retry_extra = "\n\n⚠️ الإجابة السابقة غير مقبولة. "
            if "hallucinated_articles" in str(issues):
                retry_extra += f"استخدم فقط هذه المواد: {analysis.get('facts','')}. لا تخترع أرقاماً."
            elif "hr_contamination" in issues:
                retry_extra += "لا تذكر أي إطار منهجي. فقط مواد قانونية."
            elif "legal_contamination" in issues:
                retry_extra += "لا تذكر أي مادة قانونية. فقط مناهج مهنية."
            elif "no_legal_basis" in issues:
                retry_extra += f"يجب ذكر المواد: {analysis.get('facts','')}."
            elif "no_hr_framework" in issues:
                retry_extra += f"يجب ذكر المناهج: {analysis.get('certs','')}."
            answer = _call_llm(question, reasoning_prompt + retry_extra, req_lib)
            if answer:
                return filter_response(answer, advisor_type)
        else:
            return answer
    return None

def _call_llm(question, reasoning_prompt, req_lib):
    """Call available LLM providers."""
    for prov, key_name, secret_name in [('groq','groq_api_key','groq'),('gemini','gemini_api_key','gemini'),('openrouter','openrouter_api_key','openrouter')]:
        api_key = st.session_state.get(key_name,'')
        if not api_key:
            try: api_key = st.secrets.get(secret_name,{}).get("api_key","")
            except: pass
        if not api_key: continue
        try:
            if prov == 'groq':
                resp = req_lib.post("https://api.groq.com/openai/v1/chat/completions",
                    json={"model":"llama-3.3-70b-versatile","messages":[
                        {"role":"system","content":reasoning_prompt[:3500]},
                        {"role":"user","content":question}],"max_tokens":2500,"temperature":0.3},
                    headers={'Authorization':f'Bearer {api_key}'},timeout=30)
            elif prov == 'gemini':
                resp = req_lib.post(f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent?key={api_key}",
                    json={"contents":[{"parts":[{"text":f"{reasoning_prompt[:3000]}\n\nالسؤال: {question}"}]}],
                        "generationConfig":{"maxOutputTokens":2000,"temperature":0.3}},timeout=30)
            else:
                resp = req_lib.post("https://openrouter.ai/api/v1/chat/completions",
                    json={"model":"meta-llama/llama-3.3-70b-instruct:free",
                        "messages":[{"role":"user","content":f"{reasoning_prompt[:2500]}\n\n{question}"}],"max_tokens":1500},
                    headers={'Authorization':f'Bearer {api_key}','HTTP-Referer':'https://hr-analytics-risal.streamlit.app'},timeout=30)
            if resp.status_code == 200:
                if prov == 'gemini':
                    text = resp.json().get('candidates',[{}])[0].get('content',{}).get('parts',[{}])[0].get('text','')
                else:
                    text = resp.json().get('choices',[{}])[0].get('message',{}).get('content','')
                if text and len(text) > 20: return text
        except: continue
    return None

def get_best_kb_answer(question, system_prompt=None):
    """Main entry: ALL questions go through full advisory reasoning pipeline."""
    consultant_type = "legal" if system_prompt and 'المستشار القانوني' in system_prompt else "hr"

    # Full reasoning pipeline (no shallow fallback)
    answer = generate_advisor_answer(question, consultant_type, system_prompt)
    if answer and len(answer) > 30:
        auto_learn_from_answer(question, answer, consultant_type)
        return answer

    # Diagnostic only (no shallow answers)
    k_status = []
    for p in ['groq','gemini','openrouter']:
        has = bool(st.session_state.get(f'{p}_api_key',''))
        k_status.append(f"{p}: {'✅' if has else '❌'}")
    return (f"**لم يتمكن النظام من توليد إجابة تحليلية**\n\n"
            f"المزودين: {' | '.join(k_status)}\n\n"
            f"**الحلول:**\n"
            f"1. تأكد من إضافة مفتاح Groq في الإعدادات\n"
            f"2. أعد المحاولة بعد لحظات")


def auto_learn_from_answer(question, answer, consultant_type="legal"):
    """Save good AI answers and improve them over time."""
    if not answer or len(answer) < 50: return
    skip_phrases = ["لم أتمكن", "يرجى استخدام", "المستشار القانوني ⚖️", "مستشار الموارد البشرية 📚", "المفاتيح:", "الأخطاء:"]
    if any(p in answer for p in skip_phrases): return
    try:
        conn = get_conn(); c = conn.cursor()
        c.execute(f"SELECT value FROM app_config WHERE key = {_ph()}", ("rag_learned",))
        row = c.fetchone()
        learned = json.loads(row[0]) if row else []
        q_hash = hashlib.md5(question.encode()).hexdigest()[:10]

        # Check if similar question exists - IMPROVE the answer
        for i, item in enumerate(learned):
            if item.get('hash') == q_hash and item.get('type') == consultant_type:
                # Merge: keep longer/better answer
                old_len = len(item.get('a', ''))
                if len(answer) > old_len:
                    learned[i]['a'] = answer[:1500]
                    learned[i]['updated'] = datetime.now().strftime("%Y-%m-%d")
                    learned[i]['improvements'] = item.get('improvements', 0) + 1
                _upsert_config(c, "rag_learned", json.dumps(learned, ensure_ascii=False))
                conn.commit(); conn.close()
                return

        # New question - save it
        learned.append({
            "q": question[:200], "a": answer[:1500], "hash": q_hash,
            "type": consultant_type, "date": datetime.now().strftime("%Y-%m-%d"),
            "score": 0, "improvements": 0
        })
        if len(learned) > 500: learned = learned[-500:]
        _upsert_config(c, "rag_learned", json.dumps(learned, ensure_ascii=False))
        conn.commit(); conn.close()
    except: pass

def rate_answer(q_hash, rating, consultant_type):
    """User rates answer: +1 good, -1 bad. Bad answers get removed."""
    try:
        conn = get_conn(); c = conn.cursor()
        c.execute(f"SELECT value FROM app_config WHERE key = {_ph()}", ("rag_learned",))
        row = c.fetchone()
        if row:
            learned = json.loads(row[0])
            for i, item in enumerate(learned):
                if item.get('hash') == q_hash and item.get('type') == consultant_type:
                    new_score = item.get('score', 0) + rating
                    if new_score <= -2:
                        learned.pop(i)  # Remove bad answers
                    else:
                        learned[i]['score'] = new_score
                    break
            _upsert_config(c, "rag_learned", json.dumps(learned, ensure_ascii=False))
            conn.commit()
        conn.close()
    except: pass

def get_learning_stats():
    """Get learning system statistics."""
    try:
        conn = get_conn(); c = conn.cursor()
        c.execute(f"SELECT value FROM app_config WHERE key = {_ph()}", ("rag_learned",))
        row = c.fetchone(); conn.close()
        if row:
            learned = json.loads(row[0])
            legal = [l for l in learned if l.get('type') == 'legal']
            hr = [l for l in learned if l.get('type') == 'hr']
            good = [l for l in learned if l.get('score', 0) > 0]
            improved = [l for l in learned if l.get('improvements', 0) > 0]
            return {
                "total": len(learned), "legal": len(legal), "hr": len(hr),
                "good_rated": len(good), "improved": len(improved),
                "avg_score": sum(l.get('score',0) for l in learned) / max(len(learned),1)
            }
    except: pass
    return {"total":0,"legal":0,"hr":0,"good_rated":0,"improved":0,"avg_score":0}

def export_widget(dataframes, title="تقرير", key_prefix="exp"):
    """Universal export: Excel (with charts) + CSV + PDF (with interactive charts)"""
    if dataframes is None: return
    if isinstance(dataframes, pd.DataFrame):
        if len(dataframes) == 0: return
        all_dfs = {"البيانات": dataframes}
    elif isinstance(dataframes, dict):
        all_dfs = {k: v for k, v in dataframes.items() if isinstance(v, pd.DataFrame) and len(v) > 0}
        if not all_dfs: return
    else: return

    st.markdown("---")
    st.markdown("### 📥 تصدير التقرير")
    ex1, ex2, ex3 = st.columns(3)
    fname = f"{title}_{datetime.now().strftime('%Y%m%d')}".replace(" ","_")

    # ===== EXCEL with auto-charts =====
    with ex1:
        try:
            ox = io.BytesIO()
            with pd.ExcelWriter(ox, engine='xlsxwriter') as w:
                for sname, df in all_dfs.items():
                    clean = str(sname)[:31].replace('/','_').replace('\\','_')
                    df.to_excel(w, sheet_name=clean, index=False)
                    ws = w.sheets[clean]; ws.set_column('A:Z', 18)
                    hf = w.book.add_format({'bold':True,'bg_color':'#0F4C5C','font_color':'white','border':1})
                    for ci, col in enumerate(df.columns): ws.write(0, ci, str(col), hf)
                    # Auto-charts
                    num_c = df.select_dtypes('number').columns.tolist()
                    cat_c = [c for c in df.columns if df[c].dtype=='object' and 1<df[c].nunique()<20]
                    if cat_c and num_c:
                        grp = df.groupby(cat_c[0])[num_c[0]].sum().head(10)
                        csn = f'{clean[:26]}_ch'
                        cws = w.book.add_worksheet(csn)
                        cws.write(0,0,cat_c[0],hf); cws.write(0,1,num_c[0],hf)
                        for j,(idx,val) in enumerate(grp.items()): cws.write(j+1,0,str(idx)); cws.write(j+1,1,val)
                        ch = w.book.add_chart({'type':'bar'})
                        ch.add_series({'categories':[csn,1,0,len(grp),0],'values':[csn,1,1,len(grp),1],'fill':{'color':'#0F4C5C'}})
                        ch.set_title({'name':f'{cat_c[0]} vs {num_c[0]}'}); ch.set_size({'width':520,'height':320}); ch.set_legend({'none':True})
                        cws.insert_chart('D1',ch)
                        ch2 = w.book.add_chart({'type':'pie'})
                        counts = df[cat_c[0]].value_counts().head(8)
                        for j,(idx,val) in enumerate(counts.items()): cws.write(j+1,3,str(idx)); cws.write(j+1,4,val)
                        ch2.add_series({'categories':[csn,1,3,len(counts),3],'values':[csn,1,4,len(counts),4],'data_labels':{'percentage':True}})
                        ch2.set_title({'name':f'توزيع {cat_c[0]}'}); ch2.set_size({'width':480,'height':320})
                        cws.insert_chart('D20',ch2)
            st.download_button("📊 Excel", data=ox.getvalue(), file_name=f"{fname}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                type="primary", use_container_width=True, key=f"{key_prefix}_xl")
        except Exception as e: st.error(f"خطأ Excel: {e}")

    # ===== CSV =====
    with ex2:
        try:
            parts = []
            for sname, df in all_dfs.items():
                if len(all_dfs)>1: parts.append(f"\n=== {sname} ===\n")
                parts.append(df.to_csv(index=False))
            st.download_button("📄 CSV", data="\n".join(parts).encode('utf-8-sig'), file_name=f"{fname}.csv",
                mime="text/csv", use_container_width=True, key=f"{key_prefix}_csv")
        except Exception as e: st.error(f"خطأ CSV: {e}")

    # ===== PDF with interactive Plotly charts =====
    with ex3:
        try:
            charts_html = ""; tables_html = ""
            kpi_html = "<div style='display:flex;gap:12px;flex-wrap:wrap;margin:15px 0'>"
            for sname, df in all_dfs.items():
                num_c = df.select_dtypes('number').columns.tolist()
                cat_c = [c for c in df.columns if df[c].dtype=='object' and 1<df[c].nunique()<20]
                kpi_html += f"<div class='kb'><div class='kv'>{len(df):,}</div><div class='kl'>سجلات {sname}</div></div>"
                for nc in num_c[:2]:
                    kpi_html += f"<div class='kb'><div class='kv'>{df[nc].mean():,.0f}</div><div class='kl'>متوسط {nc[:15]}</div></div>"
                    kpi_html += f"<div class='kb'><div class='kv'>{df[nc].sum():,.0f}</div><div class='kl'>إجمالي {nc[:15]}</div></div>"
                # Bar chart
                if cat_c and num_c:
                    grp = df.groupby(cat_c[0])[num_c[0]].sum().head(12).sort_values()
                    fig = px.bar(x=grp.values,y=grp.index,orientation='h',title=f'{cat_c[0]} vs {num_c[0]}',color=grp.values,color_continuous_scale='teal')
                    fig.update_layout(height=350,margin=dict(l=10,r=10,t=40,b=10),showlegend=False,coloraxis_showscale=False)
                    charts_html += fig.to_html(full_html=False,include_plotlyjs="cdn")
                # Pie chart
                if cat_c:
                    counts = df[cat_c[0]].value_counts().head(8)
                    fig = px.pie(values=counts.values,names=counts.index,title=f'توزيع {cat_c[0]}',hole=0.4)
                    fig.update_layout(height=350,margin=dict(l=10,r=10,t=40,b=10))
                    charts_html += fig.to_html(full_html=False,include_plotlyjs='cdn')
                # Histogram
                if num_c:
                    fig = px.histogram(df,x=num_c[0],nbins=20,title=f'توزيع {num_c[0]}',color_discrete_sequence=['#0F4C5C'])
                    fig.add_vline(x=df[num_c[0]].mean(),line_dash="dash",line_color="red",annotation_text=f"المتوسط: {df[num_c[0]].mean():,.0f}")
                    fig.update_layout(height=350,margin=dict(l=10,r=10,t=40,b=10))
                    charts_html += fig.to_html(full_html=False,include_plotlyjs='cdn')
                # Box plot
                if len(num_c)>=1 and cat_c:
                    fig = px.box(df,x=cat_c[0],y=num_c[0],title=f'{num_c[0]} حسب {cat_c[0]}',color_discrete_sequence=['#E9C46A'])
                    fig.update_layout(height=350,margin=dict(l=10,r=10,t=40,b=10))
                    charts_html += fig.to_html(full_html=False,include_plotlyjs='cdn')
                tables_html += f"<h2>{sname}</h2>{df.to_html(index=False,classes='tbl')}"
                break  # Charts from first df only
            kpi_html += "</div>"
            pdf_html = f"""<!DOCTYPE html><html><head><meta charset="UTF-8">
            <style>body{{font-family:Arial,sans-serif;padding:20px;direction:rtl;color:#333}}
            h1{{color:#0F4C5C;text-align:center;border-bottom:3px solid #E36414;padding-bottom:10px}}
            h2{{color:#264653;margin-top:25px;border-bottom:1px solid #eee;padding-bottom:5px}}
            .tbl{{border-collapse:collapse;width:100%;margin:10px 0}}
            .tbl th{{background:#0F4C5C;color:white;padding:8px;border:1px solid #ddd;text-align:center}}
            .tbl td{{padding:6px;border:1px solid #ddd;text-align:center}}
            .tbl tr:nth-child(even){{background:#f8f9fa}}
            .kb{{background:#f0f4f8;border-radius:8px;padding:12px 20px;border-right:4px solid #E36414;min-width:110px}}
            .kv{{font-size:1.3em;font-weight:700;color:#0F4C5C}} .kl{{font-size:0.7em;color:#888}}
            .footer{{margin-top:30px;color:#888;font-size:0.8em;text-align:center;border-top:1px solid #ddd;padding-top:10px}}
            @media print{{@page{{margin:0.5cm}} .js-plotly-plot{{break-inside:avoid}}}}</style>
            
            </head><body><h1>📊 {title}</h1>
            <p style="text-align:center;color:#888">{datetime.now().strftime('%Y-%m-%d %H:%M')} | HR Analytics Platform</p>
            {kpi_html}
            <h2>📊 الرسوم البيانية</h2>{charts_html}
            <h2>📋 البيانات التفصيلية</h2>{tables_html}
            <div class="footer">HR Analytics Platform | تم الإنشاء تلقائياً</div></body></html>"""
            st.download_button("📄 PDF", data=pdf_html.encode('utf-8'), file_name=f"{fname}.html",
                mime="text/html", use_container_width=True, key=f"{key_prefix}_pdf",
                help="افتح في المتصفح ← Save as PDF")
        except Exception as e: st.error(f"خطأ PDF: {e}")
def fmt(v): return f"{v:,.0f}"
def has(df,n): return df is not None and n in df.columns and len(df)>0

def audit_log(action, details="", user=None):
    """Save audit trail to database"""
    try:
        conn = get_conn(); c = conn.cursor()
        c.execute(f"SELECT value FROM app_config WHERE key = {_ph()}", ("audit_log",))
        row = c.fetchone()
        logs = json.loads(row[0]) if row else []
        logs.append({"time": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "user": user or st.session_state.get('user_name','النظام'),
            "action": action, "details": str(details)[:200]})
        if len(logs) > 1000: logs = logs[-1000:]
        _upsert_config(c, "audit_log", json.dumps(logs, ensure_ascii=False))
        conn.commit(); conn.close()
    except: pass

def save_snapshot(data, snapshot_type="monthly"):
    """Save historical snapshot of employee data"""
    try:
        conn = get_conn(); c = conn.cursor()
        key = f"snapshot_{snapshot_type}_{datetime.now().strftime('%Y%m')}"
        c.execute(f"SELECT value FROM app_config WHERE key = {_ph()}", (key,))
        if not c.fetchone():
            summary = {"date": datetime.now().strftime("%Y-%m-%d"), "count": len(data),
                "columns": list(data.columns[:20])}
            num_cols = data.select_dtypes('number').columns.tolist()
            for nc in num_cols[:5]:
                summary[f"avg_{nc}"] = round(data[nc].mean(), 2)
                summary[f"sum_{nc}"] = round(data[nc].sum(), 2)
            cat_cols = [c2 for c2 in data.columns if data[c2].dtype=='object' and 1<data[c2].nunique()<20]
            for cc in cat_cols[:3]:
                summary[f"dist_{cc}"] = data[cc].value_counts().head(10).to_dict()
            _upsert_config(c, key, json.dumps(summary, ensure_ascii=False, default=str))
            conn.commit()
        conn.close()
    except: pass

# ===== DATA LOADER =====
COL_MAP = {
    'emp id':'رقم الموظف','employee id':'رقم الموظف','name (english)':'الاسم الإنجليزي',
    'name (arabic)':'الاسم','name':'الاسم','department':'القسم','division':'القطاع',
    'job title':'المسمى الوظيفي','position':'المسمى الوظيفي','join date':'تاريخ التعيين',
    'hiring date':'تاريخ التعيين','location':'الموقع','city':'الموقع',
    'tenure (yrs)':'سنوات الخدمة','basic salary':'الراتب الأساسي',
    'nationality group':'الجنسية','nationality':'الجنسية','gender':'الجنس',
    'gross salary':'الراتب الإجمالي','net salary':'صافي الراتب',
    'housing allowance':'بدل السكن','transportation allowance':'بدل النقل',
    'grade':'الدرجة','level':'المستوى','age':'العمر','age group':'الفئة العمرية',
    'generation':'الجيل','employment type':'نوع التوظيف',
    'salary month':'شهر الراتب','quarter':'الربع','salary year':'سنة الراتب',
    'gosi deduction':'خصم التأمينات','overtime hours':'ساعات إضافية',
    'overtime cost':'تكلفة الإضافي','special allowance':'بدل خاص',
    'mobile allowance':'بدل جوال','living cost  allowance':'بدل معيشة',
    'salary range':'شريحة الراتب','other deduction':'خصومات أخرى',
    'hourly rate (total salary)':'سعر الساعة الإجمالي',
    'gross salary with overtime':'الإجمالي مع الإضافي',
}

def smart_read(xl, sheet):
    df_raw = pd.read_excel(xl, sheet_name=sheet, header=None)
    best_row, best_score = 0, 0
    for i in range(min(5, len(df_raw))):
        score = sum(1 for v in df_raw.iloc[i] if isinstance(v, str) and len(str(v).strip())>1 and not str(v).startswith('Unnamed') and not str(v).startswith('Total'))
        if score > best_score: best_score, best_row = score, i
    df = pd.read_excel(xl, sheet_name=sheet, header=best_row)
    df = df[[c for c in df.columns if not str(c).startswith('Unnamed')]].dropna(how='all').reset_index(drop=True)
    return df

def norm_cols(df):
    new = {}
    for c in df.columns:
        k = str(c).strip().lower()
        new[c] = COL_MAP.get(k, c)
    return df.rename(columns=new)


# ===== END-OF-SERVICE CALCULATOR (Saudi Labor Law Art 84/85) =====
def calc_eos(monthly_salary, start_date, end_date, is_resignation=False):
    """حاسبة مستحقات نهاية الخدمة - نظام العمل السعودي"""
    delta = relativedelta(end_date, start_date)
    total_days = (end_date - start_date).days
    total_years = total_days / 365.25

    # Article 84: Base calculation
    if total_years <= 5:
        eos_84 = (monthly_salary / 2) * total_years
    else:
        first_5 = (monthly_salary / 2) * 5
        remaining = monthly_salary * (total_years - 5)
        eos_84 = first_5 + remaining

    # Article 85: Resignation adjustments
    if is_resignation:
        if total_years < 2:
            eos_final = 0
            eos_pct = 0
            note = "لا يستحق مكافأة (أقل من سنتين)"
        elif total_years < 5:
            eos_final = eos_84 / 3
            eos_pct = 33.3
            note = "ثلث المكافأة (من 2 إلى 5 سنوات)"
        elif total_years < 10:
            eos_final = eos_84 * 2 / 3
            eos_pct = 66.7
            note = "ثلثا المكافأة (من 5 إلى 10 سنوات)"
        else:
            eos_final = eos_84
            eos_pct = 100
            note = "المكافأة كاملة (أكثر من 10 سنوات)"
    else:
        eos_final = eos_84
        eos_pct = 100
        note = "المكافأة كاملة (إنهاء من صاحب العمل / انتهاء العقد)"

    # Vacation balance calculation (21 days for first 5 years, 30 after)
    daily_salary = monthly_salary / 30
    vac_days_per_year = 30 if total_years >= 5 else 21

    return {
        "years": delta.years, "months": delta.months, "days": delta.days,
        "total_years": round(total_years, 2),
        "total_days": total_days,
        "eos_art84": round(eos_84, 2),
        "eos_final": round(eos_final, 2),
        "eos_pct": eos_pct,
        "note": note,
        "daily_salary": round(daily_salary, 2),
        "vac_days_per_year": vac_days_per_year,
        "is_resignation": is_resignation,
    }


# ===== TRAINING DATA (from v4) =====
PROVIDERS = {
    "السعودية": [
        {"name":"معهد الإدارة العامة","spec":"الإدارة والقيادة","type":"حكومي","url":"ipa.edu.sa"},
        {"name":"غرفة جدة","spec":"المهارات المهنية","type":"شبه حكومي","url":"jcci.org.sa"},
        {"name":"KPMG Academy","spec":"المالية والمحاسبة","type":"خاص","url":"kpmg.com/sa"},
        {"name":"PwC Academy","spec":"التحول الرقمي","type":"خاص","url":"pwcacademy.me"},
        {"name":"Misk Academy","spec":"التقنية والابتكار","type":"غير ربحي","url":"misk.org.sa"},
        {"name":"بكه للتعليم","spec":"إدارة المشاريع","type":"خاص","url":"bakkah.com"},
        {"name":"BIBF","spec":"الخدمات المالية","type":"خاص","url":"bibf.com"},
        {"name":"Udacity MENA","spec":"AI والبيانات","type":"أونلاين","url":"udacity.com"},
    ],
    "الخليج": [
        {"name":"Informa Connect","spec":"القيادة","type":"خاص","url":"informaconnect.com"},
        {"name":"London Business School ME","spec":"MBA","type":"خاص","url":"lbs.ac.uk"},
        {"name":"Dale Carnegie UAE","spec":"المهارات القيادية","type":"خاص","url":"dalecarnegie.com"},
    ],
    "مصر": [
        {"name":"الجامعة الأمريكية بالقاهرة","spec":"إدارة الأعمال","type":"أكاديمي","url":"aucegypt.edu"},
        {"name":"Sprints","spec":"البرمجة والتقنية","type":"خاص","url":"sprints.ai"},
        {"name":"Digital Egypt Pioneers","spec":"التحول الرقمي","type":"حكومي","url":"mcit.gov.eg"},
    ],
    "أونلاين": [
        {"name":"Coursera for Business","spec":"متعدد","type":"أونلاين","url":"coursera.org"},
        {"name":"LinkedIn Learning","spec":"مهارات مهنية","type":"أونلاين","url":"linkedin.com/learning"},
        {"name":"Google Certificates","spec":"التقنية","type":"أونلاين","url":"grow.google"},
    ]
}

DEFAULT_BUDGET = [
    {"dept":"المبيعات","budget":16000,"pct":22.9,"priority":"حرج","cat":"محرك إيرادات"},
    {"dept":"التسويق","budget":13000,"pct":18.6,"priority":"حرج","cat":"محرك إيرادات"},
    {"dept":"تطوير الأعمال","budget":11000,"pct":15.7,"priority":"عالي","cat":"محرك إيرادات"},
    {"dept":"عمليات المنتجات","budget":9000,"pct":12.9,"priority":"عالي","cat":"ممكّن نمو"},
    {"dept":"البيانات والذكاء","budget":7000,"pct":10.0,"priority":"عالي","cat":"ممكّن نمو"},
    {"dept":"المالية","budget":5000,"pct":7.1,"priority":"متوسط","cat":"بنية تحتية"},
    {"dept":"الموارد البشرية","budget":4000,"pct":5.7,"priority":"متوسط","cat":"بنية تحتية"},
    {"dept":"الحوكمة","budget":3000,"pct":4.3,"priority":"متوسط","cat":"بنية تحتية"},
    {"dept":"القانونية","budget":2000,"pct":2.9,"priority":"أساسي","cat":"بنية تحتية"},
]

Q_SPLIT = {"Q1":0.35,"Q2":0.30,"Q3":0.20,"Q4":0.15}

TRAINING_PROGRAMS = {
    "المبيعات": [
        {"program":"Advanced Consultative Selling Skills","budget":3600,"source":"External","timing":"Q1-Q2","impact":"Increase conversion rate & close major deals"},
        {"program":"Customer Success & CLV Optimization","budget":3000,"source":"External","timing":"Q1","impact":"Improve customer retention & increase upsell"},
        {"program":"Strategic Key Account Management","budget":2800,"source":"External","timing":"Q2","impact":"Maximize revenue from key accounts"},
        {"program":"Customer Experience & Satisfaction (NPS/CSAT)","budget":2400,"source":"Internal","timing":"Q1-Q3","impact":"Enhance CX to increase loyalty"},
        {"program":"Negotiation & Objection Handling","budget":2200,"source":"External","timing":"Q2-Q3","impact":"Improve contract terms & profit margins"},
        {"program":"CRM & Sales Automation Tools","budget":2000,"source":"Internal","timing":"Q1","impact":"Increase sales team efficiency"},
    ],
    "التسويق": [
        {"program":"Advanced Digital Marketing & Campaign Mgmt","budget":3000,"source":"External","timing":"Q1","impact":"Increase ROAS"},
        {"program":"Growth Hacking & Growth Strategy","budget":2800,"source":"External","timing":"Q1-Q2","impact":"Accelerate user acquisition at lower cost"},
        {"program":"SEO & Content Marketing Mastery","budget":2000,"source":"Online","timing":"Q2","impact":"Increase organic traffic & conversions"},
        {"program":"Marketing Analytics & Performance","budget":2000,"source":"Online","timing":"Q1-Q2","impact":"Data-driven marketing decisions"},
        {"program":"Brand Management & Positioning","budget":1600,"source":"External","timing":"Q3","impact":"Strengthen market positioning"},
        {"program":"Partnership & Alliance Marketing","budget":1600,"source":"Internal","timing":"Q2-Q3","impact":"Expand commercial partner network"},
    ],
    "تطوير الأعمال": [
        {"program":"BD Strategy & Strategic Partnerships","budget":3000,"source":"External","timing":"Q1","impact":"Build strategic partnerships"},
        {"program":"New Market Expansion - MENA","budget":2400,"source":"External","timing":"Q1-Q2","impact":"Support regional expansion"},
        {"program":"B2B Enterprise Sales","budget":2000,"source":"External","timing":"Q2","impact":"Win corporate clients"},
        {"program":"Deal Structuring & Revenue Models","budget":1800,"source":"External","timing":"Q3","impact":"Optimize deal economics"},
        {"program":"Industry Analysis & Competitive Intelligence","budget":1800,"source":"Online","timing":"Q2-Q3","impact":"Informed strategic decisions"},
    ],
    "عمليات المنتجات": [
        {"program":"Agile Product Management","budget":2500,"source":"External","timing":"Q1","impact":"Faster product delivery"},
        {"program":"Product Analytics & User Research","budget":2000,"source":"Online","timing":"Q1-Q2","impact":"Data-driven product decisions"},
        {"program":"UX/UI Design Principles","budget":1800,"source":"External","timing":"Q2","impact":"Improve user experience"},
        {"program":"API Integration & Platform Architecture","budget":1500,"source":"Online","timing":"Q2-Q3","impact":"Better system integration"},
        {"program":"Product Roadmap & Prioritization","budget":1200,"source":"Internal","timing":"Q3","impact":"Strategic product planning"},
    ],
    "البيانات والذكاء": [
        {"program":"Advanced Data Analytics & ML","budget":2500,"source":"External","timing":"Q1","impact":"Predictive business insights"},
        {"program":"Power BI / Tableau Mastery","budget":1800,"source":"Online","timing":"Q1-Q2","impact":"Self-service analytics"},
        {"program":"Data Engineering & Pipeline Design","budget":1500,"source":"Online","timing":"Q2","impact":"Scalable data infrastructure"},
        {"program":"AI/ML for Business Applications","budget":1200,"source":"External","timing":"Q3","impact":"AI-driven automation"},
    ],
    "المالية": [
        {"program":"Advanced Financial Modeling & Analysis","budget":1800,"source":"External","timing":"Q1","impact":"Better investment decisions"},
        {"program":"IFRS & Regulatory Compliance","budget":1400,"source":"External","timing":"Q2","impact":"Full regulatory compliance"},
        {"program":"Treasury & Cash Flow Management","budget":1000,"source":"Online","timing":"Q3","impact":"Optimize cash management"},
        {"program":"Budgeting & Forecasting Excellence","budget":800,"source":"Internal","timing":"Q1-Q2","impact":"Accurate financial planning"},
    ],
    "الموارد البشرية": [
        {"program":"Strategic HR Business Partnering","budget":1200,"source":"External","timing":"Q1","impact":"Align HR with business strategy"},
        {"program":"Talent Acquisition & Employer Branding","budget":1000,"source":"External","timing":"Q2","impact":"Attract top talent"},
        {"program":"Performance Management & OKRs","budget":1000,"source":"Online","timing":"Q1-Q3","impact":"Drive accountability"},
        {"program":"HR Analytics & People Insights","budget":800,"source":"Online","timing":"Q2","impact":"Data-driven HR decisions"},
    ],
    "الحوكمة": [
        {"program":"Enterprise Risk Management (ERM)","budget":1200,"source":"External","timing":"Q2","impact":"Proactive risk identification"},
        {"program":"Regulatory Compliance (SAMA/PCI-DSS)","budget":1000,"source":"Online","timing":"Q1","impact":"Ensure regulatory compliance"},
        {"program":"Information Security & Cybersecurity","budget":800,"source":"Online","timing":"Q1-Q3","impact":"Protect data & systems"},
    ],
    "القانونية": [
        {"program":"Digital Commercial Contracts & SLAs","budget":1000,"source":"External","timing":"Q2","impact":"Protect company interests"},
        {"program":"IP Protection & Trademark Law","budget":600,"source":"Online","timing":"Q3","impact":"Protect IP assets"},
        {"program":"E-Commerce & Payments Regulations","budget":400,"source":"Internal","timing":"Q1","impact":"Compliance with regulations"},
    ],
}

ROI_INDICATORS = [
    "Expected 15-20% increase in sales revenue within 12 months",
    "Improve customer retention rate by 10-15%",
    "Reduce CAC by 10% through improved marketing efficiency",
    "Increase CLV by 20% through customer success programs",
    "Improve proposal win rate by 15%",
]

TRAINING_KPIS = [
    "Training completion rate: 90%+ for all departments",
    "Trainee satisfaction score: 4.2/5 minimum",
    "Skills application within 30 days: 75%+ of trainees",
    "Professional certification pass rate: 85%+ first attempt",
    "Training hours per employee: 20 hours annually minimum",
]

# ===== SALARY BENCHMARK DATA (KSA + Egypt) =====
MARKET_BENCHMARKS = {
    "السعودية": {
        "currency": "SAR", "gosi_employer": 0.1175, "gosi_employee": 0.0975, "gosi_nsa_employer": 0.02,
        "housing_pct": 25, "transport": 500, "iqama_cost": 2400, "visa_cost": 3500, "med_insurance_avg": 500,
        "end_service_factor": 0.5, "saudization_required": True,
        "positions": {
            "الإدارة العليا": [
                {"title":"Chief Executive Officer (CEO)","title_ar":"الرئيس التنفيذي","min":45000,"mid":65000,"max":95000,"level":"C-Level","demand":"متوسط"},
                {"title":"Chief Technology Officer (CTO)","title_ar":"رئيس قطاع التقنية","min":35000,"mid":50000,"max":75000,"level":"C-Level","demand":"عالي"},
                {"title":"Chief Financial Officer (CFO)","title_ar":"المدير المالي","min":30000,"mid":45000,"max":65000,"level":"C-Level","demand":"متوسط"},
                {"title":"VP of Sales","title_ar":"نائب رئيس المبيعات","min":25000,"mid":38000,"max":55000,"level":"VP","demand":"عالي"},
                {"title":"VP of Marketing","title_ar":"نائب رئيس التسويق","min":22000,"mid":35000,"max":50000,"level":"VP","demand":"عالي"},
            ],
            "التقنية": [
                {"title":"Engineering Manager","title_ar":"مدير هندسة البرمجيات","min":20000,"mid":30000,"max":42000,"level":"Management","demand":"عالي جداً"},
                {"title":"Senior Software Engineer","title_ar":"مهندس برمجيات أول","min":15000,"mid":22000,"max":32000,"level":"Senior","demand":"عالي جداً"},
                {"title":"Software Engineer","title_ar":"مهندس برمجيات","min":10000,"mid":15000,"max":22000,"level":"Mid","demand":"عالي"},
                {"title":"Junior Developer","title_ar":"مطور مبتدئ","min":6000,"mid":8500,"max":12000,"level":"Junior","demand":"متوسط"},
                {"title":"DevOps Engineer","title_ar":"مهندس DevOps","min":14000,"mid":20000,"max":28000,"level":"Senior","demand":"عالي جداً"},
                {"title":"Data Engineer","title_ar":"مهندس بيانات","min":14000,"mid":20000,"max":28000,"level":"Senior","demand":"عالي جداً"},
                {"title":"Data Analyst","title_ar":"محلل بيانات","min":8000,"mid":12000,"max":18000,"level":"Mid","demand":"عالي"},
                {"title":"QA Engineer","title_ar":"مهندس ضمان الجودة","min":8000,"mid":12000,"max":17000,"level":"Mid","demand":"متوسط"},
                {"title":"UI/UX Designer","title_ar":"مصمم تجربة المستخدم","min":9000,"mid":14000,"max":20000,"level":"Mid","demand":"عالي"},
                {"title":"Product Manager","title_ar":"مدير المنتج","min":15000,"mid":22000,"max":32000,"level":"Senior","demand":"عالي جداً"},
            ],
            "المبيعات": [
                {"title":"Sales Director","title_ar":"مدير المبيعات","min":18000,"mid":25000,"max":38000,"level":"Director","demand":"عالي"},
                {"title":"Senior Account Manager","title_ar":"مدير حسابات أول","min":12000,"mid":16000,"max":22000,"level":"Senior","demand":"عالي"},
                {"title":"Account Manager","title_ar":"مدير حسابات","min":8000,"mid":11000,"max":16000,"level":"Mid","demand":"عالي"},
                {"title":"Sales Representative","title_ar":"مندوب مبيعات","min":5000,"mid":7500,"max":11000,"level":"Junior","demand":"متوسط"},
                {"title":"Business Development Manager","title_ar":"مدير تطوير الأعمال","min":14000,"mid":20000,"max":28000,"level":"Senior","demand":"عالي"},
            ],
            "التسويق": [
                {"title":"Marketing Director","title_ar":"مدير التسويق","min":16000,"mid":24000,"max":35000,"level":"Director","demand":"عالي"},
                {"title":"Digital Marketing Manager","title_ar":"مدير التسويق الرقمي","min":10000,"mid":15000,"max":22000,"level":"Senior","demand":"عالي"},
                {"title":"Content Creator","title_ar":"صانع محتوى","min":6000,"mid":9000,"max":14000,"level":"Mid","demand":"متوسط"},
                {"title":"SEO Specialist","title_ar":"أخصائي SEO","min":7000,"mid":10000,"max":15000,"level":"Mid","demand":"عالي"},
                {"title":"Graphic Designer","title_ar":"مصمم جرافيك","min":5000,"mid":8000,"max":13000,"level":"Mid","demand":"متوسط"},
            ],
            "الموارد البشرية": [
                {"title":"HR Director","title_ar":"مدير الموارد البشرية","min":16000,"mid":22000,"max":32000,"level":"Director","demand":"متوسط"},
                {"title":"HR Business Partner","title_ar":"شريك أعمال الموارد البشرية","min":10000,"mid":15000,"max":22000,"level":"Senior","demand":"عالي"},
                {"title":"Recruitment Specialist","title_ar":"أخصائي توظيف","min":7000,"mid":10000,"max":14000,"level":"Mid","demand":"متوسط"},
                {"title":"HR Coordinator","title_ar":"منسق موارد بشرية","min":5000,"mid":7000,"max":10000,"level":"Junior","demand":"متوسط"},
            ],
            "المالية": [
                {"title":"Finance Director","title_ar":"مدير مالي","min":18000,"mid":28000,"max":40000,"level":"Director","demand":"متوسط"},
                {"title":"Senior Accountant","title_ar":"محاسب أول","min":8000,"mid":12000,"max":17000,"level":"Senior","demand":"متوسط"},
                {"title":"Accountant","title_ar":"محاسب","min":5000,"mid":7500,"max":11000,"level":"Mid","demand":"منخفض"},
                {"title":"Financial Analyst","title_ar":"محلل مالي","min":9000,"mid":13000,"max":18000,"level":"Mid","demand":"عالي"},
            ],
            "العمليات": [
                {"title":"Operations Director","title_ar":"مدير العمليات","min":16000,"mid":24000,"max":35000,"level":"Director","demand":"متوسط"},
                {"title":"Project Manager","title_ar":"مدير مشاريع","min":12000,"mid":17000,"max":25000,"level":"Senior","demand":"عالي"},
                {"title":"Operations Coordinator","title_ar":"منسق عمليات","min":5500,"mid":8000,"max":12000,"level":"Junior","demand":"متوسط"},
                {"title":"Customer Service Manager","title_ar":"مدير خدمة العملاء","min":10000,"mid":14000,"max":20000,"level":"Senior","demand":"متوسط"},
            ],
        }
    },
    "مصر": {
        "currency": "EGP", "social_insurance_employer": 0.184, "social_insurance_employee": 0.11,
        "housing_pct": 0, "transport": 1000, "med_insurance_avg": 800,
        "end_service_factor": 0.5, "sar_to_egp": 13.2,
        "positions": {
            "الإدارة العليا": [
                {"title":"Chief Executive Officer","title_ar":"الرئيس التنفيذي","min":120000,"mid":180000,"max":300000,"level":"C-Level","demand":"متوسط"},
                {"title":"CTO","title_ar":"رئيس قطاع التقنية","min":80000,"mid":130000,"max":200000,"level":"C-Level","demand":"عالي"},
                {"title":"CFO","title_ar":"المدير المالي","min":70000,"mid":110000,"max":170000,"level":"C-Level","demand":"متوسط"},
            ],
            "التقنية": [
                {"title":"Engineering Manager","title_ar":"مدير هندسة البرمجيات","min":45000,"mid":70000,"max":100000,"level":"Management","demand":"عالي جداً"},
                {"title":"Senior Software Engineer","title_ar":"مهندس برمجيات أول","min":30000,"mid":50000,"max":80000,"level":"Senior","demand":"عالي جداً"},
                {"title":"Software Engineer","title_ar":"مهندس برمجيات","min":18000,"mid":28000,"max":45000,"level":"Mid","demand":"عالي"},
                {"title":"Junior Developer","title_ar":"مطور مبتدئ","min":8000,"mid":13000,"max":20000,"level":"Junior","demand":"متوسط"},
                {"title":"Data Analyst","title_ar":"محلل بيانات","min":12000,"mid":20000,"max":35000,"level":"Mid","demand":"عالي"},
                {"title":"UI/UX Designer","title_ar":"مصمم تجربة المستخدم","min":12000,"mid":22000,"max":35000,"level":"Mid","demand":"عالي"},
                {"title":"Product Manager","title_ar":"مدير المنتج","min":30000,"mid":50000,"max":75000,"level":"Senior","demand":"عالي جداً"},
            ],
            "المبيعات": [
                {"title":"Sales Director","title_ar":"مدير المبيعات","min":35000,"mid":55000,"max":85000,"level":"Director","demand":"عالي"},
                {"title":"Account Manager","title_ar":"مدير حسابات","min":12000,"mid":20000,"max":32000,"level":"Mid","demand":"عالي"},
                {"title":"Sales Representative","title_ar":"مندوب مبيعات","min":6000,"mid":10000,"max":16000,"level":"Junior","demand":"متوسط"},
                {"title":"BD Manager","title_ar":"مدير تطوير أعمال","min":25000,"mid":40000,"max":60000,"level":"Senior","demand":"عالي"},
            ],
            "التسويق": [
                {"title":"Marketing Director","title_ar":"مدير التسويق","min":30000,"mid":50000,"max":80000,"level":"Director","demand":"عالي"},
                {"title":"Digital Marketing Manager","title_ar":"مدير تسويق رقمي","min":15000,"mid":25000,"max":40000,"level":"Senior","demand":"عالي"},
                {"title":"Content Creator","title_ar":"صانع محتوى","min":7000,"mid":12000,"max":20000,"level":"Mid","demand":"متوسط"},
                {"title":"Graphic Designer","title_ar":"مصمم جرافيك","min":6000,"mid":10000,"max":18000,"level":"Mid","demand":"متوسط"},
            ],
            "الموارد البشرية": [
                {"title":"HR Director","title_ar":"مدير الموارد البشرية","min":28000,"mid":45000,"max":70000,"level":"Director","demand":"متوسط"},
                {"title":"HR Specialist","title_ar":"أخصائي موارد بشرية","min":8000,"mid":14000,"max":22000,"level":"Mid","demand":"متوسط"},
                {"title":"Recruiter","title_ar":"أخصائي توظيف","min":7000,"mid":12000,"max":18000,"level":"Mid","demand":"متوسط"},
            ],
            "المالية": [
                {"title":"Finance Director","title_ar":"مدير مالي","min":35000,"mid":55000,"max":85000,"level":"Director","demand":"متوسط"},
                {"title":"Senior Accountant","title_ar":"محاسب أول","min":10000,"mid":18000,"max":28000,"level":"Senior","demand":"متوسط"},
                {"title":"Accountant","title_ar":"محاسب","min":5000,"mid":9000,"max":15000,"level":"Mid","demand":"منخفض"},
            ],
        }
    }
}

BENCHMARK_SOURCES = [
    {"name":"Hays Salary Guide 2025 - GCC","url":"hays.com","region":"KSA/GCC"},
    {"name":"Bayt.com Salary Survey","url":"bayt.com","region":"MENA"},
    {"name":"GulfTalent Salary Report","url":"gulftalent.com","region":"GCC"},
    {"name":"Glassdoor Saudi Arabia","url":"glassdoor.com","region":"KSA"},
    {"name":"LinkedIn Salary Insights","url":"linkedin.com/salary","region":"Global"},
    {"name":"Robert Half Salary Guide","url":"roberthalf.ae","region":"GCC"},
    {"name":"Mercer Total Remuneration Survey","url":"mercer.com","region":"MEA"},
    {"name":"Wuzzuf Salary Explorer","url":"wuzzuf.net","region":"Egypt"},
    {"name":"Forasna Salary Data","url":"forasna.com","region":"Egypt"},
]

def calc_roi(budget, rev_inc_pct, current_rev, ret_pct, avg_sal, hc, prod_pct):
    rev_gain = current_rev * rev_inc_pct / 100
    ret_save = ret_pct / 100 * hc * avg_sal * 0.5
    prod_val = prod_pct / 100 * hc * avg_sal * 0.1
    total = rev_gain + ret_save + prod_val
    return {"rev":rev_gain,"ret":ret_save,"prod":prod_val,"total":total,
            "roi":((total-budget)/max(budget,1))*100,"bcr":total/max(budget,1),
            "payback":budget/max(total/12,1)}


# ===== ACCESS CONTROL SYSTEM =====
import hashlib

def hash_pw(pw):
    return hashlib.sha256(pw.encode()).hexdigest()

# Default users (can be managed in-app)
DEFAULT_USERS = {
    "admin": {"password": hash_pw("Rsl@Adm2026"), "role": "مدير", "name": "مدير النظام", "email": "HR@resal.me", "dept": "الإدارة", "sections": "all"},
    "analyst": {"password": hash_pw("Rsl@Anl2026"), "role": "محلل", "name": "محلل البيانات", "email": "", "dept": "التحليلات",
        "sections": "📊 التحليلات العامة,🎁 Total Rewards,👥 Headcount,🔍 التحليل العام,📤 التقارير والتصدير"},
    "viewer": {"password": hash_pw("Rsl@Vwr2026"), "role": "عارض", "name": "عارض", "email": "", "dept": "",
        "sections": "📊 التحليلات العامة,📤 التقارير والتصدير"},
    "emp1": {"password": hash_pw("Rsl@Emp2026"), "role": "موظف", "name": "أحمد محمد", "email": "", "dept": "تقنية المعلومات",
        "sections": "🧠 اختبارات الشخصية"},
    "emp2": {"password": hash_pw("Rsl@Emp2026"), "role": "موظف", "name": "سارة أحمد", "email": "", "dept": "الموارد البشرية",
        "sections": "🧠 اختبارات الشخصية"},
}

ROLE_DESCRIPTIONS = {
    "مدير": "وصول كامل لجميع الأقسام + إدارة المستخدمين",
    "محلل": "وصول للتحليلات والتقارير بدون إدارة المستخدمين",
    "موظف": "أداء الاختبارات المعيّنة + عرض نتائجه فقط",
    "عارض": "عرض التقارير فقط بدون تعديل",
}

ALL_SECTIONS = ["📊 التحليلات العامة","🎁 Total Rewards","👥 Headcount","⚖️ حاسبة المستحقات",
    "📚 التدريب والتطوير","🎯 التوظيف","🚀 Onboarding","📜 العقود","🤖 المستشار الذكي",
    "🏗️ التطوير المؤسسي OD","📈 التحليلات المتقدمة","🔍 التحليل العام","📝 الاستبيانات","🧠 اختبارات الشخصية","📤 التقارير والتصدير"]

# Email sending function
def send_test_email(to_email, emp_name, tests, deadline, assigned_by, app_url=""):
    """Send email notification about assigned tests - supports Gmail, Outlook, Yahoo, custom SMTP"""
    try:
        smtp_cfg = st.session_state.get('smtp_config', {})
        if not smtp_cfg.get('email'):
            smtp_cfg = load_smtp_config()
            if smtp_cfg:
                st.session_state.smtp_config = smtp_cfg
        if not smtp_cfg.get('server') or not smtp_cfg.get('email') or not smtp_cfg.get('password'):
            return False, "لم يتم تكوين إعدادات البريد الإلكتروني. اذهب إلى إدارة المستخدمين > إعدادات SMTP"

        sender_email = smtp_cfg['email']
        sender_name = smtp_cfg.get('sender_name', 'إدارة الموارد البشرية - رسال الود')

        tests_list = "\n".join([f"  - {t}" for t in tests])
        msg = MIMEMultipart('alternative')
        msg['From'] = f"{sender_name} <{sender_email}>"
        msg['To'] = to_email
        msg['Subject'] = f"تعيين اختبارات شخصية جديدة - {emp_name}"
        msg['X-Priority'] = '2'

        # Professional HTML email template
        app_link = app_url or smtp_cfg.get('app_url', '')
        btn_html = f'<div style="text-align:center;margin:20px 0;"><a href="{app_link}" style="display:inline-block;background:linear-gradient(135deg,#E36414,#E9C46A);color:white;padding:14px 40px;border-radius:8px;text-decoration:none;font-weight:bold;font-size:15px;">الدخول إلى المنصة وبدء الاختبارات</a></div>' if app_link else ''

        html_body = f"""<!DOCTYPE html>
<html dir="rtl" lang="ar">
<head><meta charset="UTF-8"></head>
<body style="font-family:'Segoe UI',Arial,sans-serif;direction:rtl;text-align:right;margin:0;padding:0;background:#f4f4f4;">
<div style="max-width:600px;margin:20px auto;background:white;border-radius:12px;overflow:hidden;box-shadow:0 4px 12px rgba(0,0,0,0.1);">
    <div style="background:linear-gradient(135deg,#0F4C5C,#1A1A2E);color:white;padding:24px;text-align:center;">
        <div style="background:rgba(255,255,255,0.15);width:60px;height:60px;border-radius:12px;display:inline-flex;align-items:center;justify-content:center;font-size:24px;font-weight:800;margin-bottom:10px;">HR</div>
        <h2 style="margin:8px 0 0;color:white;font-size:20px;">منصة تحليلات الموارد البشرية</h2>
        <p style="margin:4px 0 0;opacity:0.7;font-size:13px;">رسال الود لتقنية المعلومات</p>
    </div>
    <div style="padding:28px 24px;">
        <h3 style="color:#0F4C5C;margin:0 0 15px;font-size:18px;">مرحباً {emp_name},</h3>
        <p style="color:#333;line-height:1.7;font-size:14px;">تم تعيين اختبارات شخصية جديدة لك من قبل <strong style="color:#E36414;">{assigned_by}</strong>.</p>

        <div style="background:linear-gradient(135deg,#fff8f3,#fff);border-right:4px solid #E36414;padding:18px;margin:20px 0;border-radius:8px;">
            <p style="margin:0 0 12px;font-weight:bold;color:#E36414;font-size:15px;">📋 الاختبارات المطلوبة:</p>
            {"".join([f'<div style="margin:8px 0;padding:8px 15px;background:white;border-radius:6px;border:1px solid #f0f0f0;"><span style="color:#27AE60;margin-left:8px;">✅</span> <strong>{t}</strong></div>' for t in tests])}
        </div>

        <div style="background:#FFF3CD;border-radius:8px;padding:14px 18px;margin:15px 0;">
            <p style="margin:0;font-size:14px;">⏰ <strong>الموعد النهائي:</strong> <span style="color:#E74C3C;font-weight:bold;font-size:15px;">{deadline}</span></p>
        </div>

        <p style="color:#555;line-height:1.7;font-size:14px;">يرجى تسجيل الدخول إلى المنصة وإكمال الاختبارات المطلوبة قبل الموعد النهائي.</p>

        {btn_html}

        <hr style="border:none;border-top:1px solid #eee;margin:24px 0;">
        <p style="font-size:11px;color:#999;line-height:1.6;">هذه رسالة تلقائية من منصة تحليلات الموارد البشرية. الاختبارات إجبارية ولا يمكن تأجيلها بدون موافقة المدير المباشر.</p>
        <p style="font-size:11px;color:#bbb;">تاريخ الإرسال: {datetime.now().strftime('%Y-%m-%d %H:%M')}</p>
    </div>
</div>
</body></html>"""

        text_body = f"""مرحباً {emp_name},

تم تعيين اختبارات شخصية جديدة لك:
{tests_list}

الموعد النهائي: {deadline}
معيّن بواسطة: {assigned_by}

يرجى تسجيل الدخول إلى المنصة وإكمال الاختبارات المطلوبة.
{f'رابط المنصة: {app_link}' if app_link else ''}

هذه رسالة تلقائية من منصة تحليلات الموارد البشرية."""

        msg.attach(MIMEText(text_body, 'plain', 'utf-8'))
        msg.attach(MIMEText(html_body, 'html', 'utf-8'))

        port = int(smtp_cfg.get('port', 587))
        use_ssl = smtp_cfg.get('use_ssl', port == 465)

        if use_ssl:
            server = smtplib.SMTP_SSL(smtp_cfg['server'], port, timeout=30)
        else:
            server = smtplib.SMTP(smtp_cfg['server'], port, timeout=30)
            server.ehlo()
            server.starttls()
            server.ehlo()

        server.login(sender_email, smtp_cfg['password'])
        server.send_message(msg)
        server.quit()

        # Log successful send
        log_email_send(to_email, emp_name, tests, "success")
        return True, f"تم الإرسال بنجاح إلى {to_email}"
    except smtplib.SMTPAuthenticationError:
        return False, "خطأ في المصادقة: تأكد من البريد وكلمة مرور التطبيق (App Password). لـ Gmail: اذهب إلى myaccount.google.com/apppasswords"
    except smtplib.SMTPRecipientsRefused:
        return False, f"البريد المستلم غير صالح: {to_email}"
    except smtplib.SMTPConnectError:
        return False, f"لا يمكن الاتصال بالخادم {smtp_cfg.get('server','')}:{port}. تأكد من الخادم والمنفذ"
    except TimeoutError:
        return False, "انتهت مهلة الاتصال. تأكد من إعدادات الشبكة والخادم"
    except Exception as e:
        return False, f"خطأ: {str(e)}"

def log_email_send(to_email, emp_name, tests, status):
    """Log email sending to database"""
    try:
        conn = get_conn()
        c = conn.cursor()
        p = _ph()
        c.execute(f"INSERT INTO email_log (to_email, emp_name, tests, status, sent_at, sent_by) VALUES ({p},{p},{p},{p},{p},{p})",
            (to_email, emp_name, ", ".join(tests), status,
             datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
             st.session_state.get('user_name','')))
        conn.commit()
        conn.close()
    except: pass

def save_smtp_config(cfg):
    """Save SMTP config to database for persistence"""
    try:
        conn = get_conn()
        c = conn.cursor()
        _upsert_config(c, "smtp_config", json.dumps(cfg, ensure_ascii=False))
        conn.commit()
        conn.close()
    except: pass

def load_smtp_config():
    """Load SMTP config from database"""
    try:
        conn = get_conn()
        c = conn.cursor()
        p = _ph()
        c.execute(f"SELECT value FROM app_config WHERE key = {p}", ("smtp_config",))
        row = c.fetchone()
        conn.close()
        if row:
            return json.loads(row[0])
    except: pass
    return {}

def get_email_log():
    """Get email sending log"""
    try:
        conn = get_conn()
        c = conn.cursor()
        c.execute("SELECT * FROM email_log ORDER BY sent_at DESC LIMIT 50")
        rows = c.fetchall()
        cols = [d[0] for d in c.description]
        conn.close()
        return [dict(zip(cols, row)) for row in rows]
    except: return []

# SMTP Provider presets
SMTP_PROVIDERS = {
    "رسال الود (resal.me)": {"server": "smtp.resal.me", "port": 587, "use_ssl": False,
        "email": "HR@resal.me",
        "help": "البريد الرسمي للشركة. أدخل كلمة مرور حساب HR@resal.me أو تواصل مع مسؤول IT للحصول على كلمة مرور التطبيق"},
    "Gmail": {"server": "smtp.gmail.com", "port": 587, "use_ssl": False,
        "help": "استخدم App Password من: myaccount.google.com/apppasswords (تأكد من تفعيل المصادقة الثنائية أولاً)"},
    "Outlook/Hotmail": {"server": "smtp-mail.outlook.com", "port": 587, "use_ssl": False,
        "help": "استخدم كلمة مرور حسابك أو App Password"},
    "Yahoo": {"server": "smtp.mail.yahoo.com", "port": 465, "use_ssl": True,
        "help": "استخدم App Password من إعدادات أمان Yahoo"},
    "Office 365": {"server": "smtp.office365.com", "port": 587, "use_ssl": False,
        "help": "استخدم بيانات حساب Office 365"},
    "Zoho": {"server": "smtp.zoho.com", "port": 465, "use_ssl": True,
        "help": "استخدم بيانات حساب Zoho Mail"},
    "خادم مخصص": {"server": "", "port": 587, "use_ssl": False,
        "help": "أدخل إعدادات SMTP الخاصة بخادمك"},
}

def init_users():
    if 'users_db' not in st.session_state:
        if '_users_loaded' not in st.session_state:
            st.session_state._users_loaded = True
            cloud_users = db_load_users()
            if cloud_users:
                st.session_state.users_db = cloud_users
            else:
                st.session_state.users_db = DEFAULT_USERS.copy()
                db_save_users(DEFAULT_USERS)
        else:
            st.session_state.users_db = DEFAULT_USERS.copy()

def _save_login_token(username):
    """Save login token to database + query_params for persistence across reruns"""
    try:
        import hashlib
        token = hashlib.sha256(f"{username}_{datetime.now().strftime('%Y%m%d')}_{os.urandom(8).hex()}".encode()).hexdigest()[:32]
        conn = get_conn()
        c = conn.cursor()
        _upsert_config(c, f"login_token_{token}", json.dumps({"user": username, "time": datetime.now().strftime("%Y-%m-%d %H:%M:%S")}))
        conn.commit()
        conn.close()
        st.session_state['_login_token'] = token
        # Save token in URL query params - survives full page reloads
        try: st.query_params["tk"] = token
        except: pass
    except: pass

def _restore_login():
    """Try to restore login session from query_params or session_state"""
    # Skip if already restored this session
    if st.session_state.get('_login_restored'):
        return False
    st.session_state._login_restored = True
    try:
        token = None
        try: token = st.query_params.get("tk")
        except: pass
        if not token: token = st.session_state.get('_login_token')
        if not token: return False

        conn = get_conn()
        c = conn.cursor()
        c.execute(f"SELECT value FROM app_config WHERE key = {_ph()}", (f"login_token_{token}",))
        row = c.fetchone()
        conn.close()

        if row:
            data = json.loads(row[0])
            username = data.get("user")
            init_users()
            users = st.session_state.users_db
            if username in users:
                st.session_state.logged_in = True
                st.session_state.current_user = username
                st.session_state.user_role = users[username]["role"]
                st.session_state.user_name = users[username]["name"]
                st.session_state.user_sections = users[username]["sections"]
                st.session_state.user_email = users[username].get("email", "")
                audit_log("تسجيل دخول", f"المستخدم: {username}")
                st.session_state.user_dept = users[username].get("dept", "")
                st.session_state['_login_token'] = token
                return True
        try: del st.query_params["tk"]
        except: pass
    except: pass
    return False

def login_page():
    st.markdown("<div style='text-align:center;padding:40px 0;'><div style='background:linear-gradient(135deg,#E36414,#E9C46A);width:80px;height:80px;border-radius:16px;display:flex;align-items:center;justify-content:center;margin:0 auto 16px;font-size:32px;font-weight:800;color:white;'>HR</div><h1 style='color:#1A1A2E;'>منصة تحليلات الموارد البشرية</h1><p style='color:#64748B;'>رسال الود لتقنية المعلومات</p></div>", unsafe_allow_html=True)
    col1, col2, col3 = st.columns([1,2,1])
    with col2:
        login_tab, forgot_tab = st.tabs(["🔐 تسجيل الدخول", "🔑 استرجاع كلمة السر"])

        with login_tab:
            with st.form("login_form", clear_on_submit=False):
                username = st.text_input("اسم المستخدم:", key="login_user")
                password = st.text_input("كلمة المرور:", type="password", key="login_pass")
                lc1, lc2 = st.columns(2)
                with lc1:
                    login_btn = st.form_submit_button("🔓 دخول", type="primary", use_container_width=True)
                with lc2:
                    guest_btn = st.form_submit_button("👤 دخول بدون حساب", use_container_width=True)

                if login_btn:
                    init_users()
                    users = st.session_state.users_db
                    if username in users and users[username]["password"] == hash_pw(password):
                        st.session_state.logged_in = True
                        st.session_state.current_user = username
                        st.session_state.user_role = users[username]["role"]
                        st.session_state.user_name = users[username]["name"]
                        st.session_state.user_sections = users[username]["sections"]
                        st.session_state.user_email = users[username].get("email", "")
                        st.session_state.user_dept = users[username].get("dept", "")
                        _save_login_token(username)
                        st.rerun()
                    else:
                        st.error("❌ اسم المستخدم أو كلمة المرور غير صحيحة")

                if guest_btn:
                    st.session_state.logged_in = True
                    st.session_state.current_user = "guest"
                    st.session_state.user_role = "عارض"
                    st.session_state.user_name = "زائر"
                    st.session_state.user_sections = "all"
                    st.session_state.user_email = ""
                    st.session_state.user_dept = ""
                    st.rerun()

        with forgot_tab:
            st.markdown("#### 🔑 استرجاع كلمة المرور عبر البريد الإلكتروني")
            with st.form("forgot_form", clear_on_submit=False):
                reset_email = st.text_input("أدخل بريدك الإلكتروني المسجل:", key="reset_email", placeholder="you@company.com")
                reset_btn = st.form_submit_button("📧 إرسال رابط الاسترجاع", type="primary", use_container_width=True)

                if reset_btn and reset_email:
                    init_users()
                    users = st.session_state.users_db
                    # Find user by email
                    found_user = None
                    for uname, udata in users.items():
                        if udata.get('email','').lower().strip() == reset_email.lower().strip():
                            found_user = uname
                            break

                    if found_user:
                        # Generate temp password
                        import random, string
                        temp_pass = ''.join(random.choices(string.ascii_letters + string.digits, k=8))
                        # Update password
                        st.session_state.users_db[found_user]['password'] = hash_pw(temp_pass)
                        db_save_users(st.session_state.users_db)
                        # Send email
                        smtp_cfg = st.session_state.get('smtp_config', {})
                        if not smtp_cfg.get('email'):
                            smtp_cfg = load_smtp_config()
                        if smtp_cfg.get('email'):
                            ok, msg = send_test_email(reset_email, users[found_user]['name'],
                                [f"اسم المستخدم: {found_user}",
                                 f"كلمة المرور المؤقتة: {temp_pass}",
                                 "يرجى تغيير كلمة المرور فور تسجيل الدخول"],
                                datetime.now().strftime('%Y-%m-%d'),
                                smtp_cfg.get('sender_name', 'HR System'))
                            if ok:
                                st.success(f"✅ تم إرسال كلمة مرور مؤقتة إلى {reset_email}")
                            else:
                                st.warning(f"⚠️ لم يتم الإرسال: {msg}")
                                st.info(f"كلمة المرور المؤقتة: **{temp_pass}** (احفظها الآن)")
                        else:
                            st.info(f"⚠️ إعدادات البريد غير مكوّنة. كلمة المرور المؤقتة: **{temp_pass}**")
                    else:
                        st.error("❌ لم يتم العثور على حساب بهذا البريد الإلكتروني")

        st.markdown("---")
        with st.expander("📋 الحسابات الافتراضية"):
            st.markdown("| المستخدم | كلمة المرور | الدور |")
            st.markdown("|---|---|---|")
            st.markdown("| admin | Rsl@Adm2026 | مدير |")
            st.markdown("| analyst | Rsl@Anl2026 | محلل |")
            st.markdown("| viewer | Rsl@Vwr2026 | عارض |")
            st.markdown("| emp1 | Rsl@Emp2026 | موظف (أحمد محمد) |")
            st.markdown("| emp2 | Rsl@Emp2026 | موظف (سارة أحمد) |")

def check_section_access(section_name):
    if not st.session_state.get('logged_in'): return False
    user_sections = st.session_state.get('user_sections', 'all')
    if user_sections == "all": return True
    return section_name in user_sections

def user_management_page():
    hdr("👥 إدارة المستخدمين والصلاحيات", "إضافة وتعديل المستخدمين وصلاحياتهم")
    init_users()

    if st.session_state.get('user_role') != "مدير":
        st.warning("⚠️ هذه الصفحة متاحة للمدير فقط")
        return

    # Current users
    st.markdown("### 📋 المستخدمين الحاليين")
    users = st.session_state.users_db
    user_rows = []
    for uname, udata in users.items():
        user_rows.append({"المستخدم": uname, "الاسم": udata["name"], "الدور": udata["role"],
            "البريد": udata.get("email",""), "القسم": udata.get("dept",""),
            "الأقسام": "جميع الأقسام" if udata["sections"]=="all" else udata["sections"]})
    st.dataframe(pd.DataFrame(user_rows), use_container_width=True, hide_index=True)

    # SMTP Configuration
    st.markdown("### 📧 إعدادات البريد الإلكتروني (SMTP)")
    st.caption("مطلوب لإرسال إشعارات تعيين الاختبارات من التطبيق مباشرة إلى بريد الموظف")

    # Load saved config
    if 'smtp_config' not in st.session_state:
        st.session_state.smtp_config = load_smtp_config()
    smtp_cfg = st.session_state.get('smtp_config', {})

    with st.expander("⚙️ تكوين SMTP", expanded=not smtp_cfg.get('email')):
        # Provider selection
        provider = st.selectbox("اختر مزود البريد:", list(SMTP_PROVIDERS.keys()),
            index=0, key="smtp_provider")
        prov = SMTP_PROVIDERS[provider]
        st.info(f"💡 {prov['help']}")

        sm1, sm2 = st.columns(2)
        with sm1:
            smtp_server = st.text_input("خادم SMTP:", value=smtp_cfg.get('server', prov['server']), key="smtp_srv")
            smtp_port = st.number_input("المنفذ:", value=int(smtp_cfg.get('port', prov['port'])),
                min_value=1, max_value=65535, key="smtp_port")
            use_ssl = st.checkbox("استخدام SSL (بدلاً من TLS)", value=smtp_cfg.get('use_ssl', prov.get('use_ssl', False)), key="smtp_ssl")
        with sm2:
            default_email = smtp_cfg.get('email', '') or prov.get('email', '')
            smtp_email = st.text_input("البريد المرسل:", value=default_email, key="smtp_email",
                placeholder="HR@resal.me")
            smtp_pass = st.text_input("كلمة مرور التطبيق (App Password):", type="password",
                value=smtp_cfg.get('password',''), key="smtp_pass")
            sender_name = st.text_input("اسم المرسل:", value=smtp_cfg.get('sender_name','إدارة الموارد البشرية - رسال الود'), key="smtp_sender")

        app_url = st.text_input("رابط التطبيق (اختياري - يظهر في الإيميل):",
            value=smtp_cfg.get('app_url',''), key="smtp_url",
            placeholder="https://your-app.streamlit.app")

        sv1, sv2, sv3 = st.columns(3)
        with sv1:
            if st.button("💾 حفظ الإعدادات", type="primary", key="smtp_save", use_container_width=True):
                new_cfg = {'server': smtp_server, 'port': smtp_port, 'email': smtp_email,
                    'password': smtp_pass, 'sender_name': sender_name, 'app_url': app_url,
                    'use_ssl': use_ssl, 'provider': provider}
                st.session_state.smtp_config = new_cfg
                save_smtp_config(new_cfg)
                st.success("✅ تم حفظ إعدادات البريد في قاعدة البيانات")

        with sv2:
            if st.button("🧪 إرسال بريد تجريبي", key="smtp_test", use_container_width=True):
                if smtp_email and smtp_pass:
                    test_cfg = {'server': smtp_server, 'port': smtp_port, 'email': smtp_email,
                        'password': smtp_pass, 'sender_name': sender_name, 'app_url': app_url, 'use_ssl': use_ssl}
                    st.session_state.smtp_config = test_cfg
                    ok, result_msg = send_test_email(smtp_email, "اختبار النظام",
                        ["بريد تجريبي للتأكد من عمل الإعدادات"], str(date.today()),
                        st.session_state.get('user_name','المدير'))
                    if ok:
                        st.success(f"✅ {result_msg}")
                    else:
                        st.error(f"❌ {result_msg}")
                else:
                    st.error("أدخل البريد وكلمة المرور أولاً")

        with sv3:
            if smtp_cfg.get('email'):
                st.success(f"✅ مُكوّن: {smtp_cfg['email']}")
            else:
                st.warning("⚠️ غير مُكوّن")

    # Email sending log
    with st.expander("📋 سجل الإرسال"):
        email_logs = get_email_log()
        if email_logs:
            log_df = pd.DataFrame(email_logs)
            if 'id' in log_df.columns:
                log_df = log_df.drop(columns=['id'])
            st.dataframe(log_df, use_container_width=True, hide_index=True)
        else:
            st.caption("لا توجد رسائل مرسلة بعد")

    # Add new user
    st.markdown("### ➕ إضافة مستخدم جديد")
    uc1, uc2 = st.columns(2)
    with uc1:
        new_user = st.text_input("اسم المستخدم:", key="nu_user")
        new_pass = st.text_input("كلمة المرور:", type="password", key="nu_pass")
        new_name = st.text_input("الاسم الكامل:", key="nu_name")
        new_email = st.text_input("البريد الإلكتروني:", key="nu_email", placeholder="example@company.com")
        new_dept = st.text_input("القسم:", key="nu_dept")
    with uc2:
        new_role = st.selectbox("الدور:", list(ROLE_DESCRIPTIONS.keys()), key="nu_role")
        st.info(f"📋 {ROLE_DESCRIPTIONS[new_role]}")
        if new_role == "مدير":
            new_sections = "all"
        elif new_role == "موظف":
            new_sections = "🧠 اختبارات الشخصية"
            st.caption("📌 الموظف يحصل تلقائياً على صلاحية اختبارات الشخصية")
        else:
            new_sections_list = st.multiselect("الأقسام المتاحة:", ALL_SECTIONS, default=ALL_SECTIONS[:3], key="nu_sec")
            new_sections = ",".join(new_sections_list) if new_sections_list else "all"

    if st.button("➕ إضافة المستخدم", type="primary", key="nu_btn"):
        if new_user and new_pass and new_name:
            st.session_state.users_db[new_user] = {
                "password": hash_pw(new_pass), "role": new_role,
                "name": new_name, "email": new_email, "dept": new_dept, "sections": new_sections}
            db_save_users(st.session_state.users_db)
            st.success(f"✅ تم إضافة {new_name} بدور {new_role}")
            # Send welcome email with credentials
            if new_email:
                smtp_cfg = st.session_state.get('smtp_config', {})
                if not smtp_cfg.get('email'): smtp_cfg = load_smtp_config()
                if smtp_cfg.get('email'):
                    ok, msg = send_test_email(new_email, new_name,
                        [f"تم إنشاء حسابك في منصة تحليلات الموارد البشرية",
                         f"اسم المستخدم: {new_user}",
                         f"كلمة المرور: {new_pass}",
                         f"الدور: {new_role}",
                         "يرجى تغيير كلمة المرور فور تسجيل الدخول الأول"],
                        datetime.now().strftime('%Y-%m-%d'),
                        smtp_cfg.get('sender_name', 'HR'))
                    if ok: st.success(f"📧 تم إرسال بيانات الحساب إلى {new_email}")
                    else: st.warning(f"⚠️ لم يتم إرسال البريد: {msg}")
            st.rerun()
        else:
            st.error("يرجى تعبئة الاسم واسم المستخدم وكلمة المرور على الأقل")

    # Edit user email
    st.markdown("### ✏️ تعديل بريد مستخدم")
    edit_user = st.selectbox("اختر المستخدم:", list(users.keys()), key="edit_u")
    if edit_user:
        cur_email = users[edit_user].get("email","")
        cur_dept = users[edit_user].get("dept","")
        ec1, ec2 = st.columns(2)
        with ec1: upd_email = st.text_input("البريد الجديد:", value=cur_email, key="upd_email")
        with ec2: upd_dept = st.text_input("القسم:", value=cur_dept, key="upd_dept")
        if st.button("💾 تحديث", key="upd_btn"):
            st.session_state.users_db[edit_user]["email"] = upd_email
            st.session_state.users_db[edit_user]["dept"] = upd_dept
            db_save_users(st.session_state.users_db)
            st.success(f"✅ تم تحديث بيانات {edit_user}")
            st.rerun()

    # Delete user
    st.markdown("### 🗑️ حذف مستخدم")
    del_user = st.selectbox("اختر المستخدم:", [u for u in users.keys() if u != st.session_state.current_user], key="del_u")
    if st.button("🗑️ حذف", key="del_btn"):
        if del_user in st.session_state.users_db:
            del st.session_state.users_db[del_user]
            db_save_users(st.session_state.users_db)
            st.success(f"✅ تم حذف {del_user}")
            st.rerun()

# ===== SURVEY TEMPLATES =====
SURVEY_TEMPLATES = {
    "رضا الموظفين": {
        "description": "استبيان شامل لقياس مستوى رضا الموظفين عن بيئة العمل",
        "questions": [
            {"q": "أشعر بالرضا عن عملي بشكل عام", "cat": "الرضا العام"},
            {"q": "أحصل على تقدير كافٍ لإنجازاتي", "cat": "التقدير"},
            {"q": "لدي فرص كافية للتطور المهني", "cat": "التطور"},
            {"q": "العلاقة مع مديري المباشر جيدة", "cat": "الإدارة"},
            {"q": "بيئة العمل مريحة ومحفزة", "cat": "بيئة العمل"},
            {"q": "الراتب والمزايا عادلة مقارنة بالسوق", "cat": "التعويضات"},
            {"q": "أشعر بالانتماء للشركة", "cat": "الانتماء"},
            {"q": "التواصل الداخلي في الشركة فعّال", "cat": "التواصل"},
            {"q": "لدي توازن جيد بين العمل والحياة الشخصية", "cat": "التوازن"},
            {"q": "أوصي بالعمل في هذه الشركة للآخرين", "cat": "التوصية"},
        ]
    },
    "بيئة العمل": {
        "description": "تقييم بيئة العمل المادية والتنظيمية",
        "questions": [
            {"q": "المكتب والمرافق مجهزة بشكل جيد", "cat": "المرافق"},
            {"q": "الأدوات والتقنيات المتاحة كافية لأداء العمل", "cat": "الأدوات"},
            {"q": "إجراءات السلامة المهنية مطبقة", "cat": "السلامة"},
            {"q": "ساعات العمل مناسبة", "cat": "ساعات العمل"},
            {"q": "الإضاءة والتهوية مناسبة", "cat": "البيئة المادية"},
            {"q": "مساحة العمل كافية ومريحة", "cat": "المساحة"},
            {"q": "الضوضاء في بيئة العمل مقبولة", "cat": "البيئة المادية"},
            {"q": "خدمات الطعام والمشروبات متاحة", "cat": "الخدمات"},
        ]
    },
    "المشاركة والالتزام": {
        "description": "قياس مستوى مشاركة الموظفين والتزامهم التنظيمي",
        "questions": [
            {"q": "أبذل جهداً إضافياً عندما يتطلب العمل ذلك", "cat": "الالتزام"},
            {"q": "أشعر بالحماس تجاه عملي اليومي", "cat": "الحماس"},
            {"q": "أفهم أهداف الشركة وأساهم في تحقيقها", "cat": "التوافق"},
            {"q": "أشارك بفعالية في اجتماعات الفريق", "cat": "المشاركة"},
            {"q": "أقدم أفكاراً ومقترحات لتحسين العمل", "cat": "المبادرة"},
            {"q": "أشعر أن عملي له قيمة وتأثير", "cat": "القيمة"},
            {"q": "أتعاون بشكل جيد مع زملائي", "cat": "التعاون"},
            {"q": "أفتخر بالعمل في هذه الشركة", "cat": "الفخر"},
        ]
    }
}

# ===== PERSONALITY ASSESSMENTS - COMPREHENSIVE =====

# --- Big Five (OCEAN) - 25 Questions ---
BIG5_QUESTIONS = [
    {"q": "أستمتع بالتفاعل مع مجموعات كبيرة من الناس", "trait": "E", "d": 1},
    {"q": "أبادر ببدء المحادثات مع الغرباء", "trait": "E", "d": 1},
    {"q": "أفضل العمل بمفردي على العمل الجماعي", "trait": "E", "d": -1},
    {"q": "أشعر بالطاقة والحيوية في الأماكن الاجتماعية", "trait": "E", "d": 1},
    {"q": "أميل للهدوء والتأمل أكثر من الحديث", "trait": "E", "d": -1},
    {"q": "أهتم بمشاعر الآخرين وأتعاطف معهم بعمق", "trait": "A", "d": 1},
    {"q": "أثق في نوايا الآخرين بسهولة", "trait": "A", "d": 1},
    {"q": "أسعى لمساعدة الآخرين حتى لو لم يطلبوا", "trait": "A", "d": 1},
    {"q": "أتجنب الصراعات وأفضل التوافق", "trait": "A", "d": 1},
    {"q": "أجد صعوبة في رفض طلبات الآخرين", "trait": "A", "d": 1},
    {"q": "أنظم مهامي وأخطط مسبقاً بعناية فائقة", "trait": "C", "d": 1},
    {"q": "ألتزم بالمواعيد النهائية دائماً بلا استثناء", "trait": "C", "d": 1},
    {"q": "أهتم بالتفاصيل الدقيقة في عملي", "trait": "C", "d": 1},
    {"q": "أضع أهدافاً واضحة وأتابع تحقيقها بانتظام", "trait": "C", "d": 1},
    {"q": "أميل للتسويف وتأجيل المهام أحياناً", "trait": "C", "d": -1},
    {"q": "أشعر بالقلق أو التوتر بسهولة في المواقف الضاغطة", "trait": "N", "d": 1},
    {"q": "تتقلب مشاعري بشكل كبير خلال اليوم", "trait": "N", "d": 1},
    {"q": "أجد صعوبة في التعامل مع النقد الموجه لي", "trait": "N", "d": 1},
    {"q": "أميل للتفكير السلبي في المواقف الغامضة", "trait": "N", "d": 1},
    {"q": "أتعافى بسرعة من المواقف المحبطة", "trait": "N", "d": -1},
    {"q": "أحب تجربة أشياء جديدة وغير مألوفة", "trait": "O", "d": 1},
    {"q": "أستمتع بالأفكار المجردة والنظريات المعقدة", "trait": "O", "d": 1},
    {"q": "أقدّر الفن والجمال والإبداع بعمق", "trait": "O", "d": 1},
    {"q": "أفضل الروتين والأساليب المجربة والمضمونة", "trait": "O", "d": -1},
    {"q": "أبحث دائماً عن طرق جديدة لحل المشكلات", "trait": "O", "d": 1},
]

BIG5_TRAITS = {
    "E": {"name": "الانبساطية", "en": "Extraversion", "color": "#E36414",
        "desc": "مستوى الطاقة الاجتماعية والحماس والتفاعل مع الآخرين",
        "high": "اجتماعي، نشيط، متحمس، يحب العمل الجماعي، قيادي بالتأثير",
        "low": "هادئ، متأمل، يفضل العمل الفردي، يركز بعمق، مستقل",
        "jobs_high": "المبيعات، العلاقات العامة، التدريب، إدارة الفرق",
        "jobs_low": "التحليل، البرمجة، البحث، المحاسبة"},
    "A": {"name": "القبول", "en": "Agreeableness", "color": "#2D6A4F",
        "desc": "مستوى التعاون والثقة والتعاطف مع الآخرين",
        "high": "متعاون، متسامح، يثق بالآخرين، يتجنب الصراع، داعم",
        "low": "تنافسي، مباشر، ناقد، يتحدى الأفكار، مستقل الرأي",
        "jobs_high": "خدمة العملاء، الموارد البشرية، التمريض، الإرشاد",
        "jobs_low": "المحاماة، التفاوض، إدارة المخاطر، مراقبة الجودة"},
    "C": {"name": "الإتقان", "en": "Conscientiousness", "color": "#0F4C5C",
        "desc": "مستوى التنظيم والانضباط والمسؤولية والإنجاز",
        "high": "منظم، منضبط، مسؤول، يركز على التفاصيل، موثوق",
        "low": "مرن، عفوي، يتكيف بسرعة، أقل تنظيماً، مبدع",
        "jobs_high": "إدارة المشاريع، المحاسبة، الجودة، العمليات",
        "jobs_low": "الإبداع، الفنون، ريادة الأعمال، البحث الاستكشافي"},
    "N": {"name": "العصابية", "en": "Neuroticism", "color": "#9A031E",
        "desc": "مستوى الحساسية العاطفية والاستجابة للضغوط",
        "high": "حساس، متيقظ للمخاطر، عاطفي، يحتاج دعم في الضغوط",
        "low": "هادئ تحت الضغط، مستقر عاطفياً، متزن، مرن نفسياً",
        "jobs_high": "أدوار إبداعية تحتاج حساسية، الكتابة، الفنون",
        "jobs_low": "إدارة الأزمات، الطب، القيادة العليا، التفاوض"},
    "O": {"name": "الانفتاح", "en": "Openness", "color": "#7209B7",
        "desc": "مستوى حب الاستكشاف والإبداع والفضول الفكري",
        "high": "مبدع، فضولي، يحب التجديد، يتقبل التغيير، مبتكر",
        "low": "عملي، تقليدي، يفضل المألوف، واقعي، ثابت",
        "jobs_high": "التصميم، الابتكار، الاستراتيجية، البحث والتطوير",
        "jobs_low": "العمليات الروتينية، المالية، الإدارة التنفيذية"},
}

# --- Thomas PPA (Personal Profile Analysis) - 24 Questions ---
THOMAS_QUESTIONS = [
    {"q": "أتخذ قراراتي بسرعة وحزم ولا أتردد", "scale": "D", "d": 1},
    {"q": "أسعى للسيطرة على المواقف وتوجيه النتائج", "scale": "D", "d": 1},
    {"q": "أحب التحديات الصعبة وأزدهر تحت الضغط", "scale": "D", "d": 1},
    {"q": "أواجه المشكلات مباشرة ولا أتجنبها", "scale": "D", "d": 1},
    {"q": "أحب المنافسة وأسعى للفوز دائماً", "scale": "D", "d": 1},
    {"q": "لا أرتاح حتى أحقق الأهداف المحددة", "scale": "D", "d": 1},
    {"q": "أؤثر في الآخرين بالإقناع والحماس", "scale": "I", "d": 1},
    {"q": "أبني علاقات جديدة بسهولة وسرعة", "scale": "I", "d": 1},
    {"q": "أحب البيئة الاجتماعية المرحة والتفاعلية", "scale": "I", "d": 1},
    {"q": "أعبّر عن مشاعري وأفكاري بوضوح أمام الجميع", "scale": "I", "d": 1},
    {"q": "أجيد تحفيز الآخرين ورفع معنوياتهم", "scale": "I", "d": 1},
    {"q": "أحب أن أكون محور الاهتمام والتقدير", "scale": "I", "d": 1},
    {"q": "أفضل بيئة العمل المستقرة والمتوقعة", "scale": "S", "d": 1},
    {"q": "أصبر على المهام الطويلة والمتكررة بدون ملل", "scale": "S", "d": 1},
    {"q": "أستمع للآخرين بعناية قبل أن أبدي رأيي", "scale": "S", "d": 1},
    {"q": "أفضل التغيير التدريجي المدروس على المفاجئ", "scale": "S", "d": 1},
    {"q": "أدعم فريقي وأضع مصلحة المجموعة أولاً", "scale": "S", "d": 1},
    {"q": "أحافظ على هدوئي حتى في المواقف المتوترة", "scale": "S", "d": 1},
    {"q": "ألتزم بالمعايير والإجراءات بدقة متناهية", "scale": "C", "d": 1},
    {"q": "أراجع عملي عدة مرات للتأكد من خلوه من الأخطاء", "scale": "C", "d": 1},
    {"q": "أفضل اتخاذ القرارات بناءً على بيانات وحقائق", "scale": "C", "d": 1},
    {"q": "أسأل أسئلة كثيرة قبل البدء في أي مهمة جديدة", "scale": "C", "d": 1},
    {"q": "أجد صعوبة في تقبل العمل غير المنظم", "scale": "C", "d": 1},
    {"q": "أحلل المخاطر بعناية قبل اتخاذ أي خطوة", "scale": "C", "d": 1},
]

THOMAS_SCALES = {
    "D": {"name": "الهيمنة", "en": "Dominance", "color": "#E74C3C",
        "desc": "القدرة على السيطرة والتوجيه وتحقيق النتائج",
        "high": "حاسم، طموح، تنافسي، مباشر، يركز على النتائج، يتحمل الضغط",
        "low": "تعاوني، يفضل التوافق، يتجنب المواجهة، حذر، صبور",
        "role": "القيادة التنفيذية، إدارة المشاريع الكبرى، ريادة الأعمال",
        "manage": "أعطه صلاحيات واضحة، تحديات مستمرة، نتائج قابلة للقياس"},
    "I": {"name": "التأثير", "en": "Influence", "color": "#F39C12",
        "desc": "القدرة على التواصل والإقناع وبناء العلاقات",
        "high": "اجتماعي، متفائل، ملهم، مقنع، يبني علاقات بسرعة، متحمس",
        "low": "تحليلي، يفضل الحقائق، متحفظ اجتماعياً، واقعي، موضوعي",
        "role": "المبيعات، العلاقات العامة، التسويق، التدريب، خدمة العملاء",
        "manage": "أعطه تقديراً علنياً، مهام تفاعلية، فرص للتواصل والعرض"},
    "S": {"name": "الثبات", "en": "Steadiness", "color": "#27AE60",
        "desc": "الاستقرار والصبر والدعم والعمل الجماعي",
        "high": "صبور، موثوق، داعم، يعمل بإيقاع ثابت، مستمع ممتاز، وفي",
        "low": "سريع الإيقاع، مرن، يتكيف بسرعة، متعدد المهام، لا يحب الروتين",
        "role": "الموارد البشرية، الدعم الفني، العمليات المستقرة، خدمة العملاء",
        "manage": "وفر بيئة مستقرة، أشرك في القرارات، أعلن التغييرات مبكراً"},
    "C": {"name": "الامتثال", "en": "Compliance", "color": "#2980B9",
        "desc": "الالتزام بالمعايير والجودة والدقة والتحليل",
        "high": "دقيق، تحليلي، منهجي، يلتزم بالقواعد، يهتم بالجودة، حذر",
        "low": "مستقل، يتخذ مخاطر محسوبة، مبدع، يكسر القوالب، مرن",
        "role": "المالية، المراجعة، ضمان الجودة، تحليل البيانات، البرمجة",
        "manage": "أعطه معايير واضحة، وقت كافٍ للتحليل، بيئة منظمة"},
}

# --- Hogan HPI (Personality Inventory) - 28 Questions ---
HOGAN_QUESTIONS = [
    {"q": "أتعامل مع الضغوط والإحباطات بهدوء واتزان", "scale": "ADJ", "d": 1},
    {"q": "لا أقلق كثيراً بشأن أخطاء الماضي", "scale": "ADJ", "d": 1},
    {"q": "أحافظ على تفاؤلي حتى في الأوقات الصعبة", "scale": "ADJ", "d": 1},
    {"q": "أتعافى بسرعة من الانتقادات والفشل", "scale": "ADJ", "d": 1},
    {"q": "أسعى لتحقيق أهداف طموحة باستمرار", "scale": "AMB", "d": 1},
    {"q": "أحب أن أكون في موقع قيادة وسلطة", "scale": "AMB", "d": 1},
    {"q": "أبادر بأخذ زمام المبادرة في المشاريع", "scale": "AMB", "d": 1},
    {"q": "أتحمس لتحقيق النتائج والتفوق على الأهداف", "scale": "AMB", "d": 1},
    {"q": "أستمتع بالمناسبات الاجتماعية والتجمعات", "scale": "SOC", "d": 1},
    {"q": "أجد سهولة في بدء محادثات مع أشخاص جدد", "scale": "SOC", "d": 1},
    {"q": "أحب أن أكون مركز الاهتمام في المجموعة", "scale": "SOC", "d": 1},
    {"q": "أشعر بالنشاط عندما أعمل مع فريق كبير", "scale": "SOC", "d": 1},
    {"q": "أنتبه لمشاعر الآخرين وأراعي حساسياتهم", "scale": "INT", "d": 1},
    {"q": "أحافظ على علاقات مهنية قوية ومستدامة", "scale": "INT", "d": 1},
    {"q": "أتجنب إحراج الآخرين أو جرح مشاعرهم", "scale": "INT", "d": 1},
    {"q": "أجيد قراءة الإشارات غير اللفظية للآخرين", "scale": "INT", "d": 1},
    {"q": "ألتزم بالأخلاقيات المهنية بشكل صارم", "scale": "PRU", "d": 1},
    {"q": "أخطط مسبقاً وأنظم جدول أعمالي بعناية", "scale": "PRU", "d": 1},
    {"q": "أوفي بوعودي والتزاماتي دائماً", "scale": "PRU", "d": 1},
    {"q": "أتبع القواعد والسياسات المعمول بها", "scale": "PRU", "d": 1},
    {"q": "أحب استكشاف أفكار ومفاهيم جديدة ومعقدة", "scale": "INQ", "d": 1},
    {"q": "أبحث عن حلول إبداعية بدلاً من التقليدية", "scale": "INQ", "d": 1},
    {"q": "أستمتع بفهم الأنظمة المعقدة وكيف تعمل", "scale": "INQ", "d": 1},
    {"q": "أقرأ كثيراً وأتعلم من مصادر متنوعة", "scale": "INQ", "d": 1},
    {"q": "أستمتع بالتعلم الأكاديمي والبحث المنهجي", "scale": "LRN", "d": 1},
    {"q": "أبقى على اطلاع بآخر التطورات في مجالي", "scale": "LRN", "d": 1},
    {"q": "أفضل التعلم من التجربة العملية على النظري", "scale": "LRN", "d": -1},
    {"q": "أسعى دائماً لتطوير مهاراتي ومعارفي", "scale": "LRN", "d": 1},
]

HOGAN_SCALES = {
    "ADJ": {"name": "التوازن النفسي", "en": "Adjustment", "color": "#1ABC9C",
        "desc": "القدرة على إدارة الضغوط والبقاء هادئاً ومتزناً",
        "high": "هادئ تحت الضغط، متفائل، لا ينزعج بسهولة، مرن نفسياً، مستقر عاطفياً",
        "low": "حساس للضغوط، ينتقد ذاته، قلق، يتأثر بالمواقف السلبية بسرعة",
        "impact": "يؤثر بشكل مباشر على الأداء تحت الضغط والقدرة على القيادة في الأزمات"},
    "AMB": {"name": "الطموح", "en": "Ambition", "color": "#E74C3C",
        "desc": "الرغبة في القيادة والإنجاز والتقدم المهني",
        "high": "طموح، تنافسي، يسعى للقيادة، مبادر، يحب التحديات، إنجازي",
        "low": "يفضل الدعم، لا يسعى للسلطة، تعاوني، يقبل التوجيه، متواضع",
        "impact": "يتنبأ بالقدرة القيادية والرغبة في تحمل المسؤوليات الكبيرة"},
    "SOC": {"name": "الاجتماعية", "en": "Sociability", "color": "#F39C12",
        "desc": "مدى الحاجة للتفاعل الاجتماعي والظهور",
        "high": "اجتماعي، ثرثار، يحب الاهتمام، نشيط اجتماعياً، يبني شبكات واسعة",
        "low": "متحفظ، يفضل العمق على الاتساع في العلاقات، يستمع أكثر",
        "impact": "يحدد نجاح الشخص في أدوار تتطلب تواصل مستمر وبناء علاقات"},
    "INT": {"name": "الحساسية الشخصية", "en": "Interpersonal Sensitivity", "color": "#9B59B6",
        "desc": "القدرة على فهم مشاعر الآخرين وبناء علاقات",
        "high": "لبق، دبلوماسي، يراعي المشاعر، يبني علاقات قوية، متعاطف",
        "low": "مباشر، صريح، لا يتأثر بالعواطف، موضوعي في الحكم",
        "impact": "يتنبأ بالقدرة على بناء فرق متماسكة وإدارة العلاقات الحساسة"},
    "PRU": {"name": "الحصافة", "en": "Prudence", "color": "#2980B9",
        "desc": "مستوى الالتزام والانضباط والنزاهة المهنية",
        "high": "منضبط، ملتزم، موثوق، يحترم القواعد، دقيق في المواعيد",
        "low": "مرن، عفوي، يكسر القوالب، مبتكر، أقل التزاماً بالتفاصيل",
        "impact": "يتنبأ بالموثوقية والنزاهة والالتزام التنظيمي على المدى الطويل"},
    "INQ": {"name": "الفضول", "en": "Inquisitive", "color": "#E67E22",
        "desc": "حب الاستكشاف والابتكار والتفكير الإبداعي",
        "high": "فضولي، مبتكر، يحب الأفكار الجديدة، مبدع، استراتيجي التفكير",
        "low": "عملي، واقعي، يركز على التنفيذ، يفضل المجرب والمثبت",
        "impact": "يتنبأ بالقدرة على الابتكار وحل المشكلات المعقدة والتفكير الاستراتيجي"},
    "LRN": {"name": "التوجه التعلمي", "en": "Learning Approach", "color": "#16A085",
        "desc": "الرغبة في التعلم المستمر والتطوير المعرفي",
        "high": "محب للتعلم، يسعى للمعرفة، يقرأ كثيراً، أكاديمي التوجه",
        "low": "يتعلم بالممارسة، عملي، يفضل التجربة على النظرية",
        "impact": "يتنبأ بسرعة التأقلم مع المتطلبات الجديدة والقدرة على النمو المهني"},
}

# --- MBTI - 32 Questions (8 per dimension) ---
MBTI_QUESTIONS = [
    {"q": "أستمد طاقتي من التفاعل مع الآخرين أكثر من الوحدة", "dim": "EI", "d": "E"},
    {"q": "أفكر وأنا أتحدث وأحب التفكير بصوت عالٍ", "dim": "EI", "d": "E"},
    {"q": "أحتاج وقتاً بمفردي لاستعادة طاقتي بعد التجمعات", "dim": "EI", "d": "I"},
    {"q": "أفضل العمل في بيئة هادئة بعيداً عن المشتتات", "dim": "EI", "d": "I"},
    {"q": "لدي دائرة واسعة من المعارف والأصدقاء", "dim": "EI", "d": "E"},
    {"q": "أفضل العلاقات العميقة مع عدد قليل من الأشخاص", "dim": "EI", "d": "I"},
    {"q": "أشارك أفكاري بسرعة في الاجتماعات", "dim": "EI", "d": "E"},
    {"q": "أفضل التفكير جيداً قبل مشاركة آرائي", "dim": "EI", "d": "I"},
    {"q": "أركز على الحقائق والتفاصيل الملموسة الحالية", "dim": "SN", "d": "S"},
    {"q": "أثق في خبرتي العملية أكثر من الحدس", "dim": "SN", "d": "S"},
    {"q": "أرى الصورة الكبيرة والأنماط والاحتمالات المستقبلية", "dim": "SN", "d": "N"},
    {"q": "أحب التفكير في المعاني الخفية وراء الأشياء", "dim": "SN", "d": "N"},
    {"q": "أفضل التعليمات خطوة بخطوة الواضحة والمحددة", "dim": "SN", "d": "S"},
    {"q": "أستمتع بتخيل السيناريوهات المستقبلية والبدائل", "dim": "SN", "d": "N"},
    {"q": "أهتم بما هو واقعي وقابل للتطبيق الآن", "dim": "SN", "d": "S"},
    {"q": "أميل للابتكار وإيجاد طرق جديدة تماماً", "dim": "SN", "d": "N"},
    {"q": "أتخذ قراراتي بناءً على المنطق والتحليل الموضوعي", "dim": "TF", "d": "T"},
    {"q": "أولي أهمية كبيرة للعدالة والمعايير الموحدة", "dim": "TF", "d": "T"},
    {"q": "أراعي تأثير قراراتي على مشاعر الأشخاص المعنيين", "dim": "TF", "d": "F"},
    {"q": "أسعى للانسجام في العلاقات وأتجنب الصراع", "dim": "TF", "d": "F"},
    {"q": "أنتقد الأفكار الضعيفة حتى لو أزعج ذلك أصحابها", "dim": "TF", "d": "T"},
    {"q": "أجيد فهم دوافع الآخرين والتعاطف معهم", "dim": "TF", "d": "F"},
    {"q": "الحق أهم من المشاعر عند اتخاذ القرارات المهمة", "dim": "TF", "d": "T"},
    {"q": "أعتبر القيم الشخصية والإنسانية فوق كل اعتبار", "dim": "TF", "d": "F"},
    {"q": "أحب إنهاء المهام والوصول لقرارات نهائية بسرعة", "dim": "JP", "d": "J"},
    {"q": "أخطط ليومي مسبقاً وألتزم بالجدول", "dim": "JP", "d": "J"},
    {"q": "أفضل إبقاء خياراتي مفتوحة لأطول فترة ممكنة", "dim": "JP", "d": "P"},
    {"q": "أتكيف بسهولة مع التغييرات المفاجئة في الخطط", "dim": "JP", "d": "P"},
    {"q": "أشعر بالراحة عندما تكون الأمور منظمة ومحسومة", "dim": "JP", "d": "J"},
    {"q": "أعمل بشكل أفضل تحت ضغط المواعيد النهائية", "dim": "JP", "d": "P"},
    {"q": "أفضل وضع خطة والالتزام بها حتى النهاية", "dim": "JP", "d": "J"},
    {"q": "أستمتع بالعفوية واغتنام الفرص غير المتوقعة", "dim": "JP", "d": "P"},
]

MBTI_TYPES = {
    "ISTJ": {"name": "المفتش", "desc": "مسؤول، دقيق، موثوق، يحترم التقاليد والأنظمة", "strengths": "التنظيم، الموثوقية، الالتزام، إدارة التفاصيل", "careers": "المحاسبة، المراجعة، إدارة العمليات، القانون، إدارة المشاريع"},
    "ISFJ": {"name": "الحامي", "desc": "مهتم، وفي، صبور، يحمي الآخرين ويدعمهم", "strengths": "الرعاية، الصبر، الاهتمام بالتفاصيل، الوفاء", "careers": "الموارد البشرية، التمريض، التعليم، خدمة العملاء، الإدارة"},
    "INFJ": {"name": "المستشار", "desc": "مثالي، ملهم، عميق التفكير، يسعى لإحداث فرق", "strengths": "الرؤية، التعاطف، التخطيط الاستراتيجي، الإلهام", "careers": "الاستشارات، التدريب، الكتابة، التخطيط الاستراتيجي"},
    "INTJ": {"name": "المهندس الاستراتيجي", "desc": "مبتكر، استراتيجي، مستقل، طموح بلا حدود", "strengths": "التخطيط الاستراتيجي، التحليل، الابتكار، الاستقلالية", "careers": "الاستشارات الإدارية، هندسة الأنظمة، التحليل الاستراتيجي، البحث"},
    "ISTP": {"name": "الحرفي", "desc": "عملي، تحليلي، يحل المشكلات بكفاءة وسرعة", "strengths": "حل المشكلات، المرونة، التحليل العملي، الكفاءة", "careers": "الهندسة، البرمجة، التحليل الفني، إدارة الأزمات"},
    "ISFP": {"name": "الفنان", "desc": "حساس، متعاطف، يقدر الجمال والانسجام", "strengths": "الإبداع، التعاطف، المرونة، الحس الجمالي", "careers": "التصميم، الفنون، الرعاية الصحية، العمل الاجتماعي"},
    "INFP": {"name": "المعالج", "desc": "مثالي، متعاطف، مبدع، يسعى للمعنى والهدف", "strengths": "الإبداع، التعاطف العميق، الكتابة، الرؤية", "careers": "الكتابة، الإرشاد النفسي، التعليم، التصميم، العمل غير الربحي"},
    "INTP": {"name": "المفكر", "desc": "تحليلي، منطقي، مبتكر، يحب الأفكار المعقدة", "strengths": "التحليل المنطقي، الابتكار، حل المشكلات المعقدة", "careers": "البحث العلمي، البرمجة، تحليل البيانات، الفلسفة، الهندسة"},
    "ESTP": {"name": "الريادي", "desc": "نشيط، عملي، مغامر، يتصرف بسرعة وفعالية", "strengths": "التصرف السريع، التفاوض، إدارة الأزمات، الواقعية", "careers": "ريادة الأعمال، المبيعات، التسويق، إدارة الأزمات"},
    "ESFP": {"name": "المؤدي", "desc": "حيوي، عفوي، اجتماعي، يستمتع بالحياة ويشع طاقة", "strengths": "التواصل، المرونة، رفع المعنويات، التعامل مع العملاء", "careers": "العلاقات العامة، التدريب، المبيعات، الضيافة، الترفيه"},
    "ENFP": {"name": "البطل", "desc": "متحمس، مبدع، ملهم، يرى الإمكانيات في كل مكان", "strengths": "الإلهام، الإبداع، التواصل، توليد الأفكار", "careers": "التسويق، الإبداع، ريادة الأعمال، التدريب، الاستشارات"},
    "ENTP": {"name": "المناظر", "desc": "ذكي، فضولي، يتحدى الأفكار التقليدية ويبتكر", "strengths": "الابتكار، التحليل النقدي، التفاوض، حل المشكلات", "careers": "ريادة الأعمال، الاستشارات، القانون، التسويق الاستراتيجي"},
    "ESTJ": {"name": "المشرف", "desc": "منظم، حازم، يدير الأمور بكفاءة واضحة", "strengths": "التنظيم، القيادة التنفيذية، الإدارة، وضع الأنظمة", "careers": "الإدارة التنفيذية، العمليات، المالية، إدارة المشاريع"},
    "ESFJ": {"name": "القنصل", "desc": "اجتماعي، داعم، يهتم بالآخرين ويبني المجتمعات", "strengths": "بناء العلاقات، الدعم، التنظيم الاجتماعي، الرعاية", "careers": "الموارد البشرية، التعليم، الرعاية الصحية، خدمة المجتمع"},
    "ENFJ": {"name": "المعلم", "desc": "ملهم، قيادي، يطور الآخرين ويحفزهم للتميز", "strengths": "القيادة التحفيزية، التواصل، التطوير، بناء الفرق", "careers": "إدارة الأفراد، التدريب، الاستشارات، القيادة التنظيمية"},
    "ENTJ": {"name": "القائد", "desc": "استراتيجي، حازم، طموح، يقود التغيير بقوة", "strengths": "القيادة الاستراتيجية، التخطيط، اتخاذ القرارات، التنفيذ", "careers": "الإدارة العليا، الاستشارات، ريادة الأعمال، إدارة التغيير"},
}

MBTI_DIMS = {
    "EI": {"E": "الانبساطية (E)", "I": "الانطوائية (I)", "E_desc": "يستمد طاقته من التفاعل الخارجي", "I_desc": "يستمد طاقته من التأمل الداخلي"},
    "SN": {"S": "الحسية (S)", "N": "الحدسية (N)", "S_desc": "يركز على الحقائق والواقع", "N_desc": "يركز على الأنماط والإمكانيات"},
    "TF": {"T": "التفكير (T)", "F": "الشعور (F)", "T_desc": "يقرر بالمنطق والتحليل", "F_desc": "يقرر بالقيم والمشاعر"},
    "JP": {"J": "الحكم (J)", "P": "الإدراك (P)", "J_desc": "يفضل النظام والتخطيط", "P_desc": "يفضل المرونة والانفتاح"},
}

# --- DISC Assessment - 24 Questions ---
DISC_QUESTIONS = [
    {"q": "أتخذ قرارات سريعة وحاسمة بدون تردد", "style": "D"},
    {"q": "أحب السيطرة على المواقف وتوجيه النتائج", "style": "D"},
    {"q": "أواجه التحديات بثقة وأزدهر تحت الضغط", "style": "D"},
    {"q": "أفضل الطريقة المباشرة في التعامل مع المشكلات", "style": "D"},
    {"q": "أسعى للفوز وتحقيق الأهداف بأي ثمن", "style": "D"},
    {"q": "لا أتحمل البطء أو عدم الكفاءة في العمل", "style": "D"},
    {"q": "أستمتع بإقناع الآخرين وكسب تأييدهم لأفكاري", "style": "I"},
    {"q": "أبني صداقات وعلاقات جديدة بسرعة وسهولة", "style": "I"},
    {"q": "أنا متفائل بطبعي وأرى الفرص في كل موقف", "style": "I"},
    {"q": "أحب بيئة العمل المرحة والتفاعلية والاجتماعية", "style": "I"},
    {"q": "أجيد التحدث أمام مجموعات والعرض والتقديم", "style": "I"},
    {"q": "أحب أن أكون محط الاهتمام والتقدير من الآخرين", "style": "I"},
    {"q": "أفضل الاستقرار وبيئة العمل المتوقعة والآمنة", "style": "S"},
    {"q": "أصبر على المهام الطويلة والمتكررة بدون ملل", "style": "S"},
    {"q": "أستمع بعناية واهتمام قبل أن أبدي رأيي", "style": "S"},
    {"q": "أدعم فريقي وأضع مصلحة المجموعة فوق مصلحتي", "style": "S"},
    {"q": "أفضل التغيير التدريجي المدروس على المفاجئ", "style": "S"},
    {"q": "أحافظ على هدوئي واتزاني في المواقف المتوترة", "style": "S"},
    {"q": "أهتم بالدقة والجودة في كل تفصيلة من عملي", "style": "C"},
    {"q": "أراجع عملي عدة مرات للتأكد من خلوه من الأخطاء", "style": "C"},
    {"q": "أفضل اتخاذ القرارات بناءً على بيانات وحقائق ملموسة", "style": "C"},
    {"q": "ألتزم بالقواعد والمعايير والإجراءات المحددة بصرامة", "style": "C"},
    {"q": "أحلل المخاطر والاحتمالات بعناية قبل أي خطوة", "style": "C"},
    {"q": "أجد صعوبة في تقبل العمل العشوائي وغير المنظم", "style": "C"},
]

DISC_STYLES = {
    "D": {"name": "الهيمنة", "en": "Dominance", "color": "#E74C3C",
        "desc": "حاسم، تنافسي، مباشر، يركز على النتائج والإنجاز",
        "high": "قائد طبيعي، يتخذ قرارات سريعة، يتحمل الضغط، طموح، مباشر، لا يخشى المواجهة",
        "low": "متعاون، يفضل التوافق، حذر في القرارات، صبور، يتجنب الصراع",
        "strengths": "اتخاذ القرارات الصعبة، حل المشكلات، القيادة في الأزمات، تحقيق النتائج",
        "challenges": "الصبر على الآخرين، الاستماع، التعاطف، تفويض المهام، المرونة",
        "careers": "الإدارة التنفيذية، ريادة الأعمال، إدارة المشاريع، المبيعات القيادية",
        "manage": "أعطه تحديات واضحة، صلاحيات كافية، نتائج قابلة للقياس، حرية في الأسلوب",
        "communicate": "كن مباشراً، ركز على النتائج، لا تطل في التفاصيل، قدم خيارات لا أوامر"},
    "I": {"name": "التأثير", "en": "Influence", "color": "#F39C12",
        "desc": "متحمس، اجتماعي، ملهم، يبني علاقات ويؤثر بالإقناع",
        "high": "متحدث بارع، متفائل، يحفز الفريق، يبني شبكات واسعة، مبدع اجتماعياً، ملهم",
        "low": "يفضل الحقائق على العواطف، متحفظ اجتماعياً، واقعي، يعمل بصمت",
        "strengths": "التواصل والإقناع، تحفيز الفريق، بناء العلاقات، العرض والتقديم، الإبداع",
        "challenges": "التنظيم والمتابعة، الالتزام بالتفاصيل، التركيز على مهمة واحدة، الموضوعية",
        "careers": "المبيعات، التسويق، العلاقات العامة، التدريب، خدمة العملاء VIP",
        "manage": "قدم تقديراً علنياً، مهام تفاعلية، فرص للعرض والتقديم، بيئة مرحة",
        "communicate": "كن ودوداً، شارك الحماس، استمع لأفكاره، لا تكن جافاً أو رسمياً جداً"},
    "S": {"name": "الثبات", "en": "Steadiness", "color": "#27AE60",
        "desc": "صبور، داعم، مستقر، موثوق، يعمل بإيقاع ثابت ومتزن",
        "high": "موثوق جداً، صبور، مستمع ممتاز، يدعم الفريق بإخلاص، وفي، يحب الاستقرار",
        "low": "سريع الإيقاع، متعدد المهام، يتكيف بسرعة مع التغيير، لا يحب الروتين",
        "strengths": "الاستماع الفعال، دعم الفريق، الاستقرار، الوفاء، إنجاز المهام الروتينية بتميز",
        "challenges": "التكيف مع التغييرات المفاجئة، المبادرة، اتخاذ قرارات سريعة، المواجهة",
        "careers": "الموارد البشرية، الدعم الفني، العمليات، خدمة العملاء، التمريض، التعليم",
        "manage": "وفر بيئة مستقرة، أشرك في القرارات مبكراً، أعلن التغييرات تدريجياً، قدر وفاءه",
        "communicate": "كن هادئاً وصادقاً، أعطه وقتاً للتفكير، لا تفاجئه بتغييرات، أظهر التقدير"},
    "C": {"name": "الالتزام", "en": "Conscientiousness", "color": "#2980B9",
        "desc": "دقيق، تحليلي، منهجي، يلتزم بالمعايير والجودة العالية",
        "high": "دقيق جداً، يحلل البيانات بعمق، يلتزم بالقواعد، يسعى للكمال، منظم، حذر",
        "low": "مرن، يأخذ مخاطر محسوبة، مبدع، يكسر القوالب، سريع التنفيذ",
        "strengths": "ضمان الجودة، التحليل الدقيق، الالتزام بالمعايير، حل المشكلات المعقدة، التوثيق",
        "challenges": "المرونة والسرعة، التواصل العاطفي، القبول بالحلول غير المثالية، التفويض",
        "careers": "المالية والمحاسبة، المراجعة، تحليل البيانات، البرمجة، ضمان الجودة، الهندسة",
        "manage": "أعطه معايير واضحة، وقت كافٍ للتحليل، بيئة منظمة، ملاحظات مبنية على حقائق",
        "communicate": "قدم بيانات وحقائق، كن منظماً في طرحك، لا تتسرع، احترم حاجته للدقة"},
}


# ===== MAIN APP =====
def main():
    # Initialize Enterprise Architecture
    if '_orchestrator' not in st.session_state:
        st.session_state._orchestrator = _init_orchestrator()
        st.session_state._knowledge_engine = _init_knowledge()
        st.session_state._learning_system = _init_learning()

    # Auth check with session persistence
    if 'logged_in' not in st.session_state:
        st.session_state.logged_in = False

    # Try restoring login if not logged in
    if not st.session_state.logged_in:
        if not _restore_login():
            login_page()
            return

    init_users()

    # Sidebar
    with st.sidebar:
        st.markdown(f"<div style='text-align:center;padding:20px 0;'><div style='background:linear-gradient(135deg,#E36414,#E9C46A);width:56px;height:56px;border-radius:12px;display:flex;align-items:center;justify-content:center;margin:0 auto 10px;font-size:22px;font-weight:800;color:white;'>HR</div><h2 style='margin:0;font-size:16px;'>تحليلات الموارد البشرية</h2><p style='opacity:.6;font-size:11px;'>رسال الود لتقنية المعلومات v5</p><div style='background:rgba(255,255,255,.1);border-radius:6px;padding:6px 10px;margin-top:8px;font-size:11px'>👤 {st.session_state.user_name} <span style='opacity:.6'>| {st.session_state.user_role}</span></div></div>", unsafe_allow_html=True)
        st.markdown("---")

        # Filter sections by access
        available_sections = [s for s in ALL_SECTIONS if check_section_access(s)]
        if st.session_state.user_role == "مدير":
            available_sections.append("👥 إدارة المستخدمين")


        section = st.radio("📂", available_sections, label_visibility="collapsed")
        st.markdown("---")

        if section == "📊 التحليلات العامة":
            page = st.radio("📌", ["🏠 نظرة عامة","📊 الأقسام","🤖 المحلل الذكي","📋 البيانات"], label_visibility="collapsed")
        elif section == "🎁 Total Rewards":
            page = st.radio("📌", ["🎁 لوحة Total Rewards","💰 لوحة الرواتب","📈 تحليل شهري/ربعي","🏷️ تحليل حسب الفئات","📊 سلم الرواتب","💰 هيكل الرواتب","🏥 المزايا والتأمينات","📊 تحليل التنافسية","📥 تصدير TR"], label_visibility="collapsed")
        elif section == "👥 Headcount":
            page = st.radio("📌", ["👥 Headcount Report","📊 تحليل الأداء","📋 بيانات الموظفين","📥 تصدير Headcount"], label_visibility="collapsed")
        elif section == "⚖️ حاسبة المستحقات":
            page = "⚖️ حاسبة المستحقات"
        elif section == "🎯 التوظيف":
            page = st.radio("📌", ["📋 تخطيط التوظيف","🤖 Benchmark ذكاء اصطناعي","🌍 مقارنة الأسواق","📊 متابعة التوظيف","📄 تحليل السير الذاتية","🎤 تحليل المقابلات","📋 ATS تتبع المتقدمين","📥 تصدير التوظيف"], label_visibility="collapsed")
        elif section == "🚀 Onboarding":
            page = st.radio("📌", ["🚀 إنشاء Onboarding","📋 خطة 30/60/90","👥 متابعة الموظفين الجدد","📊 تحليلات Onboarding","🎬 عرض تقديمي AI","🏢 معلومات الشركة","📥 تصدير Onboarding"], label_visibility="collapsed")
        elif section == "📜 العقود":
            page = st.radio("📌", ["📜 إنشاء عقد","🔍 تحليل العقود","📋 العقود المحفوظة","📥 تصدير العقود"], label_visibility="collapsed")
        elif section == "🤖 المستشار الذكي":
            page = st.radio("📌", ["⚖️ المستشار القانوني","📚 مستشار الموارد البشرية","🧠 قاعدة المعرفة RAG","📊 التعلم والتحسين","📋 إدارة المراجع"], label_visibility="collapsed")
        elif section == "🏗️ التطوير المؤسسي OD":
            page = st.radio("📌", ["🔍 تشخيص المنظمة","📊 تحليل OD","🎯 استراتيجية OD","📋 خطة التنفيذ","📥 تصدير OD"], label_visibility="collapsed")
        elif section == "📈 التحليلات المتقدمة":
            page = st.radio("📌", ["📊 مؤشرات HR المتقدمة","🔔 التنبيهات الذكية","🔮 سيناريوهات What-If","🤖 التحليل التنبؤي","💬 تحليل المشاعر","📋 سجل التدقيق"], label_visibility="collapsed")
        elif section == "🔍 التحليل العام":
            page = st.radio("📌", ["📊 تحليل تلقائي","🤖 أسئلة ذكية"], label_visibility="collapsed")
        elif section == "📝 الاستبيانات":
            page = st.radio("📌", ["📋 قوالب جاهزة","🔨 بناء استبيان","📊 تحليل النتائج","📥 تصدير الاستبيانات"], label_visibility="collapsed")
        elif section == "🧠 اختبارات الشخصية":
            page = st.radio("📌", ["📋 تعيين الاختبارات","🧠 Big Five (OCEAN)","📊 Thomas PPA","🔬 Hogan HPI","💡 MBTI","💎 DISC","📈 تقارير الشخصية","📥 تصدير الاختبارات"], label_visibility="collapsed")
        elif section == "📤 التقارير والتصدير":
            page = st.radio("📌", ["📊 تقرير Dashboard","📝 تقرير Word","📊 تقرير شامل"], label_visibility="collapsed")
        elif section == "👥 إدارة المستخدمين":
            page = "👥 إدارة المستخدمين"
        else:
            page = st.radio("📌", ["📚 ميزانية التدريب","💹 ROI التدريب","📋 خطة ADDIE","🏫 جهات التدريب","📥 تصدير التدريب"], label_visibility="collapsed")

        # Logout button
        st.markdown("---")
        if st.button("🚪 تسجيل الخروج", use_container_width=True):
            # Clear login token from DB
            token = st.session_state.get('_login_token')
            if token:
                try:
                    conn = get_conn()
                    c = conn.cursor()
                    c.execute(f"DELETE FROM app_config WHERE key = {_ph()}", (f"login_token_{token}",))
                    conn.commit()
                    conn.close()
                except: pass
            # Clear query_params token
            try: del st.query_params["tk"]
            except: pass
            for key in ['logged_in','current_user','user_role','user_name','user_sections','_login_token']:
                st.session_state.pop(key, None)
            st.rerun()

        st.markdown("---")
        st.markdown("##### 📁 ملف البيانات")
        file = st.file_uploader("ارفع Excel", type=["xlsx","xls","csv"], label_visibility="collapsed", key="main_uploader")
        if file:
            # Store file bytes in session_state + save to DB for persistence
            file_bytes_val = file.getvalue()
            st.session_state['uploaded_file_name'] = file.name
            st.session_state['uploaded_file_bytes'] = file_bytes_val
            # Save to cloud DB for cross-device persistence
            try:
                import base64
                if len(file_bytes_val) < 4_000_000:  # Only save files under 4MB to DB
                    encoded = base64.b64encode(file_bytes_val).decode('ascii')
                    conn = get_conn()
                    c = conn.cursor()
                    _upsert_config(c, "last_uploaded_file", json.dumps({"name": file.name, "size": len(file_bytes_val)}))
                    _upsert_config(c, "last_uploaded_data", encoded)
                    conn.commit()
                    conn.close()
            except: pass
            st.success(f"✅ {file.name}")
        elif 'uploaded_file_name' in st.session_state:
            audit_log("رفع ملف", f"{file.name}")
            st.info(f"📂 {st.session_state['uploaded_file_name']}")
            if st.button("🗑️ إزالة الملف", use_container_width=True):
                for k in ['uploaded_file_name','uploaded_file_bytes','_parsed_cache_key','_parsed_emp','_parsed_sal','_parsed_sheets']:
                    st.session_state.pop(k, None)
                # Also clear from DB
                try:
                    conn = get_conn()
                    c = conn.cursor()
                    c.execute(f"DELETE FROM app_config WHERE key IN ({_ph()},{_ph()})", ("last_uploaded_file","last_uploaded_data"))
                    conn.commit()
                    conn.close()
                except: pass
                st.rerun()
        else:
            # Try restoring file from cloud DB (only once)
            if '_cloud_file_checked' not in st.session_state:
                st.session_state._cloud_file_checked = True
                try:
                    import base64
                    conn = get_conn()
                    c = conn.cursor()
                    c.execute(f"SELECT value FROM app_config WHERE key = {_ph()}", ("last_uploaded_file",))
                    meta_row = c.fetchone()
                    if meta_row:
                        meta = json.loads(meta_row[0])
                        c.execute(f"SELECT value FROM app_config WHERE key = {_ph()}", ("last_uploaded_data",))
                        data_row = c.fetchone()
                        conn.close()
                        if data_row:
                            st.session_state['uploaded_file_name'] = meta['name']
                            st.session_state['uploaded_file_bytes'] = base64.b64decode(data_row[0])
                            st.info(f"📂 {meta['name']} (مسترجع من السحابة)")
                    else:
                        conn.close()
                except: pass


    # ===== LOAD DATA =====
    emp = pd.DataFrame()
    sal_df = pd.DataFrame()
    all_sheets = {}

    # Use session_state data if file not currently in uploader
    file_bytes = None
    file_name = None
    if file:
        file_bytes = file.getvalue()
        file_name = file.name
    elif 'uploaded_file_bytes' in st.session_state:
        file_bytes = st.session_state['uploaded_file_bytes']
        file_name = st.session_state.get('uploaded_file_name', 'data.xlsx')

    if file_bytes:
        # Check if we already parsed this file (same name + size)
        cache_key = f"{file_name}_{len(file_bytes)}"
        if st.session_state.get('_parsed_cache_key') == cache_key and '_parsed_emp' in st.session_state:
            emp = st.session_state['_parsed_emp']
            sal_df = st.session_state.get('_parsed_sal', pd.DataFrame())
            all_sheets = st.session_state.get('_parsed_sheets', {})
        else:
            try:
                if file_name.endswith('.csv'):
                    emp = norm_cols(pd.read_csv(io.BytesIO(file_bytes)))
                else:
                    xl = pd.ExcelFile(io.BytesIO(file_bytes))
                    for s in xl.sheet_names:
                        try:
                            df_s = smart_read(xl, s)
                            if len(df_s) > 500 and any(c.lower() in ['salary month','gross salary','شهر الراتب'] for c in df_s.columns):
                                sal_df = norm_cols(df_s)
                            df_s = norm_cols(df_s)
                            all_sheets[s] = df_s
                            if len(emp)==0 and len(df_s)>5:
                                name_cols = [c for c in df_s.columns if any(x in str(c).lower() for x in ['name','اسم','emp','موظف'])]
                                if name_cols: emp = df_s
                        except: pass
                    if len(emp)==0 and all_sheets: emp = list(all_sheets.values())[0]

                    if 'Salary Scale' in xl.sheet_names:
                        try: all_sheets['Salary Scale'] = pd.read_excel(xl, 'Salary Scale', header=0)
                        except: pass
                    if 'Positions' in xl.sheet_names:
                        try: all_sheets['Positions'] = pd.read_excel(xl, 'Positions', header=0)
                        except: pass

            except: pass

            # Cache parsed results
            st.session_state['_parsed_cache_key'] = cache_key
            st.session_state['_parsed_emp'] = emp
            st.session_state['_parsed_sal'] = sal_df
            st.session_state['_parsed_sheets'] = all_sheets

    if '#' in emp.columns and len(emp)>0:
        emp = emp[pd.to_numeric(emp['#'], errors='coerce').notna()].reset_index(drop=True)

    n = len(emp)

    # If salary data found, also create a snapshot (latest month)
    sal_snapshot = pd.DataFrame()
    if len(sal_df) > 0:
        if has(sal_df, 'سنة الراتب'):
            latest_year = sal_df['سنة الراتب'].max()
            yr_data = sal_df[sal_df['سنة الراتب']==latest_year]
            if has(yr_data, 'شهر الراتب'):
                months_order = ['January','February','March','April','May','June','July','August','September','October','November','December']
                available = yr_data['شهر الراتب'].unique()
                for m in reversed(months_order):
                    if m in available:
                        sal_snapshot = yr_data[yr_data['شهر الراتب']==m]
                        break
        if len(sal_snapshot)==0:
            sal_snapshot = sal_df.drop_duplicates(subset=['الاسم'] if has(sal_df,'الاسم') else sal_df.columns[0], keep='last')

    # Auto-save monthly snapshot
    if len(emp) > 0 and '_snapshot_saved' not in st.session_state:
        st.session_state._snapshot_saved = True
        save_snapshot(emp, "monthly")


    # =========================================
    #            📊 GENERAL ANALYTICS
    # =========================================
    if section == "📊 التحليلات العامة":
        if page == "🏠 نظرة عامة":
            hdr("📊 لوحة التحليلات الشاملة","تحليل متقدم وشامل لبيانات القوى العاملة")
            if n==0 and len(sal_df)==0:
                st.info("📁 ارفع ملف بيانات الموظفين أو ملف الرواتب من القائمة الجانبية")
                return

            data = sal_snapshot if len(sal_snapshot)>0 else emp
            total = len(data)

            # Auto-detect columns
            dept_col = next((c for c in data.columns if any(x in c.lower() for x in ['dept','department','قسم','القطاع'])), None)
            nat_col = next((c for c in data.columns if any(x in c.lower() for x in ['nat','جنسية','nationality'])), None)
            sal_col = next((c for c in data.select_dtypes('number').columns if any(x in c.lower() for x in ['gross','salary','net','راتب','إجمالي'])), None)
            status_col = next((c for c in data.columns if any(x in c.lower() for x in ['status','حالة'])), None)
            loc_col = next((c for c in data.columns if any(x in c.lower() for x in ['location','موقع','مدينة','city'])), None)
            gender_col = next((c for c in data.columns if any(x in c.lower() for x in ['gender','جنس'])), None)
            type_col = next((c for c in data.columns if any(x in c.lower() for x in ['employment type','نوع التوظيف','type'])), None)
            level_col = next((c for c in data.columns if any(x in c.lower() for x in ['level','مستوى','grade','درجة'])), None)
            age_col = next((c for c in data.columns if any(x in c.lower() for x in ['age','عمر','age group'])), None)
            join_col = next((c for c in data.columns if any(x in c.lower() for x in ['join','hiring','التحاق','مباشرة'])), None)

            # ===== ROW 1: KPIs =====
            sa_count = 0; sa_pct = 0; active_count = total
            if nat_col:
                sa_count = len(data[data[nat_col].isin(['Saudi','سعودي','Saudi Arabian','سعودية'])])
                sa_pct = round(sa_count/max(total,1)*100,1)
            if status_col:
                active_count = len(data[data[status_col].isin(['Active','نشط','active'])])

            k1,k2,k3,k4,k5,k6 = st.columns(6)
            with k1: kpi("👥 إجمالي السجلات", f"{total:,}")
            with k2: kpi("✅ نشط", f"{active_count:,}")
            with k3: kpi("🏢 الأقسام", str(data[dept_col].nunique()) if dept_col else "-")
            with k4: kpi("🇸🇦 السعودة", f"{sa_pct}%")
            with k5: kpi("💰 متوسط الراتب", f"{data[sal_col].mean():,.0f}" if sal_col else "-")
            with k6: kpi("💵 إجمالي الرواتب", f"{data[sal_col].sum():,.0f}" if sal_col else "-")

            st.markdown("---")

            # ===== ROW 2: Dept + Nationality =====
            r2c1, r2c2 = st.columns(2)
            with r2c1:
                if dept_col:
                    dc = data[dept_col].value_counts().head(15).reset_index()
                    dc.columns = [dept_col, 'العدد']
                    fig = px.bar(dc.sort_values('العدد'), x='العدد', y=dept_col, orientation='h',
                        color='العدد', color_continuous_scale='teal', title='📊 Headcount حسب القسم')
                    fig.update_layout(font=dict(family="Noto Sans Arabic"), height=420, showlegend=False, coloraxis_showscale=False)
                    st.plotly_chart(fig, use_container_width=True)
            with r2c2:
                if nat_col:
                    nc = data[nat_col].value_counts().head(10).reset_index()
                    nc.columns = [nat_col, 'العدد']
                    fig = px.pie(nc, values='العدد', names=nat_col, title='🌍 توزيع الجنسيات', hole=0.4,
                        color_discrete_sequence=px.colors.qualitative.Set2)
                    fig.update_layout(font=dict(family="Noto Sans Arabic"), height=420)
                    st.plotly_chart(fig, use_container_width=True)

            # ===== ROW 3: Status + Location =====
            r3c1, r3c2 = st.columns(2)
            with r3c1:
                if status_col:
                    sc = data[status_col].value_counts().reset_index()
                    sc.columns = [status_col, 'العدد']
                    fig = px.pie(sc, values='العدد', names=status_col, title='📋 حالة التوظيف', hole=0.5,
                        color_discrete_sequence=['#27AE60','#E74C3C','#F39C12','#3498DB','#95A5A6'])
                    fig.update_layout(font=dict(family="Noto Sans Arabic"), height=380)
                    st.plotly_chart(fig, use_container_width=True)
                elif type_col:
                    tc = data[type_col].value_counts().reset_index()
                    tc.columns = [type_col, 'العدد']
                    fig = px.pie(tc, values='العدد', names=type_col, title='📋 نوع التوظيف', hole=0.5,
                        color_discrete_sequence=CL['dept'])
                    fig.update_layout(font=dict(family="Noto Sans Arabic"), height=380)
                    st.plotly_chart(fig, use_container_width=True)
            with r3c2:
                if loc_col:
                    lc = data[loc_col].value_counts().head(10).reset_index()
                    lc.columns = [loc_col, 'العدد']
                    fig = px.bar(lc, x=loc_col, y='العدد', title='📍 التوزيع الجغرافي', color='العدد',
                        color_continuous_scale='teal')
                    fig.update_layout(font=dict(family="Noto Sans Arabic"), height=380, showlegend=False, coloraxis_showscale=False)
                    st.plotly_chart(fig, use_container_width=True)

            # ===== ROW 4: Salary Distribution + Gender =====
            r4c1, r4c2 = st.columns(2)
            with r4c1:
                if sal_col:
                    fig = px.histogram(data, x=sal_col, nbins=25, title='💰 توزيع الرواتب',
                        color_discrete_sequence=['#0F4C5C'])
                    fig.add_vline(x=data[sal_col].mean(), line_dash="dash", line_color="red",
                        annotation_text=f"المتوسط: {data[sal_col].mean():,.0f}")
                    fig.add_vline(x=data[sal_col].median(), line_dash="dot", line_color="blue",
                        annotation_text=f"الوسيط: {data[sal_col].median():,.0f}")
                    fig.update_layout(font=dict(family="Noto Sans Arabic"), height=380)
                    st.plotly_chart(fig, use_container_width=True)
            with r4c2:
                if gender_col:
                    gc = data[gender_col].value_counts().reset_index()
                    gc.columns = [gender_col, 'العدد']
                    fig = px.pie(gc, values='العدد', names=gender_col, title='👫 التوزيع حسب الجنس', hole=0.5,
                        color_discrete_map={'Male':'#3498DB','Female':'#E91E8F','ذكر':'#3498DB','أنثى':'#E91E8F'})
                    fig.update_layout(font=dict(family="Noto Sans Arabic"), height=380)
                    st.plotly_chart(fig, use_container_width=True)
                elif type_col and not (status_col):
                    tc = data[type_col].value_counts().reset_index()
                    tc.columns = [type_col, 'العدد']
                    fig = px.bar(tc, x=type_col, y='العدد', title='📋 نوع التوظيف', color=type_col,
                        color_discrete_sequence=CL['dept'])
                    fig.update_layout(font=dict(family="Noto Sans Arabic"), height=380, showlegend=False)
                    st.plotly_chart(fig, use_container_width=True)

            # ===== ROW 5: Salary by Department + Saudization by Dept =====
            r5c1, r5c2 = st.columns(2)
            with r5c1:
                if sal_col and dept_col:
                    top_depts = data[dept_col].value_counts().head(10).index
                    fig = px.box(data[data[dept_col].isin(top_depts)], x=dept_col, y=sal_col,
                        title='💰 نطاق الرواتب حسب القسم', color_discrete_sequence=['#E9C46A'])
                    fig.update_layout(font=dict(family="Noto Sans Arabic"), height=420, xaxis_tickangle=-45)
                    st.plotly_chart(fig, use_container_width=True)
            with r5c2:
                if nat_col and dept_col:
                    saudi_vals = ['Saudi','سعودي','Saudi Arabian','سعودية']
                    sa_dept = data.groupby(dept_col).apply(lambda x: round(len(x[x[nat_col].isin(saudi_vals)])/max(len(x),1)*100,1)).reset_index()
                    sa_dept.columns = [dept_col, 'نسبة السعودة %']
                    sa_dept = sa_dept.sort_values('نسبة السعودة %', ascending=True)
                    fig = px.bar(sa_dept, x='نسبة السعودة %', y=dept_col, orientation='h',
                        title='🇸🇦 نسبة السعودة حسب القسم', color='نسبة السعودة %',
                        color_continuous_scale='RdYlGn', range_color=[0,100])
                    fig.add_vline(x=50, line_dash="dash", line_color="red", annotation_text="الحد الأدنى 50%")
                    fig.update_layout(font=dict(family="Noto Sans Arabic"), height=420, coloraxis_showscale=False)
                    st.plotly_chart(fig, use_container_width=True)

            # ===== ROW 6: Age/Level + Employment Type =====
            r6c1, r6c2 = st.columns(2)
            with r6c1:
                if age_col:
                    age_data = data[age_col].value_counts().sort_index().reset_index()
                    age_data.columns = [age_col, 'العدد']
                    fig = px.bar(age_data, x=age_col, y='العدد', title='📊 التوزيع العمري',
                        color='العدد', color_continuous_scale='viridis')
                    fig.update_layout(font=dict(family="Noto Sans Arabic"), height=380, showlegend=False, coloraxis_showscale=False)
                    st.plotly_chart(fig, use_container_width=True)
                elif level_col:
                    lv_data = data[level_col].value_counts().reset_index()
                    lv_data.columns = [level_col, 'العدد']
                    fig = px.funnel(lv_data, x='العدد', y=level_col, title='📊 التوزيع حسب المستوى الوظيفي',
                        color_discrete_sequence=['#0F4C5C'])
                    fig.update_layout(font=dict(family="Noto Sans Arabic"), height=380)
                    st.plotly_chart(fig, use_container_width=True)
            with r6c2:
                if sal_col and nat_col:
                    fig = px.violin(data, y=sal_col, x=nat_col if data[nat_col].nunique() <= 8 else None,
                        title='📊 توزيع الرواتب (Violin)', color_discrete_sequence=['#2A9D8F'], box=True)
                    fig.update_layout(font=dict(family="Noto Sans Arabic"), height=380)
                    st.plotly_chart(fig, use_container_width=True)
                elif sal_col:
                    fig = px.violin(data, y=sal_col, title='📊 توزيع الرواتب (Violin)',
                        color_discrete_sequence=['#2A9D8F'], box=True, points='all')
                    fig.update_layout(font=dict(family="Noto Sans Arabic"), height=380)
                    st.plotly_chart(fig, use_container_width=True)

            # ===== ROW 7: Hiring Trend + Treemap =====
            r7c1, r7c2 = st.columns(2)
            with r7c1:
                if join_col:
                    data['_join_dt'] = pd.to_datetime(data[join_col], errors='coerce')
                    valid = data.dropna(subset=['_join_dt'])
                    if len(valid) > 0:
                        valid['_join_month'] = valid['_join_dt'].dt.to_period('M').astype(str)
                        monthly = valid.groupby('_join_month').size().reset_index(name='العدد')
                        fig = px.area(monthly, x='_join_month', y='العدد', title='📈 اتجاه التوظيف الشهري',
                            color_discrete_sequence=['#0F4C5C'])
                        fig.update_layout(font=dict(family="Noto Sans Arabic"), height=380, xaxis_title="الشهر")
                        st.plotly_chart(fig, use_container_width=True)
            with r7c2:
                if dept_col and nat_col:
                    tree = data.groupby([dept_col, nat_col]).size().reset_index(name='العدد')
                    tree = tree[tree['العدد'] > 0]
                    fig = px.treemap(tree, path=[dept_col, nat_col], values='العدد',
                        title='🗺️ خريطة شجرية: الأقسام × الجنسيات',
                        color='العدد', color_continuous_scale='teal')
                    fig.update_layout(font=dict(family="Noto Sans Arabic"), height=420)
                    st.plotly_chart(fig, use_container_width=True)
                elif dept_col and loc_col:
                    tree = data.groupby([dept_col, loc_col]).size().reset_index(name='العدد')
                    tree = tree[tree['العدد'] > 0]
                    fig = px.treemap(tree, path=[dept_col, loc_col], values='العدد',
                        title='🗺️ خريطة شجرية: الأقسام × المواقع',
                        color='العدد', color_continuous_scale='teal')
                    fig.update_layout(font=dict(family="Noto Sans Arabic"), height=420)
                    st.plotly_chart(fig, use_container_width=True)

            # ===== ROW 8: Heatmap + Scatter =====
            if sal_col and dept_col:
                r8c1, r8c2 = st.columns(2)
                with r8c1:
                    if nat_col:
                        hm = data.groupby([dept_col, nat_col]).size().unstack(fill_value=0)
                        fig = px.imshow(hm, title='🔥 خريطة حرارية: الأقسام × الجنسيات',
                            color_continuous_scale='YlOrRd', aspect='auto')
                        fig.update_layout(font=dict(family="Noto Sans Arabic"), height=420)
                        st.plotly_chart(fig, use_container_width=True)
                with r8c2:
                    dept_stats = data.groupby(dept_col).agg(
                        count=(sal_col,'count'), avg_sal=(sal_col,'mean'), total=(sal_col,'sum')
                    ).reset_index()
                    dept_stats['cost_pct'] = dept_stats['total'] / dept_stats['total'].sum() * 100
                    dept_stats['head_pct'] = dept_stats['count'] / dept_stats['count'].sum() * 100
                    fig = px.scatter(dept_stats, x='head_pct', y='cost_pct', size='count', color=dept_col,
                        title='⚖️ نسبة العدد مقابل التكلفة', hover_data=['avg_sal'],
                        color_discrete_sequence=CL['dept'])
                    fig.add_trace(go.Scatter(x=[0,50], y=[0,50], mode='lines',
                        line=dict(dash='dash', color='gray'), name='خط التوازن', showlegend=False))
                    fig.update_layout(font=dict(family="Noto Sans Arabic"), height=420,
                        xaxis_title="% من العدد", yaxis_title="% من التكلفة")
                    st.plotly_chart(fig, use_container_width=True)



        elif page == "🤖 المحلل الذكي":
            hdr("🤖 المحلل الذكي","يبحث في كل الأوراق")
            data = sal_snapshot if len(sal_snapshot)>0 else emp
            if len(data)==0: st.info("📁 ارفع ملف"); return
            q = st.text_input("💬 اكتب سؤالك:", placeholder="ما نسبة السعودة؟ كم عدد الأقسام؟")
            if st.button("🔍 تحليل",type="primary",use_container_width=True) and q:
                ql = q.lower()
                a = ""
                total = len(data)
                if any(w in ql for w in ['سعود','جنسي','national','saudi']):
                    if has(data,'الجنسية'):
                        sa = data[data['الجنسية'].isin(['Saudi','سعودي','سعودية'])]
                        a = f"نسبة السعودة: {round(len(sa)/total*100,1)}% ({len(sa)} من {total})\n\n"
                        for nat,cnt in data['الجنسية'].value_counts().items():
                            a += f"  - {nat}: {cnt} ({round(cnt/total*100,1)}%)\n"
                    else: a = "لا يوجد عمود جنسية. أضف Nationality أو الجنسية للملف."
                elif any(w in ql for w in ['قسم','أقسام','department','division']):
                    dc = data['القسم'].value_counts() if has(data,'القسم') else (data['القطاع'].value_counts() if has(data,'القطاع') else None)
                    if dc is not None:
                        a = f"عدد الأقسام: {len(dc)}\n\n"
                        for d,c in dc.items(): a += f"  - {d}: {c} ({round(c/total*100,1)}%)\n"
                elif any(w in ql for w in ['راتب','رواتب','salary','تكلف']):
                    if has(data,'الراتب الإجمالي'):
                        a = f"إجمالي الرواتب الشهرية: {data['الراتب الإجمالي'].sum():,.0f} ريال\nمتوسط: {data['الراتب الإجمالي'].mean():,.0f}\nالأعلى: {data['الراتب الإجمالي'].max():,.0f}\nالأقل: {data['الراتب الإجمالي'].min():,.0f}"
                    elif has(data,'الراتب الأساسي'):
                        a = f"متوسط الراتب الأساسي: {data['الراتب الأساسي'].mean():,.0f} ريال"
                    else: a = "لا يوجد بيانات رواتب."
                else:
                    a = f"الموظفين: {total}\n"
                    for c in data.columns[:10]:
                        if data[c].dtype == 'object': a += f"{c}: {data[c].nunique()} قيمة فريدة\n"
                    a += f"\nالأعمدة: {', '.join(data.columns[:15])}"
                st.info(a if a else "جرب سؤال آخر")

        elif page == "📋 البيانات":
            hdr("📋 البيانات")
            if not all_sheets and n==0: st.info("📁 ارفع ملف"); return
            if all_sheets:
                sn = st.selectbox("الورقة:", list(all_sheets.keys()))
                st.dataframe(all_sheets[sn], use_container_width=True, hide_index=True, height=600)
            elif n>0:
                st.dataframe(emp, use_container_width=True, hide_index=True, height=600)


    # =========================================
    #           💰 SALARY ANALYSIS
    # =========================================
    elif section == "🎁 Total Rewards":

        data = sal_df if len(sal_df)>0 else emp
        snap = sal_snapshot if len(sal_snapshot)>0 else (data if len(data)>0 else pd.DataFrame())

        # Auto-detect columns for Total Rewards
        sal_col_tr = next((c for c in snap.select_dtypes('number').columns if any(x in c.lower() for x in ['gross','إجمالي','total sal','net'])), None) if len(snap)>0 else None
        basic_col_tr = next((c for c in snap.select_dtypes('number').columns if any(x in c.lower() for x in ['basic','أساسي','base'])), None) if len(snap)>0 else None
        dept_col_tr = next((c for c in snap.columns if any(x in c.lower() for x in ['dept','قسم','department','القطاع'])), None) if len(snap)>0 else None

        if page == "💰 لوحة الرواتب":
            hdr("💰 لوحة تحليل الرواتب والتعويضات","تحليل شامل لتكاليف الرواتب والبدلات والاستقطاعات")

            if len(snap) == 0:
                st.info("📁 ارفع ملف الرواتب من القائمة الجانبية"); return

            # Auto-detect all salary columns
            sal_col = next((c for c in ['الراتب الإجمالي','Gross Salary','إجمالي الراتب'] if has(snap,c)), None)
            basic_col = next((c for c in ['الراتب الأساسي','Basic Salary','Basic'] if has(snap,c)), None)
            net_col = next((c for c in ['صافي الراتب','Net Salary','Net'] if has(snap,c)), None)
            dept_col = next((c for c in ['القسم','Department','القطاع','Division'] if has(snap,c)), None)
            nat_col = next((c for c in ['الجنسية','Nationality'] if has(snap,c)), None)
            gender_col = next((c for c in ['الجنس','Gender'] if has(snap,c)), None)
            name_col = next((c for c in ['الاسم','Employee Name','Name'] if has(snap,c)), None)
            total_emp = snap[name_col].nunique() if name_col else len(snap)

            # ===== ROW 1: Executive KPIs =====
            k1,k2,k3,k4,k5,k6 = st.columns(6)
            with k1: kpi("👥 الموظفين", f"{total_emp:,}")
            with k2: kpi("💵 إجمالي الرواتب", f"{snap[sal_col].sum():,.0f}" if sal_col else "-")
            with k3: kpi("📊 المتوسط", f"{snap[sal_col].mean():,.0f}" if sal_col else "-")
            with k4: kpi("📈 الأعلى", f"{snap[sal_col].max():,.0f}" if sal_col else "-")
            with k5: kpi("📉 الأقل", f"{snap[sal_col].min():,.0f}" if sal_col else "-")
            with k6: kpi("📐 الوسيط", f"{snap[sal_col].median():,.0f}" if sal_col else "-")

            # ===== ROW 2: Salary Components Breakdown =====
            st.markdown("---")
            sal_components = ['الراتب الأساسي','بدل السكن','بدل النقل','بدل خاص','بدل معيشة','بدل جوال','بدلات أخرى']
            available_components = [c for c in sal_components if has(snap,c)]

            if available_components:
                st.markdown("### 📊 تركيبة الراتب (Salary Components)")
                comp_data = {c: snap[c].sum() for c in available_components}
                comp_df = pd.DataFrame(list(comp_data.items()), columns=['المكون','الإجمالي'])
                comp_df['النسبة'] = (comp_df['الإجمالي'] / comp_df['الإجمالي'].sum() * 100).round(1)
                comp_df['المتوسط/موظف'] = (comp_df['الإجمالي'] / max(total_emp,1)).round(0)

                c1,c2 = st.columns(2)
                with c1:
                    fig = px.pie(comp_df, values='الإجمالي', names='المكون', title='توزيع مكونات الراتب', hole=.4, color_discrete_sequence=CL['sal'])
                    fig.update_layout(font=dict(family="Noto Sans Arabic"), height=380)
                    st.plotly_chart(fig, use_container_width=True)
                with c2:
                    fig = px.bar(comp_df.sort_values('الإجمالي'), x='الإجمالي', y='المكون', orientation='h',
                        color='المكون', color_discrete_sequence=CL['dept'], title='مكونات الراتب بالقيمة',
                        text=comp_df.sort_values('الإجمالي')['النسبة'].apply(lambda x: f'{x}%'))
                    fig.update_layout(font=dict(family="Noto Sans Arabic"), height=380, showlegend=False, xaxis_tickformat=',')
                    st.plotly_chart(fig, use_container_width=True)

                st.dataframe(comp_df, use_container_width=True, hide_index=True)

            # ===== ROW 3: Department Analysis =====
            if dept_col and sal_col:
                st.markdown(f"### 🏢 التحليل حسب {dept_col}")
                dept_stats = snap.groupby(dept_col).agg(
                    العدد=(sal_col, 'count'), الإجمالي=(sal_col, 'sum'),
                    المتوسط=(sal_col, 'mean'), الوسيط=(sal_col, 'median'),
                    الأقل=(sal_col, 'min'), الأعلى=(sal_col, 'max')
                ).sort_values('العدد', ascending=False).reset_index()
                dept_stats['نسبة التكلفة %'] = (dept_stats['الإجمالي'] / dept_stats['الإجمالي'].sum() * 100).round(1)
                dept_stats['نسبة العدد %'] = (dept_stats['العدد'] / dept_stats['العدد'].sum() * 100).round(1)
                for c in ['الإجمالي','المتوسط','الوسيط','الأقل','الأعلى']:
                    dept_stats[c] = dept_stats[c].round(0)

                c1,c2 = st.columns(2)
                with c1:
                    fig = px.bar(dept_stats.sort_values('الإجمالي'), x='الإجمالي', y=dept_col, orientation='h',
                        title=f'إجمالي التكلفة حسب {dept_col}', color='الإجمالي', color_continuous_scale='teal')
                    fig.update_layout(font=dict(family="Noto Sans Arabic"), height=420, xaxis_tickformat=',', coloraxis_showscale=False)
                    st.plotly_chart(fig, use_container_width=True)
                with c2:
                    fig = px.bar(dept_stats.sort_values('المتوسط'), x='المتوسط', y=dept_col, orientation='h',
                        title=f'متوسط الراتب حسب {dept_col}', color='المتوسط', color_continuous_scale='teal')
                    fig.update_layout(font=dict(family="Noto Sans Arabic"), height=420, xaxis_tickformat=',', coloraxis_showscale=False)
                    st.plotly_chart(fig, use_container_width=True)

                # Headcount vs Cost scatter
                c3,c4 = st.columns(2)
                with c3:
                    fig = px.scatter(dept_stats, x='نسبة العدد %', y='نسبة التكلفة %', size='العدد',
                        color=dept_col, title='نسبة العدد مقابل التكلفة', color_discrete_sequence=CL['dept'],
                        hover_data=['المتوسط'])
                    fig.add_trace(go.Scatter(x=[0,60],y=[0,60],mode='lines',line=dict(dash='dash',color='gray'),showlegend=False))
                    fig.update_layout(font=dict(family="Noto Sans Arabic"), height=400)
                    st.plotly_chart(fig, use_container_width=True)
                with c4:
                    fig = px.box(snap, x=dept_col, y=sal_col, title=f'نطاق الرواتب حسب {dept_col}',
                        color_discrete_sequence=['#E9C46A'])
                    fig.update_layout(font=dict(family="Noto Sans Arabic"), height=400, xaxis_tickangle=-45)
                    st.plotly_chart(fig, use_container_width=True)

                st.dataframe(dept_stats, use_container_width=True, hide_index=True)

            # ===== ROW 4: Nationality & Saudization =====
            if nat_col and sal_col:
                st.markdown("### 🌍 الرواتب حسب الجنسية والسعودة")
                nat_stats = snap.groupby(nat_col).agg(
                    العدد=(sal_col,'count'), المتوسط=(sal_col,'mean'), الإجمالي=(sal_col,'sum')
                ).sort_values('العدد', ascending=False).reset_index()
                nat_stats['النسبة %'] = (nat_stats['العدد'] / nat_stats['العدد'].sum() * 100).round(1)

                c1,c2 = st.columns(2)
                with c1:
                    fig = px.pie(nat_stats, values='العدد', names=nat_col, title='توزيع الجنسيات', hole=0.4,
                        color_discrete_sequence=px.colors.qualitative.Set2)
                    fig.update_layout(font=dict(family="Noto Sans Arabic"), height=380)
                    st.plotly_chart(fig, use_container_width=True)
                with c2:
                    fig = px.bar(nat_stats, x=nat_col, y='المتوسط', title='متوسط الراتب حسب الجنسية',
                        color='العدد', text='العدد', color_continuous_scale='teal')
                    fig.update_layout(font=dict(family="Noto Sans Arabic"), height=380, yaxis_tickformat=',', coloraxis_showscale=False)
                    st.plotly_chart(fig, use_container_width=True)

                # Saudization KPIs
                saudi_vals = ['Saudi','سعودي','Saudi Arabian','سعودية']
                sa_count = len(snap[snap[nat_col].isin(saudi_vals)])
                sa_pct = round(sa_count / max(total_emp,1) * 100, 1)
                sa_cost = snap[snap[nat_col].isin(saudi_vals)][sal_col].sum() if sal_col else 0
                non_sa_cost = snap[~snap[nat_col].isin(saudi_vals)][sal_col].sum() if sal_col else 0

                sk1,sk2,sk3,sk4 = st.columns(4)
                with sk1: kpi("🇸🇦 سعوديين", str(sa_count))
                with sk2: kpi("🌍 غير سعوديين", str(total_emp - sa_count))
                with sk3: kpi("📊 نسبة السعودة", f"{sa_pct}%")
                with sk4: kpi("💰 تكلفة السعوديين", f"{sa_cost:,.0f}")

            # ===== ROW 5: Gender + Distribution =====
            c5,c6 = st.columns(2)
            with c5:
                if gender_col and sal_col:
                    gs = snap.groupby(gender_col).agg(العدد=(sal_col,'count'), المتوسط=(sal_col,'mean')).reset_index()
                    fig = px.bar(gs, x=gender_col, y='المتوسط', title='متوسط الراتب حسب الجنس',
                        color=gender_col, text='العدد', color_discrete_map={'Male':'#3498DB','Female':'#E91E8F','ذكر':'#3498DB','أنثى':'#E91E8F'})
                    fig.update_layout(font=dict(family="Noto Sans Arabic"), height=380, yaxis_tickformat=',')
                    st.plotly_chart(fig, use_container_width=True)
            with c6:
                if sal_col:
                    fig = px.histogram(snap, x=sal_col, nbins=25, title='توزيع الرواتب الإجمالية',
                        color_discrete_sequence=[CL['p']])
                    fig.add_vline(x=snap[sal_col].mean(), line_dash="dash", line_color="red",
                        annotation_text=f"المتوسط: {snap[sal_col].mean():,.0f}")
                    fig.add_vline(x=snap[sal_col].median(), line_dash="dot", line_color="blue",
                        annotation_text=f"الوسيط: {snap[sal_col].median():,.0f}")
                    fig.update_layout(font=dict(family="Noto Sans Arabic"), height=380)
                    st.plotly_chart(fig, use_container_width=True)

            # ===== ROW 6: Deductions & Net =====
            deduction_cols = [c for c in snap.columns if any(x in c.lower() for x in ['خصم','استقطاع','deduction','gosi','تأمين'])]
            if deduction_cols and sal_col:
                st.markdown("### 📉 الاستقطاعات والصافي")
                ded_data = {c: snap[c].sum() for c in deduction_cols}
                ded_df = pd.DataFrame(list(ded_data.items()), columns=['الاستقطاع','الإجمالي'])
                c1,c2 = st.columns(2)
                with c1:
                    fig = px.pie(ded_df, values='الإجمالي', names='الاستقطاع', title='توزيع الاستقطاعات', hole=0.4,
                        color_discrete_sequence=['#E74C3C','#C0392B','#A93226','#922B21'])
                    fig.update_layout(font=dict(family="Noto Sans Arabic"), height=350)
                    st.plotly_chart(fig, use_container_width=True)
                with c2:
                    if net_col:
                        gross_total = snap[sal_col].sum()
                        net_total = snap[net_col].sum()
                        ded_total = gross_total - net_total
                        gn_df = pd.DataFrame([
                            {"البند":"الإجمالي","المبلغ":gross_total},
                            {"البند":"الاستقطاعات","المبلغ":ded_total},
                            {"البند":"الصافي","المبلغ":net_total}])
                        fig = px.bar(gn_df, x='البند', y='المبلغ', title='الإجمالي vs الصافي',
                            color='البند', color_discrete_map={'الإجمالي':'#2A9D8F','الاستقطاعات':'#E74C3C','الصافي':'#27AE60'})
                        fig.update_layout(font=dict(family="Noto Sans Arabic"), height=350, yaxis_tickformat=',', showlegend=False)
                        st.plotly_chart(fig, use_container_width=True)

            # ===== ROW 7: Top Earners =====
            if sal_col and name_col:
                st.markdown("### 🏆 أعلى 10 رواتب")
                top10 = snap.nlargest(10, sal_col)[[name_col, dept_col or sal_col, sal_col]].reset_index(drop=True)
                top10.index = range(1, len(top10)+1)
                st.dataframe(top10, use_container_width=True)

            # ===== ROW 8: Cost Efficiency =====
            if sal_col and dept_col:
                st.markdown("### ⚡ كفاءة التكلفة")
                eff = snap.groupby(dept_col).agg(
                    العدد=(sal_col,'count'), التكلفة_الشهرية=(sal_col,'sum')
                ).reset_index()
                eff['التكلفة/موظف'] = (eff['التكلفة_الشهرية'] / eff['العدد']).round(0)
                eff['التكلفة_السنوية'] = eff['التكلفة_الشهرية'] * 12
                fig = px.treemap(eff, path=[dept_col], values='التكلفة_الشهرية', color='التكلفة/موظف',
                    title='خريطة تكلفة الأقسام (الحجم = التكلفة الشهرية، اللون = التكلفة/موظف)',
                    color_continuous_scale='RdYlGn_r')
                fig.update_layout(font=dict(family="Noto Sans Arabic"), height=420)
                st.plotly_chart(fig, use_container_width=True)

            # Export salary dashboard
            if len(snap) > 0:
                export_sheets = {"الراتب الشامل": snap}
                try:
                    if dept_stats is not None and len(dept_stats) > 0:
                        export_sheets["تحليل الأقسام"] = dept_stats
                except: pass
                try:
                    if nat_stats is not None and len(nat_stats) > 0:
                        export_sheets["الجنسيات"] = nat_stats
                except: pass
                export_widget(export_sheets, "تحليل_الرواتب", "sal_dash")

        elif page == "📈 تحليل شهري/ربعي":
            hdr("📈 تحليل الرواتب الشهري والربعي")
            if len(sal_df)==0: st.info("📁 ارفع ملف رواتب شهري (من القائمة الجانبية)"); return

            if has(sal_df,'سنة الراتب'):
                year = st.selectbox("📅 السنة:", sorted(sal_df['سنة الراتب'].unique(), reverse=True))
                yr = sal_df[sal_df['سنة الراتب']==year]

                if has(yr,'شهر الراتب') and has(yr,'الراتب الإجمالي'):
                    months_order = ['January','February','March','April','May','June','July','August','September','October','November','December']
                    monthly = yr.groupby('شهر الراتب')['الراتب الإجمالي'].sum().reindex(months_order).dropna()
                    fig = go.Figure()
                    fig.add_trace(go.Bar(x=monthly.index, y=monthly.values, marker_color='#0F4C5C', text=[f"{v:,.0f}" for v in monthly.values], textposition='outside'))
                    fig.update_layout(title=f'إجمالي الرواتب الشهرية - {year}', font=dict(family="Noto Sans Arabic"), height=400, yaxis_tickformat=',')
                    st.plotly_chart(fig, use_container_width=True)

                if has(yr,'الربع') and has(yr,'الراتب الإجمالي'):
                    quarterly = yr.groupby('الربع')['الراتب الإجمالي'].sum()
                    c1,c2 = st.columns(2)
                    with c1:
                        fig = px.bar(quarterly.reset_index(), x='الربع', y='الراتب الإجمالي', title=f'الرواتب ربع السنوية - {year}', color='الربع', color_discrete_sequence=[CL['p'],CL['a'],CL['s'],'#64748B'])
                        fig.update_layout(font=dict(family="Noto Sans Arabic"),height=350,yaxis_tickformat=','); st.plotly_chart(fig,use_container_width=True)
                    with c2:
                        # Headcount trend by month
                        hc_monthly = yr.groupby('شهر الراتب')['الاسم'].nunique().reindex(months_order).dropna() if has(yr,'الاسم') else None
                        if hc_monthly is not None:
                            fig = go.Figure()
                            fig.add_trace(go.Scatter(x=hc_monthly.index, y=hc_monthly.values, mode='lines+markers', line=dict(color=CL['a'],width=3), fill='tozeroy', fillcolor='rgba(227,100,20,0.1)'))
                            fig.update_layout(title=f'عدد الموظفين شهرياً - {year}', font=dict(family="Noto Sans Arabic"),height=350)
                            st.plotly_chart(fig, use_container_width=True)

                # Overtime analysis
                if has(yr,'ساعات إضافية') and has(yr,'تكلفة الإضافي'):
                    st.markdown("### ⏰ تحليل الساعات الإضافية")
                    c1,c2 = st.columns(2)
                    with c1: st.metric("🕐 إجمالي الساعات", f"{yr['ساعات إضافية'].sum():,.0f}")
                    with c2: st.metric("💰 تكلفة الإضافي", f"{yr['تكلفة الإضافي'].sum():,.0f} ريال")



        elif page == "🏷️ تحليل حسب الفئات":
            hdr("🏷️ تحليل حسب الفئات","الجنس، الجيل، المستوى، نوع التوظيف")
            if len(sal_df)==0 and n==0: st.info("📁 ارفع ملف"); return
            data = sal_snapshot if len(sal_snapshot)>0 else emp

            tabs = st.tabs(["👫 الجنس","🎂 الأجيال","📊 المستويات","📋 نوع التوظيف"])

            with tabs[0]:
                if has(data,'الجنس'):
                    gc = data['الجنس'].value_counts().reset_index(); gc.columns=['الجنس','العدد']
                    c1,c2 = st.columns(2)
                    with c1:
                        fig = px.pie(gc, values='العدد', names='الجنس', title='التوزيع حسب الجنس', hole=.4, color_discrete_map={'Male':CL['p'],'Female':CL['a']})
                        fig.update_layout(font=dict(family="Noto Sans Arabic"),height=350); st.plotly_chart(fig,use_container_width=True)
                    with c2:
                        if has(data,'الراتب الإجمالي'):
                            gs = data.groupby('الجنس')['الراتب الإجمالي'].agg(['mean','median']).reset_index()
                            gs.columns = ['الجنس','المتوسط','الوسيط']
                            st.dataframe(gs, use_container_width=True, hide_index=True)
                else: st.info("لا يوجد عمود جنس")

            with tabs[1]:
                if has(data,'الجيل'):
                    gc2 = data['الجيل'].value_counts().reset_index(); gc2.columns=['الجيل','العدد']
                    fig = px.bar(gc2, x='الجيل', y='العدد', title='التوزيع حسب الجيل', color='الجيل', color_discrete_sequence=CL['dept'])
                    fig.update_layout(font=dict(family="Noto Sans Arabic"),height=350); st.plotly_chart(fig,use_container_width=True)
                elif has(data,'الفئة العمرية'):
                    ac = data['الفئة العمرية'].value_counts().reset_index(); ac.columns=['الفئة','العدد']
                    fig = px.bar(ac, x='الفئة', y='العدد', title='التوزيع حسب الفئة العمرية', color='الفئة', color_discrete_sequence=CL['dept'])
                    fig.update_layout(font=dict(family="Noto Sans Arabic"),height=350); st.plotly_chart(fig,use_container_width=True)
                else: st.info("لا يوجد بيانات أجيال")

            with tabs[2]:
                if has(data,'المستوى'):
                    lc = data['المستوى'].value_counts().reset_index(); lc.columns=['المستوى','العدد']
                    c1,c2 = st.columns(2)
                    with c1:
                        fig = px.pie(lc, values='العدد', names='المستوى', title='التوزيع حسب المستوى', hole=.35, color_discrete_sequence=CL['dept'])
                        fig.update_layout(font=dict(family="Noto Sans Arabic"),height=350); st.plotly_chart(fig,use_container_width=True)
                    with c2:
                        if has(data,'الراتب الإجمالي'):
                            ls = data.groupby('المستوى')['الراتب الإجمالي'].mean().reset_index().sort_values('الراتب الإجمالي',ascending=True)
                            fig = px.bar(ls, x='الراتب الإجمالي', y='المستوى', orientation='h', title='متوسط الراتب حسب المستوى', color='الراتب الإجمالي', color_continuous_scale='teal')
                            fig.update_layout(font=dict(family="Noto Sans Arabic"),height=350,xaxis_tickformat=','); st.plotly_chart(fig,use_container_width=True)
                else: st.info("لا يوجد عمود مستوى")

            with tabs[3]:
                if has(data,'نوع التوظيف'):
                    ec = data['نوع التوظيف'].value_counts().reset_index(); ec.columns=['النوع','العدد']
                    fig = px.pie(ec, values='العدد', names='النوع', title='أنواع التوظيف', hole=.35, color_discrete_sequence=CL['sal'])
                    fig.update_layout(font=dict(family="Noto Sans Arabic"),height=350); st.plotly_chart(fig,use_container_width=True)
                else: st.info("لا يوجد عمود نوع التوظيف")



        elif page == "📊 سلم الرواتب":
            hdr("📊 سلم الرواتب والدرجات")
            if 'Salary Scale' in all_sheets:
                ss = all_sheets['Salary Scale'].dropna(subset=['Grade'] if 'Grade' in all_sheets['Salary Scale'].columns else all_sheets['Salary Scale'].columns[:1])
                ss_norm = norm_cols(ss)
                st.dataframe(ss, use_container_width=True, hide_index=True)
                if 'Min Salary' in ss.columns and 'Max Salary' in ss.columns:
                    ss_clean = ss.dropna(subset=['Min Salary','Max Salary'])
                    fig = go.Figure()
                    fig.add_trace(go.Bar(name='الحد الأدنى', x=ss_clean['Grade'].astype(str), y=ss_clean['Min Salary'], marker_color=CL['s']))
                    fig.add_trace(go.Bar(name='المتوسط', x=ss_clean['Grade'].astype(str), y=ss_clean['Mid Salary'], marker_color='#E36414'))
                    fig.add_trace(go.Bar(name='الحد الأقصى', x=ss_clean['Grade'].astype(str), y=ss_clean['Max Salary'], marker_color=CL['d']))
                    fig.update_layout(title='سلم الرواتب حسب الدرجة', barmode='group', font=dict(family="Noto Sans Arabic"), height=420, yaxis_tickformat=',')
                    st.plotly_chart(fig, use_container_width=True)
            else:
                st.info("لم يتم العثور على ورقة Salary Scale في الملف المرفوع")




            export_widget(snap if len(snap)>0 else None, "سلم_الرواتب", "slsc")
        elif page == "🎁 لوحة Total Rewards":
            hdr("🎁 لوحة Total Rewards الشاملة","Compensation + Benefits + Work-Life + Performance + Development")
            if len(snap) == 0: st.info("📁 ارفع ملف بيانات الموظفين"); return

            st.markdown("### 📊 مكونات Total Rewards (WorldatWork Model)")
            ibox("**نموذج WorldatWork:** Total Rewards = Compensation + Benefits + Work-Life Effectiveness + Recognition + Development")

            trc1, trc2, trc3 = st.columns(3)
            with trc1:
                tr_comp_pct = st.number_input("💰 Compensation (% من الميزانية):", 0, 100, 60, key="tr_comp")
                tr_benefits_pct = st.number_input("🏥 Benefits (%):", 0, 100, 20, key="tr_ben")
            with trc2:
                tr_worklife_pct = st.number_input("⚖️ Work-Life (%):", 0, 100, 8, key="tr_wl")
                tr_recognition_pct = st.number_input("🏆 Recognition (%):", 0, 100, 5, key="tr_rec")
            with trc3:
                tr_development_pct = st.number_input("📚 Development (%):", 0, 100, 7, key="tr_dev")
                total_budget_tr = st.number_input("💵 إجمالي ميزانية TR (ريال/شهر):", value=int(snap[sal_col_tr].sum()) if sal_col_tr else 500000, key="tr_budget")

            n_emp = len(snap)
            avg_sal = snap[sal_col_tr].mean() if sal_col_tr else 0
            k1,k2,k3,k4,k5 = st.columns(5)
            with k1: kpi("👥 الموظفين", str(n_emp))
            with k2: kpi("💰 متوسط الراتب", f"{avg_sal:,.0f}")
            with k3: kpi("💵 الميزانية الشهرية", f"{total_budget_tr:,.0f}")
            with k4: kpi("📅 السنوية", f"{total_budget_tr*12:,.0f}")
            with k5: kpi("💵 للفرد/سنة", f"{total_budget_tr*12//max(n_emp,1):,.0f}")

            tr_data = [
                {"المكون":"💰 Compensation","النسبة":tr_comp_pct,"المبلغ":total_budget_tr*tr_comp_pct//100},
                {"المكون":"🏥 Benefits","النسبة":tr_benefits_pct,"المبلغ":total_budget_tr*tr_benefits_pct//100},
                {"المكون":"⚖️ Work-Life","النسبة":tr_worklife_pct,"المبلغ":total_budget_tr*tr_worklife_pct//100},
                {"المكون":"🏆 Recognition","النسبة":tr_recognition_pct,"المبلغ":total_budget_tr*tr_recognition_pct//100},
                {"المكون":"📚 Development","النسبة":tr_development_pct,"المبلغ":total_budget_tr*tr_development_pct//100},
            ]
            tr_df = pd.DataFrame(tr_data)
            tc1, tc2 = st.columns(2)
            with tc1:
                fig = px.pie(tr_df, values='النسبة', names='المكون', title='توزيع Total Rewards', hole=0.4,
                    color_discrete_sequence=['#E36414','#2A9D8F','#E9C46A','#264653','#F4A261'])
                fig.update_layout(font=dict(family="Noto Sans Arabic"), height=380)
                st.plotly_chart(fig, use_container_width=True)
            with tc2:
                fig = px.bar(tr_df, x='المكون', y='المبلغ', title='التوزيع بالريال', color='المكون',
                    color_discrete_sequence=['#E36414','#2A9D8F','#E9C46A','#264653','#F4A261'])
                fig.update_layout(font=dict(family="Noto Sans Arabic"), height=380, showlegend=False)
                st.plotly_chart(fig, use_container_width=True)
            st.dataframe(tr_df, use_container_width=True, hide_index=True)



            export_widget({"Total Rewards": tr_df, "بيانات الموظفين": snap} if len(snap)>0 else tr_df, "Total_Rewards", "tr1")
        elif page == "💰 هيكل الرواتب":
            hdr("💰 هيكل الرواتب والعدالة","Salary Structure & Pay Equity")
            if len(snap)==0: st.info("📁 ارفع ملف"); return
            if sal_col_tr:
                st.markdown("### 📊 نطاقات الرواتب (Salary Bands)")
                snap['_SalBand'] = pd.cut(snap[sal_col_tr], bins=[0,3000,5000,8000,12000,18000,30000,50000,999999],
                    labels=['<3K','3-5K','5-8K','8-12K','12-18K','18-30K','30-50K','50K+'])
                band_counts = snap['_SalBand'].value_counts().sort_index()
                fig = px.bar(x=band_counts.index.astype(str), y=band_counts.values, title='توزيع الموظفين حسب نطاق الراتب',
                    color=band_counts.values, color_continuous_scale='teal')
                fig.update_layout(font=dict(family="Noto Sans Arabic"), height=380, showlegend=False, coloraxis_showscale=False)
                st.plotly_chart(fig, use_container_width=True)

                if dept_col_tr:
                    st.markdown("### 📊 Compa-Ratio حسب القسم")
                    market_mid = snap[sal_col_tr].median()
                    dept_compa = snap.groupby(dept_col_tr)[sal_col_tr].mean().reset_index()
                    dept_compa.columns = [dept_col_tr, 'متوسط']
                    dept_compa['Compa-Ratio'] = (dept_compa['متوسط'] / market_mid * 100).round(1)
                    dept_compa['الحالة'] = dept_compa['Compa-Ratio'].apply(lambda x: '🔴 أقل' if x<90 else ('🟢 تنافسي' if x<=110 else '🟡 أعلى'))
                    fig = px.bar(dept_compa.sort_values('Compa-Ratio'), x='Compa-Ratio', y=dept_col_tr, orientation='h',
                        title='Compa-Ratio (100% = وسيط السوق)', color='Compa-Ratio', color_continuous_scale='RdYlGn', range_color=[70,130])
                    fig.add_vline(x=100, line_dash="dash", line_color="black", annotation_text="Market Mid")
                    fig.update_layout(font=dict(family="Noto Sans Arabic"), height=420, coloraxis_showscale=False)
                    st.plotly_chart(fig, use_container_width=True)
                    st.dataframe(dept_compa, use_container_width=True, hide_index=True)



            export_widget(snap if len(snap)>0 else None, "هيكل_الرواتب", "tr2")
        elif page == "🏥 المزايا والتأمينات":
            hdr("🏥 المزايا والتأمينات","Benefits & Insurance")
            benefits_data = st.data_editor(
                pd.DataFrame([
                    {"المزية":"التأمين الطبي","التكلفة/موظف (شهري)":500,"المشمولين %":100,"الفئة":"تأمين"},
                    {"المزية":"التأمينات الاجتماعية","التكلفة/موظف (شهري)":800,"المشمولين %":100,"الفئة":"تأمين"},
                    {"المزية":"بدل سكن","التكلفة/موظف (شهري)":2500,"المشمولين %":85,"الفئة":"بدل"},
                    {"المزية":"بدل مواصلات","التكلفة/موظف (شهري)":500,"المشمولين %":90,"الفئة":"بدل"},
                    {"المزية":"تذاكر سفر","التكلفة/موظف (شهري)":300,"المشمولين %":40,"الفئة":"سفر"},
                    {"المزية":"تدريب وتطوير","التكلفة/موظف (شهري)":200,"المشمولين %":70,"الفئة":"تطوير"},
                    {"المزية":"مكافآت أداء","التكلفة/موظف (شهري)":1000,"المشمولين %":60,"الفئة":"مكافآت"},
                ]),
                column_config={'الفئة': st.column_config.SelectboxColumn('الفئة', options=['تأمين','بدل','سفر','تطوير','مكافآت','رفاهية'])},
                use_container_width=True, hide_index=True, num_rows="dynamic", key="benefits_editor"
            )
            if len(benefits_data) > 0:
                fig = px.pie(benefits_data, values='التكلفة/موظف (شهري)', names='المزية', title='توزيع تكلفة المزايا', hole=0.4)
                fig.update_layout(font=dict(family="Noto Sans Arabic"), height=380)
                st.plotly_chart(fig, use_container_width=True)
                kpi("💵 إجمالي المزايا/موظف/شهر", f"{benefits_data['التكلفة/موظف (شهري)'].sum():,} ريال")




            export_widget(benefits_data if len(benefits_data)>0 else None, "المزايا_والتأمينات", "bnft")
        elif page == "📊 تحليل التنافسية":
            hdr("📊 تحليل التنافسية","Market Competitiveness")
            if len(snap)==0 or not sal_col_tr: st.info("📁 ارفع ملف رواتب"); return
            fig = go.Figure()
            fig.add_trace(go.Indicator(mode="gauge+number", value=snap[sal_col_tr].mean(),
                title={'text': "متوسط الراتب مقارنة بالسوق"},
                gauge={'axis':{'range':[0, snap[sal_col_tr].quantile(0.95)]}, 'bar':{'color':'#E36414'},
                    'steps':[{'range':[0,snap[sal_col_tr].quantile(0.25)],'color':'#E74C3C'},
                        {'range':[snap[sal_col_tr].quantile(0.25),snap[sal_col_tr].quantile(0.75)],'color':'#F39C12'},
                        {'range':[snap[sal_col_tr].quantile(0.75),snap[sal_col_tr].quantile(0.95)],'color':'#27AE60'}]}))
            fig.update_layout(height=300)
            st.plotly_chart(fig, use_container_width=True)



            export_widget(snap, "تحليل التنافسية", "tr4")
        elif page == "📥 تصدير TR":
            hdr("📥 تصدير تقرير الرواتب","Excel تقرير شامل للرواتب والتعويضات")
            data = sal_df if len(sal_df)>0 else (sal_snapshot if len(sal_snapshot)>0 else emp)
            snap = sal_snapshot if len(sal_snapshot)>0 else data
            if len(data)==0: st.info("📁 ارفع ملف"); return

            exp_format = st.radio("صيغة التصدير:", ["📊 Excel شامل","📄 CSV","📝 HTML"], horizontal=True, key="sal_exp_fmt")

            if st.button("📥 إنشاء التقرير", type="primary", use_container_width=True):
                if exp_format == "📊 Excel شامل":
                    o = io.BytesIO()
                    with pd.ExcelWriter(o, engine='xlsxwriter') as w:
                        wb = w.book
                        hdr_f = wb.add_format({'bold':True,'font_size':14,'bg_color':'#0F4C5C','font_color':'white','align':'center','border':1})
                        num_f = wb.add_format({'num_format':'#,##0.00','border':1})

                        # Sheet 1: Dashboard summary
                        ws1 = wb.add_worksheet('Dashboard')
                        ws1.set_column('A:Z', 18)
                        ws1.merge_range('B2:F2', 'Salary Analysis', hdr_f)
                        r = 4
                        sal_col = next((c for c in ['الراتب الإجمالي','Gross Salary'] if has(snap,c)), None)
                        dept_col = next((c for c in ['القسم','Department','القطاع'] if has(snap,c)), None)
                        if sal_col:
                            stats = [('Total Employees', snap.shape[0]), ('Total Monthly Payroll', f"{snap[sal_col].sum():,.2f}"),
                                ('Average Salary', f"{snap[sal_col].mean():,.2f}"), ('Median Salary', f"{snap[sal_col].median():,.2f}"),
                                ('Max Salary', f"{snap[sal_col].max():,.2f}"), ('Min Salary', f"{snap[sal_col].min():,.2f}")]
                            for i, (lbl, val) in enumerate(stats):
                                ws1.write(r+i, 1, lbl, wb.add_format({'bold':True,'border':1}))
                                ws1.write(r+i, 2, str(val), wb.add_format({'border':1}))

                        # Sheet 2: Analysis (monthly/quarterly)
                        month_col = next((c for c in ['Salary Month','الشهر'] if has(data,c)), None)
                        if sal_col and month_col:
                            monthly = data.groupby(month_col)[sal_col].sum().reset_index()
                            monthly.columns = ['Month','Total Gross Salary']
                            monthly.to_excel(w, sheet_name='Analysis', index=False, startrow=1)

                        # Sheet 3: Employee Salaries (full data)
                        data.to_excel(w, sheet_name='Employee Salaries', index=False)

                        # Sheet 4: Current Snapshot
                        snap.to_excel(w, sheet_name='Current Snapshot', index=False)

                        # Sheet 5: Salary Scale
                        if 'Salary Scale' in all_sheets:
                            all_sheets['Salary Scale'].to_excel(w, sheet_name='Salary Scale', index=False)
                        else:
                            # Default salary scale
                            scale_data = [
                                {'Grade':4,'Level':'Staff','Min':1800,'Mid':2200,'Max':2600},
                                {'Grade':5,'Level':'Staff','Min':2200,'Mid':2800,'Max':3400},
                                {'Grade':6,'Level':'Staff','Min':2800,'Mid':3500,'Max':4200},
                                {'Grade':7,'Level':'Junior','Min':3500,'Mid':4300,'Max':5100},
                                {'Grade':8,'Level':'Junior','Min':4300,'Mid':5300,'Max':6300},
                                {'Grade':9,'Level':'Junior','Min':5300,'Mid':6600,'Max':7900},
                                {'Grade':10,'Level':'Senior','Min':6600,'Mid':8200,'Max':9800},
                                {'Grade':11,'Level':'Senior','Min':8200,'Mid':10200,'Max':12200},
                                {'Grade':12,'Level':'Senior','Min':10200,'Mid':12800,'Max':15400},
                                {'Grade':13,'Level':'Supervisor','Min':12800,'Mid':16000,'Max':19200},
                                {'Grade':14,'Level':'Supervisor','Min':16000,'Mid':20000,'Max':24000},
                                {'Grade':15,'Level':'Supervisor','Min':20000,'Mid':25000,'Max':30000},
                                {'Grade':16,'Level':'Management','Min':25000,'Mid':32000,'Max':39000},
                                {'Grade':17,'Level':'Management','Min':32000,'Mid':40000,'Max':48000},
                                {'Grade':18,'Level':'Management','Min':40000,'Mid':50000,'Max':60000},
                                {'Grade':19,'Level':'Senior Management','Min':50000,'Mid':65000,'Max':80000},
                                {'Grade':20,'Level':'Senior Management','Min':65000,'Mid':95000,'Max':125000},
                            ]
                            pd.DataFrame(scale_data).to_excel(w, sheet_name='Salary Scale', index=False)

                        # Sheet 6: Positions
                        if 'Positions' in all_sheets:
                            all_sheets['Positions'].to_excel(w, sheet_name='Positions', index=False)

                        # Sheet 7: Department Summary
                        if dept_col and sal_col:
                            dept_summary = snap.groupby(dept_col).agg(
                                Count=(sal_col, 'count'),
                                Total=(sal_col, 'sum'),
                                Average=(sal_col, 'mean'),
                                Median=(sal_col, 'median'),
                                Min=(sal_col, 'min'),
                                Max=(sal_col, 'max'),
                            ).reset_index()
                            dept_summary.columns = [dept_col,'عدد','إجمالي','متوسط','وسيط','أقل','أعلى']
                            dept_summary.to_excel(w, sheet_name='Dept Summary', index=False)

                    st.download_button("📥 تحميل تقرير الرواتب الشامل", data=o.getvalue(),
                        file_name=f"Salary_Analysis_{datetime.now().strftime('%Y%m%d')}.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", type="primary", use_container_width=True)

                elif exp_format == "📄 CSV":
                    csv_data = snap.to_csv(index=False).encode('utf-8-sig')
                    st.download_button("📥 تحميل CSV", data=csv_data,
                        file_name=f"Salary_Data_{datetime.now().strftime('%Y%m%d')}.csv",
                        mime="text/csv", type="primary", use_container_width=True)

                elif exp_format == "📝 HTML":
                    html = f"<html dir='rtl'><head><meta charset='utf-8'><title>Salary Report</title><style>body{{font-family:Arial}}table{{border-collapse:collapse;width:100%}}th{{background:#0F4C5C;color:white;padding:8px;border:1px solid #ddd}}td{{padding:6px;border:1px solid #ddd}}</style></head><body><h1>تقرير تحليل الرواتب</h1><p>التاريخ: {datetime.now().strftime('%Y-%m-%d')}</p>{snap.to_html(index=False)}</body></html>"
                    st.download_button("📥 تحميل HTML", data=html.encode('utf-8'),
                        file_name=f"Salary_Report_{datetime.now().strftime('%Y%m%d')}.html",
                        mime="text/html", type="primary", use_container_width=True)


    # =========================================
    #        👥 HEADCOUNT & PERFORMANCE
    # =========================================
    elif section == "👥 Headcount":

        if page == "👥 Headcount Report":
            hdr("👥 HR Headcount Report","تقرير القوى العاملة الشهري - مطابق لنموذج Headcount 2025-2026")

            data = sal_snapshot if len(sal_snapshot)>0 else emp
            total = len(data)

            # Detect columns
            dept_col = next((c for c in ['Department','القسم','القطاع'] if has(data,c)), None)
            nat_col = next((c for c in ['Nationality','الجنسية','Nationality Group'] if has(data,c)), None)
            status_col = next((c for c in ['Status','الحالة'] if has(data,c)), None)
            join_col = next((c for c in ['Join Date','Hiring Date','تاريخ الالتحاق','Join_Date'] if has(data,c)), None)
            end_col = next((c for c in ['End Date','End_Date','تاريخ الانتهاء'] if has(data,c)), None)
            loc_col = next((c for c in ['Location','الموقع','المدينة'] if has(data,c)), None)
            type_col = next((c for c in ['Employment Type','نوع التوظيف'] if has(data,c)), None)
            name_en = next((c for c in ['Full Name (EN)','الاسم','Name','الاسم بالانجليزي'] if has(data,c)), None)
            name_ar = next((c for c in ['Full Name (AR)','الاسم بالعربي'] if has(data,c)), None)

            if total == 0 or not dept_col:
                st.info("📁 ارفع ملف بيانات الموظفين (مثل Headcount_2025_2026.xlsx)")
                st.markdown("### 📋 أو أدخل البيانات يدوياً")
                num_depts = st.number_input("عدد الأقسام", 1, 30, 10)
                manual_data = []
                for i in range(num_depts):
                    c1,c2,c3,c4 = st.columns(4)
                    with c1: dept = st.text_input(f"القسم {i+1}", f"قسم {i+1}", key=f"hd_{i}")
                    with c2: current = st.number_input("الحالي", 0, 500, 10, key=f"hc_{i}")
                    with c3: sa = st.number_input("سعودي", 0, 500, 5, key=f"hs_{i}")
                    with c4: target = st.number_input("المستهدف", 0, 500, 12, key=f"ht_{i}")
                    manual_data.append({"القسم":dept, "الحالي":current, "سعودي":sa, "غير سعودي":current-sa, "المستهدف":target, "الفرق":target-current, "السعودة %":round(sa/max(current,1)*100,1)})
                if manual_data:
                    mdf = pd.DataFrame(manual_data)
                    st.dataframe(mdf, use_container_width=True, hide_index=True)
                    k1,k2,k3,k4 = st.columns(4)
                    with k1: kpi("👥 الإجمالي", str(mdf['الحالي'].sum()))
                    with k2: kpi("🎯 المستهدف", str(mdf['المستهدف'].sum()))
                    with k3: kpi("📊 المطلوب", str(mdf['الفرق'].sum()))
                    with k4: kpi("🇸🇦 السعودة", f"{round(mdf['سعودي'].sum()/max(mdf['الحالي'].sum(),1)*100,1)}%")
                return

            # Convert dates
            if join_col:
                data[join_col] = pd.to_datetime(data[join_col], errors='coerce')
            if end_col:
                data[end_col] = pd.to_datetime(data[end_col], errors='coerce')

            # Determine active employees
            if status_col:
                active = data[data[status_col].isin(['Active','نشط','active'])]
            elif end_col:
                active = data[data[end_col].isna()]
            else:
                active = data

            total_active = len(active)
            total_all = len(data)

            # ===== SECTION 0: KPI Summary =====
            st.markdown("### 📊 ملخص القوى العاملة")
            k1,k2,k3,k4,k5,k6 = st.columns(6)
            with k1: kpi("👥 إجمالي السجلات", str(total_all))
            with k2: kpi("✅ نشط حالياً", str(total_active))
            with k3: kpi("❌ منتهي", str(total_all - total_active))
            if nat_col:
                saudi_vals = ['Saudi','سعودي','Saudi Arabian']
                sa_count = len(active[active[nat_col].isin(saudi_vals)])
                sa_pct = round(sa_count/max(total_active,1)*100,1)
                with k4: kpi("🇸🇦 سعودي", str(sa_count))
                with k5: kpi("🌍 غير سعودي", str(total_active - sa_count))
                with k6: kpi("📊 نسبة السعودة", f"{sa_pct}%")
            if dept_col:
                n_depts = active[dept_col].nunique()
            if loc_col:
                n_locs = active[loc_col].nunique()

            # ===== SECTION 1: Monthly Headcount by Department =====
            if join_col and dept_col:
                st.markdown("---")
                st.markdown("### 📅 SECTION 1: Headcount الشهري حسب القسم")

                # Generate monthly periods
                yr_sel = st.selectbox("السنة:", [2024,2025,2026,2027], index=1, key="hc_yr")
                months = pd.date_range(f'{yr_sel}-01-01', f'{yr_sel}-12-01', freq='MS')
                if yr_sel == 2025:
                    months = pd.date_range('2025-07-01', '2026-12-01', freq='MS')

                depts = sorted(data[dept_col].dropna().unique())

                # Calculate headcount per month per dept
                hc_monthly = []
                for dept in depts:
                    dept_data = data[data[dept_col]==dept]
                    row = {'القسم': dept}
                    for m in months:
                        m_end = m + pd.offsets.MonthEnd(0)
                        joined = dept_data[dept_data[join_col] <= m_end] if join_col else dept_data
                        if end_col:
                            still_active = joined[(joined[end_col].isna()) | (joined[end_col] > m_end)]
                        else:
                            still_active = joined
                        row[m.strftime('%b %Y')] = len(still_active)
                    hc_monthly.append(row)

                hc_df = pd.DataFrame(hc_monthly)
                # Add total row
                total_row = {'القسم': 'الإجمالي'}
                for col in hc_df.columns[1:]:
                    total_row[col] = hc_df[col].sum()
                hc_df = pd.concat([hc_df, pd.DataFrame([total_row])], ignore_index=True)

                st.dataframe(hc_df, use_container_width=True, hide_index=True)

                # Chart
                chart_data = hc_df[hc_df['القسم']!='الإجمالي'].set_index('القسم').T
                fig = px.area(chart_data, title='Headcount الشهري حسب القسم',
                    color_discrete_sequence=CL['dept'])
                fig.update_layout(font=dict(family="Noto Sans Arabic"), height=450,
                    xaxis_title="الشهر", yaxis_title="عدد الموظفين")
                st.plotly_chart(fig, use_container_width=True)

                # Total trend
                totals = {col: hc_df[hc_df['القسم']=='الإجمالي'][col].values[0] for col in hc_df.columns[1:]}
                fig2 = go.Figure()
                fig2.add_trace(go.Scatter(x=list(totals.keys()), y=list(totals.values()),
                    mode='lines+markers+text', text=list(totals.values()), textposition='top center',
                    line=dict(color=CL['p'], width=3), marker=dict(size=10)))
                fig2.update_layout(title='إجمالي Headcount الشهري', font=dict(family="Noto Sans Arabic"),
                    height=350, xaxis_title="الشهر", yaxis_title="الإجمالي")
                st.plotly_chart(fig2, use_container_width=True)

            # ===== SECTION 2: Saudization by Month =====
            if join_col and nat_col and dept_col:
                st.markdown("---")
                st.markdown("### 🇸🇦 SECTION 2: نسبة السعودة الشهرية")

                saudi_vals = ['Saudi','سعودي','Saudi Arabian']
                sa_monthly = []
                nsa_monthly = []
                pct_monthly = []
                for m in months:
                    m_end = m + pd.offsets.MonthEnd(0)
                    joined = data[data[join_col] <= m_end]
                    if end_col:
                        active_m = joined[(joined[end_col].isna()) | (joined[end_col] > m_end)]
                    else:
                        active_m = joined
                    sa = len(active_m[active_m[nat_col].isin(saudi_vals)])
                    tot = len(active_m)
                    sa_monthly.append(sa)
                    nsa_monthly.append(tot - sa)
                    pct_monthly.append(round(sa/max(tot,1)*100, 1))

                sa_df = pd.DataFrame({
                    'الشهر': [m.strftime('%b %Y') for m in months],
                    'سعودي': sa_monthly, 'غير سعودي': nsa_monthly,
                    'الإجمالي': [s+n for s,n in zip(sa_monthly, nsa_monthly)],
                    'نسبة السعودة %': pct_monthly
                })
                st.dataframe(sa_df, use_container_width=True, hide_index=True)

                fig3 = go.Figure()
                fig3.add_trace(go.Bar(x=sa_df['الشهر'], y=sa_df['سعودي'], name='سعودي', marker_color='#27AE60'))
                fig3.add_trace(go.Bar(x=sa_df['الشهر'], y=sa_df['غير سعودي'], name='غير سعودي', marker_color='#3498DB'))
                fig3.add_trace(go.Scatter(x=sa_df['الشهر'], y=sa_df['نسبة السعودة %'], name='السعودة %',
                    yaxis='y2', mode='lines+markers', line=dict(color='#E74C3C', width=2)))
                fig3.update_layout(barmode='stack', title='السعودة الشهرية',
                    yaxis2=dict(overlaying='y', side='right', title='%', range=[0,100]),
                    font=dict(family="Noto Sans Arabic"), height=400)
                st.plotly_chart(fig3, use_container_width=True)

            # ===== SECTION 3: Hires & Terminations =====
            if join_col:
                st.markdown("---")
                st.markdown("### 📈 SECTION 3: التوظيف والاستقالات الشهرية")

                hire_term = []
                for m in months:
                    m_start = m
                    m_end = m + pd.offsets.MonthEnd(0)
                    hires = len(data[(data[join_col] >= m_start) & (data[join_col] <= m_end)])
                    terms = 0
                    if end_col:
                        terms = len(data[(data[end_col] >= m_start) & (data[end_col] <= m_end)])
                    hire_term.append({
                        'الشهر': m.strftime('%b %Y'),
                        'تعيينات جديدة': hires,
                        'إنهاء خدمة': terms,
                        'صافي التغيير': hires - terms
                    })

                ht_df = pd.DataFrame(hire_term)
                st.dataframe(ht_df, use_container_width=True, hide_index=True)

                fig4 = go.Figure()
                fig4.add_trace(go.Bar(x=ht_df['الشهر'], y=ht_df['تعيينات جديدة'], name='تعيينات', marker_color='#27AE60'))
                fig4.add_trace(go.Bar(x=ht_df['الشهر'], y=[-t for t in ht_df['إنهاء خدمة']], name='إنهاء', marker_color='#E74C3C'))
                fig4.add_trace(go.Scatter(x=ht_df['الشهر'], y=ht_df['صافي التغيير'], name='صافي التغيير',
                    mode='lines+markers', line=dict(color=CL['p'], width=2)))
                fig4.update_layout(barmode='relative', title='التوظيف مقابل الاستقالات',
                    font=dict(family="Noto Sans Arabic"), height=400)
                st.plotly_chart(fig4, use_container_width=True)

            # ===== SECTION 4: KPIs =====
            st.markdown("---")
            st.markdown("### 🎯 مؤشرات الأداء الرئيسية (KPIs)")

            kpi_data = {}
            if 'totals' in dir() or join_col:
                try:
                    vals = list(totals.values()) if 'totals' in dir() else [total_active]
                    kpi_data['أعلى Headcount'] = max(vals)
                    kpi_data['أقل Headcount'] = min(vals)
                    if len(vals) > 1:
                        kpi_data['نمو الفترة %'] = round((vals[-1] - vals[0]) / max(vals[0],1) * 100, 1)
                except: pass
            if nat_col:
                kpi_data['متوسط السعودة %'] = round(sum(pct_monthly)/max(len(pct_monthly),1), 1) if 'pct_monthly' in dir() else sa_pct
            if 'ht_df' in dir() and len(ht_df) > 0:
                total_hc_avg = sum([t.get('الإجمالي', total_active) for t in [{}]]) or total_active
                if ht_df['إنهاء خدمة'].sum() > 0:
                    kpi_data['معدل الدوران الشهري %'] = round(ht_df['إنهاء خدمة'].mean() / max(total_active,1) * 100, 2)
                kpi_data['إجمالي التعيينات'] = ht_df['تعيينات جديدة'].sum()
                kpi_data['إجمالي الاستقالات'] = ht_df['إنهاء خدمة'].sum()

            if kpi_data:
                cols = st.columns(min(len(kpi_data), 4))
                for i, (k, v) in enumerate(kpi_data.items()):
                    with cols[i % len(cols)]: kpi(k, str(v))

            # Distribution charts
            if dept_col:
                st.markdown("---")
                c1, c2 = st.columns(2)
                with c1:
                    dept_counts = active[dept_col].value_counts().reset_index()
                    dept_counts.columns = [dept_col, 'العدد']
                    fig = px.pie(dept_counts, values='العدد', names=dept_col, title='توزيع الموظفين حسب القسم',
                        hole=.35, color_discrete_sequence=CL['dept'])
                    fig.update_layout(font=dict(family="Noto Sans Arabic"), height=400)
                    st.plotly_chart(fig, use_container_width=True)
                with c2:
                    if loc_col:
                        loc_counts = active[loc_col].value_counts().reset_index()
                        loc_counts.columns = [loc_col, 'العدد']
                        fig = px.pie(loc_counts, values='العدد', names=loc_col, title='توزيع الموظفين حسب الموقع',
                            hole=.35, color_discrete_sequence=['#E36414','#264653','#2A9D8F','#E9C46A','#F4A261'])
                        fig.update_layout(font=dict(family="Noto Sans Arabic"), height=400)
                        st.plotly_chart(fig, use_container_width=True)
                    elif nat_col:
                        nat_counts = active[nat_col].value_counts().head(10).reset_index()
                        nat_counts.columns = [nat_col, 'العدد']
                        fig = px.bar(nat_counts, x='العدد', y=nat_col, orientation='h', title='أكثر 10 جنسيات',
                            color_discrete_sequence=[CL['a']])
                        fig.update_layout(font=dict(family="Noto Sans Arabic"), height=400)
                        st.plotly_chart(fig, use_container_width=True)



            export_widget({"Headcount": data} if len(data)>0 else None, "Headcount", "hc1")
        elif page == "📋 بيانات الموظفين":
            hdr("📋 سجل بيانات الموظفين","Employee Data Registry - عرض وتصفية")
            data = sal_snapshot if len(sal_snapshot)>0 else emp
            if len(data) == 0:
                st.info("📁 ارفع ملف بيانات الموظفين")
                return
            dept_col = next((c for c in ['Department','القسم','القطاع'] if has(data,c)), None)
            status_col = next((c for c in ['Status','الحالة'] if has(data,c)), None)
            nat_col = next((c for c in ['Nationality','الجنسية'] if has(data,c)), None)

            # Filters
            fc1, fc2, fc3 = st.columns(3)
            with fc1:
                if dept_col:
                    dept_f = st.selectbox("القسم:", ["الكل"] + sorted(data[dept_col].dropna().unique().tolist()), key="hc_df")
            with fc2:
                if status_col:
                    stat_f = st.selectbox("الحالة:", ["الكل"] + sorted(data[status_col].dropna().unique().tolist()), key="hc_sf")
            with fc3:
                if nat_col:
                    nat_f = st.selectbox("الجنسية:", ["الكل"] + sorted(data[nat_col].dropna().unique().tolist()), key="hc_nf")

            filtered = data.copy()
            if dept_col and dept_f != "الكل": filtered = filtered[filtered[dept_col]==dept_f]
            if status_col and stat_f != "الكل": filtered = filtered[filtered[status_col]==stat_f]
            if nat_col and nat_f != "الكل": filtered = filtered[filtered[nat_col]==nat_f]

            st.success(f"📊 {len(filtered)} سجل من أصل {len(data)}")
            st.dataframe(filtered, use_container_width=True, hide_index=True)

        elif page == "📥 تصدير Headcount":
            hdr("📥 تصدير تقرير Headcount","تصدير Excel مطابق لنموذج HR Headcount Report")
            data = sal_snapshot if len(sal_snapshot)>0 else emp
            if len(data) == 0:
                st.info("📁 ارفع ملف بيانات الموظفين أولاً")
                return

            dept_col = next((c for c in ['Department','القسم','القطاع'] if has(data,c)), None)
            nat_col = next((c for c in ['Nationality','الجنسية'] if has(data,c)), None)
            join_col = next((c for c in ['Join Date','Hiring Date','تاريخ الالتحاق'] if has(data,c)), None)
            end_col = next((c for c in ['End Date','End_Date','تاريخ الانتهاء'] if has(data,c)), None)
            status_col = next((c for c in ['Status','الحالة'] if has(data,c)), None)

            yr_sel = st.selectbox("السنة:", [2024,2025,2026], index=1, key="exp_yr")
            if st.button("📥 إنشاء تقرير Headcount Excel", type="primary", use_container_width=True):
                ox = io.BytesIO()
                with pd.ExcelWriter(ox, engine='xlsxwriter') as w:
                    # Sheet 1: Employee Data
                    data.to_excel(w, sheet_name='Emp Data Sheet', index=False)

                    # Sheet 2: Monthly Headcount
                    if join_col and dept_col:
                        data[join_col] = pd.to_datetime(data[join_col], errors='coerce')
                        if end_col: data[end_col] = pd.to_datetime(data[end_col], errors='coerce')
                        months = pd.date_range(f'{yr_sel}-01-01', f'{yr_sel}-12-01', freq='MS')
                        depts = sorted(data[dept_col].dropna().unique())

                        hc_rows = []
                        for dept in depts:
                            dd = data[data[dept_col]==dept]
                            row = {'Department': dept}
                            for m in months:
                                m_end = m + pd.offsets.MonthEnd(0)
                                joined = dd[dd[join_col] <= m_end]
                                if end_col:
                                    act = joined[(joined[end_col].isna()) | (joined[end_col] > m_end)]
                                else:
                                    act = joined
                                row[m.strftime('%b-%Y')] = len(act)
                            hc_rows.append(row)
                        hc_export = pd.DataFrame(hc_rows)
                        total_r = {'Department': 'TOTAL'}
                        for c in hc_export.columns[1:]: total_r[c] = hc_export[c].sum()
                        hc_export = pd.concat([hc_export, pd.DataFrame([total_r])], ignore_index=True)
                        hc_export.to_excel(w, sheet_name='HR Headcount Report', index=False)

                        # Sheet 3: Saudization
                        if nat_col:
                            saudi_vals = ['Saudi','سعودي','Saudi Arabian']
                            sa_rows = []
                            for m in months:
                                m_end = m + pd.offsets.MonthEnd(0)
                                joined = data[data[join_col] <= m_end]
                                if end_col: act = joined[(joined[end_col].isna()) | (joined[end_col] > m_end)]
                                else: act = joined
                                sa = len(act[act[nat_col].isin(saudi_vals)])
                                tot = len(act)
                                sa_rows.append({'Month': m.strftime('%b-%Y'), 'Saudi': sa, 'Non-Saudi': tot-sa,
                                    'Total': tot, 'Saudization %': round(sa/max(tot,1)*100,1)})
                            pd.DataFrame(sa_rows).to_excel(w, sheet_name='Saudization', index=False)

                        # Sheet 4: Hires & Terminations
                        ht_rows = []
                        for m in months:
                            m_start = m; m_end = m + pd.offsets.MonthEnd(0)
                            h = len(data[(data[join_col] >= m_start) & (data[join_col] <= m_end)])
                            t = len(data[(data[end_col] >= m_start) & (data[end_col] <= m_end)]) if end_col else 0
                            ht_rows.append({'Month': m.strftime('%b-%Y'), 'New Hires': h, 'Terminations': t, 'Net Change': h-t})
                        pd.DataFrame(ht_rows).to_excel(w, sheet_name='Hires & Terminations', index=False)

                    # Format workbook
                    for sname in w.sheets:
                        ws = w.sheets[sname]
                        ws.set_column('A:Z', 15)

                st.download_button("📥 تحميل Headcount Report", data=ox.getvalue(),
                    file_name=f"HR_Headcount_Report_{yr_sel}.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    type="primary", use_container_width=True)

        elif page == "📊 تحليل الأداء":
            hdr("📊 تحليل الأداء","تحليل إنتاجية وأداء الموظفين")

            data = sal_snapshot if len(sal_snapshot)>0 else emp

            if len(data)==0: st.info("📁 ارفع ملف"); return

            # Productivity metrics from salary data
            if has(data,'الراتب الإجمالي'):
                st.markdown("### 💰 مؤشرات التكلفة والإنتاجية")
                dept_col = 'القسم' if has(data,'القسم') else ('القطاع' if has(data,'القطاع') else None)

                c1,c2,c3 = st.columns(3)
                with c1: st.metric("💵 متوسط تكلفة الموظف/شهر", f"{data['الراتب الإجمالي'].mean():,.0f}")
                with c2: st.metric("📊 الانحراف المعياري", f"{data['الراتب الإجمالي'].std():,.0f}")
                with c3:
                    if has(data,'ساعات إضافية'):
                        st.metric("⏰ متوسط الإضافي", f"{data['ساعات إضافية'].mean():.1f} ساعة")

                if dept_col:
                    perf = data.groupby(dept_col).agg({
                        'الراتب الإجمالي': ['mean','sum','count'],
                    }).reset_index()
                    perf.columns = [dept_col, 'متوسط الراتب', 'إجمالي الرواتب', 'عدد الموظفين']
                    perf['نسبة التكلفة %'] = (perf['إجمالي الرواتب'] / perf['إجمالي الرواتب'].sum() * 100).round(1)
                    perf['نسبة العدد %'] = (perf['عدد الموظفين'] / perf['عدد الموظفين'].sum() * 100).round(1)
                    perf['كفاءة التكلفة'] = (perf['نسبة التكلفة %'] / perf['نسبة العدد %']).round(2)

                    st.markdown("### 📊 كفاءة التكلفة حسب القسم")
                    st.dataframe(perf.sort_values('كفاءة التكلفة', ascending=False), use_container_width=True, hide_index=True)

                    c1,c2 = st.columns(2)
                    with c1:
                        fig = px.scatter(perf, x='نسبة العدد %', y='نسبة التكلفة %', size='عدد الموظفين', color=dept_col,
                            title='العدد مقابل التكلفة (الحجم = عدد الموظفين)', color_discrete_sequence=CL['dept'])
                        fig.add_trace(go.Scatter(x=[0,50], y=[0,50], mode='lines', line=dict(dash='dash',color='gray'), name='خط التوازن'))
                        fig.update_layout(font=dict(family="Noto Sans Arabic"),height=400); st.plotly_chart(fig,use_container_width=True)
                    with c2:
                        fig = px.bar(perf.sort_values('كفاءة التكلفة'), x='كفاءة التكلفة', y=dept_col, orientation='h',
                            title='مؤشر كفاءة التكلفة (1 = متوازن)', color='كفاءة التكلفة', color_continuous_scale='RdYlGn_r')
                        fig.add_vline(x=1, line_dash="dash", line_color="gray")
                        fig.update_layout(font=dict(family="Noto Sans Arabic"),height=400); st.plotly_chart(fig,use_container_width=True)

                    ibox("مؤشر كفاءة التكلفة: إذا كان أكبر من 1 فالقسم يكلف أكثر من حجمه النسبي. إذا كان أقل من 1 فالقسم فعّال من حيث التكلفة.")

                # Overtime analysis as performance indicator
                if has(data,'ساعات إضافية') and dept_col:
                    st.markdown("### ⏰ تحليل الساعات الإضافية")
                    ot = data.groupby(dept_col)['ساعات إضافية'].agg(['mean','sum']).reset_index()
                    ot.columns = [dept_col, 'المتوسط','الإجمالي']
                    fig = px.bar(ot.sort_values('المتوسط',ascending=True), x='المتوسط', y=dept_col, orientation='h',
                        title='متوسط الساعات الإضافية حسب القسم', color='المتوسط', color_continuous_scale='teal')
                    fig.update_layout(font=dict(family="Noto Sans Arabic"),height=400); st.plotly_chart(fig,use_container_width=True)
            else:
                st.warning("لا يوجد بيانات رواتب للتحليل. ارفع ملف رواتب أو أضف عمود Gross Salary.")




    # =========================================
    #         ⚖️ LABOR CALCULATOR (MOJ-MATCHING)
    # =========================================
    elif section == "⚖️ حاسبة المستحقات":
        hdr("⚖️ الحاسبة العمالية الشاملة","مطابقة لحاسبة وزارة العدل - نظام العمل السعودي")

        # ===== Auto-fill from uploaded data =====
        data_source = sal_snapshot if len(sal_snapshot)>0 else emp
        auto_filled = False
        auto_vals = {}

        if len(data_source) > 0:
            st.markdown("### 🔍 بحث عن موظف من البيانات المرفوعة")
            search_method = st.radio("طريقة البحث:", ["يدوي (بدون بحث)","بحث بالاسم","بحث بالرقم الوظيفي"], horizontal=True, key="calc_search")

            if search_method != "يدوي (بدون بحث)":
                # Detect name and ID columns
                name_col = next((c for c in data_source.columns if any(x in c.lower() for x in ['name','اسم','full name'])), None)
                id_col = next((c for c in data_source.columns if any(x in c.lower() for x in ['emp id','employee id','رقم','id','code'])), None)
                sal_basic_col = next((c for c in data_source.columns if any(x in c.lower() for x in ['basic','أساسي','base'])), None)
                sal_housing_col = next((c for c in data_source.columns if any(x in c.lower() for x in ['housing','سكن','hous'])), None)
                sal_transport_col = next((c for c in data_source.columns if any(x in c.lower() for x in ['transport','مواصلات','trans'])), None)
                sal_other_col = next((c for c in data_source.columns if any(x in c.lower() for x in ['other','أخرى','allow'])), None)
                sal_gross_col = next((c for c in data_source.columns if any(x in c.lower() for x in ['gross','إجمالي','total sal','net'])), None)
                nat_col_c = next((c for c in data_source.columns if any(x in c.lower() for x in ['nat','جنسية','nationality'])), None)
                dept_col_c = next((c for c in data_source.columns if any(x in c.lower() for x in ['dept','قسم','department'])), None)
                title_col = next((c for c in data_source.columns if any(x in c.lower() for x in ['title','مسمى','position','وظيفة'])), None)
                join_col_c = next((c for c in data_source.columns if any(x in c.lower() for x in ['join','hiring','التحاق','مباشرة','start'])), None)

                found_emp = None
                if search_method == "بحث بالاسم" and name_col:
                    search_val = st.selectbox("اختر الموظف:", [""] + sorted(data_source[name_col].dropna().unique().tolist()), key="calc_name_sel")
                    if search_val:
                        found_emp = data_source[data_source[name_col] == search_val].iloc[0]
                elif search_method == "بحث بالرقم الوظيفي" and id_col:
                    search_val = st.text_input("أدخل الرقم الوظيفي:", key="calc_id_input")
                    if search_val:
                        matches = data_source[data_source[id_col].astype(str).str.contains(str(search_val), na=False)]
                        if len(matches) > 0:
                            found_emp = matches.iloc[0]
                        else:
                            st.warning("لم يتم العثور على موظف بهذا الرقم")
                elif search_method == "بحث بالاسم" and not name_col:
                    st.warning("لا يوجد عمود اسم في البيانات المرفوعة")
                elif search_method == "بحث بالرقم الوظيفي" and not id_col:
                    st.warning("لا يوجد عمود رقم وظيفي في البيانات المرفوعة")

                if found_emp is not None:
                    auto_filled = True
                    auto_vals['name'] = str(found_emp[name_col]) if name_col else ""
                    auto_vals['id'] = str(found_emp[id_col]) if id_col else ""
                    auto_vals['basic'] = float(found_emp[sal_basic_col]) if sal_basic_col and pd.notna(found_emp.get(sal_basic_col)) else 5000.0
                    auto_vals['housing'] = float(found_emp[sal_housing_col]) if sal_housing_col and pd.notna(found_emp.get(sal_housing_col)) else 0.0
                    auto_vals['transport'] = float(found_emp[sal_transport_col]) if sal_transport_col and pd.notna(found_emp.get(sal_transport_col)) else 0.0
                    auto_vals['other'] = float(found_emp[sal_other_col]) if sal_other_col and pd.notna(found_emp.get(sal_other_col)) else 0.0
                    auto_vals['nat'] = str(found_emp[nat_col_c]) if nat_col_c else ""
                    auto_vals['dept'] = str(found_emp[dept_col_c]) if dept_col_c else ""
                    auto_vals['title'] = str(found_emp[title_col]) if title_col else ""
                    auto_vals['join'] = found_emp[join_col_c] if join_col_c else None
                    auto_vals['gross'] = float(found_emp[sal_gross_col]) if sal_gross_col and pd.notna(found_emp.get(sal_gross_col)) else 0.0
                    st.success(f"✅ تم العثور على: {auto_vals['name']} | {auto_vals.get('dept','')} | {auto_vals.get('title','')}")

        # ===== بيانات الموظف =====
        st.markdown("---")
        st.markdown("### 👤 بيانات الموظف")
        e1, e2, e3 = st.columns([2,1,1])
        with e1: emp_name = st.text_input("اسم الموظف:", value=auto_vals.get('name','') if auto_filled else "", key="empn")
        with e2: emp_id = st.text_input("رقم الموظف:", value=auto_vals.get('id','') if auto_filled else "", key="empid")
        with e3:
            nat_default = 0 if (auto_filled and auto_vals.get('nat','') in ['Saudi','سعودي','Saudi Arabian','سعودية']) else (1 if auto_filled else 0)
            worker_type = st.radio("الجنسية:", ["سعودي","غير سعودي"], horizontal=True, key="wt", index=nat_default)

        # Extra info for report
        emp_dept = auto_vals.get('dept','') if auto_filled else ""
        emp_title = auto_vals.get('title','') if auto_filled else ""

        st.markdown("### 💵 تفاصيل الأجر")
        if auto_filled and auto_vals.get('gross',0) > 0 and auto_vals.get('basic',0) == 5000.0:
            st.caption("💡 تم استيراد الراتب الإجمالي. يمكنك تعديل التفاصيل يدوياً.")
        s1, s2, s3, s4 = st.columns(4)
        with s1: basic_sal = st.number_input("الأجر الأساسي:", min_value=0.0, max_value=500000.0, value=auto_vals.get('basic',5000.0) if auto_filled else 5000.0, step=0.01, format="%.2f", key="bsal")
        with s2: housing = st.number_input("بدل السكن:", min_value=0.0, max_value=500000.0, value=auto_vals.get('housing',1250.0) if auto_filled else 1250.0, step=0.01, format="%.2f", key="hous")
        with s3: transport = st.number_input("بدل المواصلات:", min_value=0.0, max_value=100000.0, value=auto_vals.get('transport',500.0) if auto_filled else 500.0, step=0.01, format="%.2f", key="trns")
        with s4: other_allow = st.number_input("بدلات أخرى:", min_value=0.0, max_value=500000.0, value=auto_vals.get('other',0.0) if auto_filled else 0.0, step=0.01, format="%.2f", key="otha")

        # GOSI calculation - Saudi only
        is_saudi = worker_type == "سعودي"
        if is_saudi:
            gosi_pct = st.slider("🏛️ نسبة خصم التأمينات الاجتماعية (%):", 0.0, 25.0, 9.75, 0.25, key="gosi_pct",
                help="النسبة الافتراضية: 9.75% (حصة الموظف)")
            gosi_base = basic_sal + housing
            gosi_deduction = gosi_base * (gosi_pct / 100)
            net_after_gosi = gosi_base - gosi_deduction
            total_sal = net_after_gosi + transport + other_allow

            g1,g2,g3,g4 = st.columns(4)
            with g1: kpi("وعاء التأمينات (أساسي+سكن)", f"{gosi_base:,.2f}")
            with g2: kpi(f"خصم التأمينات ({gosi_pct}%)", f"{gosi_deduction:,.2f}")
            with g3: kpi("بعد خصم التأمينات", f"{net_after_gosi:,.2f}")
            with g4: kpi("💰 الأجر النهائي (صافي)", f"{total_sal:,.2f}")

            ibox(f"طريقة الحساب: (الأساسي {basic_sal:,} + السكن {housing:,}) - التأمينات {gosi_pct}% = {net_after_gosi:,.2f} + المواصلات {transport:,} + أخرى {other_allow:,} = **{total_sal:,.2f} ريال**")
            gross_sal = basic_sal + housing + transport + other_allow
        else:
            gosi_pct = 0; gosi_deduction = 0
            gross_sal = basic_sal + housing + transport + other_allow
            total_sal = gross_sal
            g1,g2 = st.columns(2)
            with g1: kpi("إجمالي الأجر", f"{total_sal:,.2f}")
            with g2: kpi("💰 الأجر النهائي", f"{total_sal:,.2f}")
            ibox("غير سعودي: لا يوجد خصم تأمينات اجتماعية.", "success")

        daily_sal = gross_sal / 30
        results_summary = []

        # ========== 1. الأجور المتأخرة ==========
        st.markdown("---")
        st.markdown("### 💰 1. الأجور المتأخرة")
        dw_method = st.radio("طريقة الإدخال:", ["بإدخال التاريخ من إلى","بإدخال عدد الأشهر والأيام"], horizontal=True, key="dwm")
        if dw_method == "بإدخال التاريخ من إلى":
            d1, d2 = st.columns(2)
            with d1: dw_from = st.date_input("من:", value=date.today(), key="dwf")
            with d2: dw_to = st.date_input("إلى:", value=date.today(), key="dwt")
            dw_total_days = (dw_to - dw_from).days
        else:
            d1, d2 = st.columns(2)
            with d1: dw_months = st.number_input("عدد الأشهر:", 0, 120, 0, key="dwmo")
            with d2: dw_extra_days = st.number_input("عدد الأيام:", 0, 30, 0, key="dwdy")
            dw_total_days = dw_months * 30 + dw_extra_days
        delayed_amount = daily_sal * dw_total_days
        if dw_total_days > 0:
            st.success(f"الأجور المتأخرة: **{delayed_amount:,.2f} ريال** ({dw_total_days} يوم x {daily_sal:,.2f})")
            results_summary.append(("الأجور المتأخرة", delayed_amount))

        # ========== 2. مكافأة نهاية الخدمة ==========
        st.markdown("---")
        st.markdown("### 📊 2. مكافأة نهاية الخدمة")
        ec1, ec2 = st.columns(2)
        with ec1: eos_method = st.radio("طريقة الاحتساب:", ["حسب المادة (84)","حسب المادة (85) - استقالة"], key="eosm")
        with ec2: unpaid_leave = st.number_input("إجمالي أيام الإجازات بدون أجر:", 0, 9999, 0, key="unp")
        ec3, ec4 = st.columns(2)
        with ec3: eos_start = st.date_input("بداية العمل:", value=date(2018,1,1), key="eoss")
        with ec4: eos_end = st.date_input("نهاية العمل:", value=date.today(), key="eose")

        eos_service_days = max((eos_end - eos_start).days - unpaid_leave, 0)
        eos_years = eos_service_days / 365.25
        eos_delta = relativedelta(eos_end, eos_start)
        eos_monthly = gross_sal
        if eos_years <= 5:
            eos_84 = (eos_monthly / 2) * eos_years
        else:
            eos_84 = (eos_monthly / 2) * 5 + eos_monthly * (eos_years - 5)
        is_85 = "85" in eos_method
        if is_85:
            if eos_years < 2: eos_final=0; eos_pct=0; eos_note="لا يستحق (أقل من سنتين)"
            elif eos_years < 5: eos_final=eos_84/3; eos_pct=33.3; eos_note="ثلث المكافأة (2-5 سنوات)"
            elif eos_years < 10: eos_final=eos_84*2/3; eos_pct=66.7; eos_note="ثلثا المكافأة (5-10 سنوات)"
            else: eos_final=eos_84; eos_pct=100; eos_note="كاملة (10+ سنوات)"
        else:
            eos_final=eos_84; eos_pct=100; eos_note="المكافأة كاملة (المادة 84)"

        ek1,ek2,ek3,ek4 = st.columns(4)
        with ek1: kpi("مدة الخدمة", f"{eos_delta.years} سنة {eos_delta.months} شهر")
        with ek2: kpi("المكافأة كاملة (84)", f"{eos_84:,.2f}")
        with ek3: kpi(f"المستحق ({eos_pct}%)", f"{eos_final:,.2f}")
        with ek4: kpi("الأجر اليومي", f"{daily_sal:,.2f}")

        calc_rows = []
        if eos_years <= 5:
            calc_rows.append({"البند": f"{eos_years:.2f} سنة x نصف شهر", "المبلغ": f"{eos_84:,.2f}"})
        else:
            f5 = (eos_monthly/2)*5; r5 = eos_monthly*(eos_years-5)
            calc_rows.append({"البند": "أول 5 سنوات x نصف شهر", "المبلغ": f"{f5:,.2f}"})
            calc_rows.append({"البند": f"ما بعد 5 سنوات ({eos_years-5:.2f}) x شهر كامل", "المبلغ": f"{r5:,.2f}"})
        if is_85: calc_rows.append({"البند": f"المستحق (مادة 85): {eos_pct}%", "المبلغ": f"{eos_final:,.2f}"})
        st.dataframe(pd.DataFrame(calc_rows), use_container_width=True, hide_index=True)
        ibox(eos_note, "success" if eos_pct==100 else ("danger" if eos_pct==0 else "warning"))
        if unpaid_leave > 0: ibox(f"تم خصم {unpaid_leave} يوم إجازة بدون أجر.")
        ibox(f"المكافأة تُحسب على أساس الأجر الإجمالي قبل خصم التأمينات: {eos_monthly:,.2f} ريال")
        results_summary.append(("مكافأة نهاية الخدمة", eos_final))

        # ========== 3. أجر الإجازة ==========
        st.markdown("---")
        st.markdown("### 🏖️ 3. أجر الإجازة")
        vac_days_input = st.number_input("عدد أيام الإجازة المستحقة:", min_value=0.0, max_value=365.0, value=0.0, step=0.01, format="%.2f", key="vacd")
        vac_amount = daily_sal * vac_days_input
        if vac_days_input > 0:
            st.success(f"أجر الإجازة: **{vac_amount:,.2f} ريال** ({vac_days_input} يوم x {daily_sal:,.2f})")
            results_summary.append(("أجر الإجازة", vac_amount))

        # ========== 4. أجر العمل الإضافي ==========
        st.markdown("---")
        st.markdown("### ⏰ 4. أجر العمل الإضافي")
        ibox("المادة 107: أجر ساعة الإضافي = أجر الساعة + 50% (150%).")
        oc1, oc2, oc3 = st.columns(3)
        with oc1: ot_work_hours = st.selectbox("ساعات اليوم الفعلية:", list(range(2,13)), index=6, key="oth")
        with oc2: ot_days = st.number_input("عدد الأيام الإضافية:", 0, 365, 0, key="otd")
        with oc3: ot_hours = st.number_input("عدد الساعات الإضافية:", 0, 9999, 0, key="othr")
        ot_hourly = basic_sal / 30 / ot_work_hours
        ot_rate = ot_hourly * 1.5
        ot_total_hours = (ot_days * ot_work_hours) + ot_hours
        ot_amount = ot_total_hours * ot_rate
        if ot_total_hours > 0:
            st.success(f"ساعة الإضافي: {ot_rate:,.2f} | الساعات: {ot_total_hours} | **الإجمالي: {ot_amount:,.2f} ريال**")
            results_summary.append(("أجر العمل الإضافي", ot_amount))

        # ========== 5. التعويض عن الإنهاء غير المشروع ==========
        st.markdown("---")
        st.markdown("### 🚫 5. التعويض عن الإنهاء لغير سبب مشروع (المادة 77)")
        ibox("""**نص المادة 77:** ما لم يتضمن العقد تعويضاً محدداً مقابل إنهائه من أحد الطرفين لسبب غير مشروع:
1. أجر 15 يوماً عن كل سنة خدمة (غير محدد المدة)
2. أجر المدة الباقية من العقد (محدد المدة)
3. **لا يقل التعويض في كلا الحالتين عن أجر شهرين**""")

        contract_type = st.radio("نوع العقد:", ["عقد محدد المدة","عقد غير محدد المدة"], key="ctype")

        comp_method = st.radio("طريقة حساب التعويض:", [
            "① أجر أساسي شهر واحد (اتفاقي)",
            "② راتبين إجماليين كاملين (اتفاقي)",
            "③ الرواتب حتى نهاية العقد / 15 يوم لكل سنة (م.77 تلقائي)",
            "④ مبلغ مخصص (يدوي)"
        ], key="comp_method", index=2)

        unfair_amount = 0
        note77 = ""

        if "①" in comp_method:
            unfair_amount = basic_sal
            note77 = "تعويض اتفاقي: أجر أساسي شهر واحد"

        elif "②" in comp_method:
            unfair_amount = gross_sal * 2
            note77 = "تعويض اتفاقي: راتبين إجماليين كاملين"

        elif "③" in comp_method:
            if contract_type == "عقد محدد المدة":
                st.markdown("**المدة المتبقية من العقد:**")
                uc1, uc2 = st.columns(2)
                with uc1: ct_from = st.date_input("تاريخ الإنهاء:", value=date.today(), key="ctf")
                with uc2: ct_to = st.date_input("تاريخ انتهاء العقد:", value=date.today(), key="ctt")
                remaining_days = (ct_to - ct_from).days
                remaining_months = remaining_days / 30.0

                if remaining_days > 0:
                    comp_remaining = daily_sal * remaining_days
                    min_comp = gross_sal * 2

                    # المادة 77: لا يقل عن شهرين
                    if comp_remaining < min_comp:
                        unfair_amount = min_comp
                        note77 = f"المتبقي: {remaining_days} يوم ({remaining_months:.1f} شهر) = {comp_remaining:,.2f} ريال. لكنها أقل من الحد الأدنى (شهرين)، لذا التعويض = **{min_comp:,.2f} ريال** (أجر شهرين كاملين)"
                    else:
                        unfair_amount = comp_remaining
                        note77 = f"المتبقي: {remaining_days} يوم ({remaining_months:.1f} شهر) | التعويض = أجر المدة المتبقية"
                else:
                    st.warning("تاريخ انتهاء العقد يجب أن يكون بعد تاريخ الإنهاء")

            else:
                # عقد غير محدد المدة
                st.markdown("**بيانات الخدمة:**")
                ibox("""**العقد غير محدد المدة (المادة 55):** يتحول العقد المحدد المدة إلى غير محدد المدة في الحالات التالية:
- إذا استمر الطرفان في تنفيذه بعد انتهاء مدته
- إذا تجدد العقد لثلاث مرات متتالية
- إذا بلغت مدة العقد الأصلية مع التجديدات أربع سنوات (أيهما أقل)

**التعويض:** 15 يوم أجر عن كل سنة خدمة، بحد أدنى شهرين""")

                uc1, uc2 = st.columns(2)
                with uc1: uct_start = st.date_input("بداية العمل:", value=date(2018,1,1), key="ucts")
                with uc2: uct_end = st.date_input("تاريخ الإنهاء:", value=date.today(), key="ucte")
                svc_yrs = (uct_end - uct_start).days / 365.25

                if svc_yrs > 0:
                    # 15 يوم عن كل سنة
                    comp_15days = (daily_sal * 15) * svc_yrs
                    min_comp = gross_sal * 2

                    if comp_15days < min_comp:
                        unfair_amount = min_comp
                        note77 = f"الخدمة: {svc_yrs:.2f} سنة | 15 يوم/سنة = {comp_15days:,.2f} ريال. لكنها أقل من الحد الأدنى، لذا التعويض = **{min_comp:,.2f} ريال** (أجر شهرين)"
                    else:
                        unfair_amount = comp_15days
                        note77 = f"الخدمة: {svc_yrs:.2f} سنة | التعويض = 15 يوم × {svc_yrs:.2f} سنة"

        elif "④" in comp_method:
            unfair_amount = st.number_input("أدخل مبلغ التعويض يدوياً (ريال):", min_value=0.0, value=0.0, format="%.2f", key="custom_comp")
            note77 = "مبلغ تعويض مخصص (يدوي)"

        if unfair_amount > 0:
            st.success(f"**التعويض عن الفسخ غير المشروع: {unfair_amount:,.2f} ريال**")
            if note77:
                st.caption(f"📌 {note77}")
            results_summary.append(("تعويض إنهاء غير مشروع (م.77)", unfair_amount))

        # ========== 6. أيام الإجازة المستحقة ==========
        st.markdown("---")
        st.markdown("### 📅 6. أيام الإجازة المستحقة في فترة الخدمة")
        ibox("المادة 109: الحد الأدنى 21 يوم في أول 5 سنوات، 30 يوم بعدها.")
        vc1, vc2 = st.columns(2)
        with vc1: vd_first5 = st.number_input("أيام الإجازة في أول 5 سنوات:", min_value=21, max_value=60, value=21, key="vd5")
        with vc2: vd_after5 = st.number_input("أيام الإجازة بعد 5 سنوات:", min_value=30, max_value=60, value=30, key="vda5")
        vc3, vc4 = st.columns(2)
        with vc3: vd_from = st.date_input("من تاريخ:", value=date(2018,1,1), key="vdf")
        with vc4: vd_to = st.date_input("إلى تاريخ:", value=date.today(), key="vdt")
        vd_yrs = (vd_to - vd_from).days / 365.25
        vd_delta = relativedelta(vd_to, vd_from)
        vd_total = (vd_yrs * vd_first5) if vd_yrs <= 5 else (5 * vd_first5) + ((vd_yrs - 5) * vd_after5)
        if vd_yrs > 0:
            st.success(f"الخدمة: {vd_delta.years} سنة {vd_delta.months} شهر | **الإجازة المستحقة: {vd_total:.1f} يوم**")

        # ========== 7. حسم الغياب والتأخر ==========
        st.markdown("---")
        st.markdown("### 📉 7. حسم الغياب والتأخر")
        ac1, ac2 = st.columns(2)
        with ac1:
            abs_hours_day = st.selectbox("ساعات العمل اليومية:", list(range(2,13)), index=6, key="absh")
            abs_days = st.number_input("عدد أيام الغياب:", 0, 365, 0, key="absd")
        with ac2:
            abs_hours = st.number_input("عدد ساعات التأخر:", 0, 999, 0, key="abshr")
            abs_minutes = st.number_input("عدد دقائق التأخر:", 0, 59, 0, key="absmin")
        abs_hourly = daily_sal / abs_hours_day
        abs_minute_rate = abs_hourly / 60
        abs_day_ded = abs_days * daily_sal
        abs_hr_ded = abs_hours * abs_hourly
        abs_min_ded = abs_minutes * abs_minute_rate
        abs_total = abs_day_ded + abs_hr_ded + abs_min_ded
        if abs_total > 0:
            parts = []
            if abs_days > 0: parts.append(f"غياب {abs_days} يوم = {abs_day_ded:,.2f}")
            if abs_hours > 0: parts.append(f"تأخر {abs_hours} ساعة = {abs_hr_ded:,.2f}")
            if abs_minutes > 0: parts.append(f"تأخر {abs_minutes} دقيقة = {abs_min_ded:,.2f}")
            st.warning(f"{' | '.join(parts)} | **إجمالي الحسم: {abs_total:,.2f} ريال**")
            results_summary.append(("حسم الغياب والتأخر (يُخصم)", abs_total))

        # ========== 8. متوسط أجر آخر سنة ==========
        st.markdown("---")
        st.markdown("### 📊 8. متوسط الأجر لآخر سنة")
        ibox("يُستخدم عندما يكون الأجر متغيراً (عمولات، مكافآت).")
        months_ar = ["يناير","فبراير","مارس","أبريل","مايو","يونيو","يوليو","أغسطس","سبتمبر","أكتوبر","نوفمبر","ديسمبر"]
        month_sals = []
        for i in range(0, 12, 6):
            cols = st.columns(6)
            for j in range(6):
                if i+j < 12:
                    with cols[j]:
                        val = st.number_input(f"{months_ar[i+j]}:", 0, 500000, 0, 100, key=f"ms{i+j}")
                        month_sals.append(val)
        non_zero = [s for s in month_sals if s > 0]
        if non_zero:
            avg_12 = sum(month_sals) / 12
            avg_actual = sum(non_zero) / len(non_zero)
            st.success(f"الإجمالي: {sum(month_sals):,.0f} | المتوسط (12 شهر): **{avg_12:,.2f}** | المتوسط ({len(non_zero)} أشهر فعلية): {avg_actual:,.2f}")

        # ========================================
        #          الملخص النهائي + التصدير
        # ========================================
        st.markdown("---")
        st.markdown("### 🟰 ملخص المستحقات النهائية")
        if emp_name or emp_id:
            st.markdown(f"**الموظف:** {emp_name or '-'} | **الرقم:** {emp_id or '-'} | **الجنسية:** {worker_type}")

        if results_summary:
            grand_total = 0
            summary_rows = []
            for label, amount in results_summary:
                is_ded = "خصم" in label or "حسم" in label
                summary_rows.append({"البند": label, "المبلغ (ريال)": f"{amount:,.2f}", "النوع": "🔴 خصم" if is_ded else "🟢 استحقاق"})
                grand_total += (-amount if is_ded else amount)
            summary_rows.append({"البند": "🟰 صافي المستحقات النهائية", "المبلغ (ريال)": f"{grand_total:,.2f}", "النوع": "💰 الإجمالي"})
            st.dataframe(pd.DataFrame(summary_rows), use_container_width=True, hide_index=True)

            k1, k2 = st.columns(2)
            with k1: kpi("💰 صافي المستحقات النهائية", f"{grand_total:,.2f} ريال")
            with k2: kpi("📋 عدد البنود", f"{len(results_summary)}")

            # ===== EXPORT =====
            st.markdown("### 📥 تصدير التقرير")
            rpt_lang = st.radio("🌐 لغة التقرير:", ["English","العربية","English + العربية"], horizontal=True, key="rpt_lang")

            # Translation maps
            L = {}
            if "العربية" in rpt_lang and "English" not in rpt_lang:
                L = {"title":"بيان تسوية مستحقات نهاية الخدمة","emp_info":"بيانات الموظف",
                    "emp_name":"اسم الموظف","emp_id":"الرقم الوظيفي","dept":"القسم","job":"المسمى الوظيفي",
                    "join_date":"تاريخ الالتحاق","last_day":"آخر يوم عمل","svc_years":"سنوات الخدمة","svc_days":"أيام الخدمة",
                    "leave_bal":"رصيد الإجازات (أيام)","total_sal":"إجمالي الراتب الشهري",
                    "sal_breakdown":"تفصيل الراتب الشهري","num":"م","item":"البند","amount":"المبلغ (ريال)","pct":"النسبة",
                    "basic":"الراتب الأساسي","housing":"بدل السكن","transport":"بدل المواصلات","other":"بدلات أخرى",
                    "gross_total":"إجمالي الراتب","benefits":"تفصيل المستحقات والبدلات",
                    "calc_method":"طريقة الحساب","details":"التفاصيل","legal":"السند النظامي",
                    "net_total":"صافي المستحقات النهائية","legal_basis":"الأساس النظامي (نظام العمل السعودي)",
                    "currency":"ريال","settlement":"تاريخ التسوية","prepared":"الإعداد","hr_dept":"إدارة الموارد البشرية",
                    "eos":"مكافأة نهاية الخدمة","leave_cash":"بدل إجازة","delayed":"أجور متأخرة",
                    "overtime":"عمل إضافي","unfair":"تعويض فسخ غير مشروع","absence":"حسم غياب وتأخر",
                    "art84":"م.84: مكافأة نهاية الخدمة - نصف راتب عن كل سنة من الخمس الأولى، وراتب كامل عن كل سنة بعدها",
                    "art85":"م.85: الاستقالة - ثلث المكافأة (2-5 سنوات)، ثلثان (5-10)، كاملة (10+)",
                    "art77":"م.77: الفسخ غير المشروع - 15 يوم/سنة (غير محدد) أو المدة المتبقية (محدد)، بحد أدنى شهرين",
                    "art109":"م.109: بدل الإجازة - تعويض نقدي عن رصيد الإجازات غير المستخدمة",
                    "art88":"م.88: يجب تصفية جميع المستحقات خلال 7 أيام من انتهاء العقد"}
            else:
                L = {"title":"End of Service Benefits Settlement Statement","emp_info":"Employee Information",
                    "emp_name":"Employee Name","emp_id":"Employee ID","dept":"Department","job":"Job Title",
                    "join_date":"Joining Date","last_day":"Last Working Day","svc_years":"Service (Years)","svc_days":"Service (Days)",
                    "leave_bal":"Leave Balance (Days)","total_sal":"Total Monthly Salary",
                    "sal_breakdown":"Monthly Salary Breakdown","num":"#","item":"Item","amount":"Amount (SAR)","pct":"Percentage",
                    "basic":"Basic Salary","housing":"Housing Allowance","transport":"Transportation Allowance","other":"Other Allowances",
                    "gross_total":"Total Gross Salary","benefits":"Benefits & Entitlements Detail",
                    "calc_method":"Calculation Method","details":"Details","legal":"Legal Basis",
                    "net_total":"NET TOTAL ENTITLEMENTS","legal_basis":"Legal Basis (Saudi Labor Law)",
                    "currency":"SAR","settlement":"Settlement Date","prepared":"Prepared by","hr_dept":"Human Resources Department",
                    "eos":"End of Service Award","leave_cash":"Leave Encashment","delayed":"Delayed Wages",
                    "overtime":"Overtime Pay","unfair":"Unfair Termination Compensation","absence":"Absence Deduction",
                    "art84":"Art. 84: End of Service Award - Half salary per year for first 5 years, full salary per year thereafter",
                    "art85":"Art. 85: Resignation - 1/3 of award (2-5 yrs), 2/3 (5-10 yrs), full (10+ yrs)",
                    "art77":"Art. 77: Unfair Termination - 15 days/year (indefinite) or remaining term (fixed), min 2 months",
                    "art109":"Art. 109: Leave Encashment - Cash compensation for unused annual leave balance",
                    "art88":"Art. 88: Employer must settle all dues within 7 days of contract termination"}

            # If bilingual, add Arabic in parentheses
            if "English + العربية" in rpt_lang:
                AR = {"title":"بيان تسوية مستحقات نهاية الخدمة","emp_info":"بيانات الموظف",
                    "emp_name":"اسم الموظف","emp_id":"الرقم الوظيفي","dept":"القسم","job":"المسمى الوظيفي",
                    "join_date":"تاريخ الالتحاق","last_day":"آخر يوم عمل","svc_years":"سنوات الخدمة","svc_days":"أيام الخدمة",
                    "leave_bal":"رصيد الإجازات","total_sal":"إجمالي الراتب",
                    "sal_breakdown":"تفصيل الراتب","basic":"الراتب الأساسي","housing":"بدل السكن",
                    "transport":"بدل المواصلات","other":"بدلات أخرى","gross_total":"إجمالي الراتب",
                    "benefits":"تفصيل المستحقات","net_total":"صافي المستحقات النهائية","legal_basis":"الأساس النظامي"}
                for k in AR:
                    if k in L:
                        L[k] = f"{L[k]} / {AR[k]}"

            export_rows = [
                {"Item": L['emp_name'], "Value": emp_name or "-"},
                {"Item": L['emp_id'], "Value": emp_id or "-"},
                {"Item": L.get('dept','Department'), "Value": emp_dept or "-"},
                {"Item": L.get('job','Job Title'), "Value": emp_title or "-"},
                {"Item": L['basic'], "Value": f"{basic_sal:,.2f}"},
                {"Item": L['housing'], "Value": f"{housing:,.2f}"},
                {"Item": L['transport'], "Value": f"{transport:,.2f}"},
                {"Item": L['other'], "Value": f"{other_allow:,.2f}"},
                {"Item": L['gross_total'], "Value": f"{gross_sal:,.2f}"},
                {"Item": "---", "Value": "---"}]
            for label, amount in results_summary:
                is_ded = "خصم" in label or "حسم" in label
                export_rows.append({"Item": label, "Value": f"{'-' if is_ded else ''}{amount:,.2f}"})
            export_rows.append({"Item": "---", "Value": "---"})
            export_rows.append({"Item": L['net_total'], "Value": f"{grand_total:,.2f}"})

            # ===== PROFESSIONAL EXCEL (matching MOJ template) =====
            ox = io.BytesIO()
            wb_exp = openpyxl.Workbook()
            ws_exp = wb_exp.active
            ws_exp.title = "تسوية المستحقات" if "العربية" in rpt_lang and "English" not in rpt_lang else "EOS Settlement"
            ws_exp.sheet_view.rightToLeft = False

            from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
            from openpyxl.utils import get_column_letter

            # Colors
            dark_blue = PatternFill("solid", fgColor="1F4E79")
            med_blue = PatternFill("solid", fgColor="2E75B6")
            hdr_blue = PatternFill("solid", fgColor="4472C4")
            light1 = PatternFill("solid", fgColor="D6E4F0")
            light2 = PatternFill("solid", fgColor="EBF1F8")

            white_font = Font(bold=True, color="FFFFFF", size=16, name="Calibri")
            white_font12 = Font(bold=True, color="FFFFFF", size=12, name="Calibri")
            white_font13 = Font(bold=True, color="FFFFFF", size=13, name="Calibri")
            white_font11 = Font(bold=True, color="FFFFFF", size=11, name="Calibri")
            bold11 = Font(bold=True, size=11, name="Calibri")
            normal11 = Font(size=11, name="Calibri")
            blue_val = Font(size=11, color="0000FF", name="Calibri")
            red_val = Font(size=11, color="FF0000", name="Calibri")
            small10 = Font(size=10, name="Calibri")
            gray10 = Font(size=10, color="808080", name="Calibri")

            thin_border = Border(
                left=Side(style='thin', color='B0B0B0'),
                right=Side(style='thin', color='B0B0B0'),
                top=Side(style='thin', color='B0B0B0'),
                bottom=Side(style='thin', color='B0B0B0'))

            center = Alignment(horizontal='center', vertical='center', wrap_text=True)
            left = Alignment(horizontal='left', vertical='center', wrap_text=True)
            right_al = Alignment(horizontal='right', vertical='center', wrap_text=True)

            # Column widths
            ws_exp.column_dimensions['A'].width = 8
            ws_exp.column_dimensions['B'].width = 35
            ws_exp.column_dimensions['C'].width = 35
            ws_exp.column_dimensions['D'].width = 20
            ws_exp.column_dimensions['E'].width = 20
            ws_exp.column_dimensions['F'].width = 18

            r = 1  # current row

            def write_merged(row, col1, col2, value, font, fill, align_style=center):
                ws_exp.merge_cells(start_row=row, start_column=col1, end_row=row, end_column=col2)
                c = ws_exp.cell(row=row, column=col1, value=value)
                c.font = font; c.fill = fill; c.alignment = align_style; c.border = thin_border
                for cc in range(col1+1, col2+1):
                    ws_exp.cell(row=row, column=cc).fill = fill
                    ws_exp.cell(row=row, column=cc).border = thin_border

            def write_cell(row, col, value, font=normal11, fill=None, align_style=center):
                c = ws_exp.cell(row=row, column=col, value=value)
                c.font = font; c.alignment = align_style; c.border = thin_border
                if fill: c.fill = fill

            # === ROW 1: Title ===
            write_merged(r, 1, 6, L['title'], white_font, dark_blue)
            ws_exp.row_dimensions[r].height = 35
            r += 1

            # === ROW 2: Employee Info Header ===
            write_merged(r, 1, 6, L['emp_info'], white_font12, med_blue)
            r += 1

            # === ROW 3: Name & ID ===
            bg = light1
            write_cell(r, 1, L['emp_name'], bold11, bg, left)
            ws_exp.merge_cells(start_row=r, start_column=2, end_row=r, end_column=3)
            write_cell(r, 2, emp_name or "-", normal11, bg, center)
            ws_exp.cell(r,3).fill=bg; ws_exp.cell(r,3).border=thin_border
            write_cell(r, 4, L['emp_id'], bold11, bg, left)
            ws_exp.merge_cells(start_row=r, start_column=5, end_row=r, end_column=6)
            write_cell(r, 5, emp_id or "-", normal11, bg, center)
            ws_exp.cell(r,6).fill=bg; ws_exp.cell(r,6).border=thin_border
            r += 1

            # === ROW 4: Dept & Title ===
            bg = light2
            write_cell(r, 1, L['dept'], bold11, bg, left)
            ws_exp.merge_cells(start_row=r, start_column=2, end_row=r, end_column=3)
            write_cell(r, 2, emp_dept or "-", normal11, bg, center)
            ws_exp.cell(r,3).fill=bg; ws_exp.cell(r,3).border=thin_border
            write_cell(r, 4, L['job'], bold11, bg, left)
            ws_exp.merge_cells(start_row=r, start_column=5, end_row=r, end_column=6)
            write_cell(r, 5, emp_title or "-", normal11, bg, center)
            ws_exp.cell(r,6).fill=bg; ws_exp.cell(r,6).border=thin_border
            r += 1

            # === ROW 5: Dates ===
            bg = light1
            eos_start_str = eos_start.strftime('%Y-%m-%d') if hasattr(eos_start, 'strftime') else str(eos_start)
            eos_end_str = eos_end.strftime('%Y-%m-%d') if hasattr(eos_end, 'strftime') else str(eos_end)
            write_cell(r, 1, L['join_date'], bold11, bg, left)
            ws_exp.merge_cells(start_row=r, start_column=2, end_row=r, end_column=3)
            write_cell(r, 2, eos_start_str, normal11, bg, center)
            ws_exp.cell(r,3).fill=bg; ws_exp.cell(r,3).border=thin_border
            write_cell(r, 4, L['last_day'], bold11, bg, left)
            ws_exp.merge_cells(start_row=r, start_column=5, end_row=r, end_column=6)
            write_cell(r, 5, eos_end_str, normal11, bg, center)
            ws_exp.cell(r,6).fill=bg; ws_exp.cell(r,6).border=thin_border
            r += 1

            # === ROW 6: Service duration ===
            bg = light2
            write_cell(r, 1, L['svc_years'], bold11, bg, left)
            ws_exp.merge_cells(start_row=r, start_column=2, end_row=r, end_column=3)
            write_cell(r, 2, round(eos_years, 2), blue_val, bg, center)
            ws_exp.cell(r,3).fill=bg; ws_exp.cell(r,3).border=thin_border
            write_cell(r, 4, L['svc_days'], bold11, bg, left)
            ws_exp.merge_cells(start_row=r, start_column=5, end_row=r, end_column=6)
            write_cell(r, 5, eos_service_days, blue_val, bg, center)
            ws_exp.cell(r,6).fill=bg; ws_exp.cell(r,6).border=thin_border
            r += 1

            # === ROW 7: Leave balance & total salary ===
            bg = light1
            write_cell(r, 1, L['leave_bal'], bold11, bg, left)
            ws_exp.merge_cells(start_row=r, start_column=2, end_row=r, end_column=3)
            write_cell(r, 2, vac_days_input, blue_val, bg, center)
            ws_exp.cell(r,3).fill=bg; ws_exp.cell(r,3).border=thin_border
            write_cell(r, 4, L['total_sal'], bold11, bg, left)
            ws_exp.merge_cells(start_row=r, start_column=5, end_row=r, end_column=6)
            write_cell(r, 5, round(gross_sal, 2), blue_val, bg, center)
            ws_exp.cell(r,6).fill=bg; ws_exp.cell(r,6).border=thin_border
            r += 1

            # === ROW 8: Salary Details Header ===
            write_merged(r, 1, 6, L['sal_breakdown'], white_font12, med_blue)
            r += 1

            # === ROW 9: Salary table header ===
            sal_headers = [L['num'], L['item'], "", L['amount'], L['pct'], ""]
            for i, h in enumerate(sal_headers, 1):
                write_cell(r, i, h, white_font11, hdr_blue, center)
            ws_exp.merge_cells(start_row=r, start_column=2, end_row=r, end_column=3)
            ws_exp.merge_cells(start_row=r, start_column=5, end_row=r, end_column=6)
            r += 1

            # === ROWS 9-12: Salary items ===
            sal_items = [
                (L['basic'], basic_sal),
                (L['housing'], housing),
                (L['transport'], transport),
                (L['other'], other_allow),
            ]
            for idx, (item, amt) in enumerate(sal_items, 1):
                bg = light1 if idx % 2 == 1 else light2
                pct = (amt / gross_sal * 100) if gross_sal > 0 else 0
                write_cell(r, 1, idx, normal11, bg, center)
                ws_exp.merge_cells(start_row=r, start_column=2, end_row=r, end_column=3)
                write_cell(r, 2, item, normal11, bg, left)
                ws_exp.cell(r,3).fill=bg; ws_exp.cell(r,3).border=thin_border
                write_cell(r, 4, round(amt, 2), blue_val, bg, center)
                ws_exp.merge_cells(start_row=r, start_column=5, end_row=r, end_column=6)
                write_cell(r, 5, f"{pct:.1f}%", normal11, bg, center)
                ws_exp.cell(r,6).fill=bg; ws_exp.cell(r,6).border=thin_border
                r += 1

            # === Total Salary ===
            ws_exp.merge_cells(start_row=r, start_column=2, end_row=r, end_column=3)
            write_cell(r, 2, L['gross_total'], white_font12, dark_blue, center)
            ws_exp.cell(r,3).fill=dark_blue; ws_exp.cell(r,3).border=thin_border
            write_cell(r, 4, round(gross_sal, 2), white_font12, dark_blue, center)
            ws_exp.merge_cells(start_row=r, start_column=5, end_row=r, end_column=6)
            write_cell(r, 5, "100%", white_font12, dark_blue, center)
            ws_exp.cell(r,6).fill=dark_blue; ws_exp.cell(r,6).border=thin_border
            write_cell(r, 1, "", normal11, dark_blue, center)
            r += 1

            # === Benefits Details Header ===
            write_merged(r, 1, 6, L['benefits'], white_font12, med_blue)
            r += 1

            # === Benefits table header ===
            ben_headers = [L['num'], L['item'], L['calc_method'], L['details'], L['amount'], L['legal']]
            for i, h in enumerate(ben_headers, 1):
                write_cell(r, i, h, white_font11, hdr_blue, center)
            r += 1

            # === Benefits rows ===
            ben_idx = 0
            for label, amount in results_summary:
                ben_idx += 1
                bg = light1 if ben_idx % 2 == 1 else light2
                is_ded = "خصم" in label or "حسم" in label

                # Determine calculation method and legal basis (English)
                if "نهاية الخدمة" in label:
                    calc_method = f"Half salary x years (first 5 yrs) + full salary x remaining years"
                    details = f"{eos_years:.2f} years | Article {'85' if is_85 else '84'} ({eos_pct}%)"
                    legal = f"Article {'85' if is_85 else '84'}"
                elif "إجازة" in label:
                    calc_method = f"Salary / 30 x leave days"
                    details = f"{vac_days_input} days x {daily_sal:,.2f} SAR/day"
                    legal = "Article 109"
                elif "متأخرة" in label:
                    calc_method = "Salary / 30 x delayed days"
                    details = f"{dw_total_days} days"
                    legal = "Article 88"
                elif "إضافي" in label:
                    calc_method = "Overtime rate (150%) x hours"
                    details = f"{ot_total_hours} hours x {ot_rate:,.2f}"
                    legal = "Article 107"
                elif "إنهاء" in label or "تعويض" in label:
                    calc_method = "Unfair termination compensation"
                    details = f"Contract type: {contract_type}"
                    legal = "Article 77"
                elif "حسم" in label or "غياب" in label:
                    calc_method = "Daily wage x absent days + late hours"
                    details = "Deducted from total"
                    legal = "Labor Law"
                else:
                    calc_method = "-"; details = "-"; legal = "-"

                write_cell(r, 1, ben_idx, normal11, bg, center)
                write_cell(r, 2, label, normal11, bg, left)
                write_cell(r, 3, calc_method, small10, bg, center)
                write_cell(r, 4, details, small10, bg, center)
                val_font = red_val if is_ded else blue_val
                write_cell(r, 5, round(-amount if is_ded else amount, 2), val_font, bg, center)
                write_cell(r, 6, legal, small10, bg, center)
                r += 1

            # === Total Benefits Row ===
            ws_exp.merge_cells(start_row=r, start_column=2, end_row=r, end_column=4)
            write_cell(r, 1, "", normal11, dark_blue, center)
            write_cell(r, 2, L['net_total'], white_font13, dark_blue, center)
            ws_exp.cell(r,3).fill=dark_blue; ws_exp.cell(r,3).border=thin_border
            ws_exp.cell(r,4).fill=dark_blue; ws_exp.cell(r,4).border=thin_border
            write_cell(r, 5, round(grand_total, 2), white_font13, dark_blue, center)
            write_cell(r, 6, L['currency'], white_font11, dark_blue, center)
            r += 1

            # === Legal Basis Section ===
            write_merged(r, 1, 6, L['legal_basis'], white_font12, med_blue)
            r += 1

            legal_notes = [L['art84'], L['art85'], L['art77'], L['art109'], L['art88']]
            for i, note in enumerate(legal_notes):
                bg = light1 if i % 2 == 0 else light2
                write_merged(r, 1, 6, note, small10, bg, left)
                r += 1

            # === Settlement Date ===
            write_cell(r, 1, f"{L['settlement']}:", gray10, None, left)
            write_cell(r, 2, datetime.now().strftime('%Y-%m-%d'), gray10, None, center)
            r += 1
            write_cell(r, 1, f"{L['prepared']}:", gray10, None, left)
            write_cell(r, 2, L['hr_dept'], gray10, None, left)

            # Set RTL for Arabic
            is_rtl = "العربية" in rpt_lang and "English" not in rpt_lang
            ws_exp.sheet_view.rightToLeft = is_rtl

            wb_exp.save(ox)

            fname = f"EOS_Settlement_{emp_name or 'Employee'}_{datetime.now().strftime('%Y%m%d')}"
            xc1, xc2 = st.columns(2)
            with xc1:
                st.download_button("📥 تحميل Excel", data=ox.getvalue(), file_name=f"{fname}.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", type="primary", use_container_width=True)
            with xc2:
                csv_df = pd.DataFrame(export_rows)
                csv_data = csv_df.to_csv(index=False).encode('utf-8-sig')
                st.download_button("📥 تحميل CSV", data=csv_data, file_name=f"{fname}.csv", mime="text/csv", use_container_width=True)
        else:
            st.info("عبّئ البيانات أعلاه وستظهر المستحقات هنا تلقائياً")

        ibox("This is an approximate advisory calculation and does not substitute for specialized legal consultation.", "warning")

    # =========================================
    #         📚 TRAINING & DEVELOPMENT
    # =========================================
    elif section == "📚 التدريب والتطوير":

        if 'budget_data' not in st.session_state:
            st.session_state.budget_data = DEFAULT_BUDGET.copy()

        if page == "📚 ميزانية التدريب":
            hdr("📚 ميزانية التدريب","خطة توزيع ميزانية التدريب السنوية - عدّل مباشرة في الجدول")
            c1,c2 = st.columns(2)
            with c1: total_budget = st.number_input("💰 إجمالي الميزانية (ريال)", 10000, 5000000, 70000, 5000, key="trn_bgt")
            with c2: fy = st.selectbox("📅 السنة", [2025,2026,2027], index=1, key="trn_yr")

            # Initialize editable data
            if 'custom_budget' not in st.session_state:
                st.session_state.custom_budget = []
                for d in DEFAULT_BUDGET:
                    row = dict(d)
                    row['trainees'] = 5
                    row['trainee_names'] = ''
                    row['job_titles'] = ''
                    st.session_state.custom_budget.append(row)

            # Build editable dataframe
            edit_data = []
            for d in st.session_state.custom_budget:
                budget_val = int(d['pct'] / 100 * total_budget)
                trainees = d.get('trainees', 5) or 5
                edit_data.append({
                    'القسم': d['dept'],
                    'الميزانية': budget_val,
                    'النسبة %': d['pct'],
                    'الأولوية': d['priority'],
                    'التصنيف': d['cat'],
                    'المتدربين': trainees,
                    'للفرد': int(budget_val / max(trainees, 1)),
                    'المسميات الوظيفية': d.get('job_titles', ''),
                    'أسماء المتدربين': d.get('trainee_names', ''),
                })

            edit_df = pd.DataFrame(edit_data)

            # Editable table
            st.markdown("### ✏️ عدّل مباشرة في الجدول (اضغط على أي خلية)")
            edited = st.data_editor(
                edit_df,
                column_config={
                    'القسم': st.column_config.TextColumn('القسم', disabled=True, width='medium'),
                    'الميزانية': st.column_config.NumberColumn('الميزانية', disabled=True, format='%d'),
                    'النسبة %': st.column_config.NumberColumn('النسبة %', min_value=0, max_value=100, step=0.5, format='%.1f'),
                    'الأولوية': st.column_config.SelectboxColumn('الأولوية', options=['حرج','عالي','متوسط','أساسي'], required=True),
                    'التصنيف': st.column_config.SelectboxColumn('التصنيف', options=['محرك إيرادات','ممكّن نمو','بنية تحتية'], required=True),
                    'المتدربين': st.column_config.NumberColumn('المتدربين', min_value=1, max_value=200, step=1),
                    'للفرد': st.column_config.NumberColumn('للفرد', disabled=True, format='%d'),
                    'المسميات الوظيفية': st.column_config.TextColumn('المسميات الوظيفية', width='large',
                        help='أدخل المسميات مفصولة بفاصلة: مدير مبيعات, مندوب, محلل'),
                    'أسماء المتدربين': st.column_config.TextColumn('أسماء المتدربين', width='large',
                        help='أدخل الأسماء مفصولة بفاصلة: أحمد محمد, سارة أحمد'),
                },
                use_container_width=True, hide_index=True, num_rows="fixed",
                key="budget_editor"
            )

            # Sync edits back to session_state
            for i in range(len(edited)):
                st.session_state.custom_budget[i]['pct'] = edited.iloc[i]['النسبة %']
                st.session_state.custom_budget[i]['priority'] = edited.iloc[i]['الأولوية']
                st.session_state.custom_budget[i]['cat'] = edited.iloc[i]['التصنيف']
                st.session_state.custom_budget[i]['trainees'] = int(edited.iloc[i]['المتدربين'])
                st.session_state.custom_budget[i]['job_titles'] = edited.iloc[i].get('المسميات الوظيفية', '')
                st.session_state.custom_budget[i]['trainee_names'] = edited.iloc[i].get('أسماء المتدربين', '')

            # Recalculate
            budget_df = pd.DataFrame(st.session_state.custom_budget)
            budget_df['budget'] = (budget_df['pct']/100*total_budget).astype(int)
            if 'trainees' not in budget_df.columns: budget_df['trainees'] = 5

            # Warnings
            total_pct = edited['النسبة %'].sum()
            if abs(total_pct - 100) > 0.5:
                st.warning(f"⚠️ مجموع النسب = {total_pct:.1f}% (يجب أن يكون 100%)")

            # KPIs
            st.markdown("---")
            k1,k2,k3,k4,k5 = st.columns(5)
            with k1: kpi("💰 الميزانية", f"{total_budget:,}")
            with k2: kpi("📌 الأقسام", str(len(budget_df)))
            with k3:
                rev = budget_df[budget_df['cat']=='محرك إيرادات']['budget'].sum()
                kpi("محركات الإيرادات", f"{round(rev/max(total_budget,1)*100)}%")
            with k4:
                total_trainees = budget_df['trainees'].sum()
                kpi("👥 المتدربين", str(total_trainees))
            with k5:
                avg_per_person = total_budget / max(total_trainees, 1)
                kpi("💵 للفرد", f"{avg_per_person:,.0f}")

            # Trainee details per department
            st.markdown("---")
            st.markdown("### 👥 تفاصيل المتدربين حسب القسم")
            for i, d in enumerate(st.session_state.custom_budget):
                names = d.get('trainee_names', '')
                titles = d.get('job_titles', '')
                trainees = d.get('trainees', 5)
                dept_budget = int(d['pct'] / 100 * total_budget)
                if names or titles:
                    name_list = [n.strip() for n in names.split(',') if n.strip()] if names else []
                    title_list = [t.strip() for t in titles.split(',') if t.strip()] if titles else []
                    max_len = max(len(name_list), len(title_list), 1)
                    detail_rows = []
                    for j in range(max_len):
                        detail_rows.append({
                            'الاسم': name_list[j] if j < len(name_list) else '',
                            'المسمى الوظيفي': title_list[j] if j < len(title_list) else '',
                            'الميزانية المخصصة': f"{int(dept_budget / max(trainees, 1)):,} ريال"
                        })
                    with st.expander(f"📌 {d['dept']} ({len(name_list)} متدرب - {dept_budget:,} ريال)"):
                        st.dataframe(pd.DataFrame(detail_rows), use_container_width=True, hide_index=True)

            # Charts
            c1,c2 = st.columns(2)
            with c1:
                fig = px.pie(budget_df, values='budget', names='dept', title='توزيع الميزانية حسب القسم', hole=.35, color_discrete_sequence=CL['dept'])
                fig.update_layout(font=dict(family="Noto Sans Arabic"),height=380); st.plotly_chart(fig,use_container_width=True)
            with c2:
                cat_df = budget_df.groupby('cat')['budget'].sum().reset_index()
                fig = px.pie(cat_df, values='budget', names='cat', title='التوزيع الاستراتيجي', hole=.35,
                    color_discrete_map={'محرك إيرادات':CL['p'],'ممكّن نمو':CL['a'],'بنية تحتية':'#64748B'})
                fig.update_layout(font=dict(family="Noto Sans Arabic"),height=380); st.plotly_chart(fig,use_container_width=True)

            # Quarterly plan
            st.markdown("### 📅 الخطة ربع السنوية")
            q_data = []
            for _, r in budget_df.iterrows():
                qr = {"القسم":r['dept']}
                for q,p in Q_SPLIT.items(): qr[q] = int(r['budget']*p)
                qr['الإجمالي'] = r['budget']
                q_data.append(qr)
            q_df = pd.DataFrame(q_data)
            totals = {"القسم":"الإجمالي"}
            for c in ['Q1','Q2','Q3','Q4','الإجمالي']: totals[c] = q_df[c].sum()
            q_df = pd.concat([q_df, pd.DataFrame([totals])], ignore_index=True)
            st.dataframe(q_df, use_container_width=True, hide_index=True)

            # Detailed programs per department
            st.markdown("### 📋 البرامج التدريبية المفصلة")
            for dept_name, programs in TRAINING_PROGRAMS.items():
                with st.expander(f"📌 {dept_name} ({len(programs)} برامج - {sum(p['budget'] for p in programs):,} ريال)"):
                    prog_df = pd.DataFrame(programs)
                    prog_df.columns = ['البرنامج','الميزانية','المصدر','التوقيت','الأثر المتوقع']
                    st.dataframe(prog_df, use_container_width=True, hide_index=True)

            # Reset button
            if st.button("🔄 إعادة التعيين للقيم الافتراضية", key="trn_reset"):
                st.session_state.custom_budget = []
                for d in DEFAULT_BUDGET:
                    row = dict(d); row['trainees'] = 5; row['trainee_names'] = ''; row['job_titles'] = ''
                    st.session_state.custom_budget.append(row)
                st.rerun()



            export_widget({"ميزانية التدريب": pd.DataFrame(st.session_state.get("budget_data",[])), "البرامج": pd.DataFrame(st.session_state.get("training_programs_list",[]))} if st.session_state.get("budget_data") else None, "ميزانية_التدريب", "trn1")
        elif page == "💹 ROI التدريب":
            hdr("💹 عائد التدريب ROI","نموذج Phillips ذو 5 مستويات")
            c1,c2 = st.columns(2)
            with c1:
                rb = st.number_input("💰 ميزانية التدريب", value=70000, step=5000)
                cr = st.number_input("📈 الإيرادات السنوية", value=5000000, step=100000)
                ri = st.slider("📊 الزيادة المتوقعة %", 1, 50, 15)
            with c2:
                hc2 = st.number_input("👥 عدد الموظفين", value=83)
                as2 = st.number_input("💵 متوسط الراتب الشهري", value=8000, step=500)
                rt = st.slider("🔄 تحسن الاحتفاظ %", 1, 30, 10)
                pg = st.slider("⚡ الإنتاجية %", 1, 30, 10)

            if st.button("📊 حساب ROI", type="primary", use_container_width=True):
                roi = calc_roi(rb, ri, cr, rt, as2*12, hc2, pg)
                k1,k2,k3 = st.columns(3)
                with k1: kpi("ROI", f"{roi['roi']:.0f}%")
                with k2: kpi("BCR", f"{roi['bcr']:.1f}x")
                with k3: kpi("الاسترداد", f"{roi['payback']:.1f} شهر")

                fig = go.Figure()
                fig.add_trace(go.Bar(x=['الإيرادات','الاحتفاظ','الإنتاجية'], y=[roi['rev'],roi['ret'],roi['prod']], marker_color=[CL['p'],CL['a'],CL['s']]))
                fig.add_hline(y=rb, line_dash="dash", line_color="red", annotation_text=f"التكلفة: {rb:,}")
                fig.update_layout(title='العوائد مقابل التكلفة', font=dict(family="Noto Sans Arabic"), height=380, yaxis_tickformat=',')
                st.plotly_chart(fig, use_container_width=True)



        elif page == "📋 خطة ADDIE":
            hdr("📋 خطة التدريب وفق نموذج ADDIE","Analysis → Design → Development → Implementation → Evaluation")

            # Skills-to-provider mapping
            SKILL_PROVIDER_MAP = {
                "بيع استشاري": ["PwC Academy","Informa Connect","Dale Carnegie UAE"],
                "CRM": ["بكه للتعليم","Coursera for Business","LinkedIn Learning"],
                "تفاوض": ["Dale Carnegie UAE","Informa Connect","معهد الإدارة العامة"],
                "تسويق رقمي": ["Google Certificates","Udacity MENA","Sprints"],
                "SEO": ["Google Certificates","LinkedIn Learning","Coursera for Business"],
                "Growth Hacking": ["Udacity MENA","Sprints","Coursera for Business"],
                "Python/SQL": ["Udacity MENA","Coursera for Business","Sprints","Google Certificates"],
                "Power BI": ["LinkedIn Learning","Coursera for Business","بكه للتعليم"],
                "AI": ["Udacity MENA","Misk Academy","Google Certificates","Coursera for Business"],
                "IFRS": ["KPMG Academy","PwC Academy","BIBF"],
                "نمذجة مالية": ["KPMG Academy","PwC Academy","Coursera for Business"],
                "استقطاب": ["معهد الإدارة العامة","LinkedIn Learning","الجامعة الأمريكية بالقاهرة"],
                "أداء": ["معهد الإدارة العامة","بكه للتعليم","Informa Connect"],
                "OKRs": ["بكه للتعليم","Coursera for Business","LinkedIn Learning"],
            }

            # ADDIE Tabs
            addie_tabs = st.tabs(["🔍 1. Analysis (تحليل)","🎨 2. Design (تصميم)","🔧 3. Development (تطوير)","🚀 4. Implementation (تنفيذ)","📊 5. Evaluation (تقييم)"])

            # Initialize ADDIE state
            if 'addie_plan' not in st.session_state:
                st.session_state.addie_plan = {'needs':[],'programs':[],'status':'Analysis'}

            # ===== PHASE 1: ANALYSIS =====
            with addie_tabs[0]:
                st.markdown("### 🔍 المرحلة 1: تحليل الاحتياجات التدريبية (Needs Analysis)")
                ibox("""**نموذج ADDIE - مرحلة التحليل:** تحديد فجوات الأداء والمهارات بين الوضع الحالي والمستهدف.
تشمل: تحليل المنظمة (Organizational Analysis) + تحليل المهام (Task Analysis) + تحليل الأفراد (Person Analysis)""")

                # Organizational Analysis
                st.markdown("#### 🏢 1.1 تحليل المنظمة (Organizational Analysis)")
                oa1, oa2 = st.columns(2)
                with oa1:
                    org_goals = st.text_area("الأهداف الاستراتيجية للشركة:", placeholder="مثال: زيادة المبيعات 30%، التوسع في السوق المصري", key="addie_goals", height=80)
                    org_challenges = st.text_area("التحديات الرئيسية:", placeholder="مثال: نقص الكفاءات التقنية، ارتفاع الدوران الوظيفي", key="addie_challenges", height=80)
                with oa2:
                    org_budget = st.number_input("ميزانية التدريب المتاحة (ريال):", value=70000, step=5000, key="addie_budget")
                    org_timeline = st.selectbox("الإطار الزمني:", ["Q1 (يناير-مارس)","Q2 (أبريل-يونيو)","Q3 (يوليو-سبتمبر)","Q4 (أكتوبر-ديسمبر)","السنة كاملة"], key="addie_timeline")

                # Task & Person Analysis
                st.markdown("#### 📋 1.2 تحليل المهام والأفراد (Task & Person Analysis)")
                cats = {"المبيعات":["بيع استشاري","CRM","تفاوض"],"التسويق":["تسويق رقمي","SEO","Growth Hacking"],
                        "التقنية":["Python/SQL","Power BI","AI"],"المالية":["IFRS","نمذجة مالية"],"الموارد البشرية":["استقطاب","أداء","OKRs"]}
                depts = st.multiselect("📌 الأقسام المستهدفة", list(cats.keys()), default=list(cats.keys())[:3], key="addie_depts")
                needs = []
                for d in depts:
                    with st.expander(f"📌 {d}", expanded=True):
                        skills = st.multiselect(f"المهارات المطلوبة", cats[d], default=cats[d][:2], key=f"as_{d}")
                        for s in skills:
                            c1,c2,c3 = st.columns(3)
                            with c1: lv = st.select_slider(f"الحالي: {s}", ["مبتدئ","أساسي","متوسط","متقدم","خبير"], value="أساسي", key=f"al_{d}_{s}")
                            with c2: tg = st.select_slider(f"المستهدف: {s}", ["مبتدئ","أساسي","متوسط","متقدم","خبير"], value="متقدم", key=f"at_{d}_{s}")
                            with c3: n_ppl = st.number_input(f"عدد المتدربين:", 1, 50, 5, key=f"an_{d}_{s}")
                            levels = ["مبتدئ","أساسي","متوسط","متقدم","خبير"]
                            gap = levels.index(tg) - levels.index(lv)
                            providers = SKILL_PROVIDER_MAP.get(s, ["Coursera for Business","LinkedIn Learning"])
                            needs.append({"القسم":d, "المهارة":s, "الحالي":lv, "المستهدف":tg, "الفجوة":gap,
                                "المتدربين": n_ppl,
                                "الأولوية": "حرج" if gap >= 3 else ("عالي" if gap >= 2 else "متوسط"),
                                "جهات التدريب": " | ".join(providers[:3])})

                if needs:
                    st.session_state.addie_plan['needs'] = needs
                    ndf = pd.DataFrame(needs)
                    st.markdown("#### 📊 1.3 خريطة الفجوات")
                    st.dataframe(ndf, use_container_width=True, hide_index=True)

                    nc1, nc2 = st.columns(2)
                    with nc1:
                        fig = px.bar(ndf, x='المهارة', y='الفجوة', color='القسم', title='فجوات المهارات حسب القسم', color_discrete_sequence=CL['dept'])
                        fig.update_layout(font=dict(family="Noto Sans Arabic"), height=380)
                        st.plotly_chart(fig, use_container_width=True)
                    with nc2:
                        fig = px.scatter(ndf, x='الفجوة', y='المتدربين', size='المتدربين', color='الأولوية',
                            title='مصفوفة الأولويات (الفجوة × العدد)',
                            color_discrete_map={'حرج':'#E74C3C','عالي':'#E36414','متوسط':'#F39C12'})
                        fig.update_layout(font=dict(family="Noto Sans Arabic"), height=380)
                        st.plotly_chart(fig, use_container_width=True)

            # ===== PHASE 2: DESIGN =====
            with addie_tabs[1]:
                st.markdown("### 🎨 المرحلة 2: تصميم البرنامج التدريبي (Instructional Design)")
                ibox("""**نموذج ADDIE - مرحلة التصميم:** تحديد أهداف التعلم، المحتوى، الأساليب، والتقييم.
تشمل: أهداف SMART + اختيار أساليب التدريب + تصميم المحتوى + معايير النجاح""")

                needs = st.session_state.addie_plan.get('needs', [])
                if not needs:
                    st.info("أكمل مرحلة التحليل أولاً")
                else:
                    programs = []
                    for i, need in enumerate(needs):
                        if need['الفجوة'] > 0:
                            with st.expander(f"🎯 {need['المهارة']} ({need['القسم']}) - أولوية: {need['الأولوية']}", expanded=i<3):
                                dc1, dc2 = st.columns(2)
                                with dc1:
                                    obj = st.text_input(f"هدف التعلم (SMART):",
                                        value=f"تطوير مهارة {need['المهارة']} من {need['الحالي']} إلى {need['المستهدف']}",
                                        key=f"dobj_{i}")
                                    method = st.selectbox("أسلوب التدريب:", [
                                        "Instructor-Led Training (ILT)",
                                        "Virtual Instructor-Led (VILT)",
                                        "E-Learning / Online",
                                        "Blended Learning (مدمج)",
                                        "On-the-Job Training (OJT)",
                                        "Coaching & Mentoring",
                                        "Workshop / ورشة عمل",
                                        "Self-Paced Learning",
                                    ], key=f"dmethod_{i}")
                                with dc2:
                                    duration = st.selectbox("المدة:", ["يوم واحد","يومين","3 أيام","أسبوع","أسبوعين","شهر","3 أشهر"], key=f"ddur_{i}")
                                    timing = st.selectbox("التوقيت:", ["Q1","Q2","Q3","Q4"], key=f"dtiming_{i}")
                                    provider = st.selectbox("الجهة:", need['جهات التدريب'].split(' | ') + ["أخرى"], key=f"dprov_{i}")
                                    est_cost = st.number_input("التكلفة التقديرية (ريال):", value=3000, step=500, key=f"dcost_{i}")

                                programs.append({
                                    "المهارة": need['المهارة'], "القسم": need['القسم'],
                                    "الهدف": obj, "الأسلوب": method, "المدة": duration,
                                    "التوقيت": timing, "الجهة": provider, "التكلفة": est_cost,
                                    "المتدربين": need['المتدربين'], "الأولوية": need['الأولوية']
                                })

                    if programs:
                        st.session_state.addie_plan['programs'] = programs
                        st.markdown("#### 📊 ملخص التصميم")
                        prog_df = pd.DataFrame(programs)
                        st.dataframe(prog_df[['المهارة','القسم','الأسلوب','المدة','التوقيت','الجهة','التكلفة','المتدربين']], use_container_width=True, hide_index=True)
                        st.success(f"💰 إجمالي التكلفة التقديرية: {sum(p['التكلفة'] for p in programs):,} ريال | 👥 إجمالي المتدربين: {sum(p['المتدربين'] for p in programs)}")

            # ===== PHASE 3: DEVELOPMENT =====
            with addie_tabs[2]:
                st.markdown("### 🔧 المرحلة 3: تطوير المحتوى التدريبي (Content Development)")
                ibox("""**نموذج ADDIE - مرحلة التطوير:** إعداد المواد التدريبية والموارد.
تشمل: إعداد المحتوى + تطوير المواد + إنتاج الوسائط + مراجعة الجودة""")

                programs = st.session_state.addie_plan.get('programs', [])
                if not programs:
                    st.info("أكمل مرحلة التصميم أولاً")
                else:
                    st.markdown("#### 📝 Checklist تطوير المحتوى")
                    dev_checklist = {
                        "إعداد المحتوى": ["تحديد المراجع والمصادر العلمية","كتابة الأهداف التعليمية التفصيلية","إعداد خطة الجلسات (Session Plan)","تصميم أنشطة التعلم التفاعلية"],
                        "المواد التدريبية": ["دليل المدرب (Facilitator Guide)","دليل المتدرب (Participant Workbook)","عروض تقديمية (PowerPoint/Slides)","حالات دراسية (Case Studies)","تمارين وأنشطة عملية"],
                        "الوسائط": ["فيديوهات تعليمية","إنفوجرافيك ملخصات","اختبارات ذاتية (Self-Assessment)","محتوى E-Learning (SCORM)"],
                        "ضمان الجودة": ["مراجعة المحتوى من خبير المجال (SME)","اختبار تجريبي (Pilot Test)","جمع ملاحظات وتحسين","اعتماد نهائي من الإدارة"],
                    }
                    for section, items in dev_checklist.items():
                        st.markdown(f"**{section}:**")
                        for item in items:
                            st.checkbox(item, key=f"dev_{section}_{item}")

                    # Resource allocation
                    st.markdown("---")
                    st.markdown("#### 📊 توزيع الموارد حسب البرنامج")
                    for i, prog in enumerate(programs):
                        with st.expander(f"📌 {prog['المهارة']} ({prog['القسم']})"):
                            st.markdown(f"**الأسلوب:** {prog['الأسلوب']} | **الجهة:** {prog['الجهة']} | **المدة:** {prog['المدة']}")
                            rc1, rc2 = st.columns(2)
                            with rc1:
                                content_ready = st.slider("جاهزية المحتوى %:", 0, 100, 0, key=f"dev_ready_{i}")
                            with rc2:
                                dev_notes = st.text_input("ملاحظات:", key=f"dev_notes_{i}")

            # ===== PHASE 4: IMPLEMENTATION =====
            with addie_tabs[3]:
                st.markdown("### 🚀 المرحلة 4: التنفيذ (Implementation)")
                ibox("""**نموذج ADDIE - مرحلة التنفيذ:** تقديم البرنامج التدريبي للمتدربين.
تشمل: الجدولة + التنسيق اللوجستي + تنفيذ التدريب + المتابعة اليومية""")

                programs = st.session_state.addie_plan.get('programs', [])
                if not programs:
                    st.info("أكمل المراحل السابقة أولاً")
                else:
                    st.markdown("#### 📅 الجدول الزمني للتنفيذ")
                    impl_data = []
                    for p in programs:
                        q_map = {"Q1":"يناير-مارس","Q2":"أبريل-يونيو","Q3":"يوليو-سبتمبر","Q4":"أكتوبر-ديسمبر"}
                        impl_data.append({
                            "البرنامج": p['المهارة'], "القسم": p['القسم'], "الجهة": p['الجهة'],
                            "الأسلوب": p['الأسلوب'], "المدة": p['المدة'],
                            "الفترة": q_map.get(p['التوقيت'],''), "المتدربين": p['المتدربين'],
                            "الحالة": "مجدول"
                        })
                    impl_df = pd.DataFrame(impl_data)
                    edited_impl = st.data_editor(impl_df, column_config={
                        'الحالة': st.column_config.SelectboxColumn('الحالة', options=['مجدول','جاري','مكتمل','مؤجل','ملغي'])
                    }, use_container_width=True, hide_index=True, key="impl_editor")

                    # Implementation checklist
                    st.markdown("#### ✅ Checklist التنفيذ")
                    impl_checks = [
                        "تأكيد الحجوزات والقاعات/المنصات","إرسال دعوات للمتدربين","تجهيز المواد والأدوات",
                        "تأكيد حضور المدرب","توفير الدعم التقني","جمع توقيعات الحضور",
                        "متابعة يومية للبرنامج","تصوير وتوثيق الجلسات"
                    ]
                    for ic in impl_checks:
                        st.checkbox(ic, key=f"impl_{ic}")

                    # Gantt-like visualization
                    q_start = {'Q1':1,'Q2':4,'Q3':7,'Q4':10}
                    q_end = {'Q1':3,'Q2':6,'Q3':9,'Q4':12}
                    gantt_data = []
                    for p in programs:
                        s_month = q_start.get(p['التوقيت'], 1)
                        e_month = q_end.get(p['التوقيت'], 3)
                        gantt_data.append({"Task":p['المهارة'], "Start":f"2026-{s_month:02d}-01",
                            "Finish":f"2026-{e_month:02d}-28", "القسم":p['القسم']})
                    if gantt_data:
                        fig = px.timeline(pd.DataFrame(gantt_data),
                            x_start="Start", x_end="Finish", y="Task", color="القسم",
                            title="📅 الجدول الزمني (Gantt Chart)", color_discrete_sequence=CL['dept'])
                        fig.update_layout(font=dict(family="Noto Sans Arabic"), height=350)
                        st.plotly_chart(fig, use_container_width=True)

            # ===== PHASE 5: EVALUATION =====
            with addie_tabs[4]:
                st.markdown("### 📊 المرحلة 5: التقييم (Evaluation) - Kirkpatrick's 4 Levels")
                ibox("""**نموذج ADDIE + Kirkpatrick - مرحلة التقييم:**
**Level 1 - Reaction:** رضا المتدربين عن البرنامج
**Level 2 - Learning:** مدى اكتساب المعرفة والمهارات
**Level 3 - Behavior:** تطبيق المهارات في بيئة العمل
**Level 4 - Results:** الأثر على نتائج الأعمال (KPIs)
**Level 5 - ROI (Phillips):** العائد المالي من الاستثمار في التدريب""")

                programs = st.session_state.addie_plan.get('programs', [])
                if not programs:
                    st.info("أكمل المراحل السابقة أولاً")
                else:
                    for i, prog in enumerate(programs):
                        with st.expander(f"📊 تقييم: {prog['المهارة']} ({prog['القسم']})", expanded=i==0):
                            ev1, ev2 = st.columns(2)
                            with ev1:
                                st.markdown("**Level 1: Reaction (رد الفعل)**")
                                l1_satisfaction = st.slider("رضا المتدربين (1-5):", 1.0, 5.0, 4.0, 0.1, key=f"ev1_{i}")
                                l1_relevance = st.slider("ملاءمة المحتوى (1-5):", 1.0, 5.0, 4.0, 0.1, key=f"ev1r_{i}")

                                st.markdown("**Level 2: Learning (التعلم)**")
                                l2_pre = st.number_input("نتيجة الاختبار القبلي %:", 0, 100, 40, key=f"ev2pre_{i}")
                                l2_post = st.number_input("نتيجة الاختبار البعدي %:", 0, 100, 75, key=f"ev2post_{i}")

                            with ev2:
                                st.markdown("**Level 3: Behavior (السلوك)**")
                                l3_apply = st.slider("نسبة تطبيق المهارات في العمل %:", 0, 100, 60, key=f"ev3_{i}")
                                l3_timeline = st.selectbox("فترة القياس:", ["30 يوم","60 يوم","90 يوم"], key=f"ev3t_{i}")

                                st.markdown("**Level 4: Results (النتائج)**")
                                l4_kpi = st.text_input("KPI المستهدف:", placeholder="مثال: زيادة المبيعات 15%", key=f"ev4_{i}")
                                l4_achieved = st.slider("نسبة تحقيق KPI %:", 0, 100, 50, key=f"ev4a_{i}")

                            # Results summary
                            learning_gain = l2_post - l2_pre
                            st.markdown("**📊 ملخص التقييم:**")
                            ek1,ek2,ek3,ek4 = st.columns(4)
                            with ek1: kpi("⭐ الرضا", f"{l1_satisfaction}/5")
                            with ek2: kpi("📈 التعلم", f"+{learning_gain}%")
                            with ek3: kpi("🔄 التطبيق", f"{l3_apply}%")
                            with ek4: kpi("🎯 النتائج", f"{l4_achieved}%")

                    # Overall ADDIE Summary
                    st.markdown("---")
                    st.markdown("### 📊 ملخص خطة ADDIE الشاملة")
                    total_programs = len(programs)
                    total_cost = sum(p['التكلفة'] for p in programs)
                    total_trainees = sum(p['المتدربين'] for p in programs)

                    sk1,sk2,sk3,sk4 = st.columns(4)
                    with sk1: kpi("📚 البرامج", str(total_programs))
                    with sk2: kpi("👥 المتدربين", str(total_trainees))
                    with sk3: kpi("💰 التكلفة", f"{total_cost:,}")
                    with sk4: kpi("💵 للفرد", f"{total_cost//max(total_trainees,1):,}")

                    # Export full ADDIE plan
                    if st.button("📥 تصدير خطة ADDIE الكاملة", type="primary", use_container_width=True, key="exp_addie"):
                        ox = io.BytesIO()
                        with pd.ExcelWriter(ox, engine='xlsxwriter') as w:
                            if needs: pd.DataFrame(needs).to_excel(w, sheet_name='1-Analysis', index=False)
                            if programs:
                                pd.DataFrame(programs).to_excel(w, sheet_name='2-Design', index=False)
                                pd.DataFrame(programs).to_excel(w, sheet_name='4-Implementation', index=False)
                            # Summary
                            summary = pd.DataFrame([
                                {"المرحلة":"Analysis","الحالة":"مكتمل","التفاصيل":f"{len(needs)} فجوة محددة"},
                                {"المرحلة":"Design","الحالة":"مكتمل","التفاصيل":f"{len(programs)} برنامج مصمم"},
                                {"المرحلة":"Development","الحالة":"جاري","التفاصيل":"إعداد المحتوى"},
                                {"المرحلة":"Implementation","الحالة":"مجدول","التفاصيل":f"{total_trainees} متدرب"},
                                {"المرحلة":"Evaluation","الحالة":"مخطط","التفاصيل":"Kirkpatrick 4 Levels"},
                            ])
                            summary.to_excel(w, sheet_name='ADDIE Summary', index=False)
                        st.download_button("📥 تحميل", data=ox.getvalue(),
                            file_name=f"ADDIE_Training_Plan_{datetime.now().strftime('%Y%m%d')}.xlsx",
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

        elif page == "🏫 جهات التدريب":
            hdr("🏫 دليل جهات التدريب")
            market = st.selectbox("🌍 السوق:", list(PROVIDERS.keys()))
            for p in PROVIDERS[market]:
                st.markdown(f"**{p['name']}** | {p['spec']} | {p['type']} | [{p['url']}](https://{p['url']})")
                st.markdown("---")

        elif page == "📥 تصدير التدريب":
            hdr("📥 تصدير ميزانية التدريب","Excel مطابق لنموذج Resal Training Budget")

            budget_df = pd.DataFrame(st.session_state.budget_data)
            total_budget = budget_df['budget'].sum()
            fy = st.selectbox("السنة المالية:", [2025,2026,2027], index=1, key="trn_fy")

            if st.button("📥 إنشاء ملف ميزانية التدريب", type="primary", use_container_width=True):
                ox = io.BytesIO()
                with pd.ExcelWriter(ox, engine='xlsxwriter') as w:
                    wb = w.book

                    # Header formats
                    hdr_fmt = wb.add_format({'bold':True,'font_size':14,'bg_color':'#0F4C5C','font_color':'white','align':'center','valign':'vcenter','border':1})
                    sub_fmt = wb.add_format({'bold':True,'font_size':11,'bg_color':'#E36414','font_color':'white','align':'center','border':1})
                    col_fmt = wb.add_format({'bold':True,'bg_color':'#264653','font_color':'white','align':'center','border':1,'text_wrap':True})
                    num_fmt = wb.add_format({'num_format':'#,##0','align':'center','border':1})
                    pct_fmt = wb.add_format({'num_format':'0.0%','align':'center','border':1})
                    txt_fmt = wb.add_format({'align':'right','border':1,'text_wrap':True})
                    tot_fmt = wb.add_format({'bold':True,'bg_color':'#F0F0F0','num_format':'#,##0','align':'center','border':1})

                    # ===== Sheet 1: Executive Summary =====
                    ws1 = wb.add_worksheet('Executive Summary')
                    ws1.set_column('A:A', 5)
                    ws1.set_column('B:B', 35)
                    ws1.set_column('C:G', 15)
                    ws1.merge_range('B1:G1', f'resal - Training Budget Allocation FY {fy}', hdr_fmt)
                    ws1.merge_range('B3:G3', 'Infrastructure for Non-Cash Value | Training Budget Allocation Plan', sub_fmt)
                    ws1.merge_range('B5:G5', 'BUDGET ALLOCATION BY DEPARTMENT', wb.add_format({'bold':True,'font_size':12,'align':'center'}))

                    headers = ['Department','Budget (SAR)','Allocation %','Priority','Strategic Fit','Category']
                    for i, h in enumerate(headers):
                        ws1.write(5, i+1, h, col_fmt)

                    for j, (_, r) in enumerate(budget_df.iterrows()):
                        ws1.write(6+j, 1, r['dept'], txt_fmt)
                        ws1.write(6+j, 2, r['budget'], num_fmt)
                        ws1.write(6+j, 3, r['pct']/100, pct_fmt)
                        ws1.write(6+j, 4, r['priority'], txt_fmt)
                        ws1.write(6+j, 5, 'Direct' if r['cat'] in ['محرك إيرادات','ممكّن نمو'] else 'Support', txt_fmt)
                        ws1.write(6+j, 6, r['cat'], txt_fmt)

                    tr = 6 + len(budget_df)
                    ws1.write(tr, 1, 'TOTAL', wb.add_format({'bold':True,'border':1}))
                    ws1.write(tr, 2, total_budget, tot_fmt)
                    ws1.write(tr, 3, 1.0, wb.add_format({'bold':True,'num_format':'0%','align':'center','border':1}))

                    # Strategic split
                    tr += 2
                    ws1.merge_range(tr, 1, tr, 6, 'STRATEGIC BUDGET SPLIT', sub_fmt)
                    rev_b = budget_df[budget_df['cat']=='محرك إيرادات']['budget'].sum()
                    grow_b = budget_df[budget_df['cat']=='ممكّن نمو']['budget'].sum()
                    inf_b = budget_df[budget_df['cat']=='بنية تحتية']['budget'].sum()
                    for i, (lbl, val) in enumerate([
                        ('Revenue-Generating Departments', rev_b),
                        ('Growth Enabler Departments', grow_b),
                        ('Infrastructure Departments', inf_b)]):
                        ws1.write(tr+1+i, 1, lbl, txt_fmt)
                        ws1.write(tr+1+i, 2, val, num_fmt)
                        ws1.write(tr+1+i, 3, val/total_budget, pct_fmt)

                    # ===== Sheet 2: Detailed Programs =====
                    ws2 = wb.add_worksheet('Detailed Programs')
                    ws2.set_column('A:A', 5)
                    ws2.set_column('B:B', 50)
                    ws2.set_column('C:C', 15)
                    ws2.set_column('D:E', 12)
                    ws2.set_column('F:F', 50)
                    ws2.merge_range('B1:F1', 'DETAILED TRAINING PROGRAMS BY DEPARTMENT', hdr_fmt)
                    row = 2
                    for dept, programs in TRAINING_PROGRAMS.items():
                        ws2.merge_range(row, 1, row, 5, dept, sub_fmt)
                        row += 1
                        for i, h in enumerate(['Training Program','Budget (SAR)','Source','Timing','Expected Impact']):
                            ws2.write(row, i+1, h, col_fmt)
                        row += 1
                        for p in programs:
                            ws2.write(row, 1, p['program'], txt_fmt)
                            ws2.write(row, 2, p['budget'], num_fmt)
                            ws2.write(row, 3, p['source'], txt_fmt)
                            ws2.write(row, 4, p['timing'], txt_fmt)
                            ws2.write(row, 5, p['impact'], txt_fmt)
                            row += 1
                        ws2.write(row, 1, 'Subtotal', wb.add_format({'bold':True,'border':1}))
                        ws2.write(row, 2, sum(p['budget'] for p in programs), tot_fmt)
                        row += 2

                    # ===== Sheet 3: Quarterly Plan =====
                    ws3 = wb.add_worksheet('Quarterly Plan')
                    ws3.set_column('A:A', 5)
                    ws3.set_column('B:B', 35)
                    ws3.set_column('C:G', 15)
                    ws3.merge_range('B1:G1', f'QUARTERLY TRAINING BUDGET PLAN FY {fy}', hdr_fmt)
                    for i, h in enumerate(['Department','Q1 Jan-Mar','Q2 Apr-Jun','Q3 Jul-Sep','Q4 Oct-Dec','Annual Total']):
                        ws3.write(2, i+1, h, col_fmt)
                    for j, (_, r) in enumerate(budget_df.iterrows()):
                        ws3.write(3+j, 1, r['dept'], txt_fmt)
                        for qi, (q, p) in enumerate(Q_SPLIT.items()):
                            ws3.write(3+j, 2+qi, int(r['budget']*p), num_fmt)
                        ws3.write(3+j, 6, r['budget'], num_fmt)
                    tr = 3 + len(budget_df)
                    ws3.write(tr, 1, 'TOTAL', wb.add_format({'bold':True,'border':1}))
                    for qi, (q, p) in enumerate(Q_SPLIT.items()):
                        ws3.write(tr, 2+qi, int(total_budget*p), tot_fmt)
                    ws3.write(tr, 6, total_budget, tot_fmt)

                    # ===== Sheet 4: Charts Data =====
                    chart_rows = [[r['dept'], r['budget']] for _, r in budget_df.iterrows()]
                    cdf = pd.DataFrame(chart_rows, columns=['Department','Budget'])
                    cdf.to_excel(w, sheet_name='Charts & Analytics', index=False, startrow=1)

                    # ===== Sheet 5: ROI & KPIs =====
                    ws5 = wb.add_worksheet('ROI & KPIs')
                    ws5.set_column('A:A', 5)
                    ws5.set_column('B:C', 50)
                    ws5.merge_range('B1:C1', 'EXPECTED ROI & KEY PERFORMANCE INDICATORS', hdr_fmt)
                    ws5.write(3, 1, 'Training ROI Indicators', sub_fmt)
                    for i, ind in enumerate(ROI_INDICATORS):
                        ws5.write(4+i, 1, f'{i+1}.', txt_fmt)
                        ws5.write(4+i, 2, ind, txt_fmt)
                    r = 5 + len(ROI_INDICATORS)
                    ws5.write(r, 1, 'Training KPIs', sub_fmt)
                    for i, kp in enumerate(TRAINING_KPIS):
                        ws5.write(r+1+i, 1, f'{i+1}.', txt_fmt)
                        ws5.write(r+1+i, 2, kp, txt_fmt)
                    r2 = r + 2 + len(TRAINING_KPIS)
                    ws5.write(r2, 1, f'Total Budget: SAR {total_budget:,}', txt_fmt)
                    ws5.write(r2+2, 1, 'Prepared by: Human Resources Department | resal', txt_fmt)

                st.download_button("📥 تحميل ميزانية التدريب الكاملة", data=ox.getvalue(),
                    file_name=f"Resal_Training_Budget_{fy}.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    type="primary", use_container_width=True)

            # Also show programs summary
            st.markdown("---")
            st.markdown("### 📋 ملخص البرامج التدريبية")
            all_programs = []
            for dept, progs in TRAINING_PROGRAMS.items():
                for p in progs:
                    all_programs.append({"القسم": dept, "البرنامج": p['program'], "الميزانية": p['budget'], "المصدر": p['source'], "التوقيت": p['timing']})
            st.dataframe(pd.DataFrame(all_programs), use_container_width=True, hide_index=True)
            st.caption(f"📊 إجمالي: {len(all_programs)} برنامج تدريبي | {sum(p['budget'] for p in all_programs)} ريال")


    # =========================================
    #         🎯 RECRUITMENT MODULE
    # =========================================
    elif section == "🎯 التوظيف":

        # Initialize recruitment data
        if 'recruit_plans' not in st.session_state:
            st.session_state.recruit_plans = []
        if 'recruit_tracking' not in st.session_state:
            st.session_state.recruit_tracking = []

        if page == "📋 تخطيط التوظيف":
            hdr("📋 تخطيط ميزانية التوظيف", "تخطيط تكاليف التوظيف الجديد وتقدير الميزانية السنوية")

            st.markdown("### ➕ إضافة وظيفة جديدة للخطة")
            rc1, rc2, rc3 = st.columns(3)
            with rc1:
                rp_title = st.text_input("المسمى الوظيفي:", key="rpt")
                rp_dept = st.text_input("القسم:", key="rpd")
                rp_count = st.number_input("عدد المطلوب:", 1, 50, 1, key="rpc")
            with rc2:
                rp_salary = st.number_input("الراتب الشهري المتوقع:", 0.0, 200000.0, 5000.0, 100.0, format="%.2f", key="rps")
                rp_housing_pct = st.number_input("% بدل السكن:", 0.0, 100.0, 25.0, key="rph")
                rp_transport = st.number_input("بدل المواصلات:", 0.0, 10000.0, 500.0, format="%.2f", key="rptr")
            with rc3:
                rp_agency_fee = st.number_input("رسوم الاستقدام/التوظيف:", 0.0, 100000.0, 0.0, format="%.2f", key="rpaf")
                rp_visa = st.number_input("تكلفة التأشيرة/الإقامة:", 0.0, 50000.0, 0.0, format="%.2f", key="rpv")
                rp_training = st.number_input("تكلفة التدريب الأولي:", 0.0, 50000.0, 0.0, format="%.2f", key="rptrn")
                rp_nationality = st.selectbox("الجنسية:", ["سعودي","غير سعودي"], key="rpnat")

            rp_housing = rp_salary * (rp_housing_pct / 100)
            rp_monthly_total = rp_salary + rp_housing + rp_transport
            rp_gosi = (rp_salary + rp_housing) * 0.1175 if rp_nationality == "سعودي" else (rp_salary + rp_housing) * 0.02
            rp_annual_per = (rp_monthly_total + rp_gosi) * 12 + rp_agency_fee + rp_visa + rp_training
            rp_annual_total = rp_annual_per * rp_count

            st.info(f"💰 التكلفة الشهرية للفرد: **{rp_monthly_total + rp_gosi:,.2f}** | السنوية للفرد: **{rp_annual_per:,.2f}** | الإجمالي ({rp_count}): **{rp_annual_total:,.2f} ريال**")

            if st.button("➕ إضافة للخطة", type="primary", key="rpbtn"):
                st.session_state.recruit_plans.append({
                    "المسمى": rp_title, "القسم": rp_dept, "العدد": rp_count,
                    "الجنسية": rp_nationality, "الراتب": rp_salary,
                    "السكن": rp_housing, "المواصلات": rp_transport,
                    "التأمينات (صاحب العمل)": round(rp_gosi, 2),
                    "الشهري/فرد": round(rp_monthly_total + rp_gosi, 2),
                    "رسوم التوظيف": rp_agency_fee, "التأشيرة": rp_visa,
                    "التدريب": rp_training, "السنوي/فرد": round(rp_annual_per, 2),
                    "الإجمالي السنوي": round(rp_annual_total, 2)
                })
                st.success(f"✅ تمت إضافة {rp_title} ({rp_count})")
                st.rerun()

            # Display plan
            if st.session_state.recruit_plans:
                st.markdown("---")
                st.markdown("### 📊 خطة التوظيف الحالية")
                plan_df = pd.DataFrame(st.session_state.recruit_plans)
                st.dataframe(plan_df, use_container_width=True, hide_index=True)

                # Summary
                total_headcount = plan_df["العدد"].sum()
                total_annual = plan_df["الإجمالي السنوي"].sum()
                total_monthly = plan_df.apply(lambda r: r["الشهري/فرد"] * r["العدد"], axis=1).sum()
                total_onetime = plan_df.apply(lambda r: (r["رسوم التوظيف"] + r["التأشيرة"] + r["التدريب"]) * r["العدد"], axis=1).sum()

                k1,k2,k3,k4 = st.columns(4)
                with k1: kpi("👥 إجمالي المطلوب", f"{total_headcount}")
                with k2: kpi("💰 التكلفة الشهرية", f"{total_monthly:,.0f}")
                with k3: kpi("📅 التكلفة السنوية", f"{total_annual:,.0f}")
                with k4: kpi("🔑 تكاليف لمرة واحدة", f"{total_onetime:,.0f}")

                # Charts
                ch1, ch2 = st.columns(2)
                with ch1:
                    fig = px.pie(plan_df, names="القسم", values="الإجمالي السنوي", title="توزيع الميزانية حسب القسم")
                    fig.update_layout(font=dict(family="Noto Sans Arabic"), height=350)
                    st.plotly_chart(fig, use_container_width=True)
                with ch2:
                    fig = px.bar(plan_df, x="المسمى", y="الإجمالي السنوي", color="الجنسية", title="التكلفة حسب الوظيفة", text_auto=True)
                    fig.update_layout(font=dict(family="Noto Sans Arabic"), height=350)
                    st.plotly_chart(fig, use_container_width=True)

                if st.button("🗑️ مسح الخطة بالكامل", key="rpclr"):
                    st.session_state.recruit_plans = []
                    st.rerun()



            export_widget(pd.DataFrame(st.session_state.get("recruit_plans",[])) if st.session_state.get("recruit_plans") else None, "خطة_التوظيف", "rec1")
        # ===== AI Salary Benchmark =====
        elif page == "🤖 Benchmark ذكاء اصطناعي":
            hdr("🤖 Benchmark الرواتب بالذكاء الاصطناعي","بيانات مرجعية من مصادر حية للسوق السعودي والمصري")

            bm_market = st.selectbox("🌍 اختر السوق:", list(MARKET_BENCHMARKS.keys()), key="bm_market")
            market = MARKET_BENCHMARKS[bm_market]
            cur = market['currency']

            bm_dept = st.selectbox("📌 القسم:", list(market['positions'].keys()), key="bm_dept")
            positions = market['positions'][bm_dept]

            st.markdown(f"### 💰 Salary Benchmark - {bm_dept} ({bm_market})")
            st.caption(f"العملة: {cur} | المصادر: Hays, Bayt, GulfTalent, Glassdoor, LinkedIn, Mercer")

            # Display benchmark table
            bm_rows = []
            for p in positions:
                bm_rows.append({
                    "المسمى (EN)": p['title'], "المسمى (AR)": p['title_ar'],
                    "المستوى": p['level'], f"الحد الأدنى ({cur})": f"{p['min']:,}",
                    f"المتوسط ({cur})": f"{p['mid']:,}", f"الحد الأعلى ({cur})": f"{p['max']:,}",
                    "الطلب": p['demand']
                })
            st.dataframe(pd.DataFrame(bm_rows), use_container_width=True, hide_index=True)

            # Visual benchmark
            bm_df = pd.DataFrame(positions)
            fig = go.Figure()
            fig.add_trace(go.Bar(name='الحد الأدنى', x=bm_df['title_ar'], y=bm_df['min'], marker_color='#3498DB'))
            fig.add_trace(go.Bar(name='المتوسط', x=bm_df['title_ar'], y=bm_df['mid'], marker_color='#E36414'))
            fig.add_trace(go.Bar(name='الحد الأعلى', x=bm_df['title_ar'], y=bm_df['max'], marker_color='#264653'))
            fig.update_layout(barmode='group', title=f'نطاقات الرواتب - {bm_dept} ({bm_market})',
                font=dict(family="Noto Sans Arabic"), height=450, yaxis_tickformat=',', yaxis_title=cur)
            st.plotly_chart(fig, use_container_width=True)

            # Full cost calculator
            st.markdown("---")
            st.markdown("### 🧮 حاسبة التكلفة الإجمالية للتوظيف")
            sel_pos = st.selectbox("اختر الوظيفة:", [p['title_ar'] for p in positions], key="bm_sel")
            sel_data = next(p for p in positions if p['title_ar'] == sel_pos)

            bc1, bc2 = st.columns(2)
            with bc1:
                sal_level = st.select_slider("مستوى الراتب:", ["الحد الأدنى","المتوسط","الحد الأعلى","مخصص"], value="المتوسط", key="bm_sl")
                if sal_level == "الحد الأدنى": base_sal = sel_data['min']
                elif sal_level == "المتوسط": base_sal = sel_data['mid']
                elif sal_level == "الحد الأعلى": base_sal = sel_data['max']
                else: base_sal = st.number_input("الراتب الأساسي:", value=sel_data['mid'], key="bm_custom")

                num_hires = st.number_input("عدد المطلوب:", 1, 50, 1, key="bm_num")
                is_saudi = st.radio("الجنسية:", ["سعودي","غير سعودي"], horizontal=True, key="bm_nat") == "سعودي"

            with bc2:
                housing = base_sal * market['housing_pct'] / 100
                transport = market['transport']

                if bm_market == "السعودية":
                    gosi = (base_sal + housing) * (market['gosi_employer'] if is_saudi else market['gosi_nsa_employer'])
                    med_ins = market['med_insurance_avg']
                    visa = 0 if is_saudi else market.get('visa_cost', 0)
                    iqama = 0 if is_saudi else market.get('iqama_cost', 0)
                else:
                    gosi = base_sal * market.get('social_insurance_employer', 0.184)
                    med_ins = market.get('med_insurance_avg', 800)
                    visa = 0; iqama = 0

                monthly_total = base_sal + housing + transport + gosi + med_ins
                annual_per = monthly_total * 12 + visa + iqama
                annual_all = annual_per * num_hires

                st.markdown("#### 📊 تفصيل التكلفة الشهرية")
                cost_items = [
                    ("الراتب الأساسي", base_sal), ("بدل السكن", housing),
                    ("بدل المواصلات", transport), ("التأمينات (صاحب العمل)", round(gosi,2)),
                    ("التأمين الطبي", med_ins)
                ]
                for label, val in cost_items:
                    st.markdown(f"**{label}:** {val:,.2f} {cur}")
                st.markdown(f"**الإجمالي الشهري:** **{monthly_total:,.2f} {cur}**")

            # Summary KPIs
            st.markdown("---")
            k1,k2,k3,k4 = st.columns(4)
            with k1: kpi(f"💰 الشهري/فرد ({cur})", f"{monthly_total:,.0f}")
            with k2: kpi(f"📅 السنوي/فرد ({cur})", f"{annual_per:,.0f}")
            with k3: kpi(f"👥 الإجمالي ({num_hires} موظف)", f"{annual_all:,.0f}")
            with k4:
                if bm_market == "مصر":
                    kpi("💵 المعادل SAR/شهر", f"{monthly_total/market.get('sar_to_egp',13.2):,.0f}")

            # Add to recruitment plan
            if st.button("➕ إضافة للخطة التوظيفية", type="primary", use_container_width=True, key="bm_add"):
                for _ in range(num_hires):
                    st.session_state.recruit_plans.append({
                        "المسمى": sel_pos, "القسم": bm_dept, "العدد": 1,
                        "الجنسية": "سعودي" if is_saudi else "غير سعودي",
                        "الراتب": base_sal, "السكن": housing, "المواصلات": transport,
                        "التأمينات (صاحب العمل)": round(gosi,2),
                        "الشهري/فرد": round(monthly_total,2),
                        "رسوم التوظيف": 0, "التأشيرة": visa, "التدريب": 0,
                        "السنوي/فرد": round(annual_per,2), "الإجمالي السنوي": round(annual_per,2),
                        "السوق": bm_market
                    })
                st.success(f"✅ تمت إضافة {num_hires} × {sel_pos} من سوق {bm_market}")
                st.rerun()

            # Sources
            st.markdown("---")
            st.markdown("### 📚 المصادر المرجعية")
            for s in BENCHMARK_SOURCES:
                region_match = (bm_market == "السعودية" and s['region'] in ['KSA','KSA/GCC','GCC','MENA','Global','MEA']) or \
                               (bm_market == "مصر" and s['region'] in ['Egypt','MENA','Global','MEA'])
                if region_match:
                    st.markdown(f"🔗 **{s['name']}** | {s['region']} | [{s['url']}](https://{s['url']})")



            export_widget(None, "Benchmark", "rec2")
        # ===== Market Comparison =====
        elif page == "🌍 مقارنة الأسواق":
            hdr("🌍 مقارنة تكاليف التوظيف: السعودية مقابل مصر","تحليل مقارن شامل للتكاليف والمزايا")

            st.markdown("### 📊 اختر الوظائف للمقارنة")
            # Find common positions between markets
            ksa_positions = MARKET_BENCHMARKS["السعودية"]["positions"]
            egy_positions = MARKET_BENCHMARKS["مصر"]["positions"]
            common_depts = set(ksa_positions.keys()) & set(egy_positions.keys())

            comp_dept = st.selectbox("القسم:", sorted(common_depts), key="comp_dept")
            ksa_pos = {p['title_ar']: p for p in ksa_positions[comp_dept]}
            egy_pos = {p['title_ar']: p for p in egy_positions[comp_dept]}
            common_titles = sorted(set(ksa_pos.keys()) & set(egy_pos.keys()))

            if not common_titles:
                st.warning("لا توجد وظائف مشتركة للمقارنة في هذا القسم")
            else:
                comp_titles = st.multiselect("الوظائف:", common_titles, default=common_titles[:3], key="comp_titles")

                if comp_titles:
                    sar_to_egp = MARKET_BENCHMARKS["مصر"].get("sar_to_egp", 13.2)

                    # Build comparison table
                    comp_rows = []
                    for title in comp_titles:
                        kp = ksa_pos[title]
                        ep = egy_pos[title]

                        # KSA costs
                        ksa_housing = kp['mid'] * 0.25
                        ksa_gosi = (kp['mid'] + ksa_housing) * 0.1175
                        ksa_monthly = kp['mid'] + ksa_housing + 500 + ksa_gosi + 500
                        ksa_annual = ksa_monthly * 12

                        # Egypt costs (converted to SAR)
                        egy_monthly_egp = ep['mid'] + ep['mid'] * 0.184 + 800
                        egy_monthly_sar = egy_monthly_egp / sar_to_egp
                        egy_annual_sar = egy_monthly_sar * 12

                        saving = ksa_annual - egy_annual_sar
                        saving_pct = saving / max(ksa_annual, 1) * 100

                        comp_rows.append({
                            "الوظيفة": title,
                            "🇸🇦 الراتب (SAR)": f"{kp['mid']:,}",
                            "🇸🇦 التكلفة الشهرية (SAR)": f"{ksa_monthly:,.0f}",
                            "🇸🇦 السنوية (SAR)": f"{ksa_annual:,.0f}",
                            "🇪🇬 الراتب (EGP)": f"{ep['mid']:,}",
                            "🇪🇬 التكلفة الشهرية (SAR≈)": f"{egy_monthly_sar:,.0f}",
                            "🇪🇬 السنوية (SAR≈)": f"{egy_annual_sar:,.0f}",
                            "💰 الوفر السنوي (SAR)": f"{saving:,.0f}",
                            "📊 نسبة الوفر": f"{saving_pct:.0f}%"
                        })

                    st.dataframe(pd.DataFrame(comp_rows), use_container_width=True, hide_index=True)

                    # Visual comparison
                    ksa_costs = []
                    egy_costs = []
                    for title in comp_titles:
                        kp = ksa_pos[title]
                        ep = egy_pos[title]
                        ksa_housing = kp['mid'] * 0.25
                        ksa_costs.append(kp['mid'] + ksa_housing + 500 + (kp['mid']+ksa_housing)*0.1175 + 500)
                        egy_costs.append((ep['mid'] + ep['mid']*0.184 + 800) / sar_to_egp)

                    fig = go.Figure()
                    fig.add_trace(go.Bar(name='🇸🇦 السعودية (SAR)', x=comp_titles, y=ksa_costs, marker_color='#27AE60'))
                    fig.add_trace(go.Bar(name='🇪🇬 مصر (SAR معادل)', x=comp_titles, y=egy_costs, marker_color='#3498DB'))
                    fig.update_layout(barmode='group', title='مقارنة التكلفة الشهرية الإجمالية (SAR)',
                        font=dict(family="Noto Sans Arabic"), height=450, yaxis_tickformat=',')
                    st.plotly_chart(fig, use_container_width=True)

                    # Summary analysis
                    st.markdown("---")
                    st.markdown("### 📈 ملخص التحليل")
                    total_ksa = sum(c * 12 for c in ksa_costs)
                    total_egy = sum(c * 12 for c in egy_costs)
                    total_saving = total_ksa - total_egy

                    k1,k2,k3,k4 = st.columns(4)
                    with k1: kpi("🇸🇦 إجمالي سنوي (KSA)", f"{total_ksa:,.0f} SAR")
                    with k2: kpi("🇪🇬 إجمالي سنوي (EGY≈)", f"{total_egy:,.0f} SAR")
                    with k3: kpi("💰 الوفر السنوي", f"{total_saving:,.0f} SAR")
                    with k4: kpi("📊 نسبة الوفر", f"{total_saving/max(total_ksa,1)*100:.0f}%")

                    # Pros and cons
                    c1, c2 = st.columns(2)
                    with c1:
                        st.markdown("#### 🇸🇦 مزايا التوظيف في السعودية")
                        ibox("نسبة السعودة وبرنامج نطاقات\nقرب من العمليات الأساسية\nلا حاجة لتأشيرات عمل للسعوديين\nدعم صندوق الموارد البشرية (هدف)", "success")
                    with c2:
                        st.markdown("#### 🇪🇬 مزايا التوظيف في مصر")
                        ibox("تكلفة أقل بنسبة 50-70%\nتوفر الكفاءات التقنية\nتوقيت عمل متقارب\nسهولة التواصل (نفس اللغة)", "success")

                    # Recommendation
                    st.markdown("---")
                    st.markdown("### 💡 التوصية الذكية")
                    if total_saving > 0:
                        ibox(f"""بناءً على التحليل، يمكن تحقيق وفر سنوي يقدر بـ **{total_saving:,.0f} SAR** ({total_saving/max(total_ksa,1)*100:.0f}%) عند توظيف الوظائف المختارة من السوق المصري.

**التوصية:** النموذج المختلط (Hybrid Model)
- الوظائف القيادية والمبيعات المحلية: **السعودية** (قرب من السوق + السعودة)
- الوظائف التقنية والداعمة: **مصر** (تكلفة أقل + كفاءات متوفرة)
- يُنصح بتخصيص 60% من الميزانية للسعودية و40% لمصر""")

                    # Export comparison
                    if st.button("📥 تصدير المقارنة Excel", key="comp_exp"):
                        ox = io.BytesIO()
                        with pd.ExcelWriter(ox, engine='xlsxwriter') as w:
                            pd.DataFrame(comp_rows).to_excel(w, sheet_name='Market Comparison', index=False)
                            # KSA benchmarks
                            ksa_all = []
                            for dept, poss in ksa_positions.items():
                                for p in poss:
                                    ksa_all.append({"Dept":dept, "Title":p['title'], "Title_AR":p['title_ar'],
                                        "Min":p['min'], "Mid":p['mid'], "Max":p['max'], "Level":p['level'], "Demand":p['demand']})
                            pd.DataFrame(ksa_all).to_excel(w, sheet_name='KSA Benchmarks', index=False)
                            # Egypt benchmarks
                            egy_all = []
                            for dept, poss in egy_positions.items():
                                for p in poss:
                                    egy_all.append({"Dept":dept, "Title":p['title'], "Title_AR":p['title_ar'],
                                        "Min":p['min'], "Mid":p['mid'], "Max":p['max'], "Level":p['level'], "Demand":p['demand']})
                            pd.DataFrame(egy_all).to_excel(w, sheet_name='Egypt Benchmarks', index=False)
                        st.download_button("📥 تحميل", data=ox.getvalue(),
                            file_name=f"Market_Comparison_{datetime.now().strftime('%Y%m%d')}.xlsx",
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

        elif page == "📊 متابعة التوظيف":
            hdr("📊 متابعة عمليات التوظيف", "تتبع مراحل التوظيف والتكاليف والوقت المستغرق")

            STAGES = ["طلب توظيف","إعلان","فرز السير الذاتية","مقابلة أولية","مقابلة نهائية","عرض وظيفي","قبول","مباشرة"]

            st.markdown("### ➕ إضافة عملية توظيف")
            tc1, tc2, tc3 = st.columns(3)
            with tc1:
                tr_title = st.text_input("المسمى الوظيفي:", key="trt")
                tr_dept = st.text_input("القسم:", key="trd")
            with tc2:
                tr_stage = st.selectbox("المرحلة الحالية:", STAGES, key="trs")
                tr_candidates = st.number_input("عدد المرشحين:", 0, 500, 0, key="trc")
            with tc3:
                tr_start = st.date_input("تاريخ البدء:", value=date.today(), key="trst")
                tr_budget = st.number_input("الميزانية المخصصة:", 0.0, 500000.0, 0.0, format="%.2f", key="trb")
                tr_spent = st.number_input("المصروف حتى الآن:", 0.0, 500000.0, 0.0, format="%.2f", key="trsp")

            if st.button("➕ إضافة", type="primary", key="trbtn"):
                days_elapsed = (date.today() - tr_start).days
                st.session_state.recruit_tracking.append({
                    "المسمى": tr_title, "القسم": tr_dept, "المرحلة": tr_stage,
                    "المرشحين": tr_candidates, "تاريخ البدء": str(tr_start),
                    "الأيام": days_elapsed, "الميزانية": tr_budget,
                    "المصروف": tr_spent, "المتبقي": round(tr_budget - tr_spent, 2),
                    "التقدم %": round((STAGES.index(tr_stage) + 1) / len(STAGES) * 100)
                })
                st.success(f"✅ تمت إضافة {tr_title}")
                st.rerun()

            if st.session_state.recruit_tracking:
                st.markdown("---")
                st.markdown("### 📋 العمليات الجارية")
                track_df = pd.DataFrame(st.session_state.recruit_tracking)
                st.dataframe(track_df, use_container_width=True, hide_index=True)

                # KPIs
                avg_days = track_df["الأيام"].mean()
                total_budget = track_df["الميزانية"].sum()
                total_spent = track_df["المصروف"].sum()
                total_candidates = track_df["المرشحين"].sum()
                open_positions = len(track_df[track_df["المرحلة"] != "مباشرة"])

                k1,k2,k3,k4,k5 = st.columns(5)
                with k1: kpi("📋 العمليات", f"{len(track_df)}")
                with k2: kpi("⏱️ متوسط الأيام", f"{avg_days:.0f}")
                with k3: kpi("👥 المرشحين", f"{total_candidates}")
                with k4: kpi("💰 المصروف/الميزانية", f"{total_spent:,.0f}/{total_budget:,.0f}")
                with k5: kpi("📂 مفتوحة", f"{open_positions}")

                # Pipeline chart
                stage_counts = track_df["المرحلة"].value_counts().reindex(STAGES, fill_value=0)
                fig = go.Figure(go.Funnel(y=stage_counts.index, x=stage_counts.values, textinfo="value+percent initial"))
                fig.update_layout(title="مسار التوظيف (Funnel)", font=dict(family="Noto Sans Arabic"), height=400)
                st.plotly_chart(fig, use_container_width=True)

                if st.button("🗑️ مسح المتابعة", key="trclr"):
                    st.session_state.recruit_tracking = []
                    st.rerun()



            export_widget(pd.DataFrame(st.session_state.get("recruit_tracking",[])) if st.session_state.get("recruit_tracking") else None, "متابعة_التوظيف", "rec3")

        # ===== CV ANALYSIS =====
        elif page == "📄 تحليل السير الذاتية":
            hdr("📄 تحليل السير الذاتية بالذكاء الاصطناعي","تحليل شامل للمرشح وفق أفضل الممارسات العالمية")

            # === CV Upload (REQUIRED) ===
            st.markdown("### 📄 السيرة الذاتية (مطلوب)")
            cv_file = st.file_uploader("ارفع السيرة الذاتية:", type=["pdf","docx","txt"], key="cv_file")
            cv_text = ""
            if cv_file:
                if cv_file.name.endswith('.txt'):
                    cv_text = cv_file.read().decode('utf-8',errors='ignore')
                elif cv_file.name.endswith('.docx'):
                    try:
                        from docx import Document
                        doc = Document(io.BytesIO(cv_file.read()))
                        cv_text = "\n".join([p.text for p in doc.paragraphs])
                    except: pass
                elif cv_file.name.endswith('.pdf'):
                    try:
                        import pdfplumber
                        with pdfplumber.open(io.BytesIO(cv_file.read())) as pdf:
                            cv_text = "\n".join([p.extract_text() or '' for p in pdf.pages])
                    except: pass
                if cv_text: st.success(f"✅ تم قراءة {len(cv_text)} حرف من السيرة الذاتية")
                else: st.error("❌ تعذر قراءة الملف. جرب صيغة أخرى.")

            # === JD Upload (OPTIONAL) ===
            with st.expander("📋 الوصف الوظيفي (اختياري - يزيد دقة الموائمة)", expanded=False):
                jd_method = st.radio("مصدر الوصف:", ["كتابة يدوية","رفع ملف"], horizontal=True, key="jd_src")
                jd_text = ""
                if jd_method == "كتابة يدوية":
                    jd_text = st.text_area("الوصف الوظيفي:", height=150, key="jd_txt",
                        placeholder="المسمى: محلل موارد بشرية\nالمتطلبات: خبرة 3 سنوات، شهادة PHRi...")
                else:
                    jd_file = st.file_uploader("ارفع الوصف الوظيفي:", type=["pdf","docx","txt"], key="jd_file")
                    if jd_file:
                        if jd_file.name.endswith('.txt'):
                            jd_text = jd_file.read().decode('utf-8',errors='ignore')
                        elif jd_file.name.endswith('.docx'):
                            try:
                                from docx import Document
                                doc = Document(io.BytesIO(jd_file.read()))
                                jd_text = "\n".join([p.text for p in doc.paragraphs])
                            except: pass
                        elif jd_file.name.endswith('.pdf'):
                            try:
                                import pdfplumber
                                with pdfplumber.open(io.BytesIO(jd_file.read())) as pdf:
                                    jd_text = "\n".join([p.extract_text() or '' for p in pdf.pages])
                            except: pass
                        if jd_text: st.success(f"✅ تم قراءة الوصف الوظيفي ({len(jd_text)} حرف)")

            # Company context (OPTIONAL)
            company_goals = st.text_area("🏢 أهداف واحتياجات الشركة (اختياري):", height=80, key="cv_goals",
                placeholder="شركة تقنية معلومات، تحتاج تعزيز فريق HR بالتحليلات والأتمتة...")

            # === Analysis Button ===
            if st.button("🤖 تحليل السيرة الذاتية", type="primary", use_container_width=True, key="cv_btn"):
                if not cv_text:
                    st.error("❌ يرجى رفع السيرة الذاتية أولاً")
                else:
                    with st.spinner("جاري التحليل الذكي بالذكاء الاصطناعي..."):
                        # Company info from config
                        company_info = ""
                        try:
                            ci = st.session_state.get('company_info', {})
                            if ci:
                                company_info = f"الشركة: {ci.get('name','رسال الود لتقنية المعلومات')}\nالنشاط: {ci.get('activity','تقنية المعلومات والمدفوعات الرقمية')}\nالمقر: {ci.get('location','جدة')}"
                        except: pass

                        has_jd = bool(jd_text and jd_text.strip())

                        if has_jd:
                            # === WITH JD: Matching analysis ===
                            prompt = f"""أنت خبير توظيف وتحليل سير ذاتية بخبرة 15 سنة. حلل السيرة الذاتية مقابل الوصف الوظيفي.

**السيرة الذاتية:**
{cv_text[:4000]}

**الوصف الوظيفي:**
{jd_text[:2000]}

**بيانات الشركة:**
{company_info if company_info else company_goals if company_goals else 'غير محدد'}

قدم التحليل بالعربية:

## 1. ملخص المرشح (3 أسطر)

## 2. نقاط الخبرة المتوافقة مع الوصف الوظيفي
- الخبرة: ... | التوافق: .../10

## 3. المتطلبات الناقصة عند المرشح
- المتطلب: ...

## 4. تقييم المهارات (كل مهارة من 10)
- المهارة: ... | الدرجة: .../10

## 5. نقاط القوة (أقوى 5)

## 6. نقاط الضعف والفجوات

## 7. التقييم الرقمي
- مطابقة الوصف الوظيفي: XX من 100
- جودة الخبرات: XX من 100
- المهارات التقنية: XX من 100
- المهارات القيادية: XX من 100
- الدرجة النهائية: XX من 100

## 8. التوصية (مناسب جداً / مناسب / يحتاج تطوير / غير مناسب)

## 9. أسئلة مقترحة للمقابلة (5 أسئلة)"""

                        else:
                            # === WITHOUT JD: General best-practice analysis ===
                            prompt = f"""أنت خبير توظيف دولي بخبرة 15 سنة. حلل السيرة الذاتية التالية وفق أفضل الممارسات العالمية في تقييم المرشحين.

**السيرة الذاتية:**
{cv_text[:4000]}

**معلومات إضافية:**
{company_info if company_info else ''} {company_goals if company_goals else ''}

لا يوجد وصف وظيفي محدد. حلل السيرة بشكل شامل ومستقل.

قدم التحليل بالعربية:

## 1. ملخص المرشح (3 أسطر عن الخبرة والتخصص والمستوى)

## 2. المسار المهني وتطوره
(تحليل المسار الوظيفي: هل هو تصاعدي؟ متنوع؟ متخصص؟)

## 3. تقييم المهارات (كل مهارة موجودة في السيرة من 10)
- المهارة: ... | الدرجة: .../10

## 4. نقاط القوة (أقوى 5 نقاط مع شرح)

## 5. نقاط الضعف والفجوات (3-5 نقاط)

## 6. الوظائف المناسبة (أفضل 5 مسميات وظيفية تناسب المرشح مع السبب)
1. المسمى: ... | السبب: ...

## 7. القطاعات المناسبة (أفضل 3 قطاعات)

## 8. التقييم الرقمي
- جودة الخبرات: XX من 100
- المهارات التقنية: XX من 100
- المهارات القيادية: XX من 100
- جودة التعليم والشهادات: XX من 100
- القيمة السوقية للمرشح: XX من 100
- الدرجة النهائية: XX من 100

## 9. توصيات للمرشح (5 نصائح لتطوير سيرته الذاتية ومساره المهني)

## 10. نطاق الراتب المتوقع في السوق السعودي (بالريال)"""

                        response, error = call_ai_api(prompt, prompt, model_type="hr")
                        if response:
                            st.session_state['_cv_result'] = response
                            st.session_state['_cv_name'] = cv_file.name if cv_file else "مرشح"
                            st.session_state['_cv_has_jd'] = has_jd
                        elif error:
                            st.error(f"❌ {error}")
                            st.info("💡 تأكد من إعداد مفتاح API في ⚙️ الإعدادات (Gemini مجاني أو Groq أو Claude)")

                    response, error = call_ai_api(prompt, prompt, model_type="hr")
                    if response:
                        st.session_state['_cv_result'] = response
                        st.session_state['_cv_name'] = cv_file.name if cv_file else "مرشح"

            # Display results
            if '_cv_result' in st.session_state:
                result = st.session_state['_cv_result']
                st.markdown("### 📊 نتائج التحليل")
                st.markdown(result)

                # Extract scores for charts
                import re
                scores = re.findall(r'(\d{1,3})\s*(?:من\s*100|/100|%)', result)
                scores = [int(s) for s in scores if 0 <= int(s) <= 100]

                if scores:
                    st.markdown("### 📊 الرسوم البيانية")
                    labels = ["المطابقة للوظيفة","التوافق مع الشركة","الدرجة النهائية"][:len(scores)]
                    sc1, sc2 = st.columns(2)
                    with sc1:
                        score_df = pd.DataFrame({"المعيار":labels,"الدرجة":scores[:len(labels)]})
                        fig = px.bar(score_df, x='المعيار', y='الدرجة', title='تقييم المرشح',
                            color='الدرجة', color_continuous_scale='RdYlGn', range_y=[0,100])
                        fig.update_layout(font=dict(family="Noto Sans Arabic"), height=350, coloraxis_showscale=False)
                        st.plotly_chart(fig, use_container_width=True)
                    with sc2:
                        fig = go.Figure(go.Indicator(mode="gauge+number", value=scores[-1] if scores else 0,
                            title={'text':'الدرجة النهائية'},
                            gauge={'axis':{'range':[0,100]},'bar':{'color':'#E36414'},
                                'steps':[{'range':[0,40],'color':'#EF4444'},{'range':[40,70],'color':'#F59E0B'},{'range':[70,100],'color':'#22C55E'}]}))
                        fig.update_layout(height=350)
                        st.plotly_chart(fig, use_container_width=True)

                # Skills radar
                skill_scores = re.findall(r'(\w[\w\s]+?)[\s:]+(\d{1,2})\s*(?:من\s*10|/10)', result)
                if skill_scores:
                    sk_df = pd.DataFrame(skill_scores, columns=['المهارة','الدرجة'])
                    sk_df['الدرجة'] = sk_df['الدرجة'].astype(int)
                    fig = go.Figure()
                    vals = sk_df['الدرجة'].tolist() + [sk_df['الدرجة'].iloc[0]]
                    cats = sk_df['المهارة'].tolist() + [sk_df['المهارة'].iloc[0]]
                    fig.add_trace(go.Scatterpolar(r=vals, theta=cats, fill='toself', line=dict(color='#E36414')))
                    fig.update_layout(polar=dict(radialaxis=dict(range=[0,10])), title='تقييم المهارات',
                        font=dict(family="Noto Sans Arabic"), height=400)
                    st.plotly_chart(fig, use_container_width=True)

                # Save to session for interview cross-reference
                st.session_state['_last_cv_analysis'] = result

                export_widget(pd.DataFrame([{"التحليل": result}]), "تحليل_السيرة_الذاتية", "cv1")

        # ===== INTERVIEW ANALYSIS =====
        elif page == "🎤 تحليل المقابلات":
            hdr("🎤 تحليل المقابلات والاختبارات بالذكاء الاصطناعي","تحليل ذكي للأسئلة والإجابات مع موائمة السيرة الذاتية")

            st.markdown("### ❓ الأسئلة")
            questions_text = st.text_area("أدخل الأسئلة (سؤال واحد في كل سطر):", height=150, key="intv_q",
                placeholder="1. لماذا تريد العمل في مجال الموارد البشرية؟\n2. كيف تتعامل مع موظف ذو أداء ضعيف؟\n3. ما خبرتك في أنظمة الرواتب؟")

            st.markdown("### 💬 إجابات المرشح")
            answers_text = st.text_area("أدخل الإجابات (إجابة واحدة في كل سطر بنفس ترتيب الأسئلة):", height=150, key="intv_a",
                placeholder="1. لدي شغف بتطوير بيئة العمل...\n2. أبدأ بمحادثة خاصة لفهم الأسباب...\n3. عملت 3 سنوات على نظام مدد...")

            # Company context
            intv_jd = st.text_area("📋 الوصف الوظيفي:", height=80, key="intv_jd", placeholder="المسمى والمتطلبات...")
            intv_goals = st.text_area("🏢 أهداف الشركة:", height=60, key="intv_goals", placeholder="شركة تقنية تسعى لأتمتة HR...")

            # Cross-reference with CV
            has_cv = '_last_cv_analysis' in st.session_state
            if has_cv:
                ibox("✅ تم العثور على تحليل سيرة ذاتية سابق. سيتم موائمته مع تحليل المقابلة.", "success")

            if st.button("🤖 تحليل المقابلة", type="primary", use_container_width=True, key="intv_btn") and questions_text and answers_text:
                with st.spinner("جاري التحليل الذكي..."):
                    cv_ref = f"\n\n**تحليل السيرة الذاتية السابق:**\n{st.session_state.get('_last_cv_analysis','')[:1500]}" if has_cv else ""

                    prompt = f"""حلل المقابلة الشخصية التالية وقدم تقريراً مفصلاً بالعربية:

**الأسئلة:**
{questions_text[:2000]}

**الإجابات:**
{answers_text[:2000]}

**الوصف الوظيفي:** {intv_jd[:800] if intv_jd else 'غير محدد'}
**أهداف الشركة:** {intv_goals[:500] if intv_goals else 'غير محدد'}
{cv_ref}

قدم التحليل:
1. **تقييم كل إجابة** (من 10) مع وصف
2. **نقاط القوة في المقابلة** (3-5 نقاط)
3. **نقاط الضعف** (2-3 نقاط)
4. **مدى تطابق الإجابات مع متطلبات الوظيفة** من 100
5. **مدى توافق المرشح مع أهداف الشركة** من 100
6. **التقييم العام للمقابلة** من 100
7. **التوصية النهائية** (قبول / قبول مشروط / رفض)
أعط كل درجة كرقم."""

                    response, error = call_ai_api(prompt, prompt, model_type="hr")
                    if response:
                        st.session_state['_intv_result'] = response

            if '_intv_result' in st.session_state:
                result = st.session_state['_intv_result']
                st.markdown("### 📊 نتائج تحليل المقابلة")
                st.markdown(result)

                import re
                scores = re.findall(r'(\d{1,3})\s*(?:من\s*100|/100)', result)
                scores = [int(s) for s in scores if 0 <= int(s) <= 100]
                q_scores = re.findall(r'(\d{1,2})\s*(?:من\s*10|/10)', result)
                q_scores = [int(s) for s in q_scores if 0 <= int(s) <= 10]

                ic1, ic2 = st.columns(2)
                with ic1:
                    if scores:
                        labels = ["مطابقة الوظيفة","توافق الشركة","التقييم العام"][:len(scores)]
                        fig = go.Figure(go.Indicator(mode="gauge+number", value=scores[-1],
                            title={'text':'التقييم العام للمقابلة'},
                            gauge={'axis':{'range':[0,100]},'bar':{'color':'#2A9D8F'},
                                'steps':[{'range':[0,40],'color':'#EF4444'},{'range':[40,70],'color':'#F59E0B'},{'range':[70,100],'color':'#22C55E'}]}))
                        fig.update_layout(height=350)
                        st.plotly_chart(fig, use_container_width=True)
                with ic2:
                    if q_scores:
                        q_df = pd.DataFrame({"السؤال":[f"س{i+1}" for i in range(len(q_scores))],"الدرجة":q_scores})
                        fig = px.bar(q_df, x='السؤال', y='الدرجة', title='تقييم الإجابات',
                            color='الدرجة', color_continuous_scale='RdYlGn', range_y=[0,10])
                        fig.update_layout(font=dict(family="Noto Sans Arabic"), height=350, coloraxis_showscale=False)
                        st.plotly_chart(fig, use_container_width=True)

                # Combined CV + Interview score
                if has_cv and scores:
                    st.markdown("### 🎯 التقييم المتكامل (سيرة ذاتية + مقابلة)")
                    cv_score = scores[0] if scores else 70
                    intv_score = scores[-1] if scores else 70
                    combined = round(cv_score * 0.4 + intv_score * 0.6)
                    ck1,ck2,ck3 = st.columns(3)
                    with ck1: kpi("📄 السيرة الذاتية (40%)", f"{cv_score}/100")
                    with ck2: kpi("🎤 المقابلة (60%)", f"{intv_score}/100")
                    with ck3: kpi("🎯 الدرجة النهائية", f"{combined}/100")

                    rec = "✅ مناسب جداً" if combined >= 80 else ("✅ مناسب" if combined >= 65 else ("⚠️ يحتاج تطوير" if combined >= 50 else "❌ غير مناسب"))
                    st.markdown(f"### التوصية: **{rec}**")

                export_widget(pd.DataFrame([{"التحليل": result}]), "تحليل_المقابلة", "intv1")

        # ===== ATS (Applicant Tracking System) =====
        elif page == "📋 ATS تتبع المتقدمين":
            hdr("📋 نظام تتبع المتقدمين ATS","Applicant Tracking System مع تقييم ذكي")

            if 'ats_candidates' not in st.session_state:
                st.session_state.ats_candidates = []

            # Job posting
            st.markdown("### 📋 الوظيفة المطلوبة")
            ac1, ac2 = st.columns(2)
            with ac1:
                ats_title = st.text_input("المسمى الوظيفي:", key="ats_title", placeholder="محلل موارد بشرية")
                ats_dept = st.text_input("القسم:", key="ats_dept", placeholder="الموارد البشرية")
            with ac2:
                ats_level = st.selectbox("المستوى:", ["مبتدئ","متوسط","أول","مدير","تنفيذي"], key="ats_level")
                ats_type = st.selectbox("نوع التوظيف:", ["دوام كامل","دوام جزئي","عقد مؤقت","عن بُعد"], key="ats_type")

            ats_req = st.text_area("المتطلبات والمهارات:", height=100, key="ats_req",
                placeholder="خبرة 3+ سنوات في HR\nشهادة PHRi أو SHRM\nإجادة Excel وPower BI\nمعرفة نظام العمل السعودي")
            ats_company = st.text_area("أهداف الشركة:", height=60, key="ats_company2",
                placeholder="شركة تقنية تسعى لبناء فريق HR يعتمد على البيانات")

            st.markdown("---")
            st.markdown("### 👥 إضافة مرشح")
            with st.form("ats_form", clear_on_submit=True):
                af1, af2, af3 = st.columns(3)
                with af1:
                    cand_name = st.text_input("اسم المرشح:", key="cand_name")
                    cand_email = st.text_input("البريد:", key="cand_email")
                with af2:
                    cand_phone = st.text_input("الجوال:", key="cand_phone")
                    cand_source = st.selectbox("مصدر التقديم:", ["LinkedIn","موقع الشركة","تزكية","معرض وظائف","أخرى"], key="cand_source")
                with af3:
                    cand_exp = st.number_input("سنوات الخبرة:", 0, 40, 3, key="cand_exp")
                    cand_salary = st.number_input("الراتب المتوقع:", 0, 100000, 10000, key="cand_sal")

                cand_cv = st.file_uploader("السيرة الذاتية:", type=["pdf","docx","txt"], key="cand_cv")
                cand_notes = st.text_input("ملاحظات:", key="cand_notes")
                add_cand = st.form_submit_button("➕ إضافة المرشح", type="primary", use_container_width=True)

            if add_cand and cand_name:
                cv_text = ""
                if cand_cv:
                    if cand_cv.name.endswith('.txt'): cv_text = cand_cv.read().decode('utf-8',errors='ignore')
                    elif cand_cv.name.endswith('.docx'):
                        try:
                            from docx import Document
                            doc = Document(io.BytesIO(cand_cv.read()))
                            cv_text = "\n".join([p.text for p in doc.paragraphs])
                        except: pass
                    elif cand_cv.name.endswith('.pdf'):
                        try:
                            import pdfplumber
                            with pdfplumber.open(io.BytesIO(cand_cv.read())) as pdf:
                                cv_text = "\n".join([p.extract_text() or '' for p in pdf.pages])
                        except: pass

                # Auto-score with AI
                ai_score = 50
                ai_note = ""
                if cv_text and ats_req:
                    try:
                        sc_prompt = f"قيّم هذا المرشح من 100 بناءً على السيرة الذاتية والمتطلبات. أعط رقماً واحداً فقط.\nالسيرة: {cv_text[:1500]}\nالمتطلبات: {ats_req[:500]}\nأهداف الشركة: {ats_company[:300]}"
                        sc_resp, _ = call_ai_api(sc_prompt, sc_prompt, model_type="hr")
                        if sc_resp:
                            import re
                            nums = re.findall(r'\b(\d{1,3})\b', sc_resp)
                            valid = [int(n) for n in nums if 10 <= int(n) <= 100]
                            if valid: ai_score = valid[0]
                            ai_note = sc_resp[:200]
                    except: pass

                candidate = {
                    "الاسم": cand_name, "البريد": cand_email, "الجوال": cand_phone,
                    "المصدر": cand_source, "الخبرة": cand_exp, "الراتب المتوقع": cand_salary,
                    "المرحلة": "تقديم", "الدرجة AI": ai_score,
                    "ملاحظات AI": ai_note, "ملاحظات": cand_notes,
                    "تاريخ التقديم": datetime.now().strftime("%Y-%m-%d"),
                    "الوظيفة": ats_title,
                }
                st.session_state.ats_candidates.append(candidate)
                st.success(f"✅ تم إضافة {cand_name} (درجة AI: {ai_score}/100)")
                st.rerun()

            # Display candidates
            if st.session_state.ats_candidates:
                st.markdown("### 📊 لوحة المتقدمين")
                cand_df = pd.DataFrame(st.session_state.ats_candidates)

                # KPIs
                ck1,ck2,ck3,ck4 = st.columns(4)
                with ck1: kpi("👥 إجمالي المتقدمين", str(len(cand_df)))
                with ck2: kpi("📊 متوسط الدرجة", f"{cand_df['الدرجة AI'].mean():.0f}/100")
                with ck3: kpi("⭐ مؤهلين (70+)", str(len(cand_df[cand_df['الدرجة AI']>=70])))
                with ck4: kpi("💰 متوسط الراتب", f"{cand_df['الراتب المتوقع'].mean():,.0f}")

                # Pipeline stages
                STAGES = ["تقديم","فرز أولي","مقابلة هاتفية","مقابلة شخصية","عرض وظيفي","قبول","رفض"]
                st.markdown("### 🔄 مراحل التوظيف (Pipeline)")
                for i, row in cand_df.iterrows():
                    with st.expander(f"{'⭐' if row['الدرجة AI']>=70 else '👤'} {row['الاسم']} | درجة: {row['الدرجة AI']}/100 | {row['المرحلة']}"):
                        ec1, ec2 = st.columns([2,1])
                        with ec1:
                            st.write(f"📧 {row['البريد']} | 📱 {row['الجوال']}")
                            st.write(f"💼 خبرة: {row['الخبرة']} سنة | 💰 الراتب: {row['الراتب المتوقع']:,}")
                            if row.get('ملاحظات AI'): st.caption(f"🤖 {row['ملاحظات AI'][:150]}")
                        with ec2:
                            new_stage = st.selectbox("المرحلة:", STAGES, index=STAGES.index(row['المرحلة']) if row['المرحلة'] in STAGES else 0, key=f"stg_{i}")
                            if new_stage != row['المرحلة']:
                                st.session_state.ats_candidates[i]['المرحلة'] = new_stage
                                st.rerun()

                # Charts
                ch1, ch2 = st.columns(2)
                with ch1:
                    fig = px.bar(cand_df.sort_values('الدرجة AI',ascending=True), x='الدرجة AI', y='الاسم',
                        orientation='h', title='تصنيف المرشحين حسب الدرجة', color='الدرجة AI',
                        color_continuous_scale='RdYlGn', range_x=[0,100])
                    fig.update_layout(font=dict(family="Noto Sans Arabic"), height=max(300, len(cand_df)*40), coloraxis_showscale=False)
                    st.plotly_chart(fig, use_container_width=True)
                with ch2:
                    stage_counts = cand_df['المرحلة'].value_counts()
                    fig = px.funnel(y=stage_counts.index, x=stage_counts.values, title='Pipeline التوظيف')
                    fig.update_layout(font=dict(family="Noto Sans Arabic"), height=350)
                    st.plotly_chart(fig, use_container_width=True)

                # Source analysis
                if 'المصدر' in cand_df.columns:
                    fig = px.pie(cand_df, names='المصدر', title='مصادر التقديم', hole=0.4)
                    fig.update_layout(font=dict(family="Noto Sans Arabic"), height=300)
                    st.plotly_chart(fig, use_container_width=True)

                export_widget(cand_df, "ATS_المتقدمين", "ats1")

                if st.button("🗑️ مسح كل المتقدمين", key="ats_clear"):
                    st.session_state.ats_candidates = []; st.rerun()

        elif page == "📥 تصدير التوظيف":
            hdr("📥 تصدير بيانات التوظيف")
            ox = io.BytesIO()
            with pd.ExcelWriter(ox, engine='xlsxwriter') as w:
                if st.session_state.recruit_plans:
                    pd.DataFrame(st.session_state.recruit_plans).to_excel(w, sheet_name='خطة التوظيف', index=False)
                    ws = w.sheets['خطة التوظيف']; ws.right_to_left()
                if st.session_state.recruit_tracking:
                    pd.DataFrame(st.session_state.recruit_tracking).to_excel(w, sheet_name='متابعة التوظيف', index=False)
                    ws = w.sheets['متابعة التوظيف']; ws.right_to_left()
                if not st.session_state.recruit_plans and not st.session_state.recruit_tracking:
                    pd.DataFrame({"ملاحظة": ["لا توجد بيانات"]}).to_excel(w, sheet_name='فارغ', index=False)
            st.download_button("📥 تحميل Excel", data=ox.getvalue(),
                file_name=f"Recruitment_{datetime.now().strftime('%Y%m%d')}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", type="primary", use_container_width=True)


    # =========================================
    #         🚀 ONBOARDING MODULE
    # =========================================
    elif section == "🚀 Onboarding":

        if 'onboarding_plans' not in st.session_state:
            # Load from database
            try:
                conn = get_conn()
                c = conn.cursor()
                c.execute(f"SELECT value FROM app_config WHERE key = {_ph()}", ("onboarding_plans",))
                row = c.fetchone()
                conn.close()
                if row:
                    st.session_state.onboarding_plans = json.loads(row[0])
                else:
                    st.session_state.onboarding_plans = []
            except:
                st.session_state.onboarding_plans = []

        # Load company info from DB or use defaults
        if 'company_info' not in st.session_state:
            try:
                conn = get_conn()
                c = conn.cursor()
                c.execute(f"SELECT value FROM app_config WHERE key = {_ph()}", ("company_info",))
                row = c.fetchone()
                conn.close()
                if row:
                    st.session_state.company_info = json.loads(row[0])
                else:
                    raise Exception("No data")
            except:
                st.session_state.company_info = {
                    "name": "رسال الود لتقنية المعلومات",
                    "name_en": "Resal",
                    "tagline": "Infrastructure for Non-Cash Value",
                    "who_we_are": "Resal, headquartered in Saudi Arabia, is a pioneering digital cards and rewards platform established in 2018. Renowned for its expertise in gift cards, rewards, and loyalty programs, Resal connects merchants, companies, and consumers through a platform that enables the issuance, management, and distribution of various prepaid products, including prepaid cards, loyalty, digital vouchers, and rewards.\n\nResal is at the forefront of revolutionizing the alternative payment and stored value landscape. Stored value refers to a prepaid monetary value stored on a device or in an account, which can be used for future transactions. This concept encompasses a variety of forms, including prepaid cards, gift cards, digital wallets, branded currency, and loyalty points.\n\nThe global shift towards digital payments represents a $2 trillion market projected to reach $2.5 trillion globally and $20 billion in MENA by 2027.",
                    "purpose": "Make people smile, make life easier, support our society and our economy",
                    "mission": "We are here to make people and businesses smile and delight by developing the prepaid cards, rewards and loyalty solutions to make it easy to send, and amazing to receive wherever and whenever.",
                    "vision": "The leading and the most innovative digital prepaid gift cards and rewards company in the MENA",
                    "values": [
                        {"name": "Integrity & Transparency", "name_ar": "النزاهة والشفافية", "desc": "We are open, direct, and honest with all. Successes and failures, good news and bad news, we share openly. It makes us all stronger."},
                        {"name": "One Team", "name_ar": "فريق واحد", "desc": "We work together, learn together, succeed together, to meet our customer needs, to help our company win, to support our community and our economy."},
                        {"name": "Customer Focus", "name_ar": "التركيز على العميل", "desc": "We are here to listen to our customers to help them and to make them happy. We treat our customers the way we want our own family to be treated."},
                        {"name": "Learning & Growth", "name_ar": "التعلم والنمو", "desc": "There is no end to knowledge. We need more, so we believe in knowledge sharing."},
                        {"name": "Innovation", "name_ar": "الابتكار", "desc": "We never give up seeking innovative ways to solve customer problems and delivering valuable solutions."},
                    ],
                    "products": [
                        {"name": "Glee", "desc": "SaaS solution that enables organizations to manage and send digital rewards to employees and customers."},
                        {"name": "Channels", "desc": "Business API and white-label technology offering digital gift cards from over 400 brands to 3rd party channel partners."},
                        {"name": "BOONUS (Loyalty)", "desc": "Loyalty management platform based on multiple rewarding programs that enable merchants to increase retention and spending."},
                        {"name": "xCard", "desc": "Prepaid digital gift cards issuance solution which enables merchants to easily issue, manage and redeem digital gift cards."},
                        {"name": "Resal App & Wallet", "desc": "Consumer mobile app and wallet to convert rewards from leading loyalty programs, enabling smarter spending across brands."},
                    ],
                    "milestones": [
                        {"year": "2016", "event": "Founded by Hatem Kameli and Fouad Al-Farhan in Jeddah"},
                        {"year": "2018", "event": "Launched digital gift cards + Resal for Business + Expansion to Egypt"},
                        {"year": "2020", "event": "10x Growth, +300 merchants, Amazon Partnership, Launched Glee"},
                        {"year": "2021", "event": "Pivoted to Fintech, Mastercard Partnership, Acquired BOONUS, Best Technology Provider Award"},
                        {"year": "2022", "event": "ISO 27001, Best Place to Work, +100M SAR total sales milestone"},
                        {"year": "2023", "event": "Unicorns Program - Monshaat, +1000 brands, +1.5M users"},
                    ],
                    "market": {
                        "global_prepaid": "$2.5 Trillion by 2027",
                        "closed_loop": "$680 Billion by 2027",
                        "loyalty_rewards": "$255 Billion by 2027",
                        "mena_tam": "$27+ Billion by 2027 (MENA)",
                        "saudi_prepaid": "$11 Billion by 2027",
                    },
                    "stats": {
                        "brands": "1,000+",
                        "clients": "1,000+",
                        "users": "1.5M+",
                        "growth": "100%+",
                    },
                    "work_policies": {
                        "hours": "8:00 AM - 5:00 PM (Sunday - Thursday)",
                        "leave": "21 days annual leave (Saudi Labor Law)",
                        "probation": "90 days",
                        "insurance": "Medical insurance from Day 1",
                        "locations": "Jeddah (HQ), Riyadh, Cairo",
                    },
                    "contacts": {
                        "hr_email": "HR@resal.me",
                        "it_email": "IT@resal.me",
                        "website": "resal.me",
                    }
                }

        ONBOARDING_CHECKLIST = {
            "قبل اليوم الأول": [
                {"task": "إعداد عرض العمل وإرساله", "owner": "التوظيف", "days": -14},
                {"task": "تجهيز عقد العمل", "owner": "الموارد البشرية", "days": -7},
                {"task": "فتح ملف التأمين الطبي", "owner": "الموارد البشرية", "days": -5},
                {"task": "تجهيز مكتب ومعدات العمل", "owner": "تقنية المعلومات", "days": -3},
                {"task": "إنشاء حسابات البريد والأنظمة", "owner": "تقنية المعلومات", "days": -3},
                {"task": "إرسال رسالة ترحيب + دليل الموظف الجديد", "owner": "الموارد البشرية", "days": -2},
                {"task": "إبلاغ الفريق بانضمام الموظف الجديد", "owner": "المدير المباشر", "days": -1},
            ],
            "اليوم الأول": [
                {"task": "استقبال الموظف وجولة في المكتب", "owner": "الموارد البشرية", "days": 1},
                {"task": "تسليم أجهزة العمل وبطاقة الدخول", "owner": "تقنية المعلومات", "days": 1},
                {"task": "التعريف بالفريق والأقسام", "owner": "المدير المباشر", "days": 1},
                {"task": "شرح سياسات الشركة والأنظمة", "owner": "الموارد البشرية", "days": 1},
                {"task": "إعداد حساب نظام الموارد البشرية (Jisr)", "owner": "الموارد البشرية", "days": 1},
                {"task": "غداء ترحيبي مع الفريق", "owner": "المدير المباشر", "days": 1},
            ],
            "الأسبوع الأول (2-5 أيام)": [
                {"task": "تدريب على أنظمة العمل الداخلية", "owner": "تقنية المعلومات", "days": 3},
                {"task": "اجتماع مع المدير المباشر - الأهداف والتوقعات", "owner": "المدير المباشر", "days": 3},
                {"task": "تدريب على المنتجات والخدمات", "owner": "فريق المنتجات", "days": 4},
                {"task": "مراجعة خطة الأداء 30/60/90", "owner": "المدير المباشر", "days": 5},
                {"task": "جلسة تعريف بثقافة الشركة وقيمها", "owner": "الموارد البشرية", "days": 5},
            ],
            "الشهر الأول (6-30 يوم)": [
                {"task": "إتمام جميع التدريبات الإلزامية", "owner": "الموظف", "days": 15},
                {"task": "اجتماع متابعة أسبوعي مع المدير", "owner": "المدير المباشر", "days": 7},
                {"task": "مشروع أولي / مهمة تجريبية", "owner": "المدير المباشر", "days": 14},
                {"task": "تقييم نهاية الشهر الأول", "owner": "المدير المباشر", "days": 30},
                {"task": "جمع ملاحظات الموظف عن تجربة Onboarding", "owner": "الموارد البشرية", "days": 30},
            ],
            "الشهر الثاني (31-60 يوم)": [
                {"task": "زيادة المسؤوليات تدريجياً", "owner": "المدير المباشر", "days": 35},
                {"task": "حضور اجتماعات القسم والشركة", "owner": "الموظف", "days": 35},
                {"task": "تقييم منتصف فترة التجربة", "owner": "المدير المباشر", "days": 45},
                {"task": "تطوير العلاقات مع الأقسام الأخرى", "owner": "الموظف", "days": 50},
                {"task": "تدريب متخصص حسب الوظيفة", "owner": "المدير المباشر", "days": 55},
            ],
            "الشهر الثالث (61-90 يوم)": [
                {"task": "العمل باستقلالية كاملة", "owner": "الموظف", "days": 65},
                {"task": "تقييم شامل نهاية فترة التجربة", "owner": "المدير المباشر + HR", "days": 85},
                {"task": "قرار تثبيت أو تمديد التجربة", "owner": "المدير المباشر", "days": 88},
                {"task": "وضع أهداف الأداء للسنة", "owner": "المدير المباشر", "days": 90},
                {"task": "استبيان رضا الموظف الجديد", "owner": "الموارد البشرية", "days": 90},
            ],
        }

        if page == "🚀 إنشاء Onboarding":
            hdr("🚀 إنشاء خطة Onboarding","نظام تهيئة الموظفين الجدد")

            st.markdown("### 👤 بيانات الموظف الجديد")
            oc1, oc2, oc3 = st.columns(3)
            with oc1:
                ob_name = st.text_input("الاسم الكامل:", key="ob_name")
                ob_email = st.text_input("البريد الإلكتروني:", key="ob_email")
            with oc2:
                ob_title = st.text_input("المسمى الوظيفي:", key="ob_title")
                ob_dept = st.text_input("القسم:", key="ob_dept")
            with oc3:
                ob_manager = st.text_input("المدير المباشر:", key="ob_mgr")
                ob_start = st.date_input("تاريخ المباشرة:", key="ob_start")

            ob_type = st.radio("نوع التوظيف:", ["دوام كامل","دوام جزئي","عن بُعد","تعاقد"], horizontal=True, key="ob_type")

            st.markdown("### 📋 تخصيص Checklist")
            selected_phases = st.multiselect("المراحل المطلوبة:", list(ONBOARDING_CHECKLIST.keys()),
                default=list(ONBOARDING_CHECKLIST.keys()), key="ob_phases")

            # Show editable checklist
            all_tasks = []
            for phase in selected_phases:
                with st.expander(f"📌 {phase} ({len(ONBOARDING_CHECKLIST[phase])} مهمة)", expanded=False):
                    for j, task in enumerate(ONBOARDING_CHECKLIST[phase]):
                        tc1, tc2 = st.columns([3,1])
                        with tc1: st.markdown(f"✅ {task['task']} | **{task['owner']}**")
                        with tc2: st.caption(f"اليوم {task['days']}")
                        all_tasks.append({**task, "phase": phase})

            if st.button("🚀 إنشاء خطة Onboarding", type="primary", use_container_width=True, key="ob_create"):
                if ob_name and ob_title:
                    plan = {
                        "id": len(st.session_state.onboarding_plans) + 1,
                        "name": ob_name, "email": ob_email, "title": ob_title,
                        "dept": ob_dept, "manager": ob_manager,
                        "start_date": str(ob_start), "type": ob_type,
                        "tasks": all_tasks, "status": "جاري",
                        "created_at": datetime.now().strftime("%Y-%m-%d %H:%M"),
                        "progress": 0
                    }
                    st.session_state.onboarding_plans.append(plan)
                    # Save to database for persistence
                    try:
                        conn = get_conn()
                        c = conn.cursor()
                        _upsert_config(c, "onboarding_plans", json.dumps(st.session_state.onboarding_plans, ensure_ascii=False))
                        conn.commit()
                        conn.close()
                    except: pass

                    # Send welcome email
                    smtp_cfg = st.session_state.get('smtp_config', {})
                    if ob_email and smtp_cfg.get('email'):
                        ok, msg = send_test_email(ob_email, ob_name,
                            [f"بدء خطة Onboarding - {ob_title}",
                             f"تاريخ المباشرة: {ob_start}",
                             f"المدير المباشر: {ob_manager}"],
                            str(ob_start), smtp_cfg.get('sender_name', 'HR'))
                        if ok: st.success(f"📧 تم إرسال رسالة ترحيب إلى {ob_email}")

                    st.success(f"✅ تم إنشاء خطة Onboarding لـ {ob_name}")
                    st.rerun()
                else:
                    st.error("يرجى إدخال الاسم والمسمى الوظيفي")

        elif page == "📋 خطة 30/60/90":
            hdr("📋 خطة 30/60/90 يوم","خطة تفصيلية لفترة التجربة")

            if not st.session_state.onboarding_plans:
                st.info("🚀 أنشئ خطة Onboarding أولاً من القائمة")
                return

            sel_emp = st.selectbox("اختر الموظف:", [f"{p['name']} - {p['title']}" for p in st.session_state.onboarding_plans], key="ob_sel30")
            sel_idx = [f"{p['name']} - {p['title']}" for p in st.session_state.onboarding_plans].index(sel_emp)
            plan = st.session_state.onboarding_plans[sel_idx]

            st.markdown(f"**{plan['name']}** | {plan['title']} | {plan['dept']} | بدأ: {plan['start_date']}")

            # 30/60/90 Tabs
            t1, t2, t3 = st.tabs(["📅 أول 30 يوم", "📅 31-60 يوم", "📅 61-90 يوم"])

            with t1:
                st.markdown("### 🎯 أهداف الشهر الأول")
                goals_30 = ["فهم ثقافة الشركة وقيمها","التعرف على الفريق والأقسام","إتمام التدريبات الإلزامية","فهم المنتجات والخدمات","إتمام مشروع أولي"]
                for g in goals_30:
                    st.checkbox(g, key=f"g30_{g}_{sel_idx}")

                phases_30 = [p for p in plan.get('tasks',[]) if p.get('days',0) <= 30]
                if phases_30:
                    st.markdown("### ✅ المهام")
                    for t in phases_30:
                        done = st.checkbox(f"{t['task']} ({t['owner']})", key=f"t30_{t['task']}_{sel_idx}")

            with t2:
                st.markdown("### 🎯 أهداف الشهر الثاني")
                goals_60 = ["زيادة المسؤوليات","المساهمة في مشاريع القسم","بناء علاقات مع الأقسام","تقديم أفكار تحسينية"]
                for g in goals_60:
                    st.checkbox(g, key=f"g60_{g}_{sel_idx}")

                phases_60 = [p for p in plan.get('tasks',[]) if 30 < p.get('days',0) <= 60]
                if phases_60:
                    st.markdown("### ✅ المهام")
                    for t in phases_60:
                        st.checkbox(f"{t['task']} ({t['owner']})", key=f"t60_{t['task']}_{sel_idx}")

            with t3:
                st.markdown("### 🎯 أهداف الشهر الثالث")
                goals_90 = ["العمل باستقلالية","تحقيق أهداف الأداء","المساهمة الفعالة في الفريق","جاهزية للتثبيت"]
                for g in goals_90:
                    st.checkbox(g, key=f"g90_{g}_{sel_idx}")

                phases_90 = [p for p in plan.get('tasks',[]) if 60 < p.get('days',0) <= 90]
                if phases_90:
                    st.markdown("### ✅ المهام")
                    for t in phases_90:
                        st.checkbox(f"{t['task']} ({t['owner']})", key=f"t90_{t['task']}_{sel_idx}")

        elif page == "👥 متابعة الموظفين الجدد":
            hdr("👥 متابعة الموظفين الجدد","لوحة تتبع حالة Onboarding لجميع الموظفين")

            plans = st.session_state.onboarding_plans
            if not plans:
                st.info("🚀 لا يوجد موظفين جدد. أنشئ خطة Onboarding أولاً.")
                return

            # KPIs
            k1,k2,k3,k4 = st.columns(4)
            with k1: kpi("👥 الموظفين الجدد", str(len(plans)))
            active = [p for p in plans if p.get('status') == 'جاري']
            with k2: kpi("🔄 جاري", str(len(active)))
            completed = [p for p in plans if p.get('status') == 'مكتمل']
            with k3: kpi("✅ مكتمل", str(len(completed)))
            depts = set(p.get('dept','') for p in plans)
            with k4: kpi("📌 الأقسام", str(len(depts)))

            # Table
            plan_rows = []
            for p in plans:
                start = datetime.strptime(p['start_date'], '%Y-%m-%d') if isinstance(p['start_date'], str) else p['start_date']
                days_elapsed = (datetime.now() - start).days if isinstance(start, datetime) else 0
                phase = "قبل المباشرة" if days_elapsed < 0 else ("30 يوم" if days_elapsed <= 30 else ("60 يوم" if days_elapsed <= 60 else ("90 يوم" if days_elapsed <= 90 else "بعد التجربة")))
                progress = min(100, max(0, int(days_elapsed / 90 * 100)))

                plan_rows.append({
                    "الاسم": p['name'], "المسمى": p['title'], "القسم": p['dept'],
                    "المدير": p.get('manager',''), "تاريخ البدء": p['start_date'],
                    "الأيام": days_elapsed, "المرحلة": phase,
                    "التقدم %": progress, "الحالة": p.get('status','جاري')
                })

            st.dataframe(pd.DataFrame(plan_rows), use_container_width=True, hide_index=True)

            # Timeline chart
            if plan_rows:
                fig = px.bar(pd.DataFrame(plan_rows), x='الاسم', y='التقدم %', color='المرحلة',
                    title='تقدم الموظفين الجدد في Onboarding',
                    color_discrete_map={'قبل المباشرة':'#95A5A6','30 يوم':'#E36414','60 يوم':'#E9C46A','90 يوم':'#27AE60','بعد التجربة':'#2980B9'})
                fig.update_layout(font=dict(family="Noto Sans Arabic"), height=400)
                st.plotly_chart(fig, use_container_width=True)

        elif page == "📊 تحليلات Onboarding":
            hdr("📊 تحليلات Onboarding","إحصائيات وتحليل فعالية برنامج تهيئة الموظفين الجدد")

            plans = st.session_state.onboarding_plans

            # Initialize survey data
            if 'ob_surveys' not in st.session_state:
                try:
                    conn = get_conn()
                    c = conn.cursor()
                    c.execute(f"SELECT value FROM app_config WHERE key = {_ph()}", ("ob_surveys",))
                    row = c.fetchone()
                    conn.close()
                    st.session_state.ob_surveys = json.loads(row[0]) if row else []
                except:
                    st.session_state.ob_surveys = []

            # ===== KPIs =====
            st.markdown("### 📊 مؤشرات الأداء الرئيسية")
            total_plans = len(plans)
            surveys = st.session_state.ob_surveys

            if total_plans > 0:
                # Calculate metrics
                active_plans = [p for p in plans if p.get('status','جاري') == 'جاري']
                completed_plans = [p for p in plans if p.get('status') == 'مكتمل']
                depts = set(p.get('dept','') for p in plans)
                avg_satisfaction = sum(s.get('overall',0) for s in surveys) / max(len(surveys),1) if surveys else 0
                completion_rate = len(completed_plans) / max(total_plans,1) * 100

                k1,k2,k3,k4,k5,k6 = st.columns(6)
                with k1: kpi("📋 إجمالي الخطط", str(total_plans))
                with k2: kpi("🔄 جاري", str(len(active_plans)))
                with k3: kpi("✅ مكتمل", str(len(completed_plans)))
                with k4: kpi("📊 معدل الإكمال", f"{completion_rate:.0f}%")
                with k5: kpi("⭐ رضا الموظفين", f"{avg_satisfaction:.1f}/5" if surveys else "N/A")
                with k6: kpi("📌 الأقسام", str(len(depts)))

                # ===== Charts =====
                st.markdown("---")
                ch1, ch2 = st.columns(2)

                with ch1:
                    # Plans by department
                    dept_counts = {}
                    for p in plans:
                        d = p.get('dept','غير محدد')
                        dept_counts[d] = dept_counts.get(d,0) + 1
                    fig = px.pie(values=list(dept_counts.values()), names=list(dept_counts.keys()),
                        title='توزيع Onboarding حسب القسم', hole=0.35, color_discrete_sequence=CL['dept'])
                    fig.update_layout(font=dict(family="Noto Sans Arabic"), height=350)
                    st.plotly_chart(fig, use_container_width=True)

                with ch2:
                    # Plans by status/phase
                    phase_counts = {"قبل المباشرة":0, "30 يوم":0, "60 يوم":0, "90 يوم":0, "بعد التجربة":0}
                    for p in plans:
                        try:
                            start = datetime.strptime(p['start_date'], '%Y-%m-%d')
                            days = (datetime.now() - start).days
                            phase = "قبل المباشرة" if days < 0 else ("30 يوم" if days <= 30 else ("60 يوم" if days <= 60 else ("90 يوم" if days <= 90 else "بعد التجربة")))
                            phase_counts[phase] += 1
                        except: pass
                    fig = px.bar(x=list(phase_counts.keys()), y=list(phase_counts.values()),
                        title='توزيع حسب المرحلة', color=list(phase_counts.keys()),
                        color_discrete_map={"قبل المباشرة":"#95A5A6","30 يوم":"#E36414","60 يوم":"#E9C46A","90 يوم":"#27AE60","بعد التجربة":"#2980B9"})
                    fig.update_layout(font=dict(family="Noto Sans Arabic"), height=350, showlegend=False)
                    st.plotly_chart(fig, use_container_width=True)

                # Monthly trend
                if total_plans > 1:
                    monthly = {}
                    for p in plans:
                        m = p.get('start_date','')[:7]
                        monthly[m] = monthly.get(m,0) + 1
                    if monthly:
                        fig = px.line(x=list(monthly.keys()), y=list(monthly.values()),
                            title='عدد Onboarding الشهري', markers=True)
                        fig.update_layout(font=dict(family="Noto Sans Arabic"), height=300,
                            xaxis_title="الشهر", yaxis_title="العدد")
                        st.plotly_chart(fig, use_container_width=True)

            else:
                st.info("لا توجد خطط Onboarding بعد. أنشئ خطة أولاً.")

            # ===== Survey Section =====
            st.markdown("---")
            st.markdown("### 📝 استبيان رضا الموظف الجديد عن Onboarding")

            if plans:
                with st.expander("➕ إضافة تقييم جديد", expanded=not bool(surveys)):
                    sv_emp = st.selectbox("الموظف:", [f"{p['name']} - {p['title']}" for p in plans], key="sv_emp")

                    st.markdown("#### قيّم تجربة Onboarding (1 = ضعيف، 5 = ممتاز)")
                    survey_questions = {
                        "overall": "التقييم العام لتجربة Onboarding",
                        "welcome": "جودة الاستقبال والترحيب",
                        "training": "كفاية التدريب والتأهيل",
                        "tools": "توفر الأدوات والأنظمة في الوقت المناسب",
                        "manager": "دعم المدير المباشر",
                        "team": "تعاون الفريق والاندماج",
                        "clarity": "وضوح المهام والتوقعات",
                        "culture": "فهم ثقافة وقيم الشركة",
                    }
                    survey_labels = {
                        "overall": "التقييم العام", "welcome": "الاستقبال والترحيب",
                        "training": "التدريب والتأهيل", "tools": "الأدوات والأنظمة",
                        "manager": "دعم المدير", "team": "تعاون الفريق",
                        "clarity": "وضوح المهام", "culture": "ثقافة الشركة",
                    }

                    scores = {}
                    for key, question in survey_questions.items():
                        scores[key] = st.slider(f"{'⭐' * 5} {question}", 1, 5, 4, key=f"sv_{key}")

                    sv_comment = st.text_area("ملاحظات إضافية:", placeholder="ما أعجبك؟ ما يمكن تحسينه؟", key="sv_comment")
                    sv_recommend = st.radio("هل توصي بتجربة Onboarding للموظفين الجدد؟", ["نعم بالتأكيد","نعم","محايد","لا"], horizontal=True, key="sv_rec")

                    if st.button("📊 حفظ التقييم", type="primary", use_container_width=True, key="sv_save"):
                        survey = {
                            "employee": sv_emp,
                            "date": datetime.now().strftime("%Y-%m-%d"),
                            "scores": scores,
                            **scores,
                            "comment": sv_comment,
                            "recommend": sv_recommend,
                            "nps": 5 if "بالتأكيد" in sv_recommend else (4 if sv_recommend=="نعم" else (3 if sv_recommend=="محايد" else 1))
                        }
                        st.session_state.ob_surveys.append(survey)
                        # Save to DB
                        try:
                            conn = get_conn()
                            c = conn.cursor()
                            _upsert_config(c, "ob_surveys", json.dumps(st.session_state.ob_surveys, ensure_ascii=False))
                            conn.commit()
                            conn.close()
                        except: pass
                        st.success("✅ تم حفظ التقييم بنجاح")
                        st.rerun()

            # ===== Survey Results =====
            if surveys:
                st.markdown("---")
                st.markdown("### 📈 نتائج استبيانات Onboarding")

                # Summary table
                sv_rows = []
                for s in surveys:
                    sv_rows.append({
                        "الموظف": s.get('employee',''), "التاريخ": s.get('date',''),
                        "التقييم العام": f"{s.get('overall',0)}/5",
                        "الاستقبال": s.get('welcome',0), "التدريب": s.get('training',0),
                        "الأدوات": s.get('tools',0), "المدير": s.get('manager',0),
                        "الفريق": s.get('team',0), "الوضوح": s.get('clarity',0),
                        "الثقافة": s.get('culture',0), "التوصية": s.get('recommend','')
                    })
                st.dataframe(pd.DataFrame(sv_rows), use_container_width=True, hide_index=True)

                # Average scores radar chart
                avg_scores = {}
                survey_labels_ar = {"overall":"التقييم العام","welcome":"الاستقبال","training":"التدريب",
                    "tools":"الأدوات","manager":"المدير","team":"الفريق","clarity":"الوضوح","culture":"الثقافة"}
                for key in ["overall","welcome","training","tools","manager","team","clarity","culture"]:
                    vals = [s.get(key,0) for s in surveys if s.get(key,0) > 0]
                    avg_scores[survey_labels_ar.get(key,key)] = sum(vals)/max(len(vals),1) if vals else 0

                rc1, rc2 = st.columns(2)
                with rc1:
                    fig = go.Figure()
                    categories = list(avg_scores.keys())
                    values = list(avg_scores.values()) + [list(avg_scores.values())[0]]
                    fig.add_trace(go.Scatterpolar(r=values, theta=categories + [categories[0]],
                        fill='toself', line=dict(color='#E36414'), fillcolor='rgba(227,100,20,0.2)', name='المتوسط'))
                    fig.update_layout(polar=dict(radialaxis=dict(range=[0,5])),
                        title='متوسط التقييمات (Radar)', font=dict(family="Noto Sans Arabic"), height=400)
                    st.plotly_chart(fig, use_container_width=True)

                with rc2:
                    # Bar chart of averages
                    fig = px.bar(x=list(avg_scores.keys()), y=list(avg_scores.values()),
                        title='متوسط التقييمات حسب المعيار', color=list(avg_scores.values()),
                        color_continuous_scale='RdYlGn', range_color=[1,5])
                    fig.add_hline(y=4, line_dash="dash", line_color="green", annotation_text="الهدف: 4/5")
                    fig.update_layout(font=dict(family="Noto Sans Arabic"), height=400, yaxis_range=[0,5])
                    st.plotly_chart(fig, use_container_width=True)

                # NPS Score
                nps_scores = [s.get('nps',3) for s in surveys]
                promoters = sum(1 for n in nps_scores if n >= 4)
                detractors = sum(1 for n in nps_scores if n <= 2)
                nps = round((promoters - detractors) / max(len(nps_scores),1) * 100)

                st.markdown("### 🎯 مؤشر صافي الترويج (eNPS)")
                nc1, nc2, nc3, nc4 = st.columns(4)
                with nc1: kpi("🟢 مروّجين", str(promoters))
                with nc2: kpi("🟡 محايدين", str(len(nps_scores) - promoters - detractors))
                with nc3: kpi("🔴 منتقدين", str(detractors))
                with nc4: kpi("📊 eNPS Score", f"{nps}")

                # Comments
                comments = [s.get('comment','') for s in surveys if s.get('comment','').strip()]
                if comments:
                    st.markdown("### 💬 ملاحظات الموظفين")
                    for c in comments:
                        st.markdown(f"> {c}")

                # Effectiveness summary
                st.markdown("---")
                st.markdown("### 📊 ملخص فعالية Onboarding")
                overall_avg = avg_scores.get('التقييم العام', 0)
                if overall_avg >= 4:
                    ibox(f"تقييم البرنامج: **{overall_avg:.1f}/5** - ممتاز ✅\n\nبرنامج Onboarding يحقق نتائج ممتازة. يُنصح بالاستمرار مع تحسينات طفيفة في المجالات الأقل تقييماً.", "success")
                elif overall_avg >= 3:
                    lowest = min(avg_scores, key=avg_scores.get)
                    ibox(f"تقييم البرنامج: **{overall_avg:.1f}/5** - جيد ⚠️\n\nيحتاج تحسين في: **{lowest}** ({avg_scores[lowest]:.1f}/5). يُنصح بمراجعة هذا الجانب وتطويره.", "warning")
                else:
                    ibox(f"تقييم البرنامج: **{overall_avg:.1f}/5** - يحتاج تطوير 🔴\n\nيُنصح بمراجعة شاملة لبرنامج Onboarding وإعادة تصميم المراحل ذات التقييم المنخفض.", "error" if hasattr(st, 'error') else "warning")



            export_widget(None, "تحليلات_Onboarding", "ob1")
        elif page == "🎬 عرض تقديمي AI":
            hdr("🎬 عرض تقديمي Onboarding بالذكاء الاصطناعي","عرض slides احترافي مع سرد صوتي تفاعلي")

            if not st.session_state.onboarding_plans:
                st.info("🚀 أنشئ خطة Onboarding أولاً")
                return

            sel_emp = st.selectbox("اختر الموظف:", [f"{p['name']} - {p['title']}" for p in st.session_state.onboarding_plans], key="ob_sel_ai")
            sel_idx = [f"{p['name']} - {p['title']}" for p in st.session_state.onboarding_plans].index(sel_emp)
            plan = st.session_state.onboarding_plans[sel_idx]
            ci = st.session_state.company_info

            st.markdown("### ⚙️ تخصيص العرض")
            slides_sel = st.multiselect("الشرائح المطلوبة:",
                ["ترحيب","من نحن","رؤيتنا ورسالتنا","قيمنا","منتجاتنا","إنجازاتنا","السوق","فريقك","خطة 30/60/90","سياسات","تواصل"],
                default=["ترحيب","من نحن","رؤيتنا ورسالتنا","قيمنا","منتجاتنا","فريقك","خطة 30/60/90"], key="ob_slides")
            voice_lang = st.selectbox("لغة السرد:", ["ar-SA","en-US","ar-EG"], key="ob_voice")

            if st.button("🎬 إنشاء العرض التقديمي", type="primary", use_container_width=True, key="ob_gen"):

                # Build narration script for TTS
                narration_parts = []
                if "ترحيب" in slides_sel:
                    narration_parts.append(f"مرحباً {plan['name']}! أهلاً بك في {ci['name']}. نحن سعيدون جداً بانضمامك لفريقنا بصفتك {plan['title']} في قسم {plan['dept']}.")
                if "من نحن" in slides_sel:
                    narration_parts.append(f"{ci['name_en']} هي منصة رائدة في البطاقات الرقمية والمكافآت تأسست في 2018 في المملكة العربية السعودية. نربط التجار والشركات والمستهلكين من خلال منصة شاملة للبطاقات مسبقة الدفع وبرامج الولاء.")
                if "رؤيتنا ورسالتنا" in slides_sel:
                    narration_parts.append(f"هدفنا هو {ci['purpose']}. رؤيتنا أن نكون {ci['vision']}.")
                if "قيمنا" in slides_sel:
                    vals_text = "، ".join([v['name_ar'] for v in ci['values']])
                    narration_parts.append(f"قيمنا الأساسية هي: {vals_text}.")
                if "منتجاتنا" in slides_sel:
                    prods = "، ".join([p['name'] for p in ci['products']])
                    narration_parts.append(f"نقدم مجموعة من المنتجات المبتكرة تشمل: {prods}.")
                if "فريقك" in slides_sel:
                    narration_parts.append(f"ستعمل في قسم {plan['dept']} تحت إشراف {plan.get('manager','')}. تاريخ مباشرتك هو {plan['start_date']}.")
                narration_parts.append("نتمنى لك التوفيق والنجاح في رحلتك معنا!")
                full_narration = " ".join(narration_parts)

                # Build values HTML
                vals_html = ""
                colors = ['#E36414','#264653','#2A9D8F','#E9C46A','#F4A261']
                for i, v in enumerate(ci.get('values',[])):
                    vals_html += f"<div style='background:white;border-radius:12px;padding:20px;border-top:4px solid {colors[i%5]};text-align:center'><h3 style='color:{colors[i%5]};margin:0'>{v['name_ar']}</h3><p style='font-size:0.85em;color:#0F4C5C;font-weight:600;margin:5px 0'>{v['name']}</p><p style='color:#555;font-size:0.85em;margin:0'>{v['desc']}</p></div>"

                # Build products HTML
                prods_html = ""
                for p in ci.get('products',[]):
                    prods_html += f"<div style='background:white;border-radius:10px;padding:18px;border-right:4px solid #E36414'><h4 style='color:#0F4C5C;margin:0 0 5px'>{p['name']}</h4><p style='color:#555;margin:0;font-size:0.9em'>{p['desc']}</p></div>"

                # Build milestones HTML
                miles_html = ""
                for m in ci.get('milestones',[]):
                    miles_html += f"<div style='display:flex;align-items:center;gap:15px;margin:10px 0'><div style='background:#E36414;color:white;padding:8px 16px;border-radius:20px;font-weight:700;min-width:60px;text-align:center'>{m['year']}</div><div style='color:#333'>{m['event']}</div></div>"

                # Build market stats HTML
                market = ci.get('market',{})
                market_html = ""
                for k, v in market.items():
                    label = k.replace('_',' ').title()
                    market_html += f"<div style='background:white;border-radius:10px;padding:15px;text-align:center'><div style='font-size:1.3em;font-weight:700;color:#E36414'>{v}</div><div style='color:#888;font-size:0.8em;margin-top:4px'>{label}</div></div>"

                # Build tasks HTML
                tasks_html = ""
                for phase, tasks in ONBOARDING_CHECKLIST.items():
                    tasks_html += f"<h4 style='color:#E36414;margin:15px 0 8px'>{phase}</h4>"
                    for t in tasks[:4]:
                        tasks_html += f"<div style='padding:6px 12px;margin:4px 0;background:#f8f9fa;border-radius:6px;font-size:0.9em'>✅ {t['task']} <span style='color:#888'>({t['owner']})</span></div>"

                stats = ci.get('stats',{})

                html = f"""<!DOCTYPE html>
<html dir="rtl" lang="ar">
<head>
<meta charset="UTF-8"><meta name="viewport" content="width=device-width,initial-scale=1">
<title>Onboarding - {plan['name']}</title>
<style>
@import url('https://fonts.googleapis.com/css2?family=Noto+Sans+Arabic:wght@300;400;600;700;800&display=swap');
*{{margin:0;padding:0;box-sizing:border-box}}
body{{font-family:'Noto Sans Arabic',sans-serif;background:#0a0a1a;color:white;overflow:hidden;height:100vh}}
.slide{{display:none;height:100vh;padding:40px 60px;flex-direction:column;justify-content:center;position:relative;overflow:hidden}}
.slide.active{{display:flex}}
.slide::before{{content:'';position:absolute;top:0;left:0;right:0;bottom:0;z-index:0}}
.slide>*{{position:relative;z-index:1}}
.s-welcome::before{{background:linear-gradient(135deg,#0F4C5C,#1A1A2E)}}
.s-who::before{{background:linear-gradient(135deg,#1A1A2E,#264653)}}
.s-vision::before{{background:linear-gradient(135deg,#264653,#0F4C5C)}}
.s-values::before{{background:linear-gradient(135deg,#f8f9fa,#e9ecef);}} .s-values *{{color:#333}}
.s-products::before{{background:linear-gradient(135deg,#1A1A2E,#0F4C5C)}}
.s-milestones::before{{background:linear-gradient(135deg,#0F4C5C,#264653)}}
.s-market::before{{background:linear-gradient(135deg,#264653,#1A1A2E)}}
.s-team::before{{background:linear-gradient(135deg,#E36414,#E9C46A);}}
.s-plan::before{{background:linear-gradient(135deg,#f8f9fa,#e9ecef);}} .s-plan *{{color:#333}}
.s-policies::before{{background:linear-gradient(135deg,#264653,#0F4C5C)}}
.s-contact::before{{background:linear-gradient(135deg,#0F4C5C,#1A1A2E)}}
.logo{{width:70px;height:70px;background:linear-gradient(135deg,#E36414,#E9C46A);border-radius:14px;display:flex;align-items:center;justify-content:center;font-size:24px;font-weight:800;margin-bottom:20px}}
h1{{font-size:2.8em;margin-bottom:10px}} h2{{font-size:2em;margin-bottom:15px}} h3{{margin-bottom:8px}}
.subtitle{{opacity:0.7;font-size:1.1em;margin-bottom:30px}}
.grid2{{display:grid;grid-template-columns:1fr 1fr;gap:15px;margin-top:15px}}
.grid3{{display:grid;grid-template-columns:1fr 1fr 1fr;gap:12px;margin-top:15px}}
.grid5{{display:grid;grid-template-columns:repeat(5,1fr);gap:10px;margin-top:15px}}
.stat{{background:rgba(255,255,255,0.1);border-radius:10px;padding:18px;text-align:center}}
.stat .num{{font-size:2em;font-weight:800;color:#E9C46A}}
.stat .lbl{{font-size:0.8em;opacity:0.7;margin-top:4px}}
.nav{{position:fixed;bottom:20px;left:50%;transform:translateX(-50%);display:flex;gap:8px;z-index:999;background:rgba(0,0,0,0.6);padding:10px 20px;border-radius:30px;backdrop-filter:blur(10px)}}
.nav button{{background:rgba(255,255,255,0.2);border:none;color:white;padding:8px 16px;border-radius:20px;cursor:pointer;font-family:inherit;font-size:0.9em}}
.nav button:hover{{background:#E36414}} .nav button.active{{background:#E36414}}
.nav .speak{{background:#27AE60}}
.counter{{position:fixed;top:20px;left:20px;background:rgba(0,0,0,0.4);padding:6px 14px;border-radius:20px;font-size:0.8em;z-index:999}}
.progress{{position:fixed;top:0;left:0;height:4px;background:#E36414;z-index:999;transition:width 0.3s}}
</style>
</head>
<body>
<div class="progress" id="prog"></div>
<div class="counter" id="counter"></div>

{"<div class='slide s-welcome active' id='s0'><div class='logo'>HR</div><h1>" + ci['name_en'] + "</h1><p class='subtitle'>" + ci['tagline'] + "</p><h2>Welcome, " + plan['name'] + "!</h2><p style=\"font-size:1.2em;opacity:0.9\">" + plan['title'] + " | " + plan['dept'] + "</p><p style=\"margin-top:20px;opacity:0.6\">Starting: " + plan['start_date'] + "</p></div>" if "ترحيب" in slides_sel else ""}

{"<div class='slide s-who'><h2>🏢 Who We Are</h2><p style=\"font-size:1.05em;line-height:1.8;max-width:800px\">" + ci.get('who_we_are','')[:500] + "</p><div class='grid2' style='margin-top:25px'><div class='stat'><div class='num'>" + stats.get('brands','1000+') + "</div><div class='lbl'>Brands</div></div><div class='stat'><div class='num'>" + stats.get('clients','1000+') + "</div><div class='lbl'>Clients</div></div><div class='stat'><div class='num'>" + stats.get('users','1.5M+') + "</div><div class='lbl'>Users</div></div><div class='stat'><div class='num'>" + stats.get('growth','100%+') + "</div><div class='lbl'>Growth</div></div></div></div>" if "من نحن" in slides_sel else ""}

{"<div class='slide s-vision'><h2>🎯 Our Purpose, Mission & Vision</h2><div style='margin:20px 0'><div style='background:rgba(255,255,255,0.1);padding:20px;border-radius:12px;margin:12px 0;border-right:4px solid #E9C46A'><h3 style='color:#E9C46A'>Purpose (Why)</h3><p>" + ci.get('purpose','') + "</p></div><div style='background:rgba(255,255,255,0.1);padding:20px;border-radius:12px;margin:12px 0;border-right:4px solid #E36414'><h3 style='color:#E36414'>Mission (What)</h3><p>" + ci.get('mission','') + "</p></div><div style='background:rgba(255,255,255,0.1);padding:20px;border-radius:12px;margin:12px 0;border-right:4px solid #2A9D8F'><h3 style='color:#2A9D8F'>Vision (Where)</h3><p>" + ci.get('vision','') + "</p></div></div></div>" if "رؤيتنا ورسالتنا" in slides_sel else ""}

{"<div class='slide s-values'><h2 style='color:#0F4C5C'>💎 Our Values</h2><div class='grid5'>" + vals_html + "</div></div>" if "قيمنا" in slides_sel else ""}

{"<div class='slide s-products'><h2>📦 Our Products</h2><div class='grid2' style='gap:12px'>" + prods_html + "</div></div>" if "منتجاتنا" in slides_sel else ""}

{"<div class='slide s-milestones'><h2>🏆 Our Journey</h2><div style='max-width:700px'>" + miles_html + "</div></div>" if "إنجازاتنا" in slides_sel else ""}

{"<div class='slide s-market'><h2>📊 Market Opportunity</h2><div class='grid3'>" + market_html + "</div></div>" if "السوق" in slides_sel else ""}

{"<div class='slide s-team'><h2>👥 Your Team</h2><div class='grid2'><div class='stat'><div class='lbl'>Department</div><div class='num' style='font-size:1.3em'>" + plan['dept'] + "</div></div><div class='stat'><div class='lbl'>Manager</div><div class='num' style='font-size:1.3em'>" + plan.get('manager','') + "</div></div><div class='stat'><div class='lbl'>Start Date</div><div class='num' style='font-size:1.3em'>" + plan['start_date'] + "</div></div><div class='stat'><div class='lbl'>Type</div><div class='num' style='font-size:1.3em'>" + plan.get('type','Full-Time') + "</div></div></div></div>" if "فريقك" in slides_sel else ""}

{"<div class='slide s-plan'><h2 style='color:#0F4C5C'>📋 Your 30/60/90 Day Plan</h2><div style='max-height:70vh;overflow-y:auto'>" + tasks_html + "</div></div>" if "خطة 30/60/90" in slides_sel else ""}

{"<div class='slide s-policies'><h2>📌 Key Policies</h2><div class='grid2'>" + "".join([f"<div class='stat'><div class='lbl'>{k}</div><div style='font-size:1em;margin-top:6px'>{v}</div></div>" for k,v in ci.get('work_policies',{}).items()]) + "</div></div>" if "سياسات" in slides_sel else ""}

{"<div class='slide s-contact'><h2>📞 Stay Connected</h2><div class='grid2'>" + "".join([f"<div class='stat'><div class='lbl'>{k}</div><div class='num' style='font-size:1em'>{v}</div></div>" for k,v in ci.get('contacts',{}).items()]) + "</div><p style='margin-top:40px;font-size:1.3em;text-align:center'>Welcome aboard, {plan['name']}! 🎉</p></div>" if "تواصل" in slides_sel else ""}

<div class="nav">
<button onclick="prev()">◀ السابق</button>
<button onclick="next()">التالي ▶</button>
<button class="speak" onclick="speakAll()">🔊 سرد صوتي</button>
<button onclick="stopSpeak()">⏹</button>
</div>

<script>
const slides=document.querySelectorAll('.slide');
let cur=0;
function show(n){{cur=Math.max(0,Math.min(n,slides.length-1));slides.forEach((s,i)=>s.classList.toggle('active',i===cur));document.getElementById('counter').textContent=(cur+1)+'/'+slides.length;document.getElementById('prog').style.width=((cur+1)/slides.length*100)+'%';}}
function next(){{show(cur+1)}} function prev(){{show(cur-1)}}
document.addEventListener('keydown',e=>{{if(e.key==='ArrowRight'||e.key===' ')next();if(e.key==='ArrowLeft')prev();}});
show(0);
const narration=`{full_narration}`;
let utt;
function speakAll(){{stopSpeak();utt=new SpeechSynthesisUtterance(narration);utt.lang='{voice_lang}';utt.rate=0.85;speechSynthesis.speak(utt);}}
function stopSpeak(){{speechSynthesis.cancel()}}
</script>
</body></html>"""

                st.components.v1.html(html, height=650, scrolling=False)
                st.caption("⬅️ ➡️ استخدم الأسهم أو الأزرار للتنقل بين الشرائح | 🔊 اضغط 'سرد صوتي' للاستماع")

                dc1, dc2 = st.columns(2)
                with dc1:
                    st.download_button("📥 تحميل العرض HTML", data=html.encode('utf-8'),
                        file_name=f"Onboarding_Presentation_{plan['name']}.html",
                        mime="text/html", use_container_width=True)
                with dc2:
                    if plan.get('email'):
                        if st.button("📧 إرسال للموظف", use_container_width=True, key="ob_send"):
                            ok, msg = send_test_email(plan['email'], plan['name'],
                                [f"عرض Onboarding التقديمي جاهز لك", f"المسمى: {plan['title']}", f"البدء: {plan['start_date']}"],
                                plan['start_date'], ci['name'])
                            if ok: st.success(f"📧 تم الإرسال إلى {plan['email']}")
                            else: st.warning(f"⚠️ {msg}")

        elif page == "🏢 معلومات الشركة":
            hdr("🏢 إدارة معلومات الشركة","تحديث المعلومات التي تظهر في عرض Onboarding")

            if st.session_state.get('user_role') != "مدير":
                st.warning("⚠️ تعديل معلومات الشركة متاح للمدير فقط")
                # Show read-only
                ci = st.session_state.company_info
                st.markdown(f"### {ci['name']} ({ci['name_en']})")
                st.markdown(f"**الرؤية:** {ci['vision']}")
                st.markdown(f"**الرسالة:** {ci['mission']}")
                return

            ci = st.session_state.company_info

            with st.expander("🏢 المعلومات الأساسية", expanded=True):
                ci['name'] = st.text_input("اسم الشركة (عربي):", value=ci.get('name',''), key="ci_name")
                ci['name_en'] = st.text_input("اسم الشركة (English):", value=ci.get('name_en',''), key="ci_name_en")
                ci['tagline'] = st.text_input("الشعار:", value=ci.get('tagline',''), key="ci_tag")

            with st.expander("📝 من نحن"):
                ci['who_we_are'] = st.text_area("نبذة عن الشركة:", value=ci.get('who_we_are',''), height=200, key="ci_who")

            with st.expander("🎯 الغرض والرؤية والرسالة"):
                ci['purpose'] = st.text_area("Purpose (Why):", value=ci.get('purpose',''), key="ci_purpose")
                ci['mission'] = st.text_area("Mission (What):", value=ci.get('mission',''), key="ci_mission")
                ci['vision'] = st.text_area("Vision (Where):", value=ci.get('vision',''), key="ci_vision")

            with st.expander("💎 القيم"):
                for i, v in enumerate(ci.get('values',[])):
                    vc1, vc2, vc3 = st.columns([1,1,2])
                    with vc1: ci['values'][i]['name'] = st.text_input(f"القيمة {i+1} EN:", value=v['name'], key=f"cv_en_{i}")
                    with vc2: ci['values'][i]['name_ar'] = st.text_input(f"القيمة {i+1} AR:", value=v['name_ar'], key=f"cv_ar_{i}")
                    with vc3: ci['values'][i]['desc'] = st.text_input(f"الوصف:", value=v['desc'], key=f"cv_d_{i}")

            with st.expander("📦 المنتجات"):
                for i, p in enumerate(ci.get('products',[])):
                    pc1, pc2 = st.columns([1,2])
                    with pc1: ci['products'][i]['name'] = st.text_input(f"المنتج {i+1}:", value=p['name'], key=f"cp_n_{i}")
                    with pc2: ci['products'][i]['desc'] = st.text_input(f"الوصف:", value=p['desc'], key=f"cp_d_{i}")

            with st.expander("📊 إحصائيات"):
                sc1, sc2, sc3, sc4 = st.columns(4)
                with sc1: ci['stats']['brands'] = st.text_input("Brands:", value=ci.get('stats',{}).get('brands',''), key="cs_b")
                with sc2: ci['stats']['clients'] = st.text_input("Clients:", value=ci.get('stats',{}).get('clients',''), key="cs_c")
                with sc3: ci['stats']['users'] = st.text_input("Users:", value=ci.get('stats',{}).get('users',''), key="cs_u")
                with sc4: ci['stats']['growth'] = st.text_input("Growth:", value=ci.get('stats',{}).get('growth',''), key="cs_g")

            with st.expander("📞 التواصل والسياسات"):
                cc1, cc2, cc3 = st.columns(3)
                with cc1: ci['contacts']['hr_email'] = st.text_input("HR Email:", value=ci.get('contacts',{}).get('hr_email',''), key="cc_hr")
                with cc2: ci['contacts']['it_email'] = st.text_input("IT Email:", value=ci.get('contacts',{}).get('it_email',''), key="cc_it")
                with cc3: ci['contacts']['website'] = st.text_input("Website:", value=ci.get('contacts',{}).get('website',''), key="cc_web")

            if st.button("💾 حفظ معلومات الشركة", type="primary", use_container_width=True, key="ci_save"):
                st.session_state.company_info = ci
                try:
                    conn = get_conn()
                    c = conn.cursor()
                    _upsert_config(c, "company_info", json.dumps(ci, ensure_ascii=False))
                    conn.commit()
                    conn.close()
                    st.success("✅ تم حفظ معلومات الشركة في قاعدة البيانات")
                except Exception as e:
                    st.warning(f"⚠️ تم الحفظ محلياً فقط: {e}")

        elif page == "📥 تصدير Onboarding":
            hdr("📥 تصدير تقارير Onboarding")
            plans = st.session_state.onboarding_plans
            if not plans:
                st.info("لا توجد خطط Onboarding للتصدير")
                return

            if st.button("📥 تصدير Excel", type="primary", use_container_width=True, key="ob_exp"):
                ox = io.BytesIO()
                with pd.ExcelWriter(ox, engine='xlsxwriter') as w:
                    # Plans summary
                    plan_rows = [{"الاسم":p['name'],"المسمى":p['title'],"القسم":p['dept'],
                        "المدير":p.get('manager',''),"البدء":p['start_date'],"النوع":p.get('type',''),
                        "الحالة":p.get('status',''),"أُنشئ":p.get('created_at','')} for p in plans]
                    pd.DataFrame(plan_rows).to_excel(w, sheet_name='Onboarding Plans', index=False)

                    # Full checklist
                    checklist_rows = []
                    for phase, tasks in ONBOARDING_CHECKLIST.items():
                        for t in tasks:
                            checklist_rows.append({"المرحلة":phase,"المهمة":t['task'],"المسؤول":t['owner'],"اليوم":t['days']})
                    pd.DataFrame(checklist_rows).to_excel(w, sheet_name='Checklist Template', index=False)

                st.download_button("📥 تحميل", data=ox.getvalue(),
                    file_name=f"Onboarding_Report_{datetime.now().strftime('%Y%m%d')}.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")


    # =========================================
    #         📜 CONTRACTS MODULE (DOCX Templates)
    # =========================================
    elif section == "📜 العقود":

        if 'saved_contracts' not in st.session_state:
            try:
                conn = get_conn(); c = conn.cursor()
                c.execute(f"SELECT value FROM app_config WHERE key = {_ph()}", ("saved_contracts",))
                row = c.fetchone(); conn.close()
                st.session_state.saved_contracts = json.loads(row[0]) if row else []
            except:
                st.session_state.saved_contracts = []

        # Template file paths (uploaded to GitHub alongside app)
        APP_DIR = os.path.dirname(os.path.abspath(__file__))
        TEMPLATES = {
            "عقد عمل محدد المدة": {"file": "contract_employment.docx", "icon": "💼"},
            "عقد تدريب": {"file": "contract_training.docx", "icon": "🎓"},
            "عقد خدمات فنية واستشارية": {"file": "contract_consulting.docx", "icon": "🤝"},
        }

        def _replace_in_runs(paragraph, old_text, new_text):
            """Replace text across runs while preserving formatting"""
            full = paragraph.text
            if old_text not in full:
                return False
            # Try run-by-run replacement first
            for run in paragraph.runs:
                if old_text in run.text:
                    run.text = run.text.replace(old_text, str(new_text))
                    return True
            # Cross-run replacement: rebuild runs
            if old_text in full:
                new_full = full.replace(old_text, str(new_text), 1)
                # Clear all runs except first, put all text in first run
                if paragraph.runs:
                    first_run = paragraph.runs[0]
                    fmt = first_run.font
                    for run in paragraph.runs[1:]:
                        run.text = ""
                    first_run.text = new_full
                    return True
            return False

        def _fill_contract(template_path, replacements):
            """Fill a DOCX template with replacement values"""
            from docx import Document as DocxDocument
            doc = DocxDocument(template_path)
            for para in doc.paragraphs:
                for old_val, new_val in replacements.items():
                    if old_val in para.text:
                        _replace_in_runs(para, old_val, new_val)
            # Also check tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            for old_val, new_val in replacements.items():
                                if old_val in para.text:
                                    _replace_in_runs(para, old_val, new_val)
            return doc

        if page == "📜 إنشاء عقد":
            hdr("📜 إنشاء عقد جديد","نفس نموذج العقد الأصلي مع تعبئة البيانات تلقائياً")

            contract_type = st.selectbox("📝 نوع العقد:", list(TEMPLATES.keys()), key="ct_type")
            tmpl = TEMPLATES[contract_type]
            template_path = os.path.join(APP_DIR, tmpl['file'])

            # Check template exists
            if not os.path.exists(template_path):
                st.warning(f"⚠️ ملف القالب غير موجود: {tmpl['file']}")
                st.info("ارفع ملفات القوالب الثلاثة (contract_employment.docx, contract_training.docx, contract_consulting.docx) على GitHub بجانب ملف التطبيق")
                # Allow upload
                uploaded_tmpl = st.file_uploader(f"أو ارفع القالب هنا ({tmpl['file']}):", type=["docx"], key="ct_upload")
                if uploaded_tmpl:
                    template_path = os.path.join(APP_DIR, tmpl['file'])
                    with open(template_path, 'wb') as f:
                        f.write(uploaded_tmpl.getvalue())
                    st.success("✅ تم حفظ القالب")
                    st.rerun()
                else:
                    return

            st.markdown(f"### {tmpl['icon']} {contract_type}")

            # ===== EMPLOYMENT CONTRACT =====
            if contract_type == "عقد عمل محدد المدة":
                st.markdown("### 👤 بيانات الموظف")
                ec1, ec2, ec3 = st.columns(3)
                with ec1:
                    ct_name = st.text_input("اسم الموظف:", key="ct_name")
                    ct_nationality = st.text_input("الجنسية:", key="ct_nat")
                    ct_id_type = st.selectbox("نوع الهوية:", ["جواز","هوية وطنية","إقامة"], key="ct_idt")
                with ec2:
                    ct_id_num = st.text_input("رقم الهوية:", key="ct_idn")
                    ct_city = st.text_input("عنوان الموظف (المدينة):", key="ct_city")
                    ct_email = st.text_input("البريد الإلكتروني:", key="ct_email")
                with ec3:
                    ct_phone = st.text_input("رقم الجوال:", key="ct_phone")
                    ct_job = st.text_input("المسمى الوظيفي:", key="ct_job")
                    ct_work_city = st.text_input("مدينة العمل:", value="جدة", key="ct_wcity")

                st.markdown("### 📅 مدة العقد")
                dc1, dc2, dc3 = st.columns(3)
                with dc1: ct_sign_date = st.date_input("تاريخ التوقيع:", key="ct_sign")
                with dc2: ct_start = st.date_input("تاريخ المباشرة:", key="ct_start")
                with dc3: ct_duration = st.selectbox("المدة:", ["سنة واحدة","سنتين","ثلاث سنوات"], key="ct_dur")

                st.markdown("### 💰 الراتب والبدلات")
                sc1, sc2, sc3, sc4 = st.columns(4)
                with sc1: ct_basic = st.number_input("الراتب الأساسي:", value=0.0, format="%.2f", key="ct_basic")
                with sc2: ct_housing = st.number_input("بدل السكن:", value=0.0, format="%.2f", key="ct_housing")
                with sc3: ct_transport = st.number_input("بدل المواصلات:", value=0.0, format="%.2f", key="ct_trans")
                with sc4: ct_other = st.number_input("بدلات أخرى:", value=0.0, format="%.2f", key="ct_other")

                total_sal = ct_basic + ct_housing + ct_transport + ct_other
                st.success(f"💰 إجمالي الراتب الشهري: **{total_sal:,.2f}** ريال")

                # Build replacement map for employment contract
                replacements = {
                    "أنس عادل حسنين حسين": ct_name,
                    "مصري": ct_nationality,
                    "جواز رقم (A31420125)": f"{ct_id_type} رقم ({ct_id_num})",
                    "A31420125": ct_id_num,
                    "القاهرة": ct_city if ct_city else "القاهرة",
                    "anassaddel@gmail.com": ct_email,
                    "00201030698018": ct_phone,
                    "مصمم واجهة المستخدم \"UX/UI\"": ct_job,
                    "(الخميس)15/08/2024م": f"({ct_sign_date.strftime('%A')}) {ct_sign_date.strftime('%d/%m/%Y')}م",
                    "15/.08/2024م": ct_start.strftime('%d/%m/%Y') + "م",
                    "15/08/2024م": ct_start.strftime('%d/%m/%Y') + "م",
                    "سنة واحدة": ct_duration,
                    "(3,409,09)": f"({ct_basic:,.2f})",
                    "3,409,09": f"{ct_basic:,.2f}",
                    "852,27 ريال سعودي": f"{ct_housing:,.2f} ريال سعودي",
                    "511.36 ريال سعودي": f"{ct_transport:,.2f} ريال سعودي",
                    "852.27 ريال سعودي": f"{ct_other:,.2f} ريال سعودي",
                }

            # ===== TRAINING CONTRACT =====
            elif contract_type == "عقد تدريب":
                st.markdown("### 👤 بيانات المتدرب")
                tc1, tc2, tc3 = st.columns(3)
                with tc1:
                    ct_name = st.text_input("اسم المتدرب:", key="ct_name")
                    ct_nationality = st.text_input("الجنسية:", key="ct_nat")
                with tc2:
                    ct_id_num = st.text_input("رقم الهوية الوطنية:", key="ct_idn")
                    ct_city = st.text_input("المدينة:", value="جدة", key="ct_city")
                with tc3:
                    ct_email = st.text_input("البريد الإلكتروني:", key="ct_email")
                    ct_phone = st.text_input("رقم الجوال:", key="ct_phone")

                st.markdown("### 📅 تفاصيل التدريب")
                dc1, dc2, dc3 = st.columns(3)
                with dc1:
                    ct_sign_date = st.date_input("تاريخ التوقيع:", key="ct_sign")
                    ct_dept = st.text_input("قسم التدريب:", key="ct_dept")
                with dc2:
                    ct_start = st.date_input("تاريخ بداية التدريب:", key="ct_start")
                    ct_train_city = st.text_input("مدينة التدريب:", value="جدة", key="ct_tcity")
                with dc3:
                    ct_duration = st.selectbox("مدة التدريب:", ["شهر","شهرين","3 أشهر","6 أشهر","سنة"], key="ct_dur")
                    ct_hours = st.number_input("ساعات التدريب الأسبوعية:", value=40, key="ct_hours")

                ct_reward = st.number_input("المكافأة الشهرية (ريال) - اختياري:", value=0.0, format="%.2f", key="ct_reward")

                replacements = {
                    "(...............)......../......./............": f"({ct_sign_date.strftime('%A')}) {ct_sign_date.strftime('%d/%m/%Y')}",
                    " ........... ": f" {ct_city} ",
                    ": .............................، ": f": {ct_name}، ",
                    "...............  ": f"{ct_nationality}  ",
                    " (........................) ": f" ({ct_id_num}) ",
                    " .............": f" {ct_city}",
                    ".........................................": ct_email,
                    " (............................) ": f" ({ct_phone}) ",
                    " (..........................)  ": f" ({ct_dept})  ",
                    " (...................)": f" ({ct_train_city})",
                    " ................، ": f" {ct_duration}، ",
                    " (........./........../..............": f" ({ct_start.strftime('%d/%m/%Y')}",
                    " (......) ": f" ({ct_hours}) ",
                    " ...... ": f" {ct_hours // 5} ",
                    " (........) ": f" ({int(ct_hours * 0.75)}) ",
                    " (..........)": f" ({int(ct_hours * 0.75 / 5)})",
                    ":  .....................................": f":  {ct_name}",
                    ": .....................................": f": {ct_name}",
                }

            # ===== CONSULTING CONTRACT =====
            else:
                st.markdown("### 👤 بيانات المستشار")
                cc1, cc2, cc3 = st.columns(3)
                with cc1:
                    ct_name = st.text_input("اسم المستشار:", key="ct_name")
                    ct_nationality = st.text_input("الجنسية:", key="ct_nat")
                with cc2:
                    ct_id_type = st.selectbox("نوع الهوية:", ["هوية وطنية","جواز سفر","إقامة"], key="ct_idt")
                    ct_id_num = st.text_input("رقم الهوية:", key="ct_idn")
                with cc3:
                    ct_city = st.text_input("المدينة:", value="جدة", key="ct_city")
                    ct_email = st.text_input("البريد الإلكتروني:", key="ct_email")

                ct_phone = st.text_input("رقم الجوال:", key="ct_phone")

                st.markdown("### 📋 تفاصيل العقد")
                ct_service = st.text_area("وصف الخدمات المطلوبة:", key="ct_service")
                dc1, dc2, dc3 = st.columns(3)
                with dc1:
                    ct_sign_date = st.date_input("تاريخ التوقيع:", key="ct_sign")
                    ct_duration_months = st.selectbox("مدة العقد (أشهر):", [1,2,3,6,9,12], index=2, key="ct_dur")
                with dc2:
                    ct_start = st.date_input("تاريخ البداية:", key="ct_start")
                with dc3:
                    ct_fee = st.number_input("إجمالي الأتعاب (ريال):", value=0.0, format="%.2f", key="ct_fee")

                replacements = {
                    "..... /": f"{ct_sign_date.strftime('%d')} /",
                    "........ه  الموافق ../...../.....م": f"ه  الموافق {ct_sign_date.strftime('%d/%m/%Y')}م",
                    "..............................، الجنسية....................بموجب هوية (": f"{ct_name}، الجنسية {ct_nationality} بموجب {ct_id_type} (",
                    "......................) ": f"{ct_id_num}) ",
                    "...............،البريد": f"{ct_city}، البريد",
                    ".............................، جوال رقم: ........": f"{ct_email}، جوال رقم: {ct_phone}",
                    ".............": ct_service[:30] if ct_service else "",
                    ".....................................................": ct_service,
                    "...............أشهر": f"{ct_duration_months} أشهر",
                    "......../........./ ...........م": ct_start.strftime('%d/%m/%Y') + "م",
                    "......../........./...........م،": (ct_start + pd.DateOffset(months=ct_duration_months)).strftime('%d/%m/%Y') + "م،",
                    "..................................": f"{ct_fee:,.2f}",
                    "الاسم: .....................................": f"الاسم: {ct_name}",
                }

            # Generate button
            if st.button("📜 إنشاء وتحميل العقد (Word)", type="primary", use_container_width=True, key="ct_gen"):
                if not ct_name:
                    st.error("يرجى إدخال الاسم")
                else:
                    try:
                        doc = _fill_contract(template_path, replacements)
                        doc_bytes = io.BytesIO()
                        doc.save(doc_bytes)

                        # Save contract record
                        contract = {
                            "id": len(st.session_state.saved_contracts) + 1,
                            "type": contract_type, "name": ct_name,
                            "data": {k: str(v)[:100] for k, v in replacements.items() if v and not k.startswith('.')},
                            "created_at": datetime.now().strftime("%Y-%m-%d %H:%M"),
                            "created_by": st.session_state.get('user_name',''),
                            "status": "نشط"
                        }
                        st.session_state.saved_contracts.append(contract)
                        try:
                            conn = get_conn()
                            c = conn.cursor()
                            _upsert_config(c, "saved_contracts", json.dumps(st.session_state.saved_contracts, ensure_ascii=False))
                            conn.commit()
                            conn.close()
                        except: pass

                        st.download_button(f"📥 تحميل {contract_type} - {ct_name}.docx",
                            data=doc_bytes.getvalue(),
                            file_name=f"{contract_type}_{ct_name}_{ct_start.strftime('%Y%m%d')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            type="primary", use_container_width=True)
                        st.success(f"✅ تم إنشاء {contract_type} لـ {ct_name}")
                    except Exception as e:
                        st.error(f"خطأ: {e}")

        elif page == "🔍 تحليل العقود":
            hdr("🔍 تحليل العقود","ارفع عقد (Word أو PDF) لتحليله واستخراج البيانات والبنود")

            ana_file = st.file_uploader("📁 ارفع عقد للتحليل:", type=["docx","pdf","txt"], key="ct_ana_file")

            if ana_file:
                contract_text = ""
                file_type = ana_file.name.split('.')[-1].lower()

                # Extract text
                if file_type == "docx":
                    try:
                        from docx import Document as DocxDoc
                        doc = DocxDoc(io.BytesIO(ana_file.getvalue()))
                        contract_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
                    except Exception as e:
                        st.error(f"خطأ في قراءة الملف: {e}")
                elif file_type == "pdf":
                    try:
                        from fpdf import FPDF
                        # Try pdfplumber or basic extraction
                        try:
                            import pdfplumber
                            with pdfplumber.open(io.BytesIO(ana_file.getvalue())) as pdf_reader:
                                for pg in pdf_reader.pages:
                                    txt = pg.extract_text()
                                    if txt: contract_text += txt + "\n"
                        except ImportError:
                            st.warning("جاري القراءة بطريقة بديلة...")
                            contract_text = ana_file.getvalue().decode('utf-8', errors='ignore')
                    except: pass
                elif file_type == "txt":
                    contract_text = ana_file.getvalue().decode('utf-8', errors='ignore')

                if not contract_text.strip():
                    st.warning("⚠️ لم يتم استخراج نص من الملف. جرب ملف Word أو PDF نصي.")
                    return

                st.success(f"✅ تم استخراج {len(contract_text)} حرف من الملف")

                # ===== Auto Analysis =====
                st.markdown("---")
                st.markdown("### 📊 التحليل التلقائي للعقد")

                # 1. Contract type detection
                ct_type_detected = "غير محدد"
                if any(w in contract_text for w in ["تدريب","متدرب","برنامج تدريبي"]): ct_type_detected = "عقد تدريب"
                elif any(w in contract_text for w in ["استشار","خدمات فنية","أتعاب"]): ct_type_detected = "عقد خدمات استشارية"
                elif any(w in contract_text for w in ["عامل","وظيفة","راتب","بدل سكن"]): ct_type_detected = "عقد عمل"
                elif any(w in contract_text for w in ["إيجار","تأجير"]): ct_type_detected = "عقد إيجار"
                elif any(w in contract_text for w in ["بيع","شراء","مبيع"]): ct_type_detected = "عقد بيع"

                # 2. Extract key data
                import re
                extracted = {}

                # Names
                name_patterns = [
                    r'الاسم[:\s]+([^\n,،]+)',
                    r'السيد/?\s*([^\n,،]+)',
                    r'المتدرب[:\s]+([^\n,،]+)',
                ]
                for p in name_patterns:
                    m = re.search(p, contract_text)
                    if m and len(m.group(1).strip()) > 2:
                        extracted['الطرف الثاني'] = m.group(1).strip()[:50]
                        break

                # Dates
                date_patterns = [
                    r'(\d{1,2}[/\-\.]\d{1,2}[/\-\.]\d{2,4})',
                    r'(\d{4}[/\-]\d{1,2}[/\-]\d{1,2})',
                ]
                dates_found = []
                for p in date_patterns:
                    dates_found.extend(re.findall(p, contract_text))
                if dates_found:
                    extracted['التواريخ المذكورة'] = ", ".join(dates_found[:5])

                # Amounts
                amount_patterns = [
                    r'(\d[\d,\.]+)\s*ريال',
                    r'مبلغ[اً]?\s*(?:وقدره)?\s*[\(]?(\d[\d,\.]+)',
                    r'راتب[اً]?\s*(?:أساسي[اً]?)?\s*(?:قدره)?\s*[\(]?(\d[\d,\.]+)',
                ]
                amounts = []
                for p in amount_patterns:
                    amounts.extend(re.findall(p, contract_text))
                if amounts:
                    extracted['المبالغ المذكورة'] = ", ".join([f"{a} ريال" for a in amounts[:6]])

                # Duration
                dur_patterns = [
                    r'مدة العقد\s*([^\n\.]+)',
                    r'لمدة\s*([^\n\.]+)',
                ]
                for p in dur_patterns:
                    m = re.search(p, contract_text)
                    if m:
                        extracted['مدة العقد'] = m.group(1).strip()[:50]
                        break

                # ID numbers
                id_patterns = [
                    r'هوية[^\d]*(\d{10})',
                    r'جواز[^\d]*(\w{1,2}\d{6,10})',
                    r'إقامة[^\d]*(\d{10})',
                    r'سجل تجاري[^\d]*(\d{10})',
                ]
                ids_found = []
                for p in id_patterns:
                    ids_found.extend(re.findall(p, contract_text))
                if ids_found:
                    extracted['أرقام الهوية/السجلات'] = ", ".join(ids_found[:3])

                # 3. Clause analysis
                clauses_keywords = {
                    "السرية": ["سري","سرية","إفشاء","أسرار"],
                    "فترة التجربة": ["تجربة","فترة التجربة"],
                    "إنهاء العقد": ["إنهاء","فسخ","إلغاء","انتهاء"],
                    "التعويض": ["تعويض","غرامة","جزائي"],
                    "المنافسة": ["منافس","لا يعمل لدى","عدم المنافسة"],
                    "الإجازات": ["إجازة","إجازات","سنوية"],
                    "ساعات العمل": ["ساعات العمل","ساعات التدريب","دوام"],
                    "التأمين الطبي": ["تأمين","طبي","صحي"],
                    "التأمينات الاجتماعية": ["تأمينات","اجتماعية","GOSI"],
                    "مكافأة نهاية الخدمة": ["نهاية الخدمة","مكافأة","المادة 84","المادة 85"],
                    "حقوق الملكية الفكرية": ["ملكية فكرية","براءة","اختراع","حقوق"],
                    "القانون الحاكم": ["نظام العمل","محكمة","قضائي","مختصة"],
                    "البدلات": ["بدل سكن","بدل مواصلات","بدلات","علاوة"],
                }

                found_clauses = {}
                missing_clauses = []
                for clause_name, keywords in clauses_keywords.items():
                    found = any(kw in contract_text for kw in keywords)
                    if found:
                        found_clauses[clause_name] = "✅ موجود"
                    else:
                        found_clauses[clause_name] = "❌ غير موجود"
                        missing_clauses.append(clause_name)

                # Display results
                k1,k2,k3,k4 = st.columns(4)
                with k1: kpi("📄 نوع العقد", ct_type_detected)
                with k2: kpi("📝 عدد الأحرف", f"{len(contract_text):,}")
                with k3: kpi("✅ بنود موجودة", str(len(found_clauses) - len(missing_clauses)))
                with k4: kpi("⚠️ بنود ناقصة", str(len(missing_clauses)))

                # Extracted data
                if extracted:
                    st.markdown("### 📋 البيانات المستخرجة")
                    for k, v in extracted.items():
                        st.markdown(f"**{k}:** {v}")

                # Clauses table
                st.markdown("### 📊 تحليل البنود والشروط")
                clause_df = pd.DataFrame([{"البند":k, "الحالة":v} for k,v in found_clauses.items()])
                st.dataframe(clause_df, use_container_width=True, hide_index=True)

                # Warnings
                if missing_clauses:
                    st.markdown("### ⚠️ بنود مفقودة (يُنصح بإضافتها)")
                    for mc in missing_clauses:
                        st.warning(f"⚠️ **{mc}** غير موجود في العقد")

                # Compliance check
                st.markdown("### ✅ فحص التوافق مع نظام العمل السعودي (المعدل بالمرسوم م/44 - نافذ 19/2/2025)")
                compliance = []

                # 1. فترة التجربة - المادة 53 المعدلة
                if "فترة التجربة" in [k for k,v in found_clauses.items() if "✅" in v]:
                    # Check if duration exceeds 180 days
                    if any(kw in contract_text for kw in ["180","مائة وثمانين"]):
                        compliance.append({"البند":"فترة التجربة","المادة":"53","التوافق":"✅ متوافق","الملاحظة":"مذكورة بالحد الأقصى 180 يوم"})
                    elif any(kw in contract_text for kw in ["90","تسعين"]):
                        compliance.append({"البند":"فترة التجربة","المادة":"53","التوافق":"✅ متوافق","الملاحظة":"مذكورة بـ 90 يوم. ملاحظة: التمديد غير متاح بعد التعديل"})
                    else:
                        compliance.append({"البند":"فترة التجربة","المادة":"53","التوافق":"✅ متوافق","الملاحظة":"مذكورة في العقد. الحد الأقصى 180 يوم"})
                else:
                    compliance.append({"البند":"فترة التجربة","المادة":"53","التوافق":"⚠️ غير مذكورة","الملاحظة":"يجب النص عليها صراحة في العقد. الحد الأقصى 180 يوم. التمديد محظور"})

                # 2. مكافأة نهاية الخدمة - المادة 84-85
                if any(kw in contract_text for kw in ["نهاية الخدمة","المادة 84","المادة 85","مكافأة"]):
                    compliance.append({"البند":"مكافأة نهاية الخدمة","المادة":"84-85","التوافق":"✅ متوافق","الملاحظة":"نصف شهر لكل سنة من الخمس الأولى + شهر لكل سنة بعدها"})
                else:
                    compliance.append({"البند":"مكافأة نهاية الخدمة","المادة":"84-85","التوافق":"⚠️ غير مذكورة","الملاحظة":"حق مكفول بالنظام: نصف شهر/سنة (أول 5) + شهر/سنة (بعدها)"})

                # 3. الإجازة السنوية - المادة 109
                if any(kw in contract_text for kw in ["إجازة سنوية","21 يوم","22 يوم","30 يوم","واحد وعشرين"]):
                    compliance.append({"البند":"الإجازة السنوية","المادة":"109","التوافق":"✅ متوافق","الملاحظة":"21 يوم كحد أدنى، 30 يوم بعد 5 سنوات خدمة"})
                else:
                    compliance.append({"البند":"الإجازة السنوية","المادة":"109","التوافق":"⚠️ غير مذكورة","الملاحظة":"21 يوم كحد أدنى (30 يوم بعد 5 سنوات)"})

                # 4. ساعات العمل - المادة 98
                if any(kw in contract_text for kw in ["48 ساعة","ثمان ساعات","8 ساعات","ساعات العمل"]):
                    compliance.append({"البند":"ساعات العمل","المادة":"98","التوافق":"✅ متوافق","الملاحظة":"8 ساعات يومياً / 48 أسبوعياً. 6 ساعات يومياً في رمضان"})
                else:
                    compliance.append({"البند":"ساعات العمل","المادة":"98","التوافق":"⚠️ غير محددة","الملاحظة":"يجب ألا تتجاوز 8 ساعات يومياً / 48 أسبوعياً (6 في رمضان)"})

                # 5. السرية
                if any(kw in contract_text for kw in ["سري","سرية","إفشاء","أسرار"]):
                    compliance.append({"البند":"السرية","المادة":"-","التوافق":"✅ متوافق","الملاحظة":"بند السرية موجود"})
                else:
                    compliance.append({"البند":"السرية","المادة":"-","التوافق":"⚠️ غير موجود","الملاحظة":"يُنصح بشدة بإضافة بند السرية وعدم الإفشاء"})

                # 6. التعويض عن الفصل غير المشروع - المادة 77
                if any(kw in contract_text for kw in ["المادة 77","تعويض","فصل غير مشروع","إنهاء غير مشروع"]):
                    compliance.append({"البند":"التعويض عن الإنهاء","المادة":"77","التوافق":"✅ متوافق","الملاحظة":"أجر المدة المتبقية (محدد) أو 15 يوم/سنة بحد أدنى شهرين (غير محدد)"})
                else:
                    compliance.append({"البند":"التعويض عن الإنهاء","المادة":"77","التوافق":"ℹ️ غير مذكور","الملاحظة":"حق مكفول بالنظام حتى لو لم يُذكر في العقد"})

                # 7. حالات إنهاء العقد - المادة 74, 80, 81
                if any(kw in contract_text for kw in ["إنهاء","فسخ","المادة 74","المادة 80"]):
                    compliance.append({"البند":"حالات إنهاء العقد","المادة":"74,80,81","التوافق":"✅ متوافق","الملاحظة":"بند الإنهاء موجود"})
                else:
                    compliance.append({"البند":"حالات إنهاء العقد","المادة":"74,80,81","التوافق":"⚠️ غير مذكور","الملاحظة":"يُنصح بتوضيح حالات الإنهاء والفسخ"})

                # 8. فترة الإشعار - التعديل الجديد
                if any(kw in contract_text for kw in ["إشعار","إنذار","30 يوم","60 يوم","ستين يوم"]):
                    compliance.append({"البند":"فترة الإشعار","المادة":"75 معدلة","التوافق":"✅ متوافق","الملاحظة":"60 يوم لصاحب العمل / 30 يوم للعامل (التعديل الجديد)"})
                else:
                    compliance.append({"البند":"فترة الإشعار","المادة":"75 معدلة","التوافق":"⚠️ غير مذكورة","الملاحظة":"60 يوم لصاحب العمل / 30 يوم للعامل عند عدم التجديد"})

                # 9. التأمينات الاجتماعية
                if any(kw in contract_text for kw in ["تأمينات","اجتماعية","GOSI"]):
                    compliance.append({"البند":"التأمينات الاجتماعية","المادة":"نظام التأمينات","التوافق":"✅ متوافق","الملاحظة":"الاشتراك في التأمينات إلزامي"})
                else:
                    compliance.append({"البند":"التأمينات الاجتماعية","المادة":"نظام التأمينات","التوافق":"ℹ️ غير مذكور","الملاحظة":"إلزامي بحكم النظام حتى لو لم يُذكر"})

                # 10. العقد المكتوب - التعديل الجديد
                compliance.append({"البند":"العقد المكتوب","المادة":"51 معدلة","التوافق":"✅ متوافق","الملاحظة":"العقود المكتوبة أصبحت إلزامية بعد التعديل خاصة لغير السعوديين"})

                # 11. عدم المنافسة
                if any(kw in contract_text for kw in ["منافس","لا يعمل لدى","عدم المنافسة"]):
                    compliance.append({"البند":"عدم المنافسة","المادة":"83","التوافق":"✅ موجود","الملاحظة":"يجب أن يكون محدد الزمان والمكان والنشاط"})

                # 12. إجازة الوضع (للعقود مع نساء)
                if any(kw in contract_text for kw in ["وضع","أمومة","حامل"]):
                    compliance.append({"البند":"إجازة الوضع","المادة":"151 معدلة","التوافق":"✅ مذكورة","الملاحظة":"12 أسبوع بأجر كامل (6 إلزامية بعد الولادة + 6 توزع حسب الرغبة)"})

                st.dataframe(pd.DataFrame(compliance), use_container_width=True, hide_index=True)

                # Risk score
                total_checks = len(compliance)
                passed = sum(1 for c in compliance if "✅" in c['التوافق'])
                risk_score = round(passed / max(total_checks,1) * 100)

                st.markdown("### 🎯 تقييم المخاطر")
                if risk_score >= 80:
                    ibox(f"درجة التوافق: **{risk_score}%** - العقد متوافق بشكل جيد ✅", "success")
                elif risk_score >= 60:
                    ibox(f"درجة التوافق: **{risk_score}%** - يحتاج تحسينات ⚠️", "warning")
                else:
                    ibox(f"درجة التوافق: **{risk_score}%** - يحتاج مراجعة شاملة 🔴", "warning")

                # Full text preview
                with st.expander("📄 عرض النص الكامل للعقد"):
                    st.text_area("", value=contract_text, height=400, disabled=True)

                # Export analysis
                if st.button("📥 تصدير التحليل Excel", key="ct_ana_exp"):
                    ox = io.BytesIO()
                    with pd.ExcelWriter(ox, engine='xlsxwriter') as w:
                        pd.DataFrame([{"المفتاح":k,"القيمة":v} for k,v in extracted.items()]).to_excel(w, sheet_name='البيانات المستخرجة', index=False)
                        clause_df.to_excel(w, sheet_name='تحليل البنود', index=False)
                        pd.DataFrame(compliance).to_excel(w, sheet_name='فحص التوافق', index=False)
                    st.download_button("📥 تحميل", data=ox.getvalue(),
                        file_name=f"Contract_Analysis_{datetime.now().strftime('%Y%m%d')}.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

                # AI Deep Contract Analysis
                st.markdown("---")
                if st.button("🤖 تحليل العقد بالذكاء الاصطناعي", type="primary", use_container_width=True, key="ct_ai_btn"):
                    with st.spinner("جاري تحليل العقد بالذكاء الاصطناعي..."):
                        compliance_txt = "\n".join([f"- {c['البند']}: {c['التوافق']} - {c['الملاحظة']}" for c in compliance])
                        extracted_txt = "\n".join([f"- {k}: {v}" for k,v in extracted.items()])

                        ct_ai_prompt = f"""أنت محامي متخصص في نظام العمل السعودي المعدل (المرسوم الملكي م/44، نافذ 19/2/2025).

حلل العقد التالي وقدم تقريراً قانونياً شاملاً:

**نوع العقد:** {ct_type_detected}

**البيانات المستخرجة:**
{extracted_txt}

**نتائج فحص التوافق:**
{compliance_txt}

**درجة التوافق الحالية:** {risk_score}%

**نص العقد (مختصر):**
{contract_text[:3000]}

قدم تحليلاً يشمل:

## 1. التقييم القانوني العام
(هل العقد متوافق مع نظام العمل السعودي المعدل؟ 5 أسطر)

## 2. المخاطر القانونية
(أهم 5 مخاطر قانونية في العقد مع المادة المرجعية)

## 3. البنود المفقودة الضرورية
(بنود يجب إضافتها حسب نظام العمل المعدل)

## 4. توصيات التحسين
(5 توصيات محددة لتحسين العقد)

## 5. ملاحظات على التعديلات الجديدة
(هل يتوافق العقد مع تعديلات 2025؟ خاصة: فترة التجربة 180 يوم، الإشعار 60/30 يوم، إجازة الوضع 12 أسبوع)

أجب بالعربية."""

                        response, error = call_ai_api(ct_ai_prompt, ct_ai_prompt, model_type="labor_law")
                        if response:
                            st.markdown("### 🤖 التحليل القانوني بالذكاء الاصطناعي")
                            st.markdown(response)
                        elif error:
                            st.warning(f"تعذر التحليل: {error}. تأكد من إعداد مفتاح API.")

        elif page == "📋 العقود المحفوظة":
            hdr("📋 العقود المحفوظة","عرض جميع العقود المنشأة")
            contracts = st.session_state.saved_contracts
            if not contracts:
                st.info("📜 لا توجد عقود محفوظة")
                return

            k1,k2,k3,k4 = st.columns(4)
            with k1: kpi("📜 الإجمالي", str(len(contracts)))
            types = {}
            for c in contracts: types[c['type']] = types.get(c['type'],0) + 1
            with k2: kpi("💼 عقود عمل", str(types.get('عقد عمل محدد المدة',0)))
            with k3: kpi("🎓 تدريب", str(types.get('عقد تدريب',0)))
            with k4: kpi("🤝 استشاري", str(types.get('عقد خدمات فنية واستشارية',0)))

            ct_rows = [{"#":c['id'], "النوع":c['type'], "الاسم":c.get('name',''),
                "الحالة":c.get('status',''), "أُنشئ":c.get('created_at',''),
                "بواسطة":c.get('created_by','')} for c in contracts]
            st.dataframe(pd.DataFrame(ct_rows), use_container_width=True, hide_index=True)

            if st.session_state.get('user_role') == "مدير":
                if st.button("🗑️ حذف جميع العقود", key="ct_del"):
                    st.session_state.saved_contracts = []
                    try:
                        conn = get_conn()
                        c = conn.cursor()
                        _upsert_config(c, "saved_contracts", "[]")
                        conn.commit(); conn.close()
                    except: pass
                    st.rerun()



        elif page == "📥 تصدير العقود":
            hdr("📥 تصدير العقود")
            contracts = st.session_state.saved_contracts
            if not contracts:
                st.info("لا توجد عقود"); return
            if st.button("📥 تصدير Excel", type="primary", use_container_width=True):
                ox = io.BytesIO()
                with pd.ExcelWriter(ox, engine='xlsxwriter') as w:
                    rows = [{"#":c['id'],"النوع":c['type'],"الاسم":c.get('name',''),"الحالة":c.get('status',''),"أُنشئ":c.get('created_at','')} for c in contracts]
                    pd.DataFrame(rows).to_excel(w, sheet_name='العقود', index=False)
                st.download_button("📥 تحميل", data=ox.getvalue(),
                    file_name=f"Contracts_{datetime.now().strftime('%Y%m%d')}.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")


    # =========================================
    #         🤖 AI CONSULTANT MODULE (2 LLMs)
    # =========================================
    elif section == "🤖 المستشار الذكي":

        # System prompts for the 2 AI models
        LABOR_LAW_SYSTEM_PROMPT = """⚖️ أنت المستشار القانوني لنظام العمل السعودي فقط.

❌ ممنوع: لا تذكر أي مفاهيم إدارية أو نماذج أو أطر منهجية. أجب فقط بالمواد القانونية وأرقامها.

**منهجية التحليل القانوني (اتبعها في كل إجابة):**

🔍 **الخطوة 1 - تحديد المسألة:** صنّف السؤال (عقود/إنهاء/تعويض/تدريب/إجازات/أجور/تأمينات)
📋 **الخطوة 2 - الأساس النظامي:** اسرد كل المواد ذات العلاقة بأرقامها ونصوصها
⚖️ **الخطوة 3 - التحليل:** طبّق المواد على الحالة مع حسابات رقمية إن لزم
👤 **الخطوة 4 - الحقوق والالتزامات:** لكلا الطرفين (عامل وصاحب عمل)
💡 **الخطوة 5 - التوصية:** نصيحة عملية + إجراءات + مدد زمنية + جهات مختصة
⚠️ **الخطوة 6 - التحذيرات:** مخالفات محتملة وعقوبات وتبعات قانونية

أبواب نظام العمل:
الباب 4 (م42-48): التدريب | الباب 5 (م49-60): العقود | الباب 7 (م74-82): إنهاء العقد
الباب 8 (م83-88): المكافأة | الباب 9 (م89-97): الأجور | الباب 10 (م98-107): ساعات العمل
الباب 11 (م108-116): الإجازات | الباب 13 (م121-130): المرأة | الباب 15 (م146-155): إصابات

المواد: م42-48 التدريب | م50-55 العقود | م74-82 الإنهاء | م84-88 المكافأة | م89-107 الأجور والساعات | م109-115 الإجازات | م151 الوضع
GOSI: سعودي 10.5%+12.5% | غير سعودي 2% | ساند 60%+50% أقصى 9000 | تقاعد سن60+120شهر

أجب بالعربية. اتبع الخطوات الست في كل إجابة."""

        HR_EXPERT_SYSTEM_PROMPT = """You are a world-class HR professional consultant. Your answers MUST be based on these 7 certification frameworks ONLY (NOT labor law - that belongs to the legal consultant):

**المناهج المعتمدة السبعة (يجب الإشارة إليها في كل إجابة):**

**1. PHRi (Professional in Human Resources - International) by HRCI:**
- Talent Acquisition: Job analysis, workforce planning, recruitment, selection, onboarding
- HR Administration: HRIS, policies, compliance
- Talent Management: Training needs analysis, L&D, career development, succession planning
- Compensation & Benefits: Total rewards, job evaluation, pay structures
- Employee Relations: Engagement, conflict resolution, disciplinary procedures, H&S
- HR Analytics: Metrics, reporting, data-driven decisions

**2. aPHRi (Associate Professional in HR - International) by HRCI:**
- HR Operations: Day-to-day HR functions and administration
- Recruitment & Selection: Basic hiring processes
- Compensation & Benefits: Payroll basics, benefits enrollment
- HR Development: Orientation, basic training coordination
- Employee Relations: Basic employee support and communication

**3. SPHRi (Senior Professional in HR - International) by HRCI:**
- Strategic HR Management: HR as business partner
- Workforce Planning: Long-term talent strategies
- HR Development: Organizational-level learning strategy
- Total Rewards: Strategic compensation design
- Employee & Labor Relations: Complex cases
- Risk Management: Compliance, auditing

**4. SHRM (Society for Human Resource Management):**
- SHRM-BoCK: Behavioral + Technical competencies
- Leadership & Navigation, Ethical Practice, Business Acumen
- Relationship Management, Consultation, Critical Evaluation
- People, Organization, Workplace knowledge domains

**5. CIPD Level 5 (Associate):**
- People Management: Core people practices
- Organizational Performance & Culture
- Evidence-Based Practice
- Learning & Development Practice
- Reward Management Fundamentals

**6. CIPD Level 7 (Advanced/Chartered):**
- Strategic People Management & Development
- Organizational Design & Development
- Strategic Reward Management
- Leading, Managing & Developing People
- Designing High-Performance Work Systems

**7. APTD (Associate Professional in Talent Development) by ATD:**
- Instructional Design: ADDIE, SAM models
- Training Delivery: Facilitation techniques
- Learning Technologies: LMS, e-learning, microlearning
- Performance Improvement: HPT, Kirkpatrick, Phillips ROI
- Knowledge Management & Change Management
- Coaching & Mentoring

**⚠️ ممنوع ذكر أي مادة قانونية أو رقم مادة. أجب فقط بالمفاهيم العلمية والأطر المنهجية.**

**منهجية التحليل المهني (اتبعها في كل إجابة):**

🎯 **الخطوة 1 - تحديد المجال:** صنّف السؤال (تدريب/استقطاب/أداء/تعويضات/تطوير مؤسسي/قيادة/تحليلات)

📚 **الخطوة 2 - الإطار المرجعي:** حدد المناهج المعتمدة ذات العلاقة:
اذكر بوضوح: **(PHRi - المجال)** أو **(CIPD L7 - المجال)** أو **(APTD - المجال)**

🔬 **الخطوة 3 - التحليل العلمي:** طبّق النماذج والأطر المنهجية على الحالة:
- ما هو النموذج المناسب؟ (ADDIE/Kirkpatrick/Phillips/Ulrich/Burke-Litwin...)
- كيف يُطبّق عملياً؟
- ما هي المراحل والخطوات؟

📊 **الخطوة 4 - مؤشرات القياس (KPIs):** اذكر المؤشرات القابلة للقياس لتقييم النجاح

🛠️ **الخطوة 5 - التطبيق العملي:** خطوات تنفيذية واضحة + أدوات مقترحة + جدول زمني

🌍 **الخطوة 6 - أفضل الممارسات العالمية:** أمثلة من شركات عالمية أو دراسات حديثة

**النماذج والأطر:**
- ADDIE, SAM (APTD) | Kirkpatrick, Phillips ROI (APTD/SPHRi)
- 70-20-10 (CIPD) | Balanced Scorecard (SPHRi) | Competency-Based HRM (SHRM)
- Total Rewards (SHRM/SPHRi) | Ulrich Model (CIPD L7) | 9-Box Grid (PHRi/SPHRi)
- Burke-Litwin (CIPD L7) | Kotter 8-Step (CIPD L7) | HPT (APTD)
- Bloom's Taxonomy (APTD) | Gagné's 9 Events (APTD) | ARCS Model (APTD)
- 360-degree, MBO, BARS | Gallup Q12, eNPS | ADKAR, Lewin's Model
- HR Analytics: Descriptive→Diagnostic→Predictive→Prescriptive"""

        def chunk_text(text, chunk_size=500, overlap=50):
            return st.session_state._knowledge_engine._chunk_text(text, chunk_size, overlap)

        def search_knowledge_base(query, top_k=5, advisor_type=None):
            return st.session_state._knowledge_engine.search(query, top_k, advisor_type)

        def save_to_knowledge_base(text, source, doc_type="legal", advisor_type=None):
            return st.session_state._knowledge_engine.ingest(text, source, doc_type, advisor_type=advisor_type or doc_type)

        def save_qa_pair(question, answer, model_type, feedback=None):
            st.session_state._learning_system.save_interaction(question, answer, model_type, feedback)

        def get_learned_context(query, model_type, top_k=3):
            return st.session_state._learning_system.get_relevant_history(query, model_type, top_k)

        def call_ai_api(system_prompt, user_message, chat_history=None, model_type="general", provider=None):
            try:
                if '_orchestrator' not in st.session_state:
                    st.session_state._orchestrator = _init_orchestrator()
                    st.session_state._knowledge_engine = _init_knowledge()
                    st.session_state._learning_system = _init_learning()
                return st.session_state._orchestrator.call(system_prompt, user_message, chat_history, model_type)
            except Exception as e:
                return None, f"⚠️ يرجى المحاولة مرة أخرى"

        def call_claude_api(system_prompt, user_message, chat_history=None, model_type="general"):
            return call_ai_api(system_prompt, user_message, chat_history, model_type)

        # API Keys check + auto-load (single check per session)
        if '_ai_keys_loaded' not in st.session_state:
            st.session_state._ai_keys_loaded = True
            try: st.session_state.claude_api_key = st.secrets.get("anthropic", {}).get("api_key", "")
            except: st.session_state.setdefault('claude_api_key', '')
            try: st.session_state.groq_api_key = st.secrets.get("groq", {}).get("api_key", "")
            except: st.session_state.setdefault('groq_api_key', '')
            try: st.session_state.openrouter_api_key = st.secrets.get("openrouter", {}).get("api_key", "")
            except: st.session_state.setdefault('openrouter_api_key', '')
            try: st.session_state.huggingface_api_key = st.secrets.get("huggingface", {}).get("api_key", "")
            except: st.session_state.setdefault('huggingface_api_key', '')
            try: st.session_state.gemini_api_key = st.secrets.get("gemini", {}).get("api_key", "")
            except: st.session_state.setdefault('gemini_api_key', '')
            try:
                conn = get_conn(); c = conn.cursor()
                for pk in ['claude_api_key','groq_api_key','openrouter_api_key','huggingface_api_key','gemini_api_key']:
                    if not st.session_state.get(pk):
                        c.execute(f"SELECT value FROM app_config WHERE key = {_ph()}", (pk,))
                        row = c.fetchone()
                        if row: st.session_state[pk] = row[0]
                legal_ctx = ""
                for dk in ["labor_law","labor_regulations","social_insurance","health_insurance","minister_decisions"]:
                    c.execute(f"SELECT value FROM app_config WHERE key = {_ph()}", (f"legal_doc_{dk}",))
                    r2 = c.fetchone()
                    if r2: legal_ctx += r2[0][:5000] + "\n"
                conn.close()
                if legal_ctx: st.session_state['legal_docs_context'] = legal_ctx
            except: pass

        if 'ai_provider' not in st.session_state:
            st.session_state.ai_provider = 'auto'

        has_claude = bool(st.session_state.get('claude_api_key'))
        has_groq = bool(st.session_state.get('groq_api_key'))
        has_or = bool(st.session_state.get('openrouter_api_key'))
        has_hf = bool(st.session_state.get('huggingface_api_key'))
        has_gemini = bool(st.session_state.get('gemini_api_key'))

        if not has_claude and not has_groq and not has_or and not has_hf and not has_gemini:
            st.warning("⚠️ يرجى إدخال API Key واحد على الأقل")
            kc1, kc2 = st.columns(2)
            with kc1:
                st.markdown("### ✨ Google Gemini (مجاني + الأفضل)")
                gm_input = st.text_input("🔑 Gemini API Key:", type="password", key="gm_key_input",
                    help="مجاني من https://aistudio.google.com/apikey")
                if gm_input:
                    st.session_state.gemini_api_key = gm_input; st.rerun()
                st.caption("Gemini 2.0 Flash مجاني 60 طلب/دقيقة")
            with kc2:
                st.markdown("### 🟠 OpenRouter (مجاني)")
                or_input = st.text_input("🔑 OpenRouter Key:", type="password", key="or_key_input",
                    help="مجاني من https://openrouter.ai/")
                if or_input:
                    st.session_state.openrouter_api_key = or_input; st.rerun()
                st.caption("Llama 3.3 70B مجاني")
            with st.expander("🔧 مزودين إضافيين"):
                ec1, ec2, ec3 = st.columns(3)
                with ec1:
                    groq_input = st.text_input("🟢 Groq:", type="password", key="groq_key_input")
                    if groq_input: st.session_state.groq_api_key = groq_input; st.rerun()
                with ec2:
                    hf_input = st.text_input("🤗 HuggingFace:", type="password", key="hf_key_input")
                    if hf_input: st.session_state.huggingface_api_key = hf_input; st.rerun()
                with ec3:
                    claude_input = st.text_input("🔵 Claude:", type="password", key="claude_key_input")
                    if claude_input: st.session_state.claude_api_key = claude_input; st.rerun()
            st.info("💡 أسهل طريقة: احصل على مفتاح Gemini مجاناً من **aistudio.google.com/apikey**\n\nثم أضف في Secrets:\n```\n[gemini]\napi_key = \"AIza...\"\n```")
            return

        # Provider selector
        available_providers = ["🔄 تلقائي"]
        if has_gemini: available_providers.append("✨ Gemini")
        if has_groq: available_providers.append("🟢 Groq")
        if has_or: available_providers.append("🟠 OpenRouter")
        if has_hf: available_providers.append("🤗 HuggingFace")
        if has_claude: available_providers.append("🔵 Claude")
        sel_provider = st.radio("🤖 المحرك:", available_providers, horizontal=True, key="provider_sel")
        if "Gemini" in sel_provider: st.session_state.ai_provider = 'gemini'
        elif "Groq" in sel_provider: st.session_state.ai_provider = 'groq'
        elif "HuggingFace" in sel_provider: st.session_state.ai_provider = 'huggingface'
        elif "OpenRouter" in sel_provider: st.session_state.ai_provider = 'openrouter'
        elif "Claude" in sel_provider: st.session_state.ai_provider = 'claude'
        else: st.session_state.ai_provider = 'auto'

        # ===== MODEL 1: Labor Law Consultant =====
        if page == "⚖️ المستشار القانوني":
            hdr("⚖️ المستشار القانوني بالذكاء الاصطناعي",
                "مدعوم بنظام العمل السعودي + التأمينات + الضمان الصحي + قرارات وزارة الموارد البشرية")

            st.markdown("""
            **المصادر المعتمدة:**
            نظام العمل (245 مادة) | اللائحة التنفيذية | نظام التأمينات الاجتماعية (GOSI) | نظام الضمان الصحي (CCHI) | قرارات وزير الموارد البشرية
            """)

            # Chat history
            if 'labor_chat' not in st.session_state:
                st.session_state.labor_chat = []

            # Display chat history with feedback
            for idx, msg in enumerate(st.session_state.labor_chat):
                if msg['role'] == 'user':
                    st.markdown(f"<div style='background:#1e3a5f;color:white;padding:12px;border-radius:10px;margin:8px 0'>👤 {msg['content']}</div>", unsafe_allow_html=True)
                else:
                    st.markdown(f"<div style='background:#f0f4f8;color:#333;padding:12px;border-radius:10px;margin:8px 0;border-right:4px solid #0F4C5C'>⚖️ {msg['content']}</div>", unsafe_allow_html=True)
                    fc1, fc2, fc3 = st.columns([1,1,8])
                    q_text = st.session_state.labor_chat[idx-1]['content'] if idx > 0 else ""
                    q_hash = hashlib.md5(q_text.encode()).hexdigest()[:10] if q_text else ""
                    with fc1:
                        if st.button("👍", key=f"lg_{idx}", help="إجابة مفيدة"):
                            rate_answer(q_hash, 1, "legal")
                            st.toast("✅ شكراً لتقييمك")
                    with fc2:
                        if st.button("👎", key=f"lb_{idx}", help="إجابة غير دقيقة"):
                            rate_answer(q_hash, -1, "legal")
                            st.toast("📝 سيتم تحسين الإجابة")

            # Instant answers database
            INSTANT_ANSWERS = {
                "ما هي حقوقي عند الفصل وفق المادة 77؟": "**المادة 77 من نظام العمل السعودي:**\n\nإذا أُنهي العقد لسبب غير مشروع، يحق للطرف المتضرر تعويض:\n\n**عقد غير محدد المدة:** أجر 15 يوماً عن كل سنة خدمة\n**عقد محدد المدة:** أجر المدة الباقية من العقد\n**الحد الأدنى:** لا يقل التعويض عن أجر شهرين في كلا الحالتين\n\n**بالإضافة إلى:**\n- مكافأة نهاية الخدمة وفق المادة 84\n- بدل الإجازات غير المستخدمة\n- أجر الشهر الأخير كاملاً\n- شهادة خبرة\n\n**المدة:** يلتزم صاحب العمل بالتسوية خلال أسبوع من تاريخ انتهاء العلاقة (المادة 88).",
                "متى يتحول العقد المحدد لغير محدد المدة؟": "**المادة 55 من نظام العمل السعودي:**\n\nيتحول العقد محدد المدة إلى غير محدد في الحالات التالية:\n\n1. **التجديد 3 مرات متتالية** أو بلوغ مدة العقد الأصلي مع التجديد 4 سنوات، أيهما أقل\n2. **استمرار الطرفين** في تنفيذ العقد بعد انتهاء مدته دون تجديد مكتوب\n3. **اتفاق الطرفين** على تحويله\n\nيترتب على التحول أن العامل يستفيد من أحكام العقد غير المحدد في الإشعار والتعويض.",
                "كيف تُحسب مكافأة نهاية الخدمة؟": "**المادة 84 من نظام العمل السعودي:**\n\n**حساب المكافأة:**\n- **أول 5 سنوات:** نصف راتب شهري عن كل سنة\n- **ما بعد 5 سنوات:** راتب شهري كامل عن كل سنة\n- **أجزاء السنة:** تُحسب نسبياً\n\n**عند الاستقالة (المادة 85):**\n- أقل من سنتين: لا يستحق شيئاً\n- من 2 إلى 5 سنوات: ثلث المكافأة\n- من 5 إلى 10 سنوات: ثلثا المكافأة\n- أكثر من 10 سنوات: المكافأة كاملة\n\n**التسوية:** خلال أسبوع من انتهاء العلاقة (المادة 88).",
                "ما هي نسبة اشتراكات التأمينات الاجتماعية؟": "**نسب الاشتراك في التأمينات الاجتماعية (GOSI):**\n\n**السعوديون:**\n- المعاشات: 9.75% على الموظف + 9.75% على صاحب العمل\n- الأخطار المهنية: 2% على صاحب العمل\n- ساند (التعطل عن العمل): 0.75% + 0.75%\n- **إجمالي خصم الموظف السعودي: 10.5%**\n- **إجمالي تحمّل الشركة: 12.5%**\n\n**غير السعوديين:**\n- الأخطار المهنية فقط: 2% على صاحب العمل\n- لا يوجد خصم من راتب الموظف غير السعودي\n\n**الراتب الخاضع:** الراتب الأساسي + بدل السكن.",
                "ما هي فترة التجربة وشروطها؟": "**المادة 53 من نظام العمل السعودي:**\n\n- **المدة القصوى:** 90 يوماً\n- **التمديد:** يجوز تمديدها إلى 180 يوماً بموافقة مكتوبة من الطرفين\n- **لا تدخل فيها:** إجازتا عيد الفطر والأضحى والإجازات المرضية\n- **حق الإنهاء:** لكلا الطرفين إنهاء العقد خلالها بدون تعويض أو مكافأة أو إشعار مسبق\n- **عدم التكرار:** لا يجوز وضع العامل تحت التجربة أكثر من مرة لدى نفس صاحب العمل إلا بموافقة مكتوبة وفي مهنة مختلفة.",
                "ما هي حقوق المرأة العاملة في نظام العمل؟": "**حقوق المرأة العاملة في نظام العمل السعودي:**\n\n- **إجازة الوضع:** 10 أسابيع توزعها كيف تشاء، بحد أقصى 4 أسابيع قبل التاريخ المرجح (المادة 151)\n- **ساعة الرضاعة:** ساعة مدفوعة يومياً لمدة 24 شهراً بعد الوضع\n- **الحماية من الفصل:** يُحظر فصلها أثناء الحمل وإجازة الوضع وخلال 180 يوماً بعد الوضع\n- **إجازة وفاة الزوج:** 4 أشهر و10 أيام بأجر كامل (عدة)\n- **المساواة في الأجر:** أجر متساوٍ للعمل المتساوي في القيمة\n- **بيئة العمل:** توفير بيئة آمنة ومناسبة وفق اشتراطات وزارة الموارد البشرية.",
            }

            # Quick buttons - inject answer DIRECTLY (no form needed)
            st.markdown("### 💡 أسئلة شائعة")
            qc1, qc2, qc3 = st.columns(3)
            for i, q in enumerate(INSTANT_ANSWERS.keys()):
                with [qc1, qc2, qc3][i % 3]:
                    if st.button(q, key=f"lq_{i}", use_container_width=True):
                        st.session_state.labor_chat = [{"role":"user","content":q},{"role":"assistant","content":INSTANT_ANSWERS[q]}]
                        st.rerun()

            # Input for custom questions
            with st.form("labor_form", clear_on_submit=True):
                labor_q = st.text_area("اكتب سؤالك القانوني:", height=80, key="labor_q_input",
                    placeholder="مثال: تم فصلي بعد 3 سنوات خدمة بدون سبب، ما مستحقاتي؟")
                submitted = st.form_submit_button("⚖️ استشارة", type="primary", use_container_width=True)

            if submitted and labor_q:
                st.session_state.labor_chat = []
                with st.spinner("جاري التحليل القانوني..."):
                    response = get_best_kb_answer(labor_q, LABOR_LAW_SYSTEM_PROMPT)
                    auto_learn_from_answer(labor_q, response, "legal")
                    st.session_state.labor_chat = [{"role":"user","content":labor_q},{"role":"assistant","content":response}]
                    st.rerun()

            # Clear chat
            if st.session_state.labor_chat and st.button("🗑️ مسح المحادثة", key="labor_clear"):
                st.session_state.labor_chat = []
                st.rerun()

        # ===== MODEL 2: HR Expert =====
        elif page == "📚 مستشار الموارد البشرية":
            hdr("📚 مستشار الموارد البشرية بالذكاء الاصطناعي",
                "مدعوم بمناهج PHRi + aPHRi + SPHRi + SHRM + CIPD 5 + CIPD 7 + APTD")

            st.markdown("""
            **المناهج المعتمدة:**
            PHRi (HRCI) | aPHRi (HRCI) | SPHRi (HRCI) | SHRM-SCP | CIPD Level 5 | CIPD Level 7 | APTD (ATD)

            **ملاحظة:** هذا المستشار متخصص في المفاهيم العلمية والأطر المنهجية لإدارة الموارد البشرية. للاستشارات القانونية ونظام العمل، استخدم **المستشار القانوني**.
            """)

            # Chat history
            if 'hr_chat' not in st.session_state:
                st.session_state.hr_chat = []

            for idx, msg in enumerate(st.session_state.hr_chat):
                if msg['role'] == 'user':
                    st.markdown(f"<div style='background:#1e3a5f;color:white;padding:12px;border-radius:10px;margin:8px 0'>👤 {msg['content']}</div>", unsafe_allow_html=True)
                else:
                    st.markdown(f"<div style='background:#f0faf5;color:#333;padding:12px;border-radius:10px;margin:8px 0;border-right:4px solid #2A9D8F'>📚 {msg['content']}</div>", unsafe_allow_html=True)
                    fc1, fc2, fc3 = st.columns([1,1,8])
                    q_text = st.session_state.hr_chat[idx-1]['content'] if idx > 0 else ""
                    q_hash = hashlib.md5(q_text.encode()).hexdigest()[:10] if q_text else ""
                    with fc1:
                        if st.button("👍", key=f"hg_{idx}", help="إجابة مفيدة"):
                            rate_answer(q_hash, 1, "hr")
                            st.toast("✅ شكراً لتقييمك")
                    with fc2:
                        if st.button("👎", key=f"hb_{idx}", help="إجابة غير دقيقة"):
                            rate_answer(q_hash, -1, "hr")
                            st.toast("📝 سيتم تحسين الإجابة")

            # Instant HR answers database
            HR_INSTANT = {
                "كيف أبني خطة استقطاب فعالة؟": "**خطة استقطاب فعالة** *(PHRi - Talent Acquisition | SHRM - People)*\n\n**1. التحليل** *(PHRi - Workforce Planning):*\n- تحديد الاحتياج عبر Workforce Planning\n- تحليل سوق العمل والرواتب المرجعية\n\n**2. التصميم** *(CIPD L5 - Resourcing):*\n- وصف وظيفي محدد الكفاءات\n- بناء Employee Value Proposition (EVP)\n- تحديد القنوات: LinkedIn, مواقع توظيف, تزكيات\n\n**3. التنفيذ** *(SHRM - Talent Acquisition):*\n- ATS لتتبع المتقدمين\n- مقابلات منظمة (Structured Interviews)\n- تقييم الكفاءات (Competency-Based Assessment)\n\n**4. القياس** *(SPHRi - HR Metrics):*\n- Time-to-Hire | Cost-per-Hire | Quality of Hire | Source Effectiveness",
                "ما الفرق بين OKRs و KPIs؟": "**OKRs vs KPIs** *(SPHRi - Strategic HR | CIPD L7 - Performance Management)*\n\n**KPIs** *(PHRi - HR Analytics):*\n- تقيس الأداء المستمر\n- أرقام محددة مثل: معدل دوران 15%\n- ثابتة نسبياً\n\n**OKRs** *(SHRM - Business Acumen):*\n- أهداف طموحة للمستقبل\n- هدف + 3-5 نتائج قابلة للقياس\n- تتغير كل ربع سنة\n- 70% إنجاز = ممتاز\n\n**مثال HR** *(SPHRi):*\n- **KPI:** معدل الدوران = 12%\n- **OKR:** خفض الدوران إلى 10% + رفع الرضا إلى 85%",
                "ما هو نموذج Phillips ROI للتدريب؟": "**نموذج Phillips ROI** *(APTD - Performance Improvement | SPHRi - Training Evaluation)*\n\n**المستوى 1 - Reaction** *(aPHRi):* رضا المتدربين\n**المستوى 2 - Learning** *(PHRi):* المعرفة المكتسبة\n**المستوى 3 - Application** *(CIPD L5):* التطبيق في العمل\n**المستوى 4 - Impact** *(SPHRi):* التأثير على الأعمال\n**المستوى 5 - ROI** *(SPHRi):* العائد على الاستثمار\n\nROI % = (الفوائد - التكاليف) / التكاليف × 100\n\n**مثال:** تدريب 50,000 ريال → إنتاجية 150,000 ريال\nROI = **200%**",
                "كيف أصمم هيكل رواتب تنافسي؟": "**هيكل الرواتب** *(PHRi - Compensation | SPHRi - Total Rewards | CIPD L7 - Reward Management)*\n\n**1. المسح السوقي** *(SPHRi):* Salary Survey + P25/P50/P75\n**2. تقييم الوظائف** *(PHRi):* Job Evaluation + Job Grading\n**3. بناء الهيكل** *(CIPD L7):* نطاق Min-Mid-Max لكل درجة\n**4. Compa-Ratio** *(PHRi):* الراتب الفعلي / وسط النطاق\n**5. السياسات** *(SHRM):* توظيف عند Min-Mid، زيادات 3-5%، مكافآت 10-20%",
                "كيف أحسب معدل دوران الموظفين وأحسّنه؟": "**معدل الدوران** *(PHRi - HR Analytics | SPHRi - Workforce Planning)*\n\n**المعادلة** *(PHRi):*\nTurnover = (المغادرين / متوسط العدد) × 100\n\n**المعايير** *(SHRM Benchmarks):*\nأقل من 10% ممتاز | 10-15% جيد | أكثر من 25% مشكلة\n\n**التحسين** *(CIPD L7 - Retention | SPHRi):*\n1. رواتب تنافسية *(Total Rewards - SHRM)*\n2. مسار وظيفي *(Career Development - PHRi)*\n3. بيئة محفزة *(Employee Engagement - CIPD L5)*\n4. تقدير ومكافآت *(Recognition - SHRM)*\n5. تطوير القيادات *(Leadership - CIPD L7)*",
                "ما هي أفضل ممارسات تجربة الموظف (EX)؟": "**تجربة الموظف** *(SHRM - Employee Experience | CIPD L7 - People Management)*\n\n1. **الاستقطاب** *(PHRi - Talent Acquisition):* عملية سلسة وشفافة\n2. **التهيئة** *(aPHRi - Onboarding):* خطة 30/60/90 يوم + Buddy\n3. **التطوير** *(CIPD L5 - L&D):* تدريب مستمر + مسار وظيفي\n4. **الاحتفاظ** *(SPHRi - Retention):* تقدير + Total Rewards + مرونة\n5. **الانتقال** *(PHRi - Offboarding):* مقابلة خروج + شهادة خبرة + Alumni",
            }

            # Quick buttons - inject answer DIRECTLY
            st.markdown("### 💡 مواضيع شائعة")
            tc1, tc2, tc3 = st.columns(3)
            for i, t in enumerate(HR_INSTANT.keys()):
                with [tc1, tc2, tc3][i % 3]:
                    if st.button(t, key=f"ht_{i}", use_container_width=True):
                        st.session_state.hr_chat = [{"role":"user","content":t},{"role":"assistant","content":HR_INSTANT[t]}]
                        st.rerun()

            with st.form("hr_form", clear_on_submit=True):
                hr_q = st.text_area("اكتب سؤالك:", height=80, key="hr_q_input",
                    placeholder="مثال: كيف أقيس فعالية برنامج التدريب باستخدام نموذج Kirkpatrick؟")
                submitted = st.form_submit_button("📚 استشارة", type="primary", use_container_width=True)

            if submitted and hr_q:
                st.session_state.hr_chat = []
                with st.spinner("جاري التحليل المهني..."):
                    response = get_best_kb_answer(hr_q, HR_EXPERT_SYSTEM_PROMPT)
                    auto_learn_from_answer(hr_q, response, "hr")
                    st.session_state.hr_chat = [{"role":"user","content":hr_q},{"role":"assistant","content":response}]
                    st.rerun()

            if st.session_state.hr_chat and st.button("🗑️ مسح المحادثة", key="hr_clear"):
                st.session_state.hr_chat = []
                st.rerun()

        # ===== Legal Documents Management =====
        elif page == "🧠 قاعدة المعرفة RAG":
            hdr("🧠 قاعدة المعرفة RAG","Retrieval Augmented Generation - تغذية النموذج بمستنداتك")

            st.markdown("""
            **كيف يعمل RAG؟** ترفع مستنداتك → تُقسّم لأجزاء صغيرة → تُفهرس وتُخزّن → عند السؤال يُبحث عن الأجزاء المناسبة → تُرسل مع سؤالك لـ Claude → إجابة مبنية على مستنداتك
            """)

            # Upload documents
            st.markdown("### 📁 رفع مستندات للقاعدة المعرفية")

            # Advisor type selector for document tagging
            rag_advisor = st.radio("📌 نوع المستند:", ["⚖️ قانوني (نظام العمل، التأمينات، اللوائح)", "📚 موارد بشرية (مناهج، أطر منهجية، ممارسات)"],
                horizontal=True, key="rag_advisor_type")
            advisor_tag = "legal" if "قانوني" in rag_advisor else "hr"

            rag_files = st.file_uploader("ارفع ملفات (PDF, DOCX, TXT):", type=["pdf","docx","txt"],
                accept_multiple_files=True, key="rag_upload")

            if rag_files:
                for rag_file in rag_files:
                    doc_text = ""
                    if rag_file.name.endswith('.pdf'):
                        try:
                            import pdfplumber
                            with pdfplumber.open(io.BytesIO(rag_file.getvalue())) as pdf_r:
                                for pg in pdf_r.pages[:200]:
                                    txt = pg.extract_text()
                                    if txt: doc_text += txt + "\n"
                        except: pass
                    elif rag_file.name.endswith('.docx'):
                        try:
                            from docx import Document as DDoc
                            doc = DDoc(io.BytesIO(rag_file.getvalue()))
                            doc_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
                        except: pass
                    elif rag_file.name.endswith('.txt'):
                        doc_text = rag_file.getvalue().decode('utf-8', errors='ignore')

                    if doc_text:
                        n_chunks = save_to_knowledge_base(doc_text, rag_file.name, advisor_tag, advisor_tag)
                        st.success(f"✅ {rag_file.name}: {len(doc_text):,} حرف → {n_chunks} جزء مفهرس")
                    else:
                        st.warning(f"⚠️ لم يتم استخراج نص من {rag_file.name}")

            # Upload from URL (web scraping)
            st.markdown("### 🌐 إضافة من رابط ويب")
            web_url = st.text_input("رابط الصفحة:", placeholder="https://laws.boe.gov.sa/...", key="rag_url")
            if st.button("🌐 استيراد", key="rag_web") and web_url:
                try:
                    import urllib.request
                    req = urllib.request.Request(web_url, headers={'User-Agent':'Mozilla/5.0'})
                    with urllib.request.urlopen(req, timeout=15) as resp:
                        html = resp.read().decode('utf-8', errors='ignore')
                    # Simple HTML to text
                    import re
                    text = re.sub(r'<[^>]+>', ' ', html)
                    text = re.sub(r'\s+', ' ', text).strip()
                    if len(text) > 100:
                        n = save_to_knowledge_base(text, web_url[:80], advisor_tag, advisor_tag)
                        st.success(f"✅ تم استيراد {len(text):,} حرف → {n} جزء")
                    else:
                        st.warning("لم يتم استخراج محتوى كافٍ")
                except Exception as e:
                    st.error(f"خطأ: {e}")

            # Manual knowledge entry
            st.markdown("### ✏️ إضافة معرفة يدوياً")
            manual_source = st.text_input("اسم المصدر:", placeholder="مثال: سياسة الإجازات الداخلية", key="rag_msrc")
            manual_text = st.text_area("النص:", height=150, key="rag_mtxt", placeholder="الصق نص السياسة أو المعلومة هنا...")
            if st.button("➕ إضافة للقاعدة", type="primary", key="rag_madd") and manual_text and manual_source:
                n = save_to_knowledge_base(manual_text, manual_source, "manual", advisor_tag)
                st.success(f"✅ {manual_source}: {n} جزء مفهرس")

            # Knowledge base stats
            st.markdown("---")
            st.markdown("### 📊 إحصائيات القاعدة المعرفية")
            try:
                conn = get_conn()
                c = conn.cursor()
                c.execute(f"SELECT value FROM app_config WHERE key = {_ph()}", ("rag_chunks",))
                row = c.fetchone()
                conn.close()
                if row:
                    chunks = json.loads(row[0])
                    sources = {}
                    for ch in chunks:
                        src = ch.get('source','unknown')
                        sources[src] = sources.get(src, 0) + 1

                    k1,k2,k3 = st.columns(3)
                    with k1: kpi("📚 إجمالي الأجزاء", str(len(chunks)))
                    with k2: kpi("📄 المصادر", str(len(sources)))
                    with k3: kpi("📝 الكلمات (تقريباً)", f"{sum(len(ch.get('text','').split()) for ch in chunks):,}")

                    src_df = pd.DataFrame([{"المصدر":k,"الأجزاء":v,"النوع": chunks[0].get('type','') if chunks else ''} for k,v in sources.items()])
                    st.dataframe(src_df, use_container_width=True, hide_index=True)

                    if st.session_state.get('user_role') == "مدير":
                        if st.button("🗑️ مسح القاعدة المعرفية بالكامل", key="rag_clear"):
                            conn = get_conn()
                            c = conn.cursor()
                            _upsert_config(c, "rag_chunks", "[]")
                            conn.commit(); conn.close()
                            st.success("✅ تم مسح القاعدة"); st.rerun()
                else:
                    st.info("📚 القاعدة المعرفية فارغة. ارفع مستندات أو أضف معرفة يدوياً.")
            except: st.info("📚 القاعدة المعرفية فارغة")

            # Quick load legal documents
            st.markdown("---")
            st.markdown("### ⚡ تحميل سريع للأنظمة السعودية")
            legal_docs = {
                "نظام العمل السعودي": "labor_law",
                "اللائحة التنفيذية لنظام العمل": "labor_regulations",
                "نظام التأمينات الاجتماعية": "social_insurance",
                "نظام الضمان الصحي التعاوني": "health_insurance",
                "قرارات وزير الموارد البشرية": "minister_decisions",
            }
            for doc_name, doc_key in legal_docs.items():
                uploaded = st.file_uploader(f"📄 {doc_name}:", type=["pdf","docx","txt"], key=f"ql_{doc_key}")
                if uploaded:
                    doc_text = ""
                    if uploaded.name.endswith('.pdf'):
                        try:
                            import pdfplumber
                            with pdfplumber.open(io.BytesIO(uploaded.getvalue())) as pr:
                                for pg in pr.pages[:200]:
                                    txt = pg.extract_text()
                                    if txt: doc_text += txt + "\n"
                        except: pass
                    elif uploaded.name.endswith('.docx'):
                        try:
                            from docx import Document as DD
                            d = DD(io.BytesIO(uploaded.getvalue()))
                            doc_text = "\n".join([p.text for p in d.paragraphs if p.text.strip()])
                        except: pass
                    else:
                        doc_text = uploaded.getvalue().decode('utf-8', errors='ignore')
                    if doc_text:
                        n = save_to_knowledge_base(doc_text, doc_name, "legal", "legal")
                        st.success(f"✅ {doc_name}: {n} جزء")

        elif page == "📊 التعلم والتحسين":
            hdr("📊 التعلم المستمر والتحسين","نظام تعلم ذاتي يتحسن مع كل سؤال")

            # RAG Learning Stats
            stats = get_learning_stats()
            st.markdown("### 🧠 إحصائيات التعلم الذاتي")
            k1,k2,k3,k4,k5 = st.columns(5)
            with k1: kpi("📦 إجمالي المعرفة المكتسبة", str(stats['total']))
            with k2: kpi("⚖️ معرفة قانونية", str(stats['legal']))
            with k3: kpi("📚 معرفة HR", str(stats['hr']))
            with k4: kpi("👍 مُقيّمة إيجابياً", str(stats['good_rated']))
            with k5: kpi("🔄 تم تحسينها", str(stats['improved']))

            if stats['total'] > 0:
                st.progress(min(stats['avg_score'] / 5 + 0.5, 1.0), text=f"متوسط جودة الإجابات: {stats['avg_score']:.1f}/5")

            # Learning mechanism explanation
            with st.expander("🔍 كيف يتعلم النظام؟"):
                st.markdown("""
                **آلية التعلم المستمر:**

                1. **التعلم من الإجابات:** كل إجابة جديدة تُحفظ مع تصنيفها (قانوني/HR)
                2. **التقييم:** أزرار 👍/👎 ترفع أو تخفض جودة الإجابة
                3. **التحسين التلقائي:** الإجابات المحسّنة تحل محل القديمة
                4. **الحذف الذاتي:** الإجابات التي تحصل على -2 تُحذف تلقائياً
                5. **الفصل:** الإجابات القانونية لا تظهر في مستشار HR والعكس
                """)

            # Show learned Q&A pairs
            st.markdown("### 📚 المعرفة المكتسبة")
            try:
                conn = get_conn(); c = conn.cursor()
                c.execute(f"SELECT value FROM app_config WHERE key = {_ph()}", ("rag_learned",))
                row = c.fetchone(); conn.close()
                if row:
                    learned = json.loads(row[0])
                    if learned:
                        tab1, tab2 = st.tabs(["⚖️ قانوني", "📚 HR"])
                        with tab1:
                            legal_items = [l for l in learned if l.get('type') == 'legal']
                            for item in legal_items[-10:]:
                                score = item.get('score', 0)
                                score_icon = "🟢" if score > 0 else ("🔴" if score < 0 else "⚪")
                                with st.expander(f"{score_icon} {item.get('q','')[:60]}"):
                                    st.markdown(item.get('a','')[:500])
                                    st.caption(f"تاريخ: {item.get('date','')} | تقييم: {score} | تحسينات: {item.get('improvements',0)}")
                            if not legal_items:
                                st.info("لم تُكتسب معرفة قانونية بعد")
                        with tab2:
                            hr_items = [l for l in learned if l.get('type') == 'hr']
                            for item in hr_items[-10:]:
                                score = item.get('score', 0)
                                score_icon = "🟢" if score > 0 else ("🔴" if score < 0 else "⚪")
                                with st.expander(f"{score_icon} {item.get('q','')[:60]}"):
                                    st.markdown(item.get('a','')[:500])
                                    st.caption(f"تاريخ: {item.get('date','')} | تقييم: {score} | تحسينات: {item.get('improvements',0)}")
                            if not hr_items:
                                st.info("لم تُكتسب معرفة HR بعد")

                        # Cleanup button
                        if st.button("🧹 تنظيف الإجابات السيئة (تقييم أقل من 0)", key="clean_bad"):
                            clean = [l for l in learned if l.get('score', 0) >= 0]
                            removed = len(learned) - len(clean)
                            conn = get_conn(); c = conn.cursor()
                            _upsert_config(c, "rag_learned", json.dumps(clean, ensure_ascii=False))
                            conn.commit(); conn.close()
                            st.success(f"✅ تم حذف {removed} إجابة سيئة")
                            st.rerun()
            except: pass

            st.markdown("---")

            # Original Q&A history
            st.markdown("### 💬 سجل المحادثات")
            try:
                conn = get_conn(); c = conn.cursor()
                c.execute(f"SELECT value FROM app_config WHERE key = {_ph()}", ("rag_qa_history",))
                row = c.fetchone(); conn.close()
                qa_history = json.loads(row[0]) if row else []
            except:
                qa_history = []

            # KPIs
            st.markdown("### 📊 إحصائيات التعلم")
            k1,k2,k3,k4,k5 = st.columns(5)
            with k1: kpi("💬 إجمالي المحادثات", str(len(qa_history)))
            labor_qs = [h for h in qa_history if h.get('model') == 'labor']
            hr_qs = [h for h in qa_history if h.get('model') == 'hr']
            with k2: kpi("⚖️ استشارات عمالية", str(len(labor_qs)))
            with k3: kpi("📚 استشارات HR", str(len(hr_qs)))
            good = [h for h in qa_history if h.get('feedback') == 'good']
            bad = [h for h in qa_history if h.get('feedback') == 'bad']
            with k4: kpi("👍 إيجابية", str(len(good)))
            with k5: kpi("👎 سلبية", str(len(bad)))

            if qa_history:
                # Feedback rate
                total_fb = len(good) + len(bad)
                if total_fb > 0:
                    satisfaction = len(good) / total_fb * 100
                    st.progress(satisfaction / 100, text=f"معدل الرضا: {satisfaction:.0f}%")

                # Usage over time
                daily = {}
                for h in qa_history:
                    day = h.get('date','')[:10]
                    daily[day] = daily.get(day, 0) + 1
                if daily:
                    fig = px.bar(x=list(daily.keys()), y=list(daily.values()),
                        title='📈 استخدام المستشار الذكي يومياً', color_discrete_sequence=['#0F4C5C'])
                    fig.update_layout(font=dict(family="Noto Sans Arabic"), height=300)
                    st.plotly_chart(fig, use_container_width=True)

                # Topic analysis
                st.markdown("### 📋 المواضيع الأكثر تكراراً")
                all_words = " ".join([h.get('q','') for h in qa_history]).lower()
                topic_keywords = {
                    "مكافأة نهاية الخدمة": ["مكافأة","نهاية","خدمة","eos"],
                    "الفصل والتعويض": ["فصل","تعويض","77","إنهاء","فسخ"],
                    "الإجازات": ["إجازة","إجازات","سنوية","مرضية"],
                    "الرواتب والبدلات": ["راتب","رواتب","بدل","سكن","مواصلات"],
                    "التأمينات": ["تأمين","تأمينات","gosi","اجتماعية"],
                    "التوظيف والاستقطاب": ["توظيف","استقطاب","مقابلة","recruitment"],
                    "التدريب والتطوير": ["تدريب","تطوير","roi","kirkpatrick"],
                    "الأداء": ["أداء","تقييم","kpi","okr"],
                    "السعودة": ["سعودة","نطاقات","nitaqat"],
                }
                topic_counts = {}
                for topic, kws in topic_keywords.items():
                    count = sum(1 for kw in kws if kw in all_words)
                    if count > 0: topic_counts[topic] = count
                if topic_counts:
                    fig = px.bar(x=list(topic_counts.values()), y=list(topic_counts.keys()), orientation='h',
                        title='المواضيع الأكثر اهتماماً', color_discrete_sequence=['#2A9D8F'])
                    fig.update_layout(font=dict(family="Noto Sans Arabic"), height=350)
                    st.plotly_chart(fig, use_container_width=True)

                # Recent Q&A with feedback
                st.markdown("### 💬 آخر المحادثات (مع إمكانية التقييم)")
                for i, h in enumerate(reversed(qa_history[-20:])):
                    with st.expander(f"{'⚖️' if h.get('model')=='labor' else '📚'} {h.get('q','')[:80]}... | {h.get('date','')}"):
                        st.markdown(f"**السؤال:** {h.get('q','')}")
                        st.markdown(f"**الإجابة:** {h.get('a','')[:500]}...")
                        fc1, fc2 = st.columns(2)
                        with fc1:
                            if st.button("👍 إجابة جيدة", key=f"fb_good_{i}"):
                                try:
                                    idx = len(qa_history) - 1 - i
                                    qa_history[idx]['feedback'] = 'good'
                                    conn = get_conn()
                                    c = conn.cursor()
                                    _upsert_config(c, "rag_qa_history", json.dumps(qa_history, ensure_ascii=False))
                                    conn.commit(); conn.close()
                                    st.success("✅ شكراً على التقييم")
                                except: pass
                        with fc2:
                            if st.button("👎 تحتاج تحسين", key=f"fb_bad_{i}"):
                                try:
                                    idx = len(qa_history) - 1 - i
                                    qa_history[idx]['feedback'] = 'bad'
                                    conn = get_conn()
                                    c = conn.cursor()
                                    _upsert_config(c, "rag_qa_history", json.dumps(qa_history, ensure_ascii=False))
                                    conn.commit(); conn.close()
                                    st.success("✅ سنعمل على التحسين")
                                except: pass

                # Knowledge growth chart
                st.markdown("---")
                st.markdown("### 📈 نمو قاعدة المعرفة")
                try:
                    conn = get_conn()
                    c = conn.cursor()
                    c.execute(f"SELECT value FROM app_config WHERE key = {_ph()}", ("rag_chunks",))
                    row = c.fetchone()
                    conn.close()
                    chunks = json.loads(row[0]) if row else []
                    total_words = sum(len(ch.get('text','').split()) for ch in chunks)
                    total_qa = len(qa_history)

                    gc1, gc2, gc3 = st.columns(3)
                    with gc1: kpi("📚 أجزاء المعرفة", str(len(chunks)))
                    with gc2: kpi("📝 كلمات القاعدة", f"{total_words:,}")
                    with gc3: kpi("💬 تفاعلات التعلم", str(total_qa))

                    # Learning maturity
                    maturity_score = min(100, (len(chunks) * 2 + total_qa * 5 + len(good) * 10) // 10)
                    st.progress(maturity_score / 100, text=f"نضج النموذج: {maturity_score}%")
                    if maturity_score < 30:
                        ibox("🟡 المرحلة الأولى: النموذج يحتاج المزيد من المستندات والتفاعلات", "warning")
                    elif maturity_score < 70:
                        ibox("🟢 المرحلة المتوسطة: النموذج يتحسن. استمر في رفع المستندات والتقييم", "success")
                    else:
                        ibox("🔵 المرحلة المتقدمة: النموذج ناضج ويقدم إجابات عالية الجودة", "success")
                except: pass
            else:
                st.info("💬 لا توجد محادثات بعد. استخدم المستشار الذكي لبدء التعلم.")

        elif page == "📋 إدارة المراجع":
            hdr("📋 إدارة المراجع والإعدادات","API Key + تحديث الأنظمة القانونية")

            if st.session_state.get('user_role') != "مدير":
                st.warning("⚠️ إدارة المراجع متاحة للمدير فقط")
                return

            st.markdown("### 📄 رفع وتحديث المراجع القانونية")
            st.caption("ارفع ملفات PDF أو Word للأنظمة واللوائح. سيتم استخدامها كمرجع إضافي للمستشار الذكي.")

            doc_types = {
                "نظام العمل السعودي": "labor_law",
                "اللائحة التنفيذية لنظام العمل": "labor_regulations",
                "نظام التأمينات الاجتماعية": "social_insurance",
                "اللائحة التنفيذية للتأمينات": "insurance_regulations",
                "نظام الضمان الصحي التعاوني": "health_insurance",
                "قرارات وزير الموارد البشرية": "minister_decisions",
            }

            for doc_name, doc_key in doc_types.items():
                with st.expander(f"📄 {doc_name}"):
                    uploaded = st.file_uploader(f"ارفع {doc_name}:", type=["pdf","docx","txt"], key=f"legal_{doc_key}")
                    if uploaded:
                        # Extract text
                        doc_text = ""
                        if uploaded.name.endswith('.pdf'):
                            try:
                                import pdfplumber
                                with pdfplumber.open(io.BytesIO(uploaded.getvalue())) as pdf_reader:
                                    for pg in pdf_reader.pages[:100]:
                                        txt = pg.extract_text()
                                        if txt: doc_text += txt + "\n"
                            except: pass
                        elif uploaded.name.endswith('.docx'):
                            try:
                                from docx import Document as DocxDoc
                                doc = DocxDoc(io.BytesIO(uploaded.getvalue()))
                                doc_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
                            except: pass
                        elif uploaded.name.endswith('.txt'):
                            doc_text = uploaded.getvalue().decode('utf-8', errors='ignore')

                        if doc_text:
                            # Save to DB
                            try:
                                conn = get_conn()
                                c = conn.cursor()
                                _upsert_config(c, f"legal_doc_{doc_key}", doc_text[:50000])
                                _upsert_config(c, f"legal_doc_{doc_key}_date", datetime.now().strftime("%Y-%m-%d %H:%M"))
                                conn.commit()
                                conn.close()
                                st.success(f"✅ تم تحديث {doc_name} ({len(doc_text):,} حرف)")
                            except Exception as e:
                                st.error(f"خطأ: {e}")

                    # Show last update
                    try:
                        conn = get_conn()
                        c = conn.cursor()
                        c.execute(f"SELECT value FROM app_config WHERE key = {_ph()}", (f"legal_doc_{doc_key}_date",))
                        row = c.fetchone()
                        conn.close()
                        if row:
                            st.caption(f"آخر تحديث: {row[0]}")
                        else:
                            st.caption("لم يتم رفع هذا المرجع بعد")
                    except:
                        st.caption("لم يتم رفع هذا المرجع بعد")

            # Load all legal docs into context
            if st.button("🔄 تحديث سياق المستشار الذكي", type="primary", use_container_width=True, key="refresh_legal"):
                legal_context = ""
                try:
                    conn = get_conn()
                    c = conn.cursor()
                    for doc_name, doc_key in doc_types.items():
                        c.execute(f"SELECT value FROM app_config WHERE key = {_ph()}", (f"legal_doc_{doc_key}",))
                        row = c.fetchone()
                        if row:
                            legal_context += f"\n\n=== {doc_name} ===\n{row[0][:8000]}"
                    conn.close()
                    st.session_state.legal_docs_context = legal_context
                    st.success(f"✅ تم تحديث سياق المستشار ({len(legal_context):,} حرف)")
                except Exception as e:
                    st.error(f"خطأ: {e}")

            # API Key management
            st.markdown("---")
            st.markdown("### 🔑 إعدادات API Keys")
            kc1, kc2 = st.columns(2)
            with kc1:
                st.markdown("**✨ Google Gemini (مجاني + الأفضل)**")
                cur_gm = st.session_state.get('gemini_api_key','')
                st.caption(f"الحالي: {'AIza...'+cur_gm[-6:] if len(cur_gm)>8 else 'غير مُعيّن'}")
                new_gm = st.text_input("Gemini API Key:", type="password", key="new_gm_key")
                if st.button("💾 حفظ Gemini", key="save_gm"):
                    if new_gm:
                        st.session_state.gemini_api_key = new_gm
                        try:
                            conn = get_conn(); c = conn.cursor()
                            _upsert_config(c, "gemini_api_key", new_gm)
                            conn.commit(); conn.close()
                        except: pass
                        st.success("✅ تم الحفظ")
            with kc2:
                st.markdown("**🟠 OpenRouter (مجاني)**")
                cur_or = st.session_state.get('openrouter_api_key','')
                st.caption(f"الحالي: {'sk-or-...'+cur_or[-6:] if len(cur_or)>8 else 'غير مُعيّن'}")
                new_or = st.text_input("OpenRouter Key:", type="password", key="new_or_key")
                if st.button("💾 حفظ OpenRouter", key="save_or"):
                    if new_or:
                        st.session_state.openrouter_api_key = new_or
                        try:
                            conn = get_conn(); c = conn.cursor()
                            _upsert_config(c, "openrouter_api_key", new_or)
                            conn.commit(); conn.close()
                        except: pass
                        st.success("✅ تم الحفظ")
            kc3, kc4 = st.columns(2)
            with kc3:
                st.markdown("**🟢 Groq (مجاني)**")
                cur_groq = st.session_state.get('groq_api_key','')
                st.caption(f"الحالي: {'gsk_...'+cur_groq[-6:] if len(cur_groq)>8 else 'غير مُعيّن'}")
                new_groq = st.text_input("Groq API Key:", type="password", key="new_groq_key")
                if st.button("💾 حفظ Groq", key="save_groq"):
                    if new_groq:
                        st.session_state.groq_api_key = new_groq
                        try:
                            conn = get_conn(); c = conn.cursor()
                            _upsert_config(c, "groq_api_key", new_groq)
                            conn.commit(); conn.close()
                        except: pass
                        st.success("✅ تم الحفظ")
            with kc4:
                st.markdown("**🔵 Claude (مدفوع)**")
                cur_claude = st.session_state.get('claude_api_key','')
                masked_c = f"sk-ant-...{cur_claude[-8:]}" if len(cur_claude)>10 else "غير مُعيّن"
                st.caption(f"الحالي: {masked_c}")
                new_key = st.text_input("Claude API Key:", type="password", key="new_api_key")
                if st.button("💾 حفظ Claude", key="save_api"):
                    if new_key:
                        st.session_state.claude_api_key = new_key
                        try:
                            conn = get_conn(); c = conn.cursor()
                            _upsert_config(c, "claude_api_key", new_key)
                            conn.commit(); conn.close()
                        except: pass
                        st.success("✅ تم الحفظ")


    # =========================================
    #         🏗️ OD MODULE
    # =========================================
    elif section == "🏗️ التطوير المؤسسي OD":

        data = sal_snapshot if len(sal_snapshot)>0 else emp
        dept_col_od = next((c for c in data.columns if any(x in c.lower() for x in ['dept','قسم','department'])), None)
        sal_col_od = next((c for c in data.select_dtypes('number').columns if any(x in c.lower() for x in ['gross','إجمالي','total','net'])), None)

        if page == "🔍 تشخيص المنظمة":
            hdr("🔍 تشخيص المنظمة","Organizational Diagnosis - Burke-Litwin Model")
            ibox("""**نموذج Burke-Litwin للتشخيص المؤسسي:** يحلل 12 بُعداً تنظيمياً مترابطاً لفهم الحالة الحالية وتحديد فرص التحسين.""")

            st.markdown("### 📊 تقييم الأبعاد التنظيمية (1 = ضعيف، 5 = ممتاز)")
            dimensions = {
                "البيئة الخارجية": {"en":"External Environment","desc":"العوامل الخارجية المؤثرة (سوق، تقنية، تنظيمات)"},
                "القيادة": {"en":"Leadership","desc":"فعالية القيادة العليا والرؤية"},
                "الثقافة التنظيمية": {"en":"Culture","desc":"القيم والمعتقدات والسلوكيات السائدة"},
                "الاستراتيجية": {"en":"Strategy","desc":"وضوح الاستراتيجية ومواءمتها"},
                "الهيكل التنظيمي": {"en":"Structure","desc":"تصميم الهيكل وتوزيع السلطات"},
                "الأنظمة والسياسات": {"en":"Systems","desc":"السياسات والإجراءات والعمليات"},
                "ممارسات الإدارة": {"en":"Management Practices","desc":"أساليب الإدارة والإشراف"},
                "المناخ التنظيمي": {"en":"Work Climate","desc":"بيئة العمل والعلاقات"},
                "المهام والمهارات": {"en":"Task & Skills","desc":"تصميم الوظائف وملاءمة المهارات"},
                "الدافعية": {"en":"Motivation","desc":"مستوى التحفيز والالتزام"},
                "الاحتياجات الفردية": {"en":"Individual Needs","desc":"تلبية احتياجات الموظفين"},
                "الأداء المؤسسي": {"en":"Performance","desc":"النتائج والإنتاجية والجودة"},
            }

            scores = {}
            cols = st.columns(3)
            for i, (dim, info) in enumerate(dimensions.items()):
                with cols[i % 3]:
                    scores[dim] = st.slider(f"{dim}", 1, 5, 3, key=f"od_{dim}", help=info['desc'])

            # Radar chart
            fig = go.Figure()
            dims = list(scores.keys())
            vals = list(scores.values()) + [list(scores.values())[0]]
            fig.add_trace(go.Scatterpolar(r=vals, theta=dims + [dims[0]], fill='toself',
                line=dict(color='#E36414'), fillcolor='rgba(227,100,20,0.2)', name='التقييم الحالي'))
            # Add benchmark
            fig.add_trace(go.Scatterpolar(r=[4]*len(dims)+[4], theta=dims+[dims[0]], fill='toself',
                line=dict(color='#2A9D8F', dash='dash'), fillcolor='rgba(42,157,143,0.1)', name='المستوى المستهدف'))
            fig.update_layout(polar=dict(radialaxis=dict(range=[0,5])),
                title='Burke-Litwin Organizational Diagnosis', font=dict(family="Noto Sans Arabic"), height=500)
            st.plotly_chart(fig, use_container_width=True)

            # Gap analysis
            st.markdown("### 📊 تحليل الفجوات")
            gap_df = pd.DataFrame([{"البُعد":d,"التقييم":s,"المستهدف":4,"الفجوة":4-s,
                "الأولوية":"🔴 حرج" if s<=2 else ("🟡 متوسط" if s<=3 else "🟢 جيد")} for d,s in scores.items()])
            gap_df = gap_df.sort_values('الفجوة', ascending=False)
            st.dataframe(gap_df, use_container_width=True, hide_index=True)

            avg_score = sum(scores.values()) / len(scores)
            st.markdown(f"### 🎯 التقييم العام: **{avg_score:.1f}/5**")
            if avg_score >= 4: ibox("المنظمة في حالة ممتازة. ركّز على الاستدامة والابتكار.", "success")
            elif avg_score >= 3: ibox("المنظمة في حالة جيدة. ركّز على الأبعاد ذات الفجوات العالية.", "warning")
            else: ibox("المنظمة تحتاج تدخل تطويري شامل. ابدأ بالأبعاد الحرجة.", "warning")

            export_widget(gap_df, "تشخيص_المنظمة_OD", "od1")
        elif page == "📊 تحليل OD":
            hdr("📊 تحليل التطوير المؤسسي","Workforce Analytics for OD Planning")
            if len(data)==0: st.info("📁 ارفع ملف"); return

            n = len(data)
            nat_col = next((c for c in data.columns if any(x in c.lower() for x in ['nat','جنسية'])), None)
            status_col = next((c for c in data.columns if any(x in c.lower() for x in ['status','حالة'])), None)
            type_col = next((c for c in data.columns if any(x in c.lower() for x in ['type','نوع'])), None)

            # OD Metrics
            st.markdown("### 📊 مؤشرات OD الرئيسية")
            ok1,ok2,ok3,ok4 = st.columns(4)
            with ok1:
                n_depts = data[dept_col_od].nunique() if dept_col_od else 0
                kpi("🏢 الأقسام", str(n_depts))
            with ok2:
                span = n / max(n_depts,1)
                kpi("📏 Span of Control", f"{span:.1f}")
            with ok3:
                if status_col:
                    active = len(data[data[status_col].isin(['Active','نشط'])])
                    kpi("✅ نشط", str(active))
                else: kpi("👥 الإجمالي", str(n))
            with ok4:
                if sal_col_od:
                    cost_per = data[sal_col_od].mean()
                    kpi("💰 تكلفة/موظف", f"{cost_per:,.0f}")

            # Dept size distribution
            if dept_col_od:
                dc = data[dept_col_od].value_counts().reset_index()
                dc.columns = [dept_col_od, 'العدد']
                dc['النسبة'] = (dc['العدد'] / n * 100).round(1)

                oc1, oc2 = st.columns(2)
                with oc1:
                    fig = px.treemap(dc, path=[dept_col_od], values='العدد', title='الهيكل التنظيمي (Treemap)',
                        color='العدد', color_continuous_scale='teal')
                    fig.update_layout(font=dict(family="Noto Sans Arabic"), height=400)
                    st.plotly_chart(fig, use_container_width=True)
                with oc2:
                    fig = px.funnel(dc.head(10), x='العدد', y=dept_col_od, title='حجم الأقسام (Funnel)',
                        color_discrete_sequence=['#0F4C5C'])
                    fig.update_layout(font=dict(family="Noto Sans Arabic"), height=400)
                    st.plotly_chart(fig, use_container_width=True)

            # Diversity analysis
            if nat_col:
                st.markdown("### 🌍 تحليل التنوع (Diversity)")
                diversity = data[nat_col].value_counts()
                hhi = sum((cnt/n)**2 for cnt in diversity.values)
                diversity_index = round((1 - hhi) * 100, 1)
                kpi("🌍 مؤشر التنوع (HHI)", f"{diversity_index}%")
                st.caption("100% = تنوع كامل | 0% = لا تنوع")



        elif page == "🎯 استراتيجية OD":
            hdr("🎯 بناء استراتيجية التطوير المؤسسي","OD Strategy Framework")

            st.markdown("### 🎯 SWOT Analysis")
            sw1, sw2 = st.columns(2)
            with sw1:
                strengths = st.text_area("💪 نقاط القوة (Strengths):", height=100, key="od_s",
                    placeholder="فريق تقني قوي\nثقافة ابتكار\nقيادة داعمة")
                weaknesses = st.text_area("⚠️ نقاط الضعف (Weaknesses):", height=100, key="od_w",
                    placeholder="دوران وظيفي مرتفع\nضعف التدريب\nهيكل غير واضح")
            with sw2:
                opportunities = st.text_area("🚀 الفرص (Opportunities):", height=100, key="od_o",
                    placeholder="توسع السوق\nتقنيات AI جديدة\nدعم حكومي")
                threats = st.text_area("🔴 التهديدات (Threats):", height=100, key="od_t",
                    placeholder="منافسة شديدة\nتغييرات تنظيمية\nنقص الكفاءات")

            # OD Interventions
            st.markdown("---")
            st.markdown("### 🔧 التدخلات التطويرية (OD Interventions)")
            interventions = {
                "Human Process": {
                    "icon":"👥","items":["Team Building","Conflict Resolution","Process Consultation",
                        "Third-Party Intervention","Coaching & Mentoring","Survey Feedback"]},
                "Techno-Structural": {
                    "icon":"🏗️","items":["Restructuring","Job Redesign","Quality Circles",
                        "Total Quality Management","Work Design","Downsizing/Rightsizing"]},
                "HR Management": {
                    "icon":"📋","items":["Performance Management","Talent Development",
                        "Succession Planning","Career Planning","Reward Systems","Diversity & Inclusion"]},
                "Strategic": {
                    "icon":"🎯","items":["Strategic Planning","Culture Change","Organization Transformation",
                        "Merger & Acquisition Integration","Learning Organization","Knowledge Management"]},
            }

            selected_interventions = []
            for cat, info in interventions.items():
                with st.expander(f"{info['icon']} {cat}"):
                    for item in info['items']:
                        if st.checkbox(item, key=f"odi_{cat}_{item}"):
                            selected_interventions.append({"الفئة":cat, "التدخل":item})

            if selected_interventions:
                st.markdown("### ✅ التدخلات المختارة")
                st.dataframe(pd.DataFrame(selected_interventions), use_container_width=True, hide_index=True)

            # Strategic priorities
            st.markdown("---")
            st.markdown("### 🎯 الأولويات الاستراتيجية")
            priorities = st.data_editor(
                pd.DataFrame([
                    {"الأولوية":"تطوير القيادة","المدة":"6 أشهر","المسؤول":"HR","الحالة":"مخطط"},
                    {"الأولوية":"تحسين بيئة العمل","المدة":"3 أشهر","المسؤول":"HR + Operations","الحالة":"مخطط"},
                    {"الأولوية":"برنامج إدارة المواهب","المدة":"12 شهر","المسؤول":"HR","الحالة":"مخطط"},
                ]),
                column_config={'الحالة': st.column_config.SelectboxColumn('الحالة', options=['مخطط','جاري','مكتمل','مؤجل'])},
                use_container_width=True, hide_index=True, num_rows="dynamic", key="od_priorities"
            )

        elif page == "📋 خطة التنفيذ":
            hdr("📋 خطة تنفيذ OD","Kotter's 8-Step Change Model")
            ibox("""**نموذج Kotter للتغيير المؤسسي (8 خطوات):**
1. Create Urgency | 2. Form Coalition | 3. Create Vision | 4. Communicate Vision
5. Empower Action | 6. Quick Wins | 7. Build on Change | 8. Anchor in Culture""")

            kotter_steps = [
                {"step":1,"name":"خلق الإلحاح","en":"Create Urgency","desc":"بيّن لماذا التغيير ضروري الآن","duration":"أسبوعين"},
                {"step":2,"name":"تشكيل تحالف","en":"Form Coalition","desc":"حدد القادة والمؤثرين الداعمين","duration":"أسبوع"},
                {"step":3,"name":"صياغة الرؤية","en":"Create Vision","desc":"حدد الرؤية الواضحة والاستراتيجية","duration":"أسبوعين"},
                {"step":4,"name":"نشر الرؤية","en":"Communicate","desc":"تواصل الرؤية بكل الوسائل المتاحة","duration":"شهر"},
                {"step":5,"name":"تمكين العمل","en":"Empower Action","desc":"أزل العوائق ومكّن الفرق","duration":"شهر"},
                {"step":6,"name":"تحقيق مكاسب سريعة","en":"Quick Wins","desc":"حقق نتائج ملموسة مبكرة","duration":"شهرين"},
                {"step":7,"name":"البناء على التغيير","en":"Build on Change","desc":"وسّع نطاق التحسينات","duration":"3 أشهر"},
                {"step":8,"name":"ترسيخ في الثقافة","en":"Anchor","desc":"اجعل التغيير جزءاً من الثقافة","duration":"6 أشهر"},
            ]

            for step in kotter_steps:
                with st.expander(f"Step {step['step']}: {step['name']} ({step['en']})", expanded=step['step']<=3):
                    st.markdown(f"**{step['desc']}**")
                    sc1, sc2, sc3 = st.columns(3)
                    with sc1: st.text_input("الإجراءات:", key=f"k_action_{step['step']}", placeholder="ماذا ستفعل؟")
                    with sc2: st.text_input("المسؤول:", key=f"k_owner_{step['step']}")
                    with sc3: st.selectbox("الحالة:", ["لم يبدأ","جاري","مكتمل"], key=f"k_status_{step['step']}")

        elif page == "📥 تصدير OD":
            hdr("📥 تصدير خطة التطوير المؤسسي")
            if st.button("📥 تصدير Excel", type="primary", use_container_width=True, key="od_exp"):
                ox = io.BytesIO()
                with pd.ExcelWriter(ox, engine='xlsxwriter') as w:
                    kotter_df = pd.DataFrame([{"Step":s['step'],"Name":s['name'],"EN":s['en'],"Description":s['desc'],"Duration":s['duration']} for s in [
                        {"step":1,"name":"خلق الإلحاح","en":"Create Urgency","desc":"بيّن لماذا التغيير ضروري","duration":"أسبوعين"},
                        {"step":2,"name":"تشكيل تحالف","en":"Form Coalition","desc":"حدد القادة الداعمين","duration":"أسبوع"},
                        {"step":3,"name":"صياغة الرؤية","en":"Create Vision","desc":"حدد الرؤية الواضحة","duration":"أسبوعين"},
                        {"step":4,"name":"نشر الرؤية","en":"Communicate","desc":"تواصل بكل الوسائل","duration":"شهر"},
                        {"step":5,"name":"تمكين العمل","en":"Empower","desc":"أزل العوائق","duration":"شهر"},
                        {"step":6,"name":"مكاسب سريعة","en":"Quick Wins","desc":"نتائج ملموسة مبكرة","duration":"شهرين"},
                        {"step":7,"name":"البناء","en":"Build","desc":"وسّع التحسينات","duration":"3 أشهر"},
                        {"step":8,"name":"ترسيخ","en":"Anchor","desc":"جزء من الثقافة","duration":"6 أشهر"},
                    ]])
                    kotter_df.to_excel(w, sheet_name='Kotter 8 Steps', index=False)
                    if len(data)>0: data.to_excel(w, sheet_name='Workforce Data', index=False)
                st.download_button("📥 تحميل", data=ox.getvalue(),
                    file_name=f"OD_Plan_{datetime.now().strftime('%Y%m%d')}.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")


    # =========================================
    #         📈 ADVANCED ANALYTICS MODULE
    # =========================================
    elif section == "📈 التحليلات المتقدمة":

        data = sal_snapshot if len(sal_snapshot)>0 else emp
        dept_col_a = next((c for c in data.columns if any(x in c.lower() for x in ['dept','قسم','department'])), None) if len(data)>0 else None
        sal_col_a = next((c for c in data.select_dtypes('number').columns if any(x in c.lower() for x in ['gross','إجمالي','total','net'])), None) if len(data)>0 else None
        join_col_a = next((c for c in data.columns if any(x in c.lower() for x in ['join','hiring','التحاق','مباشرة','start'])), None) if len(data)>0 else None
        status_col_a = next((c for c in data.columns if any(x in c.lower() for x in ['status','حالة'])), None) if len(data)>0 else None
        name_col_a = next((c for c in data.columns if any(x in c.lower() for x in ['name','اسم'])), None) if len(data)>0 else None

        # ===== PAGE 1: HR Metrics =====
        if page == "📊 مؤشرات HR المتقدمة":
            hdr("📊 مؤشرات الموارد البشرية المتقدمة","Turnover Rate + Time-to-Hire + Cost-per-Hire + Absenteeism")
            if len(data)==0: st.info("📁 ارفع ملف بيانات"); return

            st.markdown("### ⚙️ إدخال بيانات المؤشرات")
            mc1, mc2, mc3, mc4 = st.columns(4)
            with mc1:
                m_start_count = st.number_input("عدد الموظفين بداية الفترة:", 1, 10000, len(data), key="m_start")
                m_end_count = st.number_input("عدد الموظفين نهاية الفترة:", 1, 10000, len(data), key="m_end")
            with mc2:
                m_separations = st.number_input("عدد المغادرين (Separations):", 0, 1000, 5, key="m_sep")
                m_voluntary = st.number_input("منهم استقالة (Voluntary):", 0, 1000, 3, key="m_vol")
            with mc3:
                m_hires = st.number_input("عدد التعيينات الجديدة:", 0, 1000, 8, key="m_hires")
                m_avg_days_hire = st.number_input("متوسط أيام التوظيف:", 1, 365, 35, key="m_days")
            with mc4:
                m_recruit_cost = st.number_input("تكلفة التوظيف الإجمالية (ريال):", 0, 1000000, 50000, key="m_rcost")
                m_absent_days = st.number_input("إجمالي أيام الغياب:", 0, 10000, 120, key="m_abs")
                m_work_days = st.number_input("أيام العمل في الفترة:", 1, 365, 250, key="m_wdays")

            # Calculate metrics
            avg_emp = (m_start_count + m_end_count) / 2
            turnover_rate = round(m_separations / max(avg_emp, 1) * 100, 1)
            vol_turnover = round(m_voluntary / max(avg_emp, 1) * 100, 1)
            invol_turnover = round((m_separations - m_voluntary) / max(avg_emp, 1) * 100, 1)
            retention_rate = round(100 - turnover_rate, 1)
            cost_per_hire = round(m_recruit_cost / max(m_hires, 1), 0)
            absence_rate = round(m_absent_days / (max(avg_emp, 1) * m_work_days) * 100, 2)
            turnover_cost = round(m_separations * (data[sal_col_a].mean() * 6 if sal_col_a else 30000), 0)

            st.markdown("### 📊 المؤشرات الرئيسية")
            k1,k2,k3,k4 = st.columns(4)
            with k1: kpi("📉 Turnover Rate", f"{turnover_rate}%")
            with k2: kpi("✅ Retention Rate", f"{retention_rate}%")
            with k3: kpi("⏱️ Time-to-Hire", f"{m_avg_days_hire} يوم")
            with k4: kpi("💰 Cost-per-Hire", f"{cost_per_hire:,.0f} ريال")

            k5,k6,k7,k8 = st.columns(4)
            with k5: kpi("🚶 Voluntary Turnover", f"{vol_turnover}%")
            with k6: kpi("❌ Involuntary", f"{invol_turnover}%")
            with k7: kpi("📅 Absence Rate", f"{absence_rate}%")
            with k8: kpi("💸 Turnover Cost", f"{turnover_cost:,.0f}")

            # Benchmark comparison
            st.markdown("### 📊 مقارنة بالمعايير العالمية")
            benchmarks = pd.DataFrame([
                {"المؤشر":"Turnover Rate","القيمة":turnover_rate,"المعيار العالمي":15,"الوحدة":"%","الحالة":"🟢 جيد" if turnover_rate<15 else ("🟡 متوسط" if turnover_rate<25 else "🔴 مرتفع")},
                {"المؤشر":"Voluntary Turnover","القيمة":vol_turnover,"المعيار العالمي":10,"الوحدة":"%","الحالة":"🟢" if vol_turnover<10 else "🔴"},
                {"المؤشر":"Time-to-Hire","القيمة":m_avg_days_hire,"المعيار العالمي":36,"الوحدة":"يوم","الحالة":"🟢" if m_avg_days_hire<=36 else "🔴"},
                {"المؤشر":"Cost-per-Hire","القيمة":cost_per_hire,"المعيار العالمي":15000,"الوحدة":"ريال","الحالة":"🟢" if cost_per_hire<=15000 else "🔴"},
                {"المؤشر":"Absence Rate","القيمة":absence_rate,"المعيار العالمي":2.5,"الوحدة":"%","الحالة":"🟢" if absence_rate<=2.5 else "🔴"},
                {"المؤشر":"Retention Rate","القيمة":retention_rate,"المعيار العالمي":85,"الوحدة":"%","الحالة":"🟢" if retention_rate>=85 else "🔴"},
            ])
            st.dataframe(benchmarks, use_container_width=True, hide_index=True)

            bc1, bc2 = st.columns(2)
            with bc1:
                fig = go.Figure()
                fig.add_trace(go.Bar(name='القيمة الفعلية', x=benchmarks['المؤشر'], y=benchmarks['القيمة'], marker_color='#E36414'))
                fig.add_trace(go.Bar(name='المعيار العالمي', x=benchmarks['المؤشر'], y=benchmarks['المعيار العالمي'], marker_color='#2A9D8F'))
                fig.update_layout(title='المؤشرات مقارنة بالمعايير العالمية', barmode='group', font=dict(family="Noto Sans Arabic"), height=400)
                st.plotly_chart(fig, use_container_width=True)
            with bc2:
                fig = px.pie(values=[m_voluntary, m_separations-m_voluntary, m_end_count],
                    names=['استقالة','إنهاء','باقين'], title='توزيع حالات الموظفين', hole=0.4,
                    color_discrete_sequence=['#E36414','#E74C3C','#2A9D8F'])
                fig.update_layout(font=dict(family="Noto Sans Arabic"), height=400)
                st.plotly_chart(fig, use_container_width=True)

            # Turnover by department
            if dept_col_a and status_col_a:
                st.markdown("### 🏢 الدوران حسب القسم")
                dept_counts = data[dept_col_a].value_counts()
                dept_turn = pd.DataFrame({'القسم':dept_counts.index, 'العدد':dept_counts.values,
                    'الدوران المقدر %': [round(turnover_rate * (0.8 + 0.4 * (i%3)), 1) for i in range(len(dept_counts))]})
                fig = px.bar(dept_turn, x='القسم', y='الدوران المقدر %', title='معدل الدوران المقدر حسب القسم',
                    color='الدوران المقدر %', color_continuous_scale='RdYlGn_r')
                fig.add_hline(y=15, line_dash="dash", line_color="gray", annotation_text="المعيار 15%")
                fig.update_layout(font=dict(family="Noto Sans Arabic"), height=380, coloraxis_showscale=False)
                st.plotly_chart(fig, use_container_width=True)

            export_widget(benchmarks, "مؤشرات_HR_المتقدمة", "adv1")

        # ===== PAGE 2: Smart Alerts =====
        elif page == "🔔 التنبيهات الذكية":
            hdr("🔔 نظام التنبيهات الذكية","تنبيهات تلقائية عند تجاوز العتبات المحددة")

            st.markdown("### ⚙️ إعداد العتبات")
            th1, th2, th3 = st.columns(3)
            with th1:
                thr_turnover = st.number_input("عتبة الدوران %:", 0, 100, 20, key="thr_t")
                thr_absence = st.number_input("عتبة الغياب %:", 0.0, 20.0, 3.0, 0.5, key="thr_a")
            with th2:
                thr_saudization = st.number_input("الحد الأدنى للسعودة %:", 0, 100, 30, key="thr_s")
                thr_salary_diff = st.number_input("فرق الراتب عن السوق %:", 0, 50, 20, key="thr_sd")
            with th3:
                thr_headcount_change = st.number_input("تغير العدد %:", 0, 50, 10, key="thr_hc")
                thr_overtime = st.number_input("عتبة ساعات إضافية/شهر:", 0, 100, 30, key="thr_ot")

            if len(data) > 0:
                st.markdown("### 🔔 التنبيهات النشطة")
                alerts = []
                nat_col_a = next((c for c in data.columns if any(x in c.lower() for x in ['nat','جنسية'])), None)

                # Saudization check
                if nat_col_a:
                    saudi_vals = ['Saudi','سعودي','Saudi Arabian','سعودية']
                    sa_pct = len(data[data[nat_col_a].isin(saudi_vals)]) / len(data) * 100
                    if sa_pct < thr_saudization:
                        alerts.append({"النوع":"🔴 حرج","التنبيه":f"نسبة السعودة ({sa_pct:.1f}%) أقل من الحد الأدنى ({thr_saudization}%)","الإجراء":"زيادة التوظيف السعودي أو مراجعة استراتيجية السعودة"})
                    else:
                        alerts.append({"النوع":"🟢 طبيعي","التنبيه":f"نسبة السعودة ({sa_pct:.1f}%) ضمن المعيار","الإجراء":"استمرار المتابعة"})

                # Salary distribution alerts
                if sal_col_a and dept_col_a:
                    dept_avg = data.groupby(dept_col_a)[sal_col_a].mean()
                    overall_avg = data[sal_col_a].mean()
                    for dept, avg in dept_avg.items():
                        diff_pct = abs(avg - overall_avg) / overall_avg * 100
                        if diff_pct > thr_salary_diff:
                            direction = "أعلى" if avg > overall_avg else "أقل"
                            alerts.append({"النوع":"🟡 تحذير","التنبيه":f"قسم {dept}: متوسط الراتب {direction} بنسبة {diff_pct:.0f}% عن المتوسط العام","الإجراء":"مراجعة هيكل الرواتب للقسم"})

                # Headcount concentration
                if dept_col_a:
                    max_dept_pct = data[dept_col_a].value_counts(normalize=True).iloc[0] * 100
                    if max_dept_pct > 40:
                        alerts.append({"النوع":"🟡 تحذير","التنبيه":f"تركز {max_dept_pct:.0f}% من الموظفين في قسم واحد","الإجراء":"مراجعة التوزيع التنظيمي"})

                # Overtime check
                ot_col = next((c for c in data.columns if any(x in c.lower() for x in ['overtime','إضافي','ساعات'])), None)
                if ot_col:
                    avg_ot = data[ot_col].mean()
                    if avg_ot > thr_overtime:
                        alerts.append({"النوع":"🔴 حرج","التنبيه":f"متوسط الساعات الإضافية ({avg_ot:.0f}) يتجاوز العتبة ({thr_overtime})","الإجراء":"مراجعة الأعباء الوظيفية أو زيادة التوظيف"})

                if not alerts:
                    alerts.append({"النوع":"🟢 طبيعي","التنبيه":"جميع المؤشرات ضمن النطاق الطبيعي","الإجراء":"استمرار المتابعة الدورية"})

                alerts_df = pd.DataFrame(alerts)
                critical = len([a for a in alerts if "حرج" in a["النوع"]])
                warning = len([a for a in alerts if "تحذير" in a["النوع"]])
                normal = len([a for a in alerts if "طبيعي" in a["النوع"]])

                ak1,ak2,ak3 = st.columns(3)
                with ak1: kpi("🔴 تنبيهات حرجة", str(critical))
                with ak2: kpi("🟡 تحذيرات", str(warning))
                with ak3: kpi("🟢 طبيعي", str(normal))

                st.dataframe(alerts_df, use_container_width=True, hide_index=True)

                # Policy recommendations
                st.markdown("### 📋 التوصيات المرتبطة بالسياسات")
                for alert in alerts:
                    if "حرج" in alert["النوع"] or "تحذير" in alert["النوع"]:
                        with st.expander(f"{alert['النوع']} {alert['التنبيه'][:50]}..."):
                            st.markdown(f"**التنبيه:** {alert['التنبيه']}")
                            st.markdown(f"**الإجراء المقترح:** {alert['الإجراء']}")
                            if "سعودة" in alert['التنبيه']:
                                st.markdown("**السند النظامي:** نظام نطاقات + قرارات وزارة الموارد البشرية")
                            elif "راتب" in alert['التنبيه']:
                                st.markdown("**السند النظامي:** المادة 89 + 90 من نظام العمل (المساواة في الأجور)")

                export_widget(alerts_df, "التنبيهات_الذكية", "alrt")

            else:
                st.info("📁 ارفع ملف بيانات لتفعيل التنبيهات")

        # ===== PAGE 3: What-If Scenarios =====
        elif page == "🔮 سيناريوهات What-If":
            hdr("🔮 سيناريوهات What-If","تحليل الأثر المالي والتنظيمي للقرارات")

            st.markdown("### 📊 اختر السيناريو")
            scenario = st.selectbox("السيناريو:", [
                "زيادة رواتب عامة %","تقليص عدد الموظفين %","توظيف إضافي",
                "تغيير نسبة السعودة","زيادة ميزانية التدريب"], key="wif_scn")

            if len(data) > 0 and sal_col_a:
                current_payroll = data[sal_col_a].sum()
                current_count = len(data)
                current_avg = data[sal_col_a].mean()

                if scenario == "زيادة رواتب عامة %":
                    pct_increase = st.slider("نسبة الزيادة %:", 1, 30, 10, key="wif_pct")
                    new_payroll = current_payroll * (1 + pct_increase/100)
                    monthly_diff = new_payroll - current_payroll
                    annual_diff = monthly_diff * 12

                    wc1, wc2, wc3 = st.columns(3)
                    with wc1: kpi("💰 الرواتب الحالية/شهر", f"{current_payroll:,.0f}")
                    with wc2: kpi("💰 الرواتب الجديدة/شهر", f"{new_payroll:,.0f}")
                    with wc3: kpi("📈 الفرق السنوي", f"{annual_diff:,.0f}")

                    fig = go.Figure()
                    fig.add_trace(go.Bar(name='الحالي', x=['شهري','سنوي'], y=[current_payroll, current_payroll*12], marker_color='#2A9D8F'))
                    fig.add_trace(go.Bar(name=f'بعد زيادة {pct_increase}%', x=['شهري','سنوي'], y=[new_payroll, new_payroll*12], marker_color='#E36414'))
                    fig.update_layout(title='أثر الزيادة على الرواتب', barmode='group', font=dict(family="Noto Sans Arabic"), height=400, yaxis_tickformat=',')
                    st.plotly_chart(fig, use_container_width=True)

                elif scenario == "تقليص عدد الموظفين %":
                    pct_reduce = st.slider("نسبة التقليص %:", 1, 30, 10, key="wif_red")
                    removed = int(current_count * pct_reduce / 100)
                    new_count = current_count - removed
                    saved_monthly = removed * current_avg
                    eos_cost = removed * current_avg * 2  # Estimated 2 months EOS

                    wc1, wc2, wc3, wc4 = st.columns(4)
                    with wc1: kpi("👥 الحالي", str(current_count))
                    with wc2: kpi("👥 الجديد", str(new_count))
                    with wc3: kpi("💰 توفير شهري", f"{saved_monthly:,.0f}")
                    with wc4: kpi("💸 تكلفة EOS", f"{eos_cost:,.0f}")

                    ibox(f"التوفير السنوي: {saved_monthly*12:,.0f} ريال | تكلفة نهاية الخدمة لمرة واحدة: {eos_cost:,.0f} ريال | نقطة التعادل: {round(eos_cost/max(saved_monthly,1),1)} شهر", "warning")

                elif scenario == "توظيف إضافي":
                    new_hires = st.number_input("عدد الموظفين الجدد:", 1, 100, 10, key="wif_nh")
                    avg_new_sal = st.number_input("متوسط راتب الجديد:", 3000, 50000, int(current_avg), key="wif_ns")
                    recruit_cost = st.number_input("تكلفة توظيف/فرد:", 1000, 50000, 5000, key="wif_rc")

                    total_new_monthly = new_hires * avg_new_sal
                    total_recruit = new_hires * recruit_cost

                    wc1, wc2, wc3 = st.columns(3)
                    with wc1: kpi("💰 تكلفة شهرية جديدة", f"{total_new_monthly:,.0f}")
                    with wc2: kpi("💰 تكلفة التوظيف", f"{total_recruit:,.0f}")
                    with wc3: kpi("📅 التكلفة السنوية", f"{total_new_monthly*12 + total_recruit:,.0f}")

                elif scenario == "تغيير نسبة السعودة":
                    target_sa = st.slider("نسبة السعودة المستهدفة %:", 10, 100, 50, key="wif_sa")
                    nat_col_a = next((c for c in data.columns if any(x in c.lower() for x in ['nat','جنسية'])), None)
                    if nat_col_a:
                        current_sa = len(data[data[nat_col_a].isin(['Saudi','سعودي','Saudi Arabian','سعودية'])]) / len(data) * 100
                        needed = int(len(data) * target_sa / 100 - len(data[data[nat_col_a].isin(['Saudi','سعودي','Saudi Arabian','سعودية'])]))
                        kpi("📊 السعودة الحالية", f"{current_sa:.1f}%")
                        kpi("📊 المستهدف", f"{target_sa}%")
                        kpi("👥 مطلوب توظيف سعوديين", str(max(needed, 0)))

                elif scenario == "زيادة ميزانية التدريب":
                    train_budget = st.number_input("ميزانية التدريب الحالية:", 10000, 1000000, 70000, key="wif_tb")
                    train_increase = st.slider("نسبة الزيادة %:", 10, 200, 50, key="wif_ti")
                    new_budget = train_budget * (1 + train_increase/100)
                    per_emp = new_budget / max(current_count, 1)
                    roi_estimate = new_budget * 2.5

                    wc1,wc2,wc3 = st.columns(3)
                    with wc1: kpi("💰 الميزانية الجديدة", f"{new_budget:,.0f}")
                    with wc2: kpi("💰 للفرد", f"{per_emp:,.0f}")
                    with wc3: kpi("📈 ROI المتوقع", f"{roi_estimate:,.0f}")
            else:
                st.info("📁 ارفع ملف بيانات لتفعيل السيناريوهات")

        # ===== PAGE 4: Predictive Analytics =====
        elif page == "🤖 التحليل التنبؤي":
            hdr("🤖 التحليل التنبؤي بالذكاء الاصطناعي","Predictive & Diagnostic Analytics")

            if len(data) > 0:
                analysis_type = st.selectbox("نوع التحليل:", [
                    "🔮 توقع الاستقالات (Attrition Prediction)",
                    "🔍 تشخيص أسباب الدوران (Diagnostic)",
                    "📊 تحليل Cohort (فوج التعيين)",
                    "🎯 تصنيف الموظفين (Clustering)"], key="pred_type")

                if analysis_type == "🔮 توقع الاستقالات (Attrition Prediction)":
                    st.markdown("### 🔮 عوامل خطر الاستقالة")
                    risk_factors = []
                    for _, row in data.iterrows():
                        score = 0
                        # Tenure risk
                        if join_col_a:
                            try:
                                join_date = pd.to_datetime(row.get(join_col_a), errors='coerce')
                                if pd.notna(join_date):
                                    tenure = (datetime.now() - join_date).days / 365
                                    if 1 < tenure < 3: score += 25  # High risk window
                                    elif tenure > 10: score += 5  # Low risk
                            except: pass
                        # Salary risk
                        if sal_col_a and pd.notna(row.get(sal_col_a)):
                            median = data[sal_col_a].median()
                            if row[sal_col_a] < median * 0.8: score += 30  # Below market
                        # Department risk
                        if dept_col_a and row.get(dept_col_a):
                            dept_size = len(data[data[dept_col_a]==row[dept_col_a]])
                            if dept_size < 3: score += 15  # Small team risk
                        score = min(score, 100)
                        risk_factors.append({"الاسم":row.get(name_col_a,'N/A') if name_col_a else f"#{_}",
                            "القسم":row.get(dept_col_a,'') if dept_col_a else '',
                            "مؤشر الخطر":score,
                            "المستوى":"🔴 عالي" if score>=50 else ("🟡 متوسط" if score>=25 else "🟢 منخفض")})

                    risk_df = pd.DataFrame(risk_factors).sort_values('مؤشر الخطر', ascending=False)
                    high_risk = len(risk_df[risk_df['مؤشر الخطر']>=50])
                    med_risk = len(risk_df[(risk_df['مؤشر الخطر']>=25) & (risk_df['مؤشر الخطر']<50)])

                    rk1,rk2,rk3 = st.columns(3)
                    with rk1: kpi("🔴 خطر عالي", str(high_risk))
                    with rk2: kpi("🟡 خطر متوسط", str(med_risk))
                    with rk3: kpi("🟢 مستقر", str(len(risk_df) - high_risk - med_risk))

                    st.dataframe(risk_df.head(20), use_container_width=True, hide_index=True)

                    fig = px.histogram(risk_df, x='مؤشر الخطر', nbins=10, title='توزيع مؤشر خطر الاستقالة',
                        color_discrete_sequence=['#0F4C5C'])
                    fig.update_layout(font=dict(family="Noto Sans Arabic"), height=350)
                    st.plotly_chart(fig, use_container_width=True)

                    export_widget(risk_df, "توقع_الاستقالات", "pred1")

                elif analysis_type == "🔍 تشخيص أسباب الدوران (Diagnostic)":
                    st.markdown("### 🔍 تحليل العوامل المؤثرة في الدوران")
                    factors = []
                    if sal_col_a:
                        sal_cv = round(data[sal_col_a].std() / data[sal_col_a].mean() * 100, 1) if data[sal_col_a].mean()>0 else 0
                        factors.append({"العامل":"تفاوت الرواتب (CV)","القيمة":f"{sal_cv}%","الأثر":"مرتفع" if sal_cv>30 else "منخفض","التوصية":"مراجعة هيكل الرواتب" if sal_cv>30 else "مستوى جيد"})
                    if dept_col_a:
                        dept_count = data[dept_col_a].nunique()
                        avg_size = len(data) / max(dept_count,1)
                        factors.append({"العامل":"متوسط حجم القسم","القيمة":f"{avg_size:.0f}","الأثر":"تحذير" if avg_size<5 else "طبيعي","التوصية":"دمج الأقسام الصغيرة" if avg_size<5 else "حجم مناسب"})
                    if join_col_a:
                        try:
                            data['_tenure'] = (pd.Timestamp.now() - pd.to_datetime(data[join_col_a], errors='coerce')).dt.days / 365
                            avg_tenure = data['_tenure'].mean()
                            factors.append({"العامل":"متوسط الخدمة (سنوات)","القيمة":f"{avg_tenure:.1f}","الأثر":"تحذير" if avg_tenure<2 else "جيد","التوصية":"تحسين برامج الاحتفاظ" if avg_tenure<2 else "استقرار جيد"})
                        except: pass

                    if factors:
                        st.dataframe(pd.DataFrame(factors), use_container_width=True, hide_index=True)
                        export_widget(pd.DataFrame(factors), "تشخيص_الدوران", "diag1")

                elif analysis_type == "📊 تحليل Cohort (فوج التعيين)":
                    if join_col_a:
                        try:
                            data['_join_year'] = pd.to_datetime(data[join_col_a], errors='coerce').dt.year
                            cohort = data.groupby('_join_year').agg(العدد=(data.columns[0],'count')).reset_index()
                            cohort.columns = ['سنة التعيين','العدد']
                            cohort = cohort.dropna().sort_values('سنة التعيين')

                            fig = px.bar(cohort, x='سنة التعيين', y='العدد', title='توزيع الموظفين حسب سنة التعيين (Cohort)',
                                color='العدد', color_continuous_scale='teal')
                            fig.update_layout(font=dict(family="Noto Sans Arabic"), height=400, coloraxis_showscale=False)
                            st.plotly_chart(fig, use_container_width=True)
                            export_widget(cohort, "تحليل_Cohort", "coh1")
                        except: st.warning("تعذر تحليل تاريخ التعيين")
                    else:
                        st.info("لا يوجد عمود تاريخ التعيين في البيانات")

                elif analysis_type == "🎯 تصنيف الموظفين (Clustering)":
                    num_cols = data.select_dtypes('number').columns.tolist()[:5]
                    if len(num_cols) >= 2:
                        st.markdown("### 🎯 تصنيف الموظفين حسب البيانات الرقمية")
                        try:
                            from sklearn.preprocessing import StandardScaler
                        except ImportError:
                            st.warning("مكتبة scikit-learn غير مثبتة. أضف `scikit-learn` في requirements.txt")
                            return
                        try:
                            clean = data[num_cols].dropna()
                            if len(clean) > 10:
                                scaled = StandardScaler().fit_transform(clean)
                                # Simple quantile-based clustering
                                clean['_Cluster'] = pd.qcut(scaled[:,0], q=3, labels=['مجموعة أ','مجموعة ب','مجموعة ج'], duplicates='drop')
                                fig = px.scatter(clean, x=num_cols[0], y=num_cols[1] if len(num_cols)>1 else num_cols[0],
                                    color='_Cluster', title='تصنيف الموظفين', color_discrete_sequence=['#E36414','#2A9D8F','#264653'])
                                fig.update_layout(font=dict(family="Noto Sans Arabic"), height=400)
                                st.plotly_chart(fig, use_container_width=True)
                                export_widget(clean, "تصنيف_الموظفين", "clus1")
                        except Exception as e:
                            st.warning(f"تعذر التصنيف: {e}")
                    else:
                        st.info("يلزم عمودين رقميين على الأقل")
            else:
                st.info("📁 ارفع ملف بيانات للتحليل التنبؤي")

        # ===== PAGE 5: Sentiment Analysis =====
        elif page == "💬 تحليل المشاعر":
            hdr("💬 تحليل المشاعر والنصوص","Sentiment Analysis للتقييمات والملاحظات")

            st.markdown("### 📝 أدخل نصوص للتحليل")
            text_source = st.radio("مصدر النصوص:", ["إدخال يدوي","من البيانات المرفوعة"], horizontal=True, key="sent_src")

            if text_source == "إدخال يدوي":
                texts_input = st.text_area("الصق النصوص (نص واحد في كل سطر):", height=200, key="sent_txt",
                    placeholder="الشركة ممتازة والبيئة محفزة\nالراتب غير مناسب\nالإدارة تحتاج تطوير\nفرص النمو رائعة")
                if texts_input:
                    texts = [t.strip() for t in texts_input.split('\n') if t.strip()]
                else:
                    texts = []
            else:
                text_cols = [c for c in data.columns if data[c].dtype == 'object'] if len(data)>0 else []
                if text_cols:
                    sel_col = st.selectbox("اختر العمود:", text_cols, key="sent_col")
                    texts = data[sel_col].dropna().tolist()[:50]
                else:
                    texts = []; st.info("لا توجد أعمدة نصية")

            if texts and st.button("💬 تحليل المشاعر", type="primary", key="sent_btn"):
                # Simple keyword-based sentiment
                positive_kw = ['ممتاز','رائع','جيد','محفز','فرص','نمو','تطوير','ابتكار','excellent','great','good','growth']
                negative_kw = ['سيء','ضعيف','مشكلة','غير مناسب','نقص','ضغط','تأخر','bad','poor','problem','pressure']
                neutral_kw = ['عادي','متوسط','مقبول','normal','average','ok']

                results = []
                for text in texts:
                    text_lower = text.lower()
                    pos = sum(1 for kw in positive_kw if kw in text_lower)
                    neg = sum(1 for kw in negative_kw if kw in text_lower)
                    if pos > neg: sentiment = "إيجابي 😊"; score = min(pos * 20, 100)
                    elif neg > pos: sentiment = "سلبي 😟"; score = -min(neg * 20, 100)
                    else: sentiment = "محايد 😐"; score = 0
                    results.append({"النص":text[:100],"المشاعر":sentiment,"الدرجة":score})

                res_df = pd.DataFrame(results)
                pos_count = len(res_df[res_df['الدرجة']>0])
                neg_count = len(res_df[res_df['الدرجة']<0])
                neutral_count = len(res_df[res_df['الدرجة']==0])

                sk1,sk2,sk3,sk4 = st.columns(4)
                with sk1: kpi("📊 إجمالي النصوص", str(len(res_df)))
                with sk2: kpi("😊 إيجابي", str(pos_count))
                with sk3: kpi("😟 سلبي", str(neg_count))
                with sk4: kpi("😐 محايد", str(neutral_count))

                st.dataframe(res_df, use_container_width=True, hide_index=True)

                sc1, sc2 = st.columns(2)
                with sc1:
                    fig = px.pie(values=[pos_count,neg_count,neutral_count], names=['إيجابي','سلبي','محايد'],
                        title='توزيع المشاعر', hole=0.4, color_discrete_sequence=['#27AE60','#E74C3C','#95A5A6'])
                    fig.update_layout(font=dict(family="Noto Sans Arabic"), height=350)
                    st.plotly_chart(fig, use_container_width=True)
                with sc2:
                    fig = px.histogram(res_df, x='الدرجة', nbins=10, title='توزيع درجات المشاعر',
                        color_discrete_sequence=['#0F4C5C'])
                    fig.update_layout(font=dict(family="Noto Sans Arabic"), height=350)
                    st.plotly_chart(fig, use_container_width=True)

                export_widget(res_df, "تحليل_المشاعر", "sent1")

        # ===== PAGE 6: Audit Log =====
        elif page == "📋 سجل التدقيق":
            hdr("📋 سجل التدقيق","Audit Trail - تتبع جميع العمليات")

            if st.session_state.get('user_role') != "مدير":
                st.warning("⚠️ سجل التدقيق متاح للمدير فقط"); return

            try:
                conn = get_conn(); c = conn.cursor()
                c.execute(f"SELECT value FROM app_config WHERE key = {_ph()}", ("audit_log",))
                row = c.fetchone(); conn.close()
                logs = json.loads(row[0]) if row else []
            except:
                logs = []

            if logs:
                log_df = pd.DataFrame(logs)
                log_df = log_df.sort_values('time', ascending=False).reset_index(drop=True)

                lk1,lk2,lk3 = st.columns(3)
                with lk1: kpi("📋 إجمالي العمليات", str(len(log_df)))
                with lk2: kpi("👥 المستخدمين", str(log_df['user'].nunique()))
                with lk3: kpi("📅 آخر عملية", log_df['time'].iloc[0][:16] if len(log_df)>0 else "-")

                # Filter
                fc1, fc2 = st.columns(2)
                with fc1:
                    users_filter = st.multiselect("فلتر المستخدم:", log_df['user'].unique().tolist(), key="aud_usr")
                with fc2:
                    actions_filter = st.multiselect("فلتر العملية:", log_df['action'].unique().tolist(), key="aud_act")

                filtered = log_df
                if users_filter: filtered = filtered[filtered['user'].isin(users_filter)]
                if actions_filter: filtered = filtered[filtered['action'].isin(actions_filter)]

                st.dataframe(filtered.head(100), use_container_width=True, hide_index=True)

                # Activity chart
                if len(log_df) > 1:
                    log_df['date'] = log_df['time'].str[:10]
                    daily = log_df.groupby('date').size().reset_index(name='العمليات')
                    fig = px.bar(daily, x='date', y='العمليات', title='النشاط اليومي',
                        color_discrete_sequence=['#2A9D8F'])
                    fig.update_layout(font=dict(family="Noto Sans Arabic"), height=300)
                    st.plotly_chart(fig, use_container_width=True)

                export_widget(filtered, "سجل_التدقيق", "audt")

                if st.button("🗑️ مسح السجل", key="clear_audit"):
                    try:
                        conn = get_conn(); c = conn.cursor()
                        _upsert_config(c, "audit_log", "[]")
                        conn.commit(); conn.close()
                        st.success("✅ تم مسح السجل"); st.rerun()
                    except: pass
            else:
                st.info("📋 سجل التدقيق فارغ. سيتم تسجيل العمليات تلقائياً.")
                st.caption("العمليات المُسجّلة: تسجيل الدخول، رفع الملفات، تصدير التقارير، تعديل البيانات")


    # =========================================
    #         🔍 GENERAL ANALYSIS MODULE
    # =========================================
    elif section == "🔍 التحليل العام":

        if page == "📊 تحليل تلقائي":
            hdr("📊 التحليل التلقائي للبيانات", "ارفع أي ملف Excel وسيتم تحليله تلقائياً")

            ga_file = st.file_uploader("📁 ارفع ملف Excel أو CSV:", type=["xlsx","xls","csv"], key="ga_uploader")
            ga_df = pd.DataFrame()

            if ga_file:
                try:
                    if ga_file.name.endswith('.csv'):
                        ga_df = pd.read_csv(ga_file)
                    else:
                        ga_xl = pd.ExcelFile(ga_file)
                        if len(ga_xl.sheet_names) > 1:
                            ga_sheet = st.selectbox("اختر الشيت:", ga_xl.sheet_names, key="ga_sh")
                        else:
                            ga_sheet = ga_xl.sheet_names[0]
                        ga_df = pd.read_excel(ga_xl, ga_sheet)
                except Exception as e:
                    st.error(f"خطأ في قراءة الملف: {e}")

            elif len(emp) > 0:
                ga_df = emp.copy()
                st.info("📂 يتم تحليل بيانات الموظفين المرفوعة في القائمة الجانبية")

            if len(ga_df) > 0:
                st.markdown("---")
                st.markdown("### 📊 نظرة عامة على البيانات")

                gi1, gi2, gi3, gi4 = st.columns(4)
                with gi1: kpi("📋 الصفوف", f"{len(ga_df):,}")
                with gi2: kpi("📊 الأعمدة", f"{len(ga_df.columns)}")
                with gi3: kpi("❌ القيم الفارغة", f"{ga_df.isnull().sum().sum():,}")
                with gi4: kpi("🔢 الأعمدة الرقمية", f"{len(ga_df.select_dtypes('number').columns)}")

                # Data types summary
                st.markdown("### 📋 هيكل البيانات")
                dtype_data = []
                for col in ga_df.columns:
                    dtype_data.append({
                        "العمود": col,
                        "النوع": str(ga_df[col].dtype),
                        "القيم الفريدة": ga_df[col].nunique(),
                        "الفارغة": ga_df[col].isnull().sum(),
                        "عينة": str(ga_df[col].dropna().iloc[0]) if len(ga_df[col].dropna()) > 0 else "-"
                    })
                st.dataframe(pd.DataFrame(dtype_data), use_container_width=True, hide_index=True)

                # Numeric columns analysis
                num_cols = ga_df.select_dtypes('number').columns.tolist()
                cat_cols = [c for c in ga_df.columns if ga_df[c].dtype == 'object' and ga_df[c].nunique() < 30 and ga_df[c].nunique() > 1]

                if num_cols:
                    st.markdown("### 📈 الإحصائيات الوصفية")
                    desc = ga_df[num_cols].describe().T
                    desc.columns = ["العدد","المتوسط","الانحراف","الأدنى","25%","الوسيط","75%","الأقصى"]
                    st.dataframe(desc.style.format("{:,.2f}"), use_container_width=True)

                    # Auto charts
                    st.markdown("### 📊 رسوم بيانية تلقائية")

                    # Histogram for numeric columns
                    sel_num = st.selectbox("اختر عمود رقمي:", num_cols, key="ga_num")
                    fig = px.histogram(ga_df, x=sel_num, nbins=30, title=f"توزيع: {sel_num}", color_discrete_sequence=[CL['p']])
                    fig.update_layout(font=dict(family="Noto Sans Arabic"), height=350)
                    st.plotly_chart(fig, use_container_width=True)

                    if len(num_cols) >= 2:
                        sc1, sc2 = st.columns(2)
                        with sc1: sel_x = st.selectbox("المحور X:", num_cols, index=0, key="ga_x")
                        with sc2: sel_y = st.selectbox("المحور Y:", num_cols, index=min(1, len(num_cols)-1), key="ga_y")
                        color_col = st.selectbox("التلوين حسب (اختياري):", ["بدون"] + cat_cols, key="ga_clr") if cat_cols else "بدون"
                        fig = px.scatter(ga_df, x=sel_x, y=sel_y, color=None if color_col=="بدون" else color_col,
                            title=f"{sel_y} vs {sel_x}", opacity=0.6)
                        fig.update_layout(font=dict(family="Noto Sans Arabic"), height=400)
                        st.plotly_chart(fig, use_container_width=True)

                if cat_cols:
                    st.markdown("### 📊 تحليل الأعمدة النصية")
                    sel_cat = st.selectbox("اختر عمود:", cat_cols, key="ga_cat")
                    vc = ga_df[sel_cat].value_counts().head(15)
                    fig = px.bar(x=vc.index, y=vc.values, title=f"توزيع: {sel_cat}", labels={"x": sel_cat, "y": "العدد"}, color_discrete_sequence=[CL['a']])
                    fig.update_layout(font=dict(family="Noto Sans Arabic"), height=350)
                    st.plotly_chart(fig, use_container_width=True)

                    # Cross analysis
                    if num_cols and cat_cols:
                        st.markdown("### 📊 تحليل متقاطع")
                        cx1, cx2 = st.columns(2)
                        with cx1: cross_cat = st.selectbox("التصنيف:", cat_cols, key="ga_cc")
                        with cx2: cross_num = st.selectbox("القيمة:", num_cols, key="ga_cn")
                        cross_agg = ga_df.groupby(cross_cat)[cross_num].agg(['mean','sum','count']).reset_index()
                        cross_agg.columns = [cross_cat, "المتوسط", "الإجمالي", "العدد"]
                        fig = px.bar(cross_agg, x=cross_cat, y="المتوسط", title=f"متوسط {cross_num} حسب {cross_cat}", text_auto=".1f", color_discrete_sequence=[CL['s']])
                        fig.update_layout(font=dict(family="Noto Sans Arabic"), height=350)
                        st.plotly_chart(fig, use_container_width=True)

                # Correlation heatmap
                if len(num_cols) >= 3:
                    st.markdown("### 🔥 خريطة الارتباط")
                    corr = ga_df[num_cols].corr()
                    fig = px.imshow(corr, text_auto=".2f", aspect="auto", title="مصفوفة الارتباط", color_continuous_scale="RdBu_r")
                    fig.update_layout(font=dict(family="Noto Sans Arabic"), height=500)
                    st.plotly_chart(fig, use_container_width=True)

                # Raw data
                with st.expander("📋 عرض البيانات الخام"):
                    st.dataframe(ga_df, use_container_width=True, height=400)

            export_widget(None, "تحليل_تلقائي", "ga01")

        elif page == "🤖 أسئلة ذكية":
            hdr("🤖 المحلل الذكي", "اطرح أسئلة عن بياناتك بالعربي أو الإنجليزي")

            if len(emp) > 0:
                st.success(f"📂 البيانات جاهزة: {len(emp):,} صف × {len(emp.columns)} عمود")

                q = st.text_input("💬 اسأل عن بياناتك:", placeholder="مثال: كم متوسط الرواتب حسب القسم؟ أو ما أعلى 5 رواتب؟", key="ga_q")

                if q:
                    num_cols = emp.select_dtypes('number').columns.tolist()
                    cat_cols = [c for c in emp.columns if emp[c].dtype=='object' and emp[c].nunique() < 50 and emp[c].nunique() > 1]
                    ql = q.lower()

                    try:
                        # Pattern matching for common questions
                        if any(w in ql for w in ['متوسط','average','mean']):
                            if num_cols:
                                matched_num = None
                                for nc in num_cols:
                                    if any(p in ql for p in [nc.lower(), nc.replace('_',' ').lower()]):
                                        matched_num = nc; break
                                if not matched_num: matched_num = num_cols[0]

                                matched_cat = None
                                for cc in cat_cols:
                                    if any(p in ql for p in [cc.lower(), cc.replace('_',' ').lower(), 'قسم','department','div']):
                                        matched_cat = cc; break

                                if matched_cat:
                                    result = emp.groupby(matched_cat)[matched_num].mean().sort_values(ascending=False)
                                    st.dataframe(result.reset_index().rename(columns={matched_num: f"متوسط {matched_num}"}), use_container_width=True, hide_index=True)
                                    fig = px.bar(x=result.index, y=result.values, title=f"متوسط {matched_num} حسب {matched_cat}", text_auto=".1f")
                                    fig.update_layout(font=dict(family="Noto Sans Arabic"), height=350)
                                    st.plotly_chart(fig, use_container_width=True)
                                else:
                                    for nc in num_cols:
                                        st.metric(nc, f"{emp[nc].mean():,.2f}")

                        elif any(w in ql for w in ['أعلى','top','highest','أكبر','max']):
                            n = 5
                            for w in ql.split():
                                try: n = int(w); break
                                except: pass
                            sort_col = num_cols[0] if num_cols else emp.columns[0]
                            for nc in num_cols:
                                if nc.lower() in ql: sort_col = nc; break
                            st.dataframe(emp.nlargest(n, sort_col), use_container_width=True, hide_index=True)

                        elif any(w in ql for w in ['أقل','bottom','lowest','أصغر','min']):
                            n = 5
                            for w in ql.split():
                                try: n = int(w); break
                                except: pass
                            sort_col = num_cols[0] if num_cols else emp.columns[0]
                            for nc in num_cols:
                                if nc.lower() in ql: sort_col = nc; break
                            st.dataframe(emp.nsmallest(n, sort_col), use_container_width=True, hide_index=True)

                        elif any(w in ql for w in ['عدد','count','كم','how many']):
                            matched_cat = None
                            for cc in cat_cols:
                                if any(p in ql for p in [cc.lower(), cc.replace('_',' ').lower()]):
                                    matched_cat = cc; break
                            if matched_cat:
                                vc = emp[matched_cat].value_counts()
                                st.dataframe(vc.reset_index().rename(columns={matched_cat: "الفئة", "count": "العدد"}), use_container_width=True, hide_index=True)
                                fig = px.pie(names=vc.index, values=vc.values, title=f"توزيع {matched_cat}")
                                fig.update_layout(font=dict(family="Noto Sans Arabic"), height=350)
                                st.plotly_chart(fig, use_container_width=True)
                            else:
                                st.metric("إجمالي الصفوف", f"{len(emp):,}")

                        elif any(w in ql for w in ['مجموع','total','sum','إجمالي']):
                            for nc in num_cols:
                                if nc.lower() in ql or any(p in ql for p in nc.lower().split('_')):
                                    st.metric(f"مجموع {nc}", f"{emp[nc].sum():,.2f}")
                                    break
                            else:
                                if num_cols:
                                    sums = {nc: emp[nc].sum() for nc in num_cols}
                                    st.dataframe(pd.DataFrame({"العمود": sums.keys(), "المجموع": [f"{v:,.2f}" for v in sums.values()]}), use_container_width=True, hide_index=True)

                        elif any(w in ql for w in ['توزيع','distribution','histogram']):
                            matched_col = num_cols[0] if num_cols else None
                            for nc in num_cols:
                                if nc.lower() in ql: matched_col = nc; break
                            if matched_col:
                                fig = px.histogram(emp, x=matched_col, nbins=25, title=f"توزيع {matched_col}")
                                fig.update_layout(font=dict(family="Noto Sans Arabic"), height=350)
                                st.plotly_chart(fig, use_container_width=True)

                        elif any(w in ql for w in ['أعمدة','columns','حقول','fields']):
                            st.write("📋 الأعمدة المتاحة:")
                            for i, c in enumerate(emp.columns, 1):
                                st.write(f"  {i}. **{c}** ({emp[c].dtype})")

                        else:
                            st.warning("💡 حاول صياغة السؤال بطريقة مختلفة. أمثلة:")
                            st.write("- ما متوسط الرواتب حسب القسم؟")
                            st.write("- أعلى 10 رواتب")
                            st.write("- عدد الموظفين حسب الجنسية")
                            st.write("- توزيع الأعمار")
                            st.write("- مجموع الرواتب")

                    except Exception as e:
                        st.error(f"خطأ: {e}")
                        st.info("حاول سؤال آخر أو تأكد من البيانات")
            else:
                ibox("ارفع ملف بيانات من القائمة الجانبية أولاً.", "warning")


    # =========================================
    #       📤 REPORTS & EXPORT MODULE
    # =========================================
    elif section == "📤 التقارير والتصدير":

        if page == "📊 تقرير Dashboard":
            hdr("📊 تقرير Dashboard احترافي","تقرير بأسلوب Power BI مع رسوم بيانية ورؤى ذكية بالذكاء الاصطناعي")

            if len(emp) > 0:
                st.markdown("### ⚙️ إعدادات التقرير")
                pc1, pc2 = st.columns(2)
                with pc1:
                    rpt_title = st.text_input("عنوان التقرير:", "HR Analytics Dashboard", key="rpt_t")
                    rpt_company = st.text_input("اسم الشركة:", "رسال الود لتقنية المعلومات", key="rpt_c")
                    rpt_prepared = st.text_input("إعداد:", st.session_state.get('user_name',''), key="rpt_p")
                with pc2:
                    rpt_period = st.text_input("الفترة:", datetime.now().strftime('%Y'), key="rpt_pr")
                    rpt_sections = st.multiselect("أقسام التقرير:",
                        ["ملخص تنفيذي","القوى العاملة","تحليل الرواتب","الأقسام","الجنسيات","الأداء","رؤى الذكاء الاصطناعي"],
                        default=["ملخص تنفيذي","القوى العاملة","تحليل الرواتب","رؤى الذكاء الاصطناعي"],
                        key="rpt_sec")

                if st.button("📊 إنشاء التقرير", type="primary", use_container_width=True, key="rpt_btn"):
                    with st.spinner("جاري إنشاء التقرير..."):
                        n = len(emp)
                        num_cols = emp.select_dtypes('number').columns.tolist()
                        cat_cols = [c for c in emp.columns if emp[c].dtype=='object' and 1 < emp[c].nunique() < 30]
                        sal_col = next((c for c in num_cols if any(x in c.lower() for x in ['gross','salary','net','راتب','إجمالي'])), num_cols[0] if num_cols else None)
                        dept_col = next((c for c in cat_cols if any(x in c.lower() for x in ['dept','department','قسم','القطاع'])), cat_cols[0] if cat_cols else None)
                        nat_col = next((c for c in cat_cols if any(x in c.lower() for x in ['nat','جنسية','nationality'])), None)
                        status_col = next((c for c in cat_cols if any(x in c.lower() for x in ['status','حالة'])), None)
                        loc_col = next((c for c in cat_cols if any(x in c.lower() for x in ['location','موقع','مدينة','city'])), None)

                        # Generate charts as Plotly HTML divs (no kaleido needed)
                        charts_html = ""

                        def add_chart(fig):
                            nonlocal charts_html
                            fig.update_layout(template='plotly_dark', paper_bgcolor='#1e1e2e', plot_bgcolor='#1e1e2e',
                                height=350, font=dict(color='white', size=11), margin=dict(l=10,r=10,t=40,b=10),
                                showlegend=True)
                            chart_div = fig.to_html(full_html=False, include_plotlyjs='cdn')
                            charts_html += f"<div class='chart'>{chart_div}</div>"

                        charts_html = ""

                        # Chart 1: Department distribution
                        if dept_col:
                            dept_counts = emp[dept_col].value_counts().head(12)
                            fig = px.bar(x=dept_counts.values, y=dept_counts.index, orientation='h',
                                color=dept_counts.values, color_continuous_scale='teal',
                                labels={'x':'Count','y':''})
                            fig.update_layout(title='Headcount by Department', template='plotly_dark',
                                paper_bgcolor='#1e1e2e', plot_bgcolor='#1e1e2e', height=350,
                                font=dict(color='white'), showlegend=False, coloraxis_showscale=False,
                                margin=dict(l=10,r=10,t=40,b=10))
                            add_chart(fig)

                        # Chart 2: Salary distribution
                        if sal_col:
                            fig = px.histogram(emp, x=sal_col, nbins=20, color_discrete_sequence=['#0F4C5C'])
                            fig.update_layout(title='Salary Distribution', template='plotly_dark',
                                paper_bgcolor='#1e1e2e', plot_bgcolor='#1e1e2e', height=350,
                                font=dict(color='white'), margin=dict(l=10,r=10,t=40,b=10))
                            add_chart(fig)

                        # Chart 3: Nationality pie
                        if nat_col:
                            nat_counts = emp[nat_col].value_counts().head(8)
                            fig = px.pie(values=nat_counts.values, names=nat_counts.index, hole=0.4,
                                color_discrete_sequence=px.colors.qualitative.Set2)
                            fig.update_layout(title='Nationality Distribution', template='plotly_dark',
                                paper_bgcolor='#1e1e2e', plot_bgcolor='#1e1e2e', height=350,
                                font=dict(color='white'), margin=dict(l=10,r=10,t=40,b=10))
                            add_chart(fig)

                        # Chart 4: Status
                        if status_col:
                            stat_counts = emp[status_col].value_counts()
                            fig = px.pie(values=stat_counts.values, names=stat_counts.index, hole=0.5,
                                color_discrete_sequence=['#27AE60','#E74C3C','#F39C12','#3498DB'])
                            fig.update_layout(title='Employment Status', template='plotly_dark',
                                paper_bgcolor='#1e1e2e', plot_bgcolor='#1e1e2e', height=350,
                                font=dict(color='white'), margin=dict(l=10,r=10,t=40,b=10))
                            add_chart(fig)

                        # Chart 5: Location
                        if loc_col:
                            loc_counts = emp[loc_col].value_counts().head(8)
                            fig = px.bar(x=loc_counts.index, y=loc_counts.values,
                                color_discrete_sequence=['#2A9D8F'])
                            fig.update_layout(title='Distribution by Location', template='plotly_dark',
                                paper_bgcolor='#1e1e2e', plot_bgcolor='#1e1e2e', height=350,
                                font=dict(color='white'), margin=dict(l=10,r=10,t=40,b=10))
                            add_chart(fig)

                        # Chart 6: Salary by dept box
                        if sal_col and dept_col:
                            top_depts = emp[dept_col].value_counts().head(8).index
                            fig = px.box(emp[emp[dept_col].isin(top_depts)], x=dept_col, y=sal_col,
                                color_discrete_sequence=['#E9C46A'])
                            fig.update_layout(title='Salary Range by Department', template='plotly_dark',
                                paper_bgcolor='#1e1e2e', plot_bgcolor='#1e1e2e', height=350,
                                font=dict(color='white'), margin=dict(l=10,r=10,t=40,b=10))
                            add_chart(fig)

                        # KPIs
                        active_count = len(emp[emp[status_col].isin(['Active','نشط','active'])]) if status_col else n
                        avg_sal = emp[sal_col].mean() if sal_col else 0
                        total_payroll = emp[sal_col].sum() if sal_col else 0
                        n_depts = emp[dept_col].nunique() if dept_col else 0
                        sa_pct = 0
                        if nat_col:
                            sa = len(emp[emp[nat_col].isin(['Saudi','سعودي','Saudi Arabian'])])
                            sa_pct = round(sa / max(n,1) * 100, 1)

                        # AI Insights
                        insights = []
                        if dept_col:
                            top_dept = emp[dept_col].value_counts().idxmax()
                            top_pct = emp[dept_col].value_counts().iloc[0] / n * 100
                            insights.append(f"The largest department is <strong>{top_dept}</strong> with {top_pct:.0f}% of total workforce.")
                        if sal_col:
                            median = emp[sal_col].median()
                            std = emp[sal_col].std()
                            insights.append(f"Average salary is <strong>{avg_sal:,.0f} SAR</strong> with median at <strong>{median:,.0f} SAR</strong>. Salary spread (std) is {std:,.0f}.")
                            if avg_sal > median * 1.15:
                                insights.append("Salary distribution is <strong>right-skewed</strong>, indicating a few high earners pulling the average up. Consider reviewing compensation equity.")
                        if nat_col:
                            insights.append(f"Saudization rate is <strong>{sa_pct}%</strong>." + (" Meets Nitaqat requirements." if sa_pct >= 50 else " May need improvement for Nitaqat compliance."))
                        if status_col:
                            turnover = len(emp[emp[status_col].isin(['Historical','Terminated','منتهي'])]) / max(n,1) * 100
                            insights.append(f"Historical/terminated rate is <strong>{turnover:.0f}%</strong>." + (" High turnover detected - investigate retention strategies." if turnover > 50 else ""))
                        if dept_col and sal_col:
                            dept_avg = emp.groupby(dept_col)[sal_col].mean().sort_values(ascending=False)
                            insights.append(f"Highest-paid department: <strong>{dept_avg.index[0]}</strong> ({dept_avg.iloc[0]:,.0f} SAR avg).")

                        insights_html = "".join([f"<div class='insight'><span class='dot'></span>{ins}</div>" for ins in insights])

                        # Table
                        table_html = ""
                        if dept_col:
                            dept_stats = emp.groupby(dept_col).agg(
                                Count=(dept_col, 'count'),
                                **({f'Avg_{sal_col}': (sal_col, 'mean')} if sal_col else {})
                            ).sort_values('Count', ascending=False).head(12).reset_index()
                            table_rows = ""
                            for _, r in dept_stats.iterrows():
                                pct = r['Count'] / n * 100
                                bar_w = min(pct * 3, 100)
                                sal_str = f"{r[f'Avg_{sal_col}']:,.0f}" if sal_col else "-"
                                table_rows += f"<tr><td>{r[dept_col]}</td><td>{r['Count']}</td><td>{pct:.1f}%</td><td><div class='bar' style='width:{bar_w}%'></div></td><td>{sal_str}</td></tr>"
                            table_html = f"""<table><thead><tr><th>Department</th><th>Count</th><th>%</th><th>Distribution</th><th>Avg Salary</th></tr></thead><tbody>{table_rows}</tbody></table>"""

                        # Build full HTML
                        html = f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8"><meta name="viewport" content="width=device-width,initial-scale=1">
<script src="https://cdn.plot.ly/plotly-2.27.0.min.js"></script>
<title>{rpt_title}</title>
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&display=swap');
*{{margin:0;padding:0;box-sizing:border-box}}
body{{font-family:'Inter',sans-serif;background:white;color:#e0e0e0;padding:0}}
.header{{background:linear-gradient(135deg,#0F4C5C,#1A1A2E);padding:30px 40px;border-bottom:3px solid #E36414}}
.header h1{{font-size:1.8em;color:white;margin-bottom:5px}}
.header .sub{{color:rgba(255,255,255,0.6);font-size:0.9em}}
.header .logo{{float:right;width:50px;height:50px;background:linear-gradient(135deg,#E36414,#E9C46A);border-radius:10px;display:flex;align-items:center;justify-content:center;font-weight:800;color:white;font-size:18px}}
.container{{max-width:1200px;margin:0 auto;padding:20px}}
.kpi-row{{display:grid;grid-template-columns:repeat(auto-fit,minmax(180px,1fr));gap:12px;margin:20px 0}}
.kpi{{background:linear-gradient(135deg,white,#252540);border-radius:12px;padding:18px;border-left:4px solid #E36414;position:relative;overflow:hidden}}
.kpi::after{{content:'';position:absolute;top:0;right:0;width:60px;height:100%;background:linear-gradient(90deg,transparent,rgba(227,100,20,0.05))}}
.kpi .value{{font-size:1.6em;font-weight:700;color:#E9C46A;margin-bottom:2px}}
.kpi .label{{font-size:0.75em;color:rgba(255,255,255,0.5);text-transform:uppercase;letter-spacing:1px}}
.section{{margin:25px 0}}
.section h2{{font-size:1.2em;color:#E9C46A;margin-bottom:12px;padding-bottom:6px;border-bottom:1px solid rgba(255,255,255,0.1)}}
.charts-grid{{display:grid;grid-template-columns:repeat(auto-fit,minmax(500px,1fr));gap:15px}}
.chart{{background:white;border-radius:10px;overflow:hidden;border:1px solid rgba(255,255,255,0.05)}}
.chart img{{width:100%;height:auto;display:block}}
table{{width:100%;border-collapse:collapse;background:white;border-radius:10px;overflow:hidden;margin:10px 0}}
thead{{background:#252540}} th{{padding:10px 12px;text-align:left;font-size:0.8em;text-transform:uppercase;letter-spacing:1px;color:#E9C46A;border-bottom:2px solid #E36414}}
td{{padding:8px 12px;border-bottom:1px solid rgba(255,255,255,0.05);font-size:0.85em}}
tr:hover{{background:rgba(227,100,20,0.05)}}
.bar{{height:8px;background:linear-gradient(90deg,#E36414,#E9C46A);border-radius:4px}}
.insights{{background:linear-gradient(135deg,#1a2332,white);border-radius:12px;padding:20px;border:1px solid rgba(42,157,143,0.3);margin:15px 0}}
.insights h2{{color:#2A9D8F;border-bottom-color:rgba(42,157,143,0.3)}}
.insight{{padding:8px 0;border-bottom:1px solid rgba(255,255,255,0.05);font-size:0.9em;line-height:1.6;display:flex;align-items:flex-start;gap:10px}}
.insight:last-child{{border:none}}
.dot{{width:6px;height:6px;background:#2A9D8F;border-radius:50%;margin-top:8px;flex-shrink:0}}
.insight strong{{color:#E9C46A}}
.footer{{text-align:center;padding:20px;color:rgba(255,255,255,0.3);font-size:0.75em;border-top:1px solid rgba(255,255,255,0.05);margin-top:30px}}
@media print{{body{{background:white;color:#333}} .kpi{{border-color:#0F4C5C}} .kpi .value{{color:#0F4C5C}} .section h2{{color:#0F4C5C}} th{{color:#0F4C5C}}}}
</style>
</head>
<body>
<div class="header">
<div class="logo">HR</div>
<h1>{rpt_title}</h1>
<div class="sub">{rpt_company} | {rpt_period} | Prepared by: {rpt_prepared} | Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}</div>
</div>
<div class="container">

<div class="kpi-row">
<div class="kpi"><div class="value">{n:,}</div><div class="label">Total Records</div></div>
<div class="kpi"><div class="value">{active_count:,}</div><div class="label">Active Employees</div></div>
<div class="kpi"><div class="value">{n_depts}</div><div class="label">Departments</div></div>
<div class="kpi"><div class="value">{avg_sal:,.0f}</div><div class="label">Avg Salary (SAR)</div></div>
<div class="kpi"><div class="value">{total_payroll:,.0f}</div><div class="label">Total Payroll</div></div>
<div class="kpi"><div class="value">{sa_pct}%</div><div class="label">Saudization</div></div>
</div>

{"<div class='section insights'><h2>🤖 AI-Powered Insights</h2>" + insights_html + "</div>" if "رؤى الذكاء الاصطناعي" in rpt_sections else ""}

<div class="section"><h2>📊 Workforce Analytics</h2><div class="charts-grid">{charts_html}</div></div>

{"<div class='section'><h2>📋 Department Breakdown</h2>" + table_html + "</div>" if table_html else ""}

</div>
<div class="footer">{rpt_company} | HR Analytics Platform | Confidential</div>
</body></html>"""

                        # Display in Streamlit
                        st.components.v1.html(html, height=900, scrolling=True)

                        # ===== 4 Export Formats =====
                        st.markdown("---")
                        st.markdown("### 📥 تحميل التقرير")
                        ex1, ex2, ex3, ex4 = st.columns(4)

                        # 1. HTML
                        with ex1:
                            st.download_button("🌐 HTML Dashboard", data=html.encode('utf-8'),
                                file_name=f"{rpt_title}_{datetime.now().strftime('%Y%m%d')}.html",
                                mime="text/html", type="primary", use_container_width=True)

                        # 2. PDF (via HTML with print-optimized CSS)
                        with ex2:
                            pdf_html = html.replace('white','#ffffff').replace('white','#f8f9fa').replace('#252540','#e9ecef')
                            pdf_html = pdf_html.replace("color:'white'","color:'#333'").replace("color:white","color:#333")
                            pdf_html = pdf_html.replace("color:'#e0e0e0'","color:'#333'").replace("color:#e0e0e0","color:#333")
                            # Add print CSS
                            pdf_css = """<style>@media print{body{background:white!important;color:#333!important}
                            .kpi{border-color:#0F4C5C!important;background:#f8f9fa!important}
                            .kpi .value{color:#0F4C5C!important} .kpi .label{color:#555!important}
                            .header{background:#0F4C5C!important;-webkit-print-color-adjust:exact}
                            .chart{break-inside:avoid;page-break-inside:avoid}
                            .section h2{color:#0F4C5C!important}}</style>"""
                            pdf_html = pdf_html.replace('</head>', pdf_css + '</head>')
                            # Auto-print script
                            pdf_html = pdf_html.replace('</body>', '</body>')
                            st.download_button("📄 PDF (Print)", data=pdf_html.encode('utf-8'),
                                file_name=f"{rpt_title}_{datetime.now().strftime('%Y%m%d')}_print.html",
                                mime="text/html", use_container_width=True,
                                help="افتح الملف في المتصفح وسيفتح نافذة الطباعة تلقائياً - اختر Save as PDF")

                        # 3. Excel (multi-sheet with charts matching dashboard)
                        with ex3:
                            ox = io.BytesIO()
                            with pd.ExcelWriter(ox, engine='xlsxwriter') as w:
                                wb = w.book
                                hdr_f = wb.add_format({'bold':True,'font_size':12,'bg_color':'#0F4C5C','font_color':'white','align':'center','border':1})
                                sub_f = wb.add_format({'bold':True,'font_size':10,'bg_color':'#264653','font_color':'white','align':'center','border':1})
                                num_f = wb.add_format({'num_format':'#,##0.00','border':1,'align':'center'})
                                pct_f = wb.add_format({'num_format':'0.0%','border':1,'align':'center'})
                                txt_f = wb.add_format({'border':1,'text_wrap':True})

                                # ===== Sheet 1: Dashboard =====
                                ws1 = wb.add_worksheet('Dashboard')
                                ws1.set_column('A:A', 25); ws1.set_column('B:G', 18)
                                ws1.set_tab_color('#E36414')
                                ws1.merge_range('A1:F1', rpt_title, hdr_f)
                                ws1.merge_range('A2:F2', f'{rpt_company} | {rpt_period} | {datetime.now().strftime("%Y-%m-%d")}', sub_f)

                                # KPIs row
                                kpi_items = [('Total Records', n), ('Active', active_count), ('Departments', n_depts),
                                    ('Avg Salary', round(avg_sal)), ('Total Payroll', round(total_payroll)), ('Saudization', f"{sa_pct}%")]
                                for i, (lbl, val) in enumerate(kpi_items):
                                    ws1.write(3, i, lbl, wb.add_format({'bold':True,'bg_color':'#E9C46A','align':'center','border':1}))
                                    ws1.write(4, i, str(val), wb.add_format({'bold':True,'font_size':14,'align':'center','border':1}))

                                # Dept data for chart
                                if dept_col:
                                    dept_counts = emp[dept_col].value_counts().head(12)
                                    ws1.write(6, 0, 'Department', sub_f); ws1.write(6, 1, 'Headcount', sub_f)
                                    if sal_col: ws1.write(6, 2, 'Avg Salary', sub_f)
                                    for j, (dept, cnt) in enumerate(dept_counts.items()):
                                        ws1.write(7+j, 0, dept, txt_f)
                                        ws1.write(7+j, 1, cnt, wb.add_format({'border':1,'align':'center'}))
                                        if sal_col:
                                            avg = emp[emp[dept_col]==dept][sal_col].mean()
                                            ws1.write(7+j, 2, round(avg,0), num_f)

                                    # Chart 1: Dept headcount bar
                                    chart1 = wb.add_chart({'type':'bar'})
                                    chart1.add_series({'name':'Headcount','categories':['Dashboard',7,0,6+len(dept_counts),0],
                                        'values':['Dashboard',7,1,6+len(dept_counts),1],'fill':{'color':'#0F4C5C'}})
                                    chart1.set_title({'name':'Headcount by Department'})
                                    chart1.set_style(10); chart1.set_size({'width':520,'height':320})
                                    chart1.set_legend({'none':True})
                                    ws1.insert_chart('E7', chart1)

                                    # Chart 2: Salary by dept
                                    if sal_col:
                                        chart2 = wb.add_chart({'type':'column'})
                                        chart2.add_series({'name':'Avg Salary','categories':['Dashboard',7,0,6+len(dept_counts),0],
                                            'values':['Dashboard',7,2,6+len(dept_counts),2],'fill':{'color':'#1A1A2E'}})
                                        chart2.set_title({'name':'Average Salary by Department'})
                                        chart2.set_style(10); chart2.set_size({'width':520,'height':320})
                                        chart2.set_legend({'none':True})
                                        ws1.insert_chart('E24', chart2)

                                # ===== Sheet 2: Charts Data =====
                                ws2 = wb.add_worksheet('Charts')
                                ws2.set_tab_color('#2A9D8F')
                                col_offset = 0

                                # Nationality data + pie chart
                                if nat_col:
                                    nat_counts = emp[nat_col].value_counts().head(10)
                                    ws2.write(0, 0, 'Nationality', sub_f); ws2.write(0, 1, 'Count', sub_f); ws2.write(0, 2, '%', sub_f)
                                    for j, (nat, cnt) in enumerate(nat_counts.items()):
                                        ws2.write(1+j, 0, nat, txt_f)
                                        ws2.write(1+j, 1, cnt, wb.add_format({'border':1,'align':'center'}))
                                        ws2.write(1+j, 2, cnt/n, pct_f)
                                    chart3 = wb.add_chart({'type':'pie'})
                                    chart3.add_series({'name':'Nationality','categories':['Charts',1,0,len(nat_counts),0],
                                        'values':['Charts',1,1,len(nat_counts),1],'data_labels':{'percentage':True}})
                                    chart3.set_title({'name':'Nationality Distribution'})
                                    chart3.set_size({'width':480,'height':320})
                                    ws2.insert_chart('E1', chart3)

                                # Status data + pie chart
                                if status_col:
                                    stat_counts = emp[status_col].value_counts()
                                    r_start = len(nat_counts) + 3 if nat_col else 0
                                    ws2.write(r_start, 0, 'Status', sub_f); ws2.write(r_start, 1, 'Count', sub_f)
                                    for j, (st_name, cnt) in enumerate(stat_counts.items()):
                                        ws2.write(r_start+1+j, 0, st_name, txt_f)
                                        ws2.write(r_start+1+j, 1, cnt, wb.add_format({'border':1,'align':'center'}))
                                    chart4 = wb.add_chart({'type':'pie'})
                                    chart4.add_series({'name':'Status','categories':['Charts',r_start+1,0,r_start+len(stat_counts),0],
                                        'values':['Charts',r_start+1,1,r_start+len(stat_counts),1],'data_labels':{'percentage':True}})
                                    chart4.set_title({'name':'Employment Status'})
                                    chart4.set_size({'width':480,'height':320})
                                    ws2.insert_chart('E' + str(r_start + 1), chart4)

                                # Location data + bar chart
                                if loc_col:
                                    loc_counts = emp[loc_col].value_counts().head(8)
                                    r_start2 = (r_start + len(stat_counts) + 3) if status_col else (len(nat_counts) + 3 if nat_col else 0)
                                    ws2.write(r_start2, 0, 'Location', sub_f); ws2.write(r_start2, 1, 'Count', sub_f)
                                    for j, (loc, cnt) in enumerate(loc_counts.items()):
                                        ws2.write(r_start2+1+j, 0, loc, txt_f)
                                        ws2.write(r_start2+1+j, 1, cnt, wb.add_format({'border':1,'align':'center'}))
                                    chart5 = wb.add_chart({'type':'column'})
                                    chart5.add_series({'name':'Location','categories':['Charts',r_start2+1,0,r_start2+len(loc_counts),0],
                                        'values':['Charts',r_start2+1,1,r_start2+len(loc_counts),1],'fill':{'color':'#E9C46A'}})
                                    chart5.set_title({'name':'Distribution by Location'})
                                    chart5.set_size({'width':480,'height':320}); chart5.set_legend({'none':True})
                                    ws2.insert_chart('E' + str(r_start2 + 1), chart5)

                                # Salary histogram data
                                if sal_col:
                                    ws3 = wb.add_worksheet('Salary Analysis')
                                    ws3.set_tab_color('#E36414')
                                    sal_stats = [('Average', emp[sal_col].mean()), ('Median', emp[sal_col].median()),
                                        ('Min', emp[sal_col].min()), ('Max', emp[sal_col].max()),
                                        ('Std Dev', emp[sal_col].std()), ('Total', emp[sal_col].sum())]
                                    ws3.write(0, 0, 'Salary Statistics', hdr_f); ws3.merge_range('A1:B1', 'Salary Statistics', hdr_f)
                                    for j, (lbl, val) in enumerate(sal_stats):
                                        ws3.write(1+j, 0, lbl, wb.add_format({'bold':True,'border':1}))
                                        ws3.write(1+j, 1, round(val,2), num_f)

                                    # Salary bands for chart
                                    bands = pd.cut(emp[sal_col], bins=8).value_counts().sort_index()
                                    ws3.write(8, 0, 'Salary Band', sub_f); ws3.write(8, 1, 'Count', sub_f)
                                    for j, (band, cnt) in enumerate(bands.items()):
                                        ws3.write(9+j, 0, str(band), txt_f)
                                        ws3.write(9+j, 1, cnt, wb.add_format({'border':1,'align':'center'}))
                                    chart6 = wb.add_chart({'type':'column'})
                                    chart6.add_series({'name':'Distribution','categories':['Salary Analysis',9,0,8+len(bands),0],
                                        'values':['Salary Analysis',9,1,8+len(bands),1],'fill':{'color':'#1A1A2E'}})
                                    chart6.set_title({'name':'Salary Distribution'})
                                    chart6.set_size({'width':600,'height':350}); chart6.set_legend({'none':True})
                                    ws3.insert_chart('D1', chart6)

                                    # Salary by dept box-like (min/avg/max)
                                    if dept_col:
                                        dept_sal = emp.groupby(dept_col)[sal_col].agg(['min','mean','max','count']).sort_values('count',ascending=False).head(10).reset_index()
                                        r = 9 + len(bands) + 2
                                        ws3.write(r, 0, 'Department', sub_f); ws3.write(r, 1, 'Min', sub_f)
                                        ws3.write(r, 2, 'Avg', sub_f); ws3.write(r, 3, 'Max', sub_f)
                                        for j, (_, row) in enumerate(dept_sal.iterrows()):
                                            ws3.write(r+1+j, 0, row[dept_col], txt_f)
                                            ws3.write(r+1+j, 1, round(row['min'],0), num_f)
                                            ws3.write(r+1+j, 2, round(row['mean'],0), num_f)
                                            ws3.write(r+1+j, 3, round(row['max'],0), num_f)
                                        chart7 = wb.add_chart({'type':'column'})
                                        for ci, (col_name, color) in enumerate([('Min','#3498DB'),('Avg','#E36414'),('Max','#27AE60')]):
                                            chart7.add_series({'name':col_name,'categories':['Salary Analysis',r+1,0,r+len(dept_sal),0],
                                                'values':['Salary Analysis',r+1,1+ci,r+len(dept_sal),1+ci],'fill':{'color':color}})
                                        chart7.set_title({'name':'Salary Range by Department (Min/Avg/Max)'})
                                        chart7.set_size({'width':600,'height':350})
                                        ws3.insert_chart('D' + str(r+1), chart7)

                                # ===== Sheet: AI Insights =====
                                ws_ai = wb.add_worksheet('AI Insights')
                                ws_ai.set_column('A:A', 90); ws_ai.set_tab_color('#264653')
                                ws_ai.write('A1', 'AI-Powered Insights', hdr_f)
                                for i, ins in enumerate(insights):
                                    clean = ins.replace('<strong>','').replace('</strong>','')
                                    ws_ai.write(i+2, 0, f"• {clean}", wb.add_format({'text_wrap':True,'font_size':11}))

                                # ===== Sheet: Raw Data =====
                                emp.to_excel(w, sheet_name='Raw Data', index=False)

                                # ===== Sheet: Dept Detail =====
                                if dept_col:
                                    dept_full = emp.groupby(dept_col).agg(
                                        Count=(dept_col,'count'),
                                        **({f'Avg Salary': (sal_col,'mean'), f'Total Salary': (sal_col,'sum'),
                                            f'Min': (sal_col,'min'), f'Max': (sal_col,'max'), f'Median': (sal_col,'median')} if sal_col else {})
                                    ).sort_values('Count', ascending=False).reset_index()
                                    dept_full['%'] = (dept_full['Count'] / n * 100).round(1)
                                    dept_full.to_excel(w, sheet_name='Dept Detail', index=False)

                                for sname in w.sheets: w.sheets[sname].set_column('A:Z', 18)

                            st.download_button("📊 Excel Report", data=ox.getvalue(),
                                file_name=f"{rpt_title}_{datetime.now().strftime('%Y%m%d')}.xlsx",
                                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                                use_container_width=True)

                        # 4. Power BI (structured dataset + template)
                        with ex4:
                            pbi_ox = io.BytesIO()
                            with pd.ExcelWriter(pbi_ox, engine='xlsxwriter') as w:
                                # Clean data optimized for Power BI import
                                pbi_data = emp.copy()
                                # Add calculated columns for Power BI
                                if dept_col: pbi_data['_Department'] = pbi_data[dept_col]
                                if nat_col: pbi_data['_Nationality'] = pbi_data[nat_col]
                                if status_col: pbi_data['_Status'] = pbi_data[status_col]
                                if loc_col: pbi_data['_Location'] = pbi_data[loc_col]
                                if sal_col:
                                    pbi_data['_Salary'] = pd.to_numeric(pbi_data[sal_col], errors='coerce')
                                    pbi_data['_SalaryBand'] = pd.cut(pbi_data['_Salary'], bins=[0,5000,10000,15000,25000,50000,999999],
                                        labels=['0-5K','5-10K','10-15K','15-25K','25-50K','50K+'], ordered=False)
                                pbi_data.to_excel(w, sheet_name='FactEmployees', index=False)

                                # Dimension tables for Power BI star schema
                                if dept_col:
                                    dim_dept = emp[dept_col].value_counts().reset_index()
                                    dim_dept.columns = ['Department','Headcount']
                                    if sal_col:
                                        dim_dept = dim_dept.merge(emp.groupby(dept_col)[sal_col].agg(['mean','sum','min','max']).reset_index().rename(
                                            columns={dept_col:'Department','mean':'AvgSalary','sum':'TotalSalary','min':'MinSalary','max':'MaxSalary'}), on='Department', how='left')
                                    dim_dept.to_excel(w, sheet_name='DimDepartment', index=False)

                                if nat_col:
                                    dim_nat = emp[nat_col].value_counts().reset_index()
                                    dim_nat.columns = ['Nationality','Count']
                                    dim_nat['IsSaudi'] = dim_nat['Nationality'].isin(['Saudi','سعودي','Saudi Arabian'])
                                    dim_nat.to_excel(w, sheet_name='DimNationality', index=False)

                                # KPIs table for Power BI card visuals
                                kpi_df = pd.DataFrame([
                                    {'KPI':'Total Employees','Value':n},
                                    {'KPI':'Active','Value':active_count},
                                    {'KPI':'Departments','Value':n_depts},
                                    {'KPI':'Avg Salary','Value':round(avg_sal,0)},
                                    {'KPI':'Total Payroll','Value':round(total_payroll,0)},
                                    {'KPI':'Saudization %','Value':sa_pct},
                                ])
                                kpi_df.to_excel(w, sheet_name='KPIs', index=False)

                                # Measures guide
                                measures = pd.DataFrame([
                                    {'Measure':'Total Headcount','DAX':'COUNTROWS(FactEmployees)','Table':'FactEmployees'},
                                    {'Measure':'Active Count','DAX':'CALCULATE(COUNTROWS(FactEmployees), FactEmployees[_Status]="Active")','Table':'FactEmployees'},
                                    {'Measure':'Avg Salary','DAX':'AVERAGE(FactEmployees[_Salary])','Table':'FactEmployees'},
                                    {'Measure':'Total Payroll','DAX':'SUM(FactEmployees[_Salary])','Table':'FactEmployees'},
                                    {'Measure':'Saudization %','DAX':'DIVIDE(CALCULATE(COUNTROWS(FactEmployees), DimNationality[IsSaudi]=TRUE), COUNTROWS(FactEmployees))','Table':'DimNationality'},
                                    {'Measure':'Dept Count','DAX':'DISTINCTCOUNT(FactEmployees[_Department])','Table':'FactEmployees'},
                                ])
                                measures.to_excel(w, sheet_name='DAX_Measures_Guide', index=False)

                                for sname in w.sheets: w.sheets[sname].set_column('A:Z', 18)

                            st.download_button("📈 Power BI Dataset", data=pbi_ox.getvalue(),
                                file_name=f"PowerBI_{rpt_title}_{datetime.now().strftime('%Y%m%d')}.xlsx",
                                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                                use_container_width=True,
                                help="افتح Power BI Desktop > Get Data > Excel > اختر هذا الملف")

            else:
                ibox("ارفع ملف بيانات أولاً من القائمة الجانبية.", "warning")

        elif page == "📝 تقرير Word":
            hdr("📝 تصدير تقرير Word احترافي", "مستند Word منسق بشكل احترافي")

            if len(emp) > 0:
                wc1, wc2 = st.columns(2)
                with wc1:
                    word_title = st.text_input("عنوان التقرير:", "تقرير الموارد البشرية", key="wdt")
                    word_company = st.text_input("اسم الشركة:", "رسال الود لتقنية المعلومات", key="wdc")
                with wc2:
                    word_prepared = st.text_input("إعداد:", "", key="wdp")
                    word_period = st.text_input("الفترة:", datetime.now().strftime('%Y'), key="wdpr")

                if st.button("📝 إنشاء تقرير Word", type="primary", key="wdbtn"):
                    try:
                        from docx import Document
                        from docx.shared import Inches, Pt, Cm, RGBColor
                        from docx.enum.text import WD_ALIGN_PARAGRAPH
                        from docx.enum.table import WD_TABLE_ALIGNMENT

                        doc = Document()

                        # Style adjustments
                        style = doc.styles['Normal']
                        style.font.name = 'Calibri'
                        style.font.size = Pt(11)

                        # Title
                        p = doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = p.add_run(word_company)
                        run.font.size = Pt(18)
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(15, 76, 92)

                        p = doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = p.add_run(word_title)
                        run.font.size = Pt(14)
                        run.font.color.rgb = RGBColor(46, 117, 182)

                        p = doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = p.add_run(f"Period: {word_period} | Prepared by: {word_prepared} | Date: {datetime.now().strftime('%Y-%m-%d')}")
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(128, 128, 128)

                        doc.add_paragraph("_" * 60)

                        # Summary
                        doc.add_heading('Executive Summary', level=1)
                        n = len(emp)
                        num_cols = emp.select_dtypes('number').columns.tolist()
                        doc.add_paragraph(f"Total Records: {n:,}")

                        # Category breakdown
                        cat_cols = [c for c in emp.columns if emp[c].dtype=='object' and 1 < emp[c].nunique() < 30]
                        for cc in cat_cols[:4]:
                            doc.add_heading(cc, level=2)
                            table = doc.add_table(rows=1, cols=3)
                            table.style = 'Light Grid Accent 1'
                            table.alignment = WD_TABLE_ALIGNMENT.CENTER
                            hdr_cells = table.rows[0].cells
                            hdr_cells[0].text = cc; hdr_cells[1].text = 'Count'; hdr_cells[2].text = '%'
                            for val, cnt in emp[cc].value_counts().head(8).items():
                                row_cells = table.add_row().cells
                                row_cells[0].text = str(val)
                                row_cells[1].text = str(cnt)
                                row_cells[2].text = f"{cnt/n*100:.1f}%"
                            doc.add_paragraph("")

                        # Numeric summary
                        if num_cols:
                            doc.add_heading('Numeric Summary', level=1)
                            table = doc.add_table(rows=1, cols=5)
                            table.style = 'Light Grid Accent 1'
                            table.alignment = WD_TABLE_ALIGNMENT.CENTER
                            for i, h in enumerate(['Column', 'Mean', 'Min', 'Max', 'Std']):
                                table.rows[0].cells[i].text = h
                            for nc in num_cols[:12]:
                                row_cells = table.add_row().cells
                                row_cells[0].text = nc[:30]
                                row_cells[1].text = f"{emp[nc].mean():,.2f}"
                                row_cells[2].text = f"{emp[nc].min():,.2f}"
                                row_cells[3].text = f"{emp[nc].max():,.2f}"
                                row_cells[4].text = f"{emp[nc].std():,.2f}"

                        # Footer
                        doc.add_paragraph("")
                        p = doc.add_paragraph()
                        run = p.add_run("This report is auto-generated by HR Analytics Platform - Risal Al-Wud IT")
                        run.font.size = Pt(8)
                        run.font.color.rgb = RGBColor(128, 128, 128)

                        doc_bytes = io.BytesIO()
                        doc.save(doc_bytes)

                        st.download_button("📥 تحميل Word", data=doc_bytes.getvalue(),
                            file_name=f"{word_title}_{datetime.now().strftime('%Y%m%d')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            type="primary", use_container_width=True)
                        st.success("✅ تم إنشاء التقرير بنجاح!")

                    except ImportError:
                        st.error("مكتبة python-docx غير مثبتة. أضف `python-docx` في requirements.txt")
                    except Exception as e:
                        st.error(f"خطأ: {e}")
            else:
                ibox("ارفع ملف بيانات أولاً من القائمة الجانبية.", "warning")

        elif page == "📊 تقرير شامل":
            hdr("📊 التقرير الشامل", "تصدير جميع التحليلات في ملف Excel واحد")

            if len(emp) > 0:
                st.markdown("### ⚙️ اختر محتوى التقرير")
                rpt_summary = st.checkbox("📊 ملخص تنفيذي", value=True, key="rpts")
                rpt_workforce = st.checkbox("👥 تحليل القوى العاملة", value=True, key="rptw")
                rpt_salary = st.checkbox("💰 تحليل الرواتب", value=True, key="rptsalx")
                rpt_recruit = st.checkbox("🎯 بيانات التوظيف", value=True, key="rptr2")
                rpt_training = st.checkbox("📚 بيانات التدريب", value=True, key="rptt")

                if st.button("📊 إنشاء التقرير الشامل", type="primary", key="rptbtn"):
                    ox = io.BytesIO()
                    with pd.ExcelWriter(ox, engine='xlsxwriter') as w:
                        num_cols = emp.select_dtypes('number').columns.tolist()
                        cat_cols = [c for c in emp.columns if emp[c].dtype=='object' and 1 < emp[c].nunique() < 30]

                        if rpt_summary:
                            summary_data = {"المؤشر": [], "القيمة": []}
                            summary_data["المؤشر"].append("إجمالي السجلات"); summary_data["القيمة"].append(len(emp))
                            summary_data["المؤشر"].append("الأعمدة"); summary_data["القيمة"].append(len(emp.columns))
                            for nc in num_cols[:5]:
                                summary_data["المؤشر"].append(f"متوسط {nc}"); summary_data["القيمة"].append(round(emp[nc].mean(), 2))
                                summary_data["المؤشر"].append(f"إجمالي {nc}"); summary_data["القيمة"].append(round(emp[nc].sum(), 2))
                            pd.DataFrame(summary_data).to_excel(w, sheet_name='ملخص تنفيذي', index=False)
                            w.sheets['ملخص تنفيذي'].right_to_left()

                        if rpt_workforce:
                            for cc in cat_cols[:6]:
                                safe_name = cc[:28]
                                vc = emp[cc].value_counts().reset_index()
                                vc.columns = [cc, "العدد"]
                                vc["النسبة %"] = (vc["العدد"] / len(emp) * 100).round(1)
                                vc.to_excel(w, sheet_name=safe_name, index=False)
                                try: w.sheets[safe_name].right_to_left()
                                except: pass

                        if rpt_salary and num_cols:
                            desc = emp[num_cols].describe().T.reset_index()
                            desc.columns = ["العمود","العدد","المتوسط","الانحراف","الأدنى","25%","الوسيط","75%","الأقصى"]
                            desc.to_excel(w, sheet_name='تحليل الرواتب', index=False)
                            w.sheets['تحليل الرواتب'].right_to_left()

                        if rpt_recruit:
                            if st.session_state.get('recruit_plans'):
                                pd.DataFrame(st.session_state.recruit_plans).to_excel(w, sheet_name='خطة التوظيف', index=False)
                                w.sheets['خطة التوظيف'].right_to_left()
                            if st.session_state.get('recruit_tracking'):
                                pd.DataFrame(st.session_state.recruit_tracking).to_excel(w, sheet_name='متابعة التوظيف', index=False)
                                w.sheets['متابعة التوظيف'].right_to_left()

                        if rpt_training and 'budget_data' in st.session_state:
                            pd.DataFrame(st.session_state.budget_data).to_excel(w, sheet_name='ميزانية التدريب', index=False)
                            w.sheets['ميزانية التدريب'].right_to_left()

                        # Raw data
                        emp.to_excel(w, sheet_name='البيانات الخام', index=False)

                    st.download_button("📥 تحميل التقرير الشامل", data=ox.getvalue(),
                        file_name=f"HR_Report_{datetime.now().strftime('%Y%m%d')}.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        type="primary", use_container_width=True)
                    st.success("✅ تم إنشاء التقرير الشامل!")
            else:
                ibox("ارفع ملف بيانات أولاً من القائمة الجانبية.", "warning")


    # =========================================
    #         📝 SURVEYS MODULE
    # =========================================
    elif section == "📝 الاستبيانات":

        if 'surveys_data' not in st.session_state:
            st.session_state.surveys_data = {}
        if 'survey_responses' not in st.session_state:
            st.session_state.survey_responses = []
        if 'custom_surveys' not in st.session_state:
            st.session_state.custom_surveys = {}

        if page == "📋 قوالب جاهزة":
            hdr("📋 قوالب الاستبيانات الجاهزة", "اختر قالب جاهز واملأ الاستبيان")

            template = st.selectbox("📝 اختر القالب:", list(SURVEY_TEMPLATES.keys()), key="sv_tmpl")
            tmpl = SURVEY_TEMPLATES[template]
            ibox(tmpl["description"])

            st.markdown("### 👤 بيانات المشارك")
            sv1, sv2, sv3 = st.columns(3)
            with sv1: sv_name = st.text_input("الاسم:", key="sv_name")
            with sv2: sv_dept = st.text_input("القسم:", key="sv_dept")
            with sv3: sv_date = st.date_input("التاريخ:", value=date.today(), key="sv_date")

            st.markdown(f"### 📝 {template}")
            st.info("قيّم كل عبارة من 1 (غير موافق تماماً) إلى 5 (موافق تماماً)")

            answers = {}
            for i, q_item in enumerate(tmpl["questions"]):
                answers[i] = st.slider(f"{i+1}. {q_item['q']}", 1, 5, 3, key=f"sv_q{i}")

            # Notification options
            st.markdown("### 📧 إشعار بالاستبيان")
            notif_c1, notif_c2, notif_c3 = st.columns(3)
            with notif_c1:
                sv_notify_email = st.text_input("بريد الموظف (للإشعار):", key="sv_notify_email", placeholder="employee@company.com")
            with notif_c2:
                sv_send_email = st.checkbox("📧 إرسال نتيجة بالبريد", value=False, key="sv_send_email")
            with notif_c3:
                sv_send_basecamp = st.checkbox("🏕️ إرسال إلى Basecamp", value=False, key="sv_send_bc")

            # Basecamp config
            if sv_send_basecamp:
                with st.expander("⚙️ إعدادات Basecamp"):
                    st.caption("أدخل Webhook URL من Basecamp (Chatbot Integration)")
                    bc_webhook = st.text_input("Basecamp Webhook URL:", key="bc_webhook",
                        value=st.session_state.get('basecamp_webhook',''),
                        placeholder="https://3.basecampapi.com/xxx/integrations/yyy/buckets/zzz/chats/www/lines.json")
                    if bc_webhook:
                        st.session_state.basecamp_webhook = bc_webhook

            if st.button("✅ إرسال الاستبيان", type="primary", key="sv_submit"):
                if sv_name:
                    response = {
                        "الاسم": sv_name, "القسم": sv_dept, "التاريخ": str(sv_date),
                        "القالب": template, "الإجابات": answers,
                        "المتوسط العام": round(sum(answers.values()) / len(answers), 2)
                    }
                    # Add category averages
                    cats = {}
                    for i, q_item in enumerate(tmpl["questions"]):
                        cat = q_item["cat"]
                        cats.setdefault(cat, []).append(answers[i])
                    response["التفاصيل"] = {c: round(sum(v)/len(v), 2) for c, v in cats.items()}
                    st.session_state.survey_responses.append(response)
                    st.success(f"✅ تم حفظ استبيان {sv_name} - المتوسط: {response['المتوسط العام']}/5")

                    # Send email notification
                    if sv_send_email and sv_notify_email:
                        smtp_cfg = st.session_state.get('smtp_config', {})
                        if not smtp_cfg.get('email'): smtp_cfg = load_smtp_config()
                        if smtp_cfg.get('email'):
                            details = "\n".join([f"- {c}: {v}/5" for c,v in response['التفاصيل'].items()])
                            ok, msg = send_test_email(sv_notify_email, sv_name,
                                [f"تم تسجيل نتيجة استبيان: {template}",
                                 f"المتوسط العام: {response['المتوسط العام']}/5",
                                 f"التفاصيل: {details}"],
                                str(sv_date), smtp_cfg.get('sender_name','HR'))
                            if ok: st.success(f"📧 تم إرسال النتيجة إلى {sv_notify_email}")

                    # Send to Basecamp
                    if sv_send_basecamp and st.session_state.get('basecamp_webhook'):
                        try:
                            import urllib.request
                            bc_msg = f"📝 استبيان جديد: {template}\n👤 {sv_name} ({sv_dept})\n⭐ المتوسط: {response['المتوسط العام']}/5\n📅 {sv_date}"
                            bc_data = json.dumps({"content": bc_msg}).encode('utf-8')
                            req = urllib.request.Request(st.session_state.basecamp_webhook,
                                data=bc_data, headers={'Content-Type': 'application/json'}, method='POST')
                            urllib.request.urlopen(req, timeout=10)
                            st.success("🏕️ تم الإرسال إلى Basecamp")
                        except Exception as e:
                            st.warning(f"⚠️ لم يتم الإرسال إلى Basecamp: {e}")

                    st.rerun()
                else:
                    st.error("يرجى إدخال الاسم")

        elif page == "🔨 بناء استبيان":
            hdr("🔨 بناء استبيان مخصص", "أنشئ استبيانك الخاص")

            st.markdown("### ⚙️ إعدادات الاستبيان")
            cs_name = st.text_input("اسم الاستبيان:", key="cs_name")
            cs_desc = st.text_input("الوصف:", key="cs_desc")

            st.markdown("### ➕ إضافة أسئلة")
            if 'custom_q_list' not in st.session_state:
                st.session_state.custom_q_list = []

            cq1, cq2 = st.columns([3,1])
            with cq1: new_q = st.text_input("السؤال:", key="cs_newq")
            with cq2: new_cat = st.text_input("التصنيف:", key="cs_newcat")

            if st.button("➕ إضافة سؤال", key="cs_addq"):
                if new_q:
                    st.session_state.custom_q_list.append({"q": new_q, "cat": new_cat or "عام"})
                    st.rerun()

            if st.session_state.custom_q_list:
                st.markdown("### 📋 الأسئلة المضافة")
                for i, cq in enumerate(st.session_state.custom_q_list):
                    st.write(f"{i+1}. {cq['q']} [{cq['cat']}]")

                if st.button("💾 حفظ الاستبيان", type="primary", key="cs_save"):
                    if cs_name:
                        st.session_state.custom_surveys[cs_name] = {
                            "description": cs_desc, "questions": st.session_state.custom_q_list.copy()
                        }
                        st.session_state.custom_q_list = []
                        st.success(f"✅ تم حفظ الاستبيان: {cs_name}")
                        st.rerun()

                if st.button("🗑️ مسح الأسئلة", key="cs_clear"):
                    st.session_state.custom_q_list = []
                    st.rerun()

            # Show saved custom surveys
            if st.session_state.custom_surveys:
                st.markdown("---")
                st.markdown("### 📂 الاستبيانات المخصصة المحفوظة")
                for name, survey in st.session_state.custom_surveys.items():
                    with st.expander(f"📝 {name} ({len(survey['questions'])} سؤال)"):
                        st.write(survey['description'])
                        for i, q in enumerate(survey['questions']):
                            st.write(f"{i+1}. {q['q']} [{q['cat']}]")

        elif page == "📊 تحليل النتائج":
            hdr("📊 تحليل نتائج الاستبيانات")

            if st.session_state.survey_responses:
                responses = st.session_state.survey_responses
                st.success(f"📊 إجمالي الاستجابات: {len(responses)}")

                # Summary table
                summary_rows = []
                for r in responses:
                    row = {"الاسم": r["الاسم"], "القسم": r["القسم"], "القالب": r["القالب"], "المتوسط": r["المتوسط العام"]}
                    summary_rows.append(row)
                sdf = pd.DataFrame(summary_rows)
                st.dataframe(sdf, use_container_width=True, hide_index=True)

                # KPIs
                k1,k2,k3,k4 = st.columns(4)
                avg_all = sdf["المتوسط"].mean()
                with k1: kpi("📊 المتوسط العام", f"{avg_all:.2f}/5")
                with k2: kpi("✅ الاستجابات", f"{len(responses)}")
                with k3: kpi("📈 أعلى تقييم", f"{sdf['المتوسط'].max():.2f}")
                with k4: kpi("📉 أقل تقييم", f"{sdf['المتوسط'].min():.2f}")

                # Charts
                ch1, ch2 = st.columns(2)
                with ch1:
                    fig = px.bar(sdf, x="الاسم", y="المتوسط", color="القالب", title="التقييم حسب المشارك", text_auto=".2f")
                    fig.add_hline(y=avg_all, line_dash="dash", annotation_text=f"المتوسط: {avg_all:.2f}")
                    fig.update_layout(font=dict(family="Noto Sans Arabic"), height=350)
                    st.plotly_chart(fig, use_container_width=True)
                with ch2:
                    if sdf["القسم"].nunique() > 1:
                        dept_avg = sdf.groupby("القسم")["المتوسط"].mean().reset_index()
                        fig = px.bar(dept_avg, x="القسم", y="المتوسط", title="المتوسط حسب القسم", text_auto=".2f", color_discrete_sequence=[CL['s']])
                        fig.update_layout(font=dict(family="Noto Sans Arabic"), height=350)
                        st.plotly_chart(fig, use_container_width=True)

                # Category breakdown
                st.markdown("### 📊 التحليل حسب التصنيف")
                all_cats = {}
                for r in responses:
                    if "التفاصيل" in r:
                        for cat, val in r["التفاصيل"].items():
                            all_cats.setdefault(cat, []).append(val)
                if all_cats:
                    cat_avg = {c: sum(v)/len(v) for c, v in all_cats.items()}
                    cat_df = pd.DataFrame({"التصنيف": cat_avg.keys(), "المتوسط": cat_avg.values()}).sort_values("المتوسط")
                    fig = px.bar(cat_df, x="المتوسط", y="التصنيف", orientation='h', title="المتوسط حسب التصنيف", text_auto=".2f", color_discrete_sequence=[CL['a']])
                    fig.update_layout(font=dict(family="Noto Sans Arabic"), height=400)
                    st.plotly_chart(fig, use_container_width=True)

                if st.button("🗑️ مسح جميع الاستجابات", key="sv_clr"):
                    st.session_state.survey_responses = []
                    st.rerun()
            else:
                ibox("لا توجد استجابات بعد. اذهب لصفحة القوالب الجاهزة واملأ استبيان.", "warning")




        elif page == "📥 تصدير الاستبيانات":
            hdr("📥 تصدير بيانات الاستبيانات")
            if st.session_state.survey_responses:
                ox = io.BytesIO()
                with pd.ExcelWriter(ox, engine='xlsxwriter') as w:
                    rows = []
                    for r in st.session_state.survey_responses:
                        row = {"الاسم": r["الاسم"], "القسم": r["القسم"], "التاريخ": r["التاريخ"], "القالب": r["القالب"], "المتوسط": r["المتوسط العام"]}
                        if "التفاصيل" in r:
                            row.update(r["التفاصيل"])
                        rows.append(row)
                    pd.DataFrame(rows).to_excel(w, sheet_name='الاستجابات', index=False)
                    w.sheets['الاستجابات'].right_to_left()
                st.download_button("📥 تحميل", data=ox.getvalue(), file_name=f"Surveys_{datetime.now().strftime('%Y%m%d')}.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", type="primary", use_container_width=True)
            else:
                ibox("لا توجد بيانات للتصدير.", "warning")


    # =========================================
    #       🧠 PERSONALITY TESTS MODULE
    # =========================================
    elif section == "🧠 اختبارات الشخصية":

        # Load from database
        if 'personality_results' not in st.session_state:
            st.session_state.personality_results = db_load_results()
        if 'test_assignments' not in st.session_state:
            st.session_state.test_assignments = db_load_assignments()

        ALL_TESTS = ["Big Five (OCEAN)", "Thomas PPA", "Hogan HPI", "MBTI", "DISC"]

        def save_test_result(result):
            """Save to both session_state and database"""
            created_by = st.session_state.get('user_name', '')
            rid = db_save_result(result, created_by)
            result["id"] = rid
            st.session_state.personality_results.append(result)
            st.session_state['_last_test_result'] = result
            return rid

        def calc_big5(answers):
            scores = {}; counts = {}
            for i, q in enumerate(BIG5_QUESTIONS):
                t = q["trait"]
                s = answers[i] if q["d"] == 1 else (6 - answers[i])
                scores[t] = scores.get(t, 0) + s
                counts[t] = counts.get(t, 0) + 1
            return {t: round(scores[t] / (counts[t] * 5) * 100) for t in scores}

        def calc_thomas(answers):
            scores = {}; counts = {}
            for i, q in enumerate(THOMAS_QUESTIONS):
                s = q["scale"]
                scores[s] = scores.get(s, 0) + answers[i]
                counts[s] = counts.get(s, 0) + 1
            return {s: round(scores[s] / (counts[s] * 5) * 100) for s in scores}

        def calc_hogan(answers):
            scores = {}; counts = {}
            for i, q in enumerate(HOGAN_QUESTIONS):
                s = q["scale"]
                val = answers[i] if q["d"] == 1 else (6 - answers[i])
                scores[s] = scores.get(s, 0) + val
                counts[s] = counts.get(s, 0) + 1
            return {s: round(scores[s] / (counts[s] * 5) * 100) for s in scores}

        def calc_mbti(answers):
            dim_scores = {"E": 0, "I": 0, "S": 0, "N": 0, "T": 0, "F": 0, "J": 0, "P": 0}
            for i, q in enumerate(MBTI_QUESTIONS):
                dim_scores[q["d"]] += answers[i]
                opp = {"E":"I","I":"E","S":"N","N":"S","T":"F","F":"T","J":"P","P":"J"}[q["d"]]
                dim_scores[opp] += (6 - answers[i])
            mbti_type = ""
            dims_pct = {}
            for dim_pair in ["EI","SN","TF","JP"]:
                a, b = dim_pair[0], dim_pair[1]
                total = dim_scores[a] + dim_scores[b]
                pct_a = round(dim_scores[a] / max(total,1) * 100)
                pct_b = 100 - pct_a
                winner = a if dim_scores[a] >= dim_scores[b] else b
                mbti_type += winner
                dims_pct[a] = pct_a
                dims_pct[b] = pct_b
            return mbti_type, dims_pct

        def calc_disc(answers):
            scores = {}; counts = {}
            for i, q in enumerate(DISC_QUESTIONS):
                s = q["style"]
                scores[s] = scores.get(s, 0) + answers[i]
                counts[s] = counts.get(s, 0) + 1
            return {s: round(scores[s] / (counts[s] * 5) * 100) for s in scores}

        def render_test(test_name, emp_name, emp_dept, is_mandatory=False, assigned_by=""):
            """Render any test and return result"""
            if test_name == "Big Five (OCEAN)":
                questions = BIG5_QUESTIONS
                st.info("قيّم كل عبارة: 1 (غير موافق تماماً) - 5 (موافق تماماً)")
                answers = {}
                for i, q in enumerate(questions):
                    trait_info = BIG5_TRAITS[q["trait"]]
                    answers[i] = st.slider(f'{i+1}. {q["q"]}', 1, 5, 3, key=f"t_{test_name}_{i}")
                if st.button("✅ حساب النتائج", type="primary", key=f"btn_{test_name}"):
                    scores = calc_big5(answers)
                    result = {"type": "Big Five", "الاسم": emp_name, "القسم": emp_dept,
                        "التاريخ": str(date.today()), "scores": scores,
                        "إجباري": is_mandatory, "معيّن_بواسطة": assigned_by}
                    save_test_result(result)
                    st.markdown("---")
                    cols = st.columns(5)
                    for j, (t, pct) in enumerate(scores.items()):
                        info = BIG5_TRAITS[t]
                        with cols[j]:
                            kpi(info["name"], f"{pct}%")
                            level = "مرتفع" if pct >= 70 else ("متوسط" if pct >= 40 else "منخفض")
                            st.caption(f"{level}: {info['high'] if pct>=60 else info['low']}")
                    fig = go.Figure()
                    t_list = [BIG5_TRAITS[t]["name"] for t in scores]
                    v_list = list(scores.values()) + [list(scores.values())[0]]
                    fig.add_trace(go.Scatterpolar(r=v_list, theta=t_list+[t_list[0]], fill='toself', line_color=CL['p'], fillcolor='rgba(15,76,92,0.2)'))
                    fig.update_layout(polar=dict(radialaxis=dict(range=[0,100])), title=f"Big Five - {emp_name}", font=dict(family="Noto Sans Arabic"), height=450, showlegend=False)
                    st.plotly_chart(fig, use_container_width=True)
                    st.markdown("### 🎯 التوصيات المهنية")
                    top2 = sorted(scores.items(), key=lambda x: x[1], reverse=True)[:2]
                    low2 = sorted(scores.items(), key=lambda x: x[1])[:2]
                    for t, pct in top2:
                        info = BIG5_TRAITS[t]
                        ibox(f"**{info['name']} ({pct}%):** الوظائف المناسبة: {info['jobs_high']}", "success")

                    # === ENHANCED: 5 sections ===
                    st.markdown("---")
                    # Personality Summary
                    top_names = [BIG5_TRAITS[t]["name"] for t,_ in top2]
                    low_names = [BIG5_TRAITS[t]["name"] for t,_ in low2]
                    summary_lines = [
                        f"**{emp_name}** يتميز بمستوى مرتفع في {top_names[0]} و{top_names[1]}، مما يشير إلى شخصية {'اجتماعية ومنفتحة' if 'E' in [t for t,_ in top2] else 'تحليلية ومنظمة'}.",
                        f"يميل إلى {BIG5_TRAITS[top2[0][0]]['high']}، مع قدرة على {BIG5_TRAITS[top2[1][0]]['high']}.",
                        f"المجالات التي تحتاج تطوير تشمل {low_names[0]} ({low2[0][1]}%) و{low_names[1]} ({low2[1][1]}%).",
                        f"في بيئة العمل، يُفضل المهام التي تتطلب {BIG5_TRAITS[top2[0][0]]['jobs_high'].split('،')[0]}.",
                        f"يُنصح بتوجيهه نحو أدوار تستثمر نقاط قوته مع تطوير {'المرونة والتكيف' if low2[0][0] in ['O','E'] else 'التنظيم والانضباط'}."
                    ]
                    st.markdown("### 📝 ملخص الشخصية")
                    for line in summary_lines:
                        st.markdown(f"- {line}")

                    st.markdown("### 💪 نقاط القوة")
                    for t, pct in top2:
                        info = BIG5_TRAITS[t]
                        st.success(f"✅ **{info['name']} ({pct}%):** {info['high']}")

                    st.markdown("### ⚠️ نقاط الضعف")
                    for t, pct in low2:
                        info = BIG5_TRAITS[t]
                        st.warning(f"⚠️ **{info['name']} ({pct}%):** {info['low']}")

                    st.markdown("### 💼 توصيات مهنية")
                    career_tips = {
                        "O": "ابحث عن بيئات عمل تشجع الابتكار والتجريب. انضم لمشاريع جديدة وفرق تطوير المنتجات.",
                        "C": "تولى مهام تتطلب الدقة والتنظيم. أنت مناسب لإدارة المشاريع وضمان الجودة.",
                        "E": "استثمر قدرتك على التواصل في أدوار قيادية أو مبيعات أو تدريب.",
                        "A": "أنت مناسب للموارد البشرية وحل النزاعات وبناء الفرق.",
                        "N": "استثمر حساسيتك في الأدوار الإبداعية والاستشارية. تجنب البيئات عالية الضغط."
                    }
                    for t, pct in top2:
                        st.info(f"💼 {career_tips.get(t, '')}")

                    st.markdown("### 🌱 توصيات شخصية")
                    personal_tips = {
                        "O": "جرب هوايات إبداعية جديدة. اقرأ في مجالات متنوعة خارج تخصصك.",
                        "C": "ضع أهدافاً يومية واحتفل بإنجازها. نظم بيئتك الشخصية لتعزيز إنتاجيتك.",
                        "E": "وسع شبكة علاقاتك. شارك في فعاليات اجتماعية ومهنية بانتظام.",
                        "A": "مارس التعاطف الذاتي. ساعد الآخرين لكن لا تنسَ حدودك الشخصية.",
                        "N": "مارس تقنيات الاسترخاء والتأمل. ابنِ روتيناً يومياً يعزز استقرارك النفسي."
                    }
                    for t, pct in sorted(scores.items(), key=lambda x: x[1], reverse=True)[:3]:
                        st.info(f"🌱 {personal_tips.get(t, '')}")

            elif test_name == "Thomas PPA":
                questions = THOMAS_QUESTIONS
                st.info("قيّم كل عبارة: 1 (لا تنطبق) - 5 (تنطبق تماماً)")
                answers = {}
                for i, q in enumerate(questions):
                    answers[i] = st.slider(f'{i+1}. {q["q"]}', 1, 5, 3, key=f"t_thomas_{i}")
                if st.button("✅ حساب النتائج", type="primary", key="btn_thomas"):
                    scores = calc_thomas(answers)
                    dominant = max(scores, key=scores.get)
                    result = {"type": "Thomas PPA", "الاسم": emp_name, "القسم": emp_dept,
                        "التاريخ": str(date.today()), "scores": scores, "dominant": dominant,
                        "إجباري": is_mandatory, "معيّن_بواسطة": assigned_by}
                    save_test_result(result)
                    st.markdown("---")
                    cols = st.columns(4)
                    for j, (s, pct) in enumerate(scores.items()):
                        info = THOMAS_SCALES[s]
                        with cols[j]:
                            st.markdown(f"<div style='background:{info['color']};color:white;border-radius:12px;padding:16px;text-align:center;margin-bottom:10px'><p style='font-size:11px;opacity:.7;margin:0'>{info['name']}</p><h3 style='font-size:24px;margin:6px 0;font-weight:800'>{pct}%</h3></div>", unsafe_allow_html=True)
                    dom_info = THOMAS_SCALES[dominant]
                    st.markdown(f"### 🏆 النمط السائد: {dom_info['name']} ({dom_info['en']})")
                    dc1, dc2, dc3 = st.columns(3)
                    with dc1: ibox(f"**الوصف:** {dom_info['high']}")
                    with dc2: ibox(f"**الأدوار المناسبة:** {dom_info['role']}", "success")
                    with dc3: ibox(f"**كيف تديره:** {dom_info['manage']}", "warning")
                    # Thomas profile chart
                    thomas_df = pd.DataFrame({"المقياس": [THOMAS_SCALES[s]["name"] for s in scores], "النسبة": list(scores.values())})
                    fig = px.bar(thomas_df, x="المقياس", y="النسبة", text_auto=True,
                        color="المقياس", color_discrete_map={THOMAS_SCALES[s]["name"]: THOMAS_SCALES[s]["color"] for s in THOMAS_SCALES})
                    fig.update_layout(title=f"Thomas PPA - {emp_name}", font=dict(family="Noto Sans Arabic"), height=400, showlegend=False)
                    st.plotly_chart(fig, use_container_width=True)

                    # === ENHANCED: Full analysis ===
                    st.markdown("---")
                    sorted_scores = sorted(scores.items(), key=lambda x: x[1], reverse=True)
                    top_s, sec_s = sorted_scores[0], sorted_scores[1]
                    low_s = sorted_scores[-1]
                    dom_i, sec_i, low_i = THOMAS_SCALES[top_s[0]], THOMAS_SCALES[sec_s[0]], THOMAS_SCALES[low_s[0]]

                    st.markdown("### 📝 ملخص الشخصية")
                    summary = [
                        f"**{emp_name}** يتميز بنمط {dom_i['name']} ({dom_i['en']}) كنمط سائد بنسبة {top_s[1]}%.",
                        f"يجمع بين {dom_i['desc'].split('،')[0]} و{sec_i['desc'].split('،')[0]} مما يجعله {dom_i['high'].split('،')[0]}.",
                        f"في بيئة العمل يميل إلى {dom_i['role'].split('،')[0]}، ويفضل {dom_i['communicate'].split('،')[0] if 'communicate' in dom_i else 'بيئة منظمة'}.",
                        f"المجال الأضعف هو {low_i['name']} ({low_s[1]}%)، مما يعني أنه قد يحتاج دعماً في {low_i['high'].split('،')[0]}.",
                        f"يُنصح بتوجيهه نحو أدوار {dom_i['role']} مع تطوير مهارات {low_i['name']}."
                    ]
                    for line in summary: st.markdown(f"- {line}")

                    st.markdown("### 💪 نقاط القوة")
                    st.success(f"✅ **{dom_i['name']}:** {dom_i['strengths']}")
                    st.success(f"✅ **{sec_i['name']}:** {sec_i['strengths']}")

                    st.markdown("### ⚠️ نقاط الضعف")
                    st.warning(f"⚠️ **{low_i['name']} ({low_s[1]}%):** {low_i['challenges']}")

                    st.markdown("### 💼 توصيات مهنية")
                    st.info(f"💼 **الأدوار المناسبة:** {dom_i['role']}")
                    st.info(f"💼 **الوظائف المقترحة:** {dom_i['careers']}")
                    st.info(f"💼 **أسلوب الإدارة المثالي:** {dom_i['manage']}")

                    st.markdown("### 🌱 توصيات شخصية")
                    thomas_personal = {"D":"طور مهارات الاستماع والصبر. مارس التفويض وثق بقدرات فريقك.",
                        "I":"درب نفسك على التنظيم والمتابعة. ضع قوائم مهام يومية والتزم بها.",
                        "S":"تعلم تقبل التغيير تدريجياً. مارس المبادرة واخرج من منطقة الراحة.",
                        "C":"تقبل أن الكمال ليس دائماً ممكناً. طور مهاراتك الاجتماعية والعاطفية."}
                    for s, pct in sorted_scores[:2]:
                        st.info(f"🌱 **{THOMAS_SCALES[s]['name']}:** {thomas_personal.get(s,'')}")

            elif test_name == "Hogan HPI":
                questions = HOGAN_QUESTIONS
                st.info("قيّم كل عبارة: 1 (لا تنطبق) - 5 (تنطبق تماماً)")
                answers = {}
                for i, q in enumerate(questions):
                    answers[i] = st.slider(f'{i+1}. {q["q"]}', 1, 5, 3, key=f"t_hogan_{i}")
                if st.button("✅ حساب النتائج", type="primary", key="btn_hogan"):
                    scores = calc_hogan(answers)
                    result = {"type": "Hogan HPI", "الاسم": emp_name, "القسم": emp_dept,
                        "التاريخ": str(date.today()), "scores": scores,
                        "إجباري": is_mandatory, "معيّن_بواسطة": assigned_by}
                    save_test_result(result)
                    st.markdown("---")
                    st.markdown(f"### 📊 نتائج Hogan HPI - {emp_name}")
                    # Progress bars with detailed analysis
                    for s, pct in scores.items():
                        info = HOGAN_SCALES[s]
                        level = "مرتفع" if pct >= 70 else ("متوسط" if pct >= 40 else "منخفض")
                        st.progress(pct/100, text=f"{info['name']} ({info['en']}): {pct}% - {level}")
                        with st.expander(f"📋 تفاصيل {info['name']}"):
                            st.write(f"**الوصف:** {info['desc']}")
                            st.write(f"**{'السمات (مرتفع)' if pct >= 60 else 'السمات (منخفض)'}:** {info['high'] if pct >= 60 else info['low']}")
                            st.write(f"**الأثر المهني:** {info['impact']}")
                    # Radar chart
                    fig = go.Figure()
                    h_names = [HOGAN_SCALES[s]["name"] for s in scores]
                    h_vals = list(scores.values()) + [list(scores.values())[0]]
                    fig.add_trace(go.Scatterpolar(r=h_vals, theta=h_names+[h_names[0]], fill='toself', line_color='#1ABC9C', fillcolor='rgba(26,188,156,0.2)'))
                    fig.update_layout(polar=dict(radialaxis=dict(range=[0,100])), title=f"Hogan HPI - {emp_name}", font=dict(family="Noto Sans Arabic"), height=500, showlegend=False)
                    st.plotly_chart(fig, use_container_width=True)

                    # === ENHANCED: Full analysis ===
                    st.markdown("---")
                    h_sorted = sorted(scores.items(), key=lambda x: x[1], reverse=True)
                    h_top2, h_low2 = h_sorted[:2], h_sorted[-2:]

                    st.markdown("### 📝 ملخص الشخصية")
                    ht1, ht2 = HOGAN_SCALES[h_top2[0][0]], HOGAN_SCALES[h_top2[1][0]]
                    hl1, hl2 = HOGAN_SCALES[h_low2[0][0]], HOGAN_SCALES[h_low2[1][0]]
                    h_summary = [
                        f"**{emp_name}** يتميز بمستوى مرتفع في {ht1['name']} ({h_top2[0][1]}%) و{ht2['name']} ({h_top2[1][1]}%).",
                        f"هذا يعني أنه {ht1['high'].split('،')[0]} مع قدرة على {ht2['high'].split('،')[0]}.",
                        f"الأثر المهني: {ht1['impact']}",
                        f"المجالات التي تحتاج تطوير: {hl1['name']} ({h_low2[0][1]}%) و{hl2['name']} ({h_low2[1][1]}%).",
                        f"يُنصح بالتركيز على تطوير {hl1['desc'].split('،')[0]} لتحقيق توازن أفضل في الأداء المهني."
                    ]
                    for line in h_summary: st.markdown(f"- {line}")

                    st.markdown("### 💪 نقاط القوة")
                    for s, pct in h_top2:
                        info = HOGAN_SCALES[s]
                        st.success(f"✅ **{info['name']} ({pct}%):** {info['high']}")

                    st.markdown("### ⚠️ نقاط الضعف")
                    for s, pct in h_low2:
                        info = HOGAN_SCALES[s]
                        st.warning(f"⚠️ **{info['name']} ({pct}%):** {info['low']}")

                    st.markdown("### 💼 توصيات مهنية")
                    for s, pct in h_top2:
                        info = HOGAN_SCALES[s]
                        st.info(f"💼 **{info['name']}:** {info['impact']}")
                    hogan_careers = {"ADJ":"إدارة الأزمات، القيادة تحت الضغط",
                        "AMB":"الإدارة التنفيذية، ريادة الأعمال، إدارة المشاريع",
                        "SOC":"المبيعات، العلاقات العامة، التدريب، الموارد البشرية",
                        "INT":"الدبلوماسية، التفاوض، خدمة العملاء VIP",
                        "PRU":"التحليل المالي، المراجعة، إدارة المخاطر",
                        "INQ":"البحث والتطوير، الابتكار، التحليل الاستراتيجي",
                        "LRN":"التعليم والتدريب، الاستشارات، إدارة المعرفة"}
                    st.info(f"💼 **وظائف مقترحة:** {hogan_careers.get(h_top2[0][0], 'متعدد المجالات')}")

                    st.markdown("### 🌱 توصيات شخصية")
                    hogan_personal = {"ADJ":"مارس تقنيات إدارة التوتر. ابنِ روتيناً صحياً يعزز استقرارك النفسي.",
                        "AMB":"ضع أهدافاً طويلة المدى وقسمها لخطوات. احتفل بالإنجازات الصغيرة.",
                        "SOC":"وسع شبكة علاقاتك المهنية. شارك في فعاليات ومؤتمرات بانتظام.",
                        "INT":"طور مهارات الذكاء العاطفي. تدرب على قراءة لغة الجسد.",
                        "PRU":"وازن بين الحذر والمخاطرة المحسوبة. لا تدع الخوف يمنعك من الفرص.",
                        "INQ":"خصص وقتاً للقراءة والتعلم الذاتي. جرب مجالات معرفية جديدة.",
                        "LRN":"شارك معرفتك مع الآخرين. ابدأ مدونة أو قناة تعليمية."}
                    for s, pct in h_top2:
                        st.info(f"🌱 **{HOGAN_SCALES[s]['name']}:** {hogan_personal.get(s,'')}")

            elif test_name == "MBTI":
                questions = MBTI_QUESTIONS
                st.info("قيّم كل عبارة: 1 (لا تنطبق) - 5 (تنطبق تماماً)")
                answers = {}
                for i, q in enumerate(questions):
                    answers[i] = st.slider(f'{i+1}. {q["q"]}', 1, 5, 3, key=f"t_mbti_{i}")
                if st.button("✅ حساب النتائج", type="primary", key="btn_mbti"):
                    mbti_type, dims_pct = calc_mbti(answers)
                    result = {"type": "MBTI", "الاسم": emp_name, "القسم": emp_dept,
                        "التاريخ": str(date.today()), "scores": dims_pct, "mbti_type": mbti_type,
                        "إجباري": is_mandatory, "معيّن_بواسطة": assigned_by}
                    save_test_result(result)
                    st.markdown("---")
                    type_info = MBTI_TYPES.get(mbti_type, {"name": mbti_type, "desc": "", "strengths": "", "careers": ""})
                    st.markdown(f"<div style='text-align:center;padding:20px;background:linear-gradient(135deg,#0F4C5C,#1A1A2E);border-radius:14px;color:white;margin-bottom:20px'><h1 style='color:white;font-size:48px;margin:0'>{mbti_type}</h1><h3 style='color:#E9C46A;margin:8px 0'>{type_info['name']}</h3><p style='opacity:.8;margin:0'>{type_info['desc']}</p></div>", unsafe_allow_html=True)
                    # Dimension bars
                    for dim_pair in ["EI","SN","TF","JP"]:
                        a, b = dim_pair[0], dim_pair[1]
                        dim_info = MBTI_DIMS[dim_pair]
                        pct_a = dims_pct.get(a, 50)
                        col1, col2, col3 = st.columns([2,6,2])
                        with col1: st.markdown(f"**{dim_info[a]}**\n\n{pct_a}%")
                        with col2: st.progress(pct_a/100)
                        with col3: st.markdown(f"**{dim_info[b]}**\n\n{100-pct_a}%")
                    st.markdown("### 🎯 التحليل المهني")
                    mc1, mc2 = st.columns(2)
                    with mc1: ibox(f"**نقاط القوة:** {type_info['strengths']}", "success")
                    with mc2: ibox(f"**المسارات المهنية:** {type_info['careers']}")

                    # === ENHANCED: Full analysis ===
                    st.markdown("---")
                    # Weaknesses map
                    mbti_weaknesses = {
                        "I":"قد يواجه صعوبة في المبادرة الاجتماعية والعمل الجماعي المكثف",
                        "E":"قد يفتقر للتركيز في العمل الفردي ويتشتت بالتفاعلات",
                        "S":"قد يقاوم التغيير ويفتقر للرؤية الاستراتيجية طويلة المدى",
                        "N":"قد يتجاهل التفاصيل العملية ويبالغ في التنظير",
                        "T":"قد يفتقر للحساسية العاطفية والتعاطف مع الآخرين",
                        "F":"قد يتأثر بالعواطف في اتخاذ القرارات ويتجنب المواجهة",
                        "J":"قد يكون جامداً ويقاوم التغييرات غير المخططة",
                        "P":"قد يماطل في إنجاز المهام ويفتقر للتنظيم"
                    }
                    mbti_personal_tips = {
                        "I":"خصص وقتاً لشحن طاقتك بالانعزال، لكن تحدَّ نفسك بمشاركة اجتماعية أسبوعية.",
                        "E":"تعلم فن الاستماع النشط. خصص وقتاً هادئاً للتأمل والتفكير.",
                        "S":"جرب التفكير في الصورة الكبيرة أحياناً. اقرأ عن المستقبليات.",
                        "N":"درب نفسك على الانتباه للتفاصيل. ضع قوائم مهام عملية.",
                        "T":"مارس التعبير عن مشاعرك. اسأل الآخرين عن شعورهم.",
                        "F":"تعلم فصل القرارات المهنية عن العواطف الشخصية.",
                        "J":"تقبل أن الخطط تتغير. جرب المرونة في روتينك اليومي.",
                        "P":"استخدم تقنيات إدارة الوقت. ضع مواعيد نهائية ذاتية."
                    }

                    st.markdown("### 📝 ملخص الشخصية")
                    dims_text = {"E":"منفتح اجتماعياً","I":"متأمل داخلياً","S":"واقعي عملي","N":"حدسي مبتكر",
                        "T":"منطقي تحليلي","F":"عاطفي متعاطف","J":"منظم مخطط","P":"مرن تلقائي"}
                    m_summary = [
                        f"**{emp_name}** من نمط **{mbti_type}** ({type_info['name']})، وهو {type_info['desc']}.",
                        f"يتميز بكونه {dims_text.get(mbti_type[0],'')} و{dims_text.get(mbti_type[1],'')}.",
                        f"في اتخاذ القرارات يميل لأن يكون {dims_text.get(mbti_type[2],'')}، وفي أسلوب حياته {dims_text.get(mbti_type[3],'')}.",
                        f"نقاط قوته الرئيسية: {type_info['strengths']}.",
                        f"المسارات المهنية المناسبة تشمل {type_info['careers']}."
                    ]
                    for line in m_summary: st.markdown(f"- {line}")

                    st.markdown("### 💪 نقاط القوة")
                    st.success(f"✅ {type_info['strengths']}")

                    st.markdown("### ⚠️ نقاط الضعف")
                    for letter in mbti_type:
                        if letter in mbti_weaknesses:
                            st.warning(f"⚠️ **{dims_text.get(letter,letter)}:** {mbti_weaknesses[letter]}")

                    st.markdown("### 💼 توصيات مهنية")
                    st.info(f"💼 **المسارات المهنية المثالية:** {type_info['careers']}")
                    mbti_career_advice = {"I":"ابحث عن أدوار تتيح العمل المستقل مع تفاعل محدود.",
                        "E":"اختر أدواراً تفاعلية مثل المبيعات أو التدريب أو القيادة.",
                        "T":"تميز في الأدوار التحليلية والاستشارية.",
                        "F":"ابرع في أدوار الرعاية والتطوير البشري.",
                        "J":"تفوق في إدارة المشاريع والتخطيط الاستراتيجي.",
                        "P":"اختر بيئات مرنة تسمح بالإبداع والتجريب."}
                    for letter in [mbti_type[0], mbti_type[2]]:
                        if letter in mbti_career_advice:
                            st.info(f"💼 {mbti_career_advice[letter]}")

                    st.markdown("### 🌱 توصيات شخصية")
                    for letter in mbti_type:
                        if letter in mbti_personal_tips:
                            st.info(f"🌱 **{dims_text.get(letter,letter)}:** {mbti_personal_tips[letter]}")

            elif test_name == "DISC":
                questions = DISC_QUESTIONS
                st.info("قيّم كل عبارة: 1 (لا تنطبق أبداً) - 5 (تنطبق تماماً)")
                answers = {}
                for i, q in enumerate(questions):
                    answers[i] = st.slider(f'{i+1}. {q["q"]}', 1, 5, 3, key=f"t_disc_{i}")
                if st.button("✅ حساب النتائج", type="primary", key="btn_disc"):
                    scores = calc_disc(answers)
                    dominant = max(scores, key=scores.get)
                    secondary = sorted(scores, key=scores.get, reverse=True)[1]
                    result = {"type": "DISC", "الاسم": emp_name, "القسم": emp_dept,
                        "التاريخ": str(date.today()), "scores": scores, "dominant": dominant,
                        "secondary": secondary, "إجباري": is_mandatory, "معيّن_بواسطة": assigned_by}
                    save_test_result(result)
                    st.markdown("---")
                    # Colored KPIs
                    cols = st.columns(4)
                    for j, (s, pct) in enumerate(scores.items()):
                        info = DISC_STYLES[s]
                        badge = " 🏆" if s == dominant else ""
                        with cols[j]:
                            st.markdown(f"<div style='background:{info['color']};color:white;border-radius:12px;padding:16px;text-align:center;margin-bottom:10px'><p style='font-size:11px;opacity:.7;margin:0'>{info['name']} ({info['en']}){badge}</p><h3 style='font-size:28px;margin:6px 0;font-weight:800'>{pct}%</h3></div>", unsafe_allow_html=True)
                    # Dominant + Secondary
                    dom_info = DISC_STYLES[dominant]
                    sec_info = DISC_STYLES[secondary]
                    st.markdown(f"### 🏆 النمط السائد: {dom_info['name']} ({dom_info['en']}) | الثانوي: {sec_info['name']}")
                    st.markdown(f"**النمط المركب:** {dominant}{secondary} - يجمع بين {dom_info['desc'].split('،')[0]} و{sec_info['desc'].split('،')[0]}")
                    # Detailed analysis
                    t1, t2, t3, t4 = st.tabs(["📋 الوصف","💪 نقاط القوة","⚠️ التحديات","🎯 التوصيات"])
                    with t1:
                        ibox(f"**السمات الرئيسية:** {dom_info['high']}")
                        ibox(f"**الوصف:** {dom_info['desc']}")
                    with t2:
                        ibox(f"**نقاط القوة:** {dom_info['strengths']}", "success")
                    with t3:
                        ibox(f"**التحديات:** {dom_info['challenges']}", "warning")
                    with t4:
                        ibox(f"**الوظائف المناسبة:** {dom_info['careers']}", "success")
                        ibox(f"**كيف تديره:** {dom_info['manage']}")
                        ibox(f"**كيف تتواصل معه:** {dom_info['communicate']}")

                    # === ENHANCED: Summary + Personal tips ===
                    st.markdown("---")
                    st.markdown("### 📝 ملخص الشخصية")
                    d_summary = [
                        f"**{emp_name}** يتميز بنمط {dom_info['name']} ({dom_info['en']}) كنمط سائد بنسبة {scores[dominant]}%.",
                        f"النمط الثانوي هو {sec_info['name']} ({scores[secondary]}%)، مما يشكل نمطاً مركباً {dominant}{secondary}.",
                        f"يتصف بأنه {dom_info['desc']}، مع ميل ثانوي نحو {sec_info['desc'].split('،')[0]}.",
                        f"نقاط قوته الرئيسية: {dom_info['strengths']}.",
                        f"يحتاج لتطوير: {dom_info['challenges']}."
                    ]
                    for line in d_summary: st.markdown(f"- {line}")

                    st.markdown("### 💼 توصيات مهنية")
                    st.info(f"💼 **الوظائف المثالية:** {dom_info['careers']}")
                    disc_career_extra = {"D":"ضع أهدافاً تنافسية واضحة. أنت تتميز في بيئات سريعة النتائج.",
                        "I":"استثمر قدرتك على التأثير في أدوار تتطلب إقناع وتحفيز الآخرين.",
                        "S":"ابحث عن بيئات مستقرة مع فريق داعم. أنت العمود الفقري لأي فريق.",
                        "C":"تميز في الأدوار التي تتطلب دقة وتحليل. الجودة هي ميزتك التنافسية."}
                    st.info(f"💼 {disc_career_extra.get(dominant,'')}")

                    st.markdown("### 🌱 توصيات شخصية")
                    disc_personal = {"D":"تدرب على الصبر والاستماع. الفوز ليس كل شيء - العلاقات مهمة أيضاً.",
                        "I":"نظم وقتك بشكل أفضل. الحماس رائع لكن المتابعة هي المفتاح.",
                        "S":"تقبل التغيير كفرصة وليس تهديداً. جرب شيئاً جديداً كل أسبوع.",
                        "C":"تقبل أن 80% أحياناً كافٍ. الكمال عدو الإنجاز."}
                    st.info(f"🌱 **{dom_info['name']}:** {disc_personal.get(dominant,'')}")
                    st.info(f"🌱 **{sec_info['name']}:** {disc_personal.get(secondary,'')}")
                    # Chart
                    disc_df = pd.DataFrame({"النمط": [DISC_STYLES[s]["name"] for s in scores], "النسبة": list(scores.values())})
                    fig = px.bar(disc_df, x="النمط", y="النسبة", text_auto=True,
                        color="النمط", color_discrete_map={DISC_STYLES[s]["name"]: DISC_STYLES[s]["color"] for s in DISC_STYLES})
                    fig.update_layout(title=f"DISC Profile - {emp_name}", font=dict(family="Noto Sans Arabic"), height=400, showlegend=False)
                    st.plotly_chart(fig, use_container_width=True)

            # ===== AI DEEP ANALYSIS (for all tests) =====
            # Check if any result was just saved
            if st.session_state.get('_last_test_result'):
                last_r = st.session_state['_last_test_result']
                st.markdown("---")
                if st.button("🤖 تحليل معمّق بالذكاء الاصطناعي", type="secondary", use_container_width=True, key=f"ai_deep_{test_name}"):
                    with st.spinner("جاري التحليل العميق بالذكاء الاصطناعي..."):
                        scores_txt = json.dumps(last_r.get('scores',{}), ensure_ascii=False)
                        extra = ""
                        if last_r.get('mbti_type'): extra = f"\nنمط MBTI: {last_r['mbti_type']}"
                        if last_r.get('dominant'): extra = f"\nالنمط السائد: {last_r['dominant']}"

                        ai_prompt = f"""أنت خبير في علم النفس التنظيمي وتحليل الشخصيات المهنية.

اختبار: {test_name}
الموظف: {emp_name}
القسم: {emp_dept}
النتائج: {scores_txt}{extra}

قدم تحليلاً معمقاً يشمل:

## التحليل النفسي المهني
(5 أسطر تحليل عميق للشخصية المهنية بناءً على النتائج)

## نقاط القوة الخفية
(3 نقاط قوة غير واضحة يمكن استنتاجها من تركيبة النتائج)

## المخاطر المهنية
(3 مخاطر محتملة يجب الانتباه لها بناءً على نمط الشخصية)

## خطة التطوير المقترحة
(خطة من 3 أشهر مع أهداف محددة وقابلة للقياس)

## التوافق مع الأدوار في شركة تقنية معلومات
(أفضل 3 أدوار في شركة IT وسبب التوافق)

## نصيحة للمدير المباشر
(كيف يدير هذا الموظف بفعالية في 3 نقاط)

أجب بالعربية بأسلوب مهني."""

                        response, error = call_ai_api(ai_prompt, ai_prompt, model_type="hr")
                        if response:
                            st.markdown("### 🤖 التحليل المعمّق بالذكاء الاصطناعي")
                            st.markdown(response)
                        elif error:
                            st.warning(f"تعذر التحليل: {error}. تأكد من إعداد مفتاح API في الإعدادات.")

        # ===== PAGE: Test Assignment =====
        if page == "📋 تعيين الاختبارات":
            hdr("📋 تعيين اختبارات الشخصية", "تعيين اختبارات إجبارية للموظفين مع إرسال إشعار بالبريد")

            # Get current user info
            cur_role = st.session_state.get('user_role', '')
            cur_name = st.session_state.get('user_name', '')
            cur_dept = st.session_state.get('user_dept', '')

            # === Employee Self-Service Section ===
            st.markdown("### 👤 الاختبارات الخاصة بي")
            if cur_role == "موظف":
                # Auto show employee's assigned tests
                my_tests = [a for a in st.session_state.test_assignments if a["الموظف"] == cur_name]
                if my_tests:
                    for mt in my_tests:
                        completed = any(r["الاسم"]==cur_name and mt["الاختبار"] in r.get("type","") for r in st.session_state.personality_results)
                        status = "✅ مكتمل" if completed else "⏳ مطلوب (إجباري)"
                        ibox(f"**{mt['الاختبار']}** | الموعد: {mt['الموعد النهائي']} | الحالة: {status} | معيّن بواسطة: {mt['معيّن_بواسطة']}", "success" if completed else "warning")
                        if not completed:
                            st.caption("👆 اذهب لصفحة الاختبار المطلوب من القائمة الجانبية لإتمامه")
                else:
                    ibox("لا توجد اختبارات معيّنة لك حالياً. يمكنك أداء أي اختبار اختيارياً من القائمة الجانبية.")

                st.markdown("---")
                st.markdown("### 📊 نتائجي السابقة")
                my_results = [r for r in st.session_state.personality_results if r["الاسم"] == cur_name]
                if my_results:
                    for mr in my_results:
                        st.markdown(f"- **{mr['type']}** | {mr['التاريخ']} | {'إجباري' if mr.get('إجباري') else 'اختياري'}")
                else:
                    st.caption("لم تكمل أي اختبار بعد.")

            elif cur_role != "موظف":
                # Non-employee can search
                my_name_check = st.text_input("أدخل اسم الموظف لعرض اختباراته:", key="my_name_check")
                if my_name_check and st.session_state.test_assignments:
                    my_tests = [a for a in st.session_state.test_assignments if a["الموظف"] == my_name_check]
                    if my_tests:
                        for mt in my_tests:
                            completed = any(r["الاسم"]==my_name_check and mt["الاختبار"] in r.get("type","") for r in st.session_state.personality_results)
                            status = "✅ مكتمل" if completed else "⏳ مطلوب (إجباري)"
                            ibox(f"**{mt['الاختبار']}** | الموعد: {mt['الموعد النهائي']} | الحالة: {status}", "success" if completed else "warning")
                    else:
                        ibox("لا توجد اختبارات معيّنة لهذا الموظف.")

            # === Manager Assignment Section ===
            if cur_role == "مدير" or st.session_state.get('current_user') == "guest":
                st.markdown("---")
                st.markdown("### 📌 تعيين اختبار لموظف (المدير)")

                # Get employee list from users_db
                emp_users = {uname: udata for uname, udata in st.session_state.get('users_db',{}).items()
                    if udata.get('role') == 'موظف'}
                emp_names = [udata["name"] for udata in emp_users.values()]

                ac1, ac2 = st.columns(2)
                with ac1:
                    assign_method = st.radio("طريقة التعيين:", ["اختيار من الموظفين المسجلين", "إدخال يدوي"], horizontal=True, key="assign_method")

                    if assign_method == "اختيار من الموظفين المسجلين" and emp_names:
                        assign_emp = st.selectbox("اختر الموظف:", emp_names, key="assign_emp_sel")
                        # Auto-fill dept and email
                        sel_user_data = next((ud for ud in emp_users.values() if ud["name"] == assign_emp), {})
                        assign_dept = sel_user_data.get("dept", "")
                        assign_email_default = sel_user_data.get("email", "")
                        st.caption(f"القسم: {assign_dept}" if assign_dept else "")
                    else:
                        assign_emp = st.text_input("اسم الموظف:", key="assign_emp")
                        assign_dept = st.text_input("القسم:", key="assign_dept")
                        assign_email_default = ""

                with ac2:
                    assign_tests = st.multiselect("الاختبارات المطلوبة:", ALL_TESTS, key="assign_tests")
                    assign_deadline = st.date_input("الموعد النهائي:", key="assign_dl")

                # Email section
                st.markdown("#### 📧 إرسال إشعار بالبريد")

                # Check SMTP config status
                smtp_cfg = st.session_state.get('smtp_config', {})
                if not smtp_cfg.get('email'):
                    smtp_cfg = load_smtp_config()
                    if smtp_cfg:
                        st.session_state.smtp_config = smtp_cfg

                smtp_ready = bool(smtp_cfg.get('email') and smtp_cfg.get('password') and smtp_cfg.get('server'))

                if not smtp_ready:
                    st.warning("⚠️ إعدادات البريد غير مكوّنة. أكمل الإعداد السريع أدناه:")
                    with st.expander("⚡ إعداد سريع للبريد", expanded=True):
                        qp = st.selectbox("مزود البريد:", list(SMTP_PROVIDERS.keys()), key="q_prov")
                        qprov = SMTP_PROVIDERS[qp]
                        st.caption(f"💡 {qprov['help']}")
                        q1, q2 = st.columns(2)
                        with q1:
                            q_default_email = qprov.get('email', '')
                            q_email = st.text_input("البريد المرسل:", value=q_default_email, key="q_email", placeholder="HR@resal.me")
                            q_server = st.text_input("خادم SMTP:", value=qprov['server'], key="q_server")
                        with q2:
                            q_pass = st.text_input("كلمة مرور البريد:", type="password", key="q_pass")
                            q_port = st.number_input("المنفذ:", value=qprov['port'], key="q_port")
                        q_url = st.text_input("رابط التطبيق (اختياري):", key="q_url", placeholder="https://your-app.streamlit.app")
                        if st.button("💾 حفظ وتفعيل البريد", type="primary", key="q_save"):
                            if q_email and q_pass:
                                new_cfg = {'server': q_server, 'port': int(q_port), 'email': q_email,
                                    'password': q_pass, 'sender_name': 'إدارة الموارد البشرية - رسال الود',
                                    'app_url': q_url, 'use_ssl': qprov.get('use_ssl', False), 'provider': qp}
                                st.session_state.smtp_config = new_cfg
                                save_smtp_config(new_cfg)
                                st.success("✅ تم تفعيل البريد بنجاح!")
                                st.rerun()
                            else:
                                st.error("أدخل البريد وكلمة المرور")
                else:
                    st.success(f"✅ البريد مفعّل: {smtp_cfg['email']}")

                email_c1, email_c2 = st.columns([3,1])
                with email_c1:
                    assign_email = st.text_input("البريد الإلكتروني للموظف (عمل أو شخصي):",
                        value=assign_email_default, key="assign_email",
                        placeholder="employee@company.com أو employee@gmail.com")
                with email_c2:
                    send_email_flag = st.checkbox("إرسال إشعار", value=smtp_ready, key="send_flag")

                bc1, bc2 = st.columns(2)
                with bc1:
                    if st.button("📌 تعيين الاختبارات", type="primary", key="assign_btn", use_container_width=True):
                        if assign_emp and assign_tests:
                            for t in assign_tests:
                                assignment = {
                                    "الموظف": assign_emp, "القسم": assign_dept,
                                    "الاختبار": t, "الموعد النهائي": str(assign_deadline),
                                    "الحالة": "لم يبدأ", "معيّن_بواسطة": st.session_state.get('user_name', 'المدير'),
                                    "إجباري": True, "البريد": assign_email
                                }
                                db_save_assignment(assignment)
                                st.session_state.test_assignments.append(assignment)
                            st.success(f"✅ تم تعيين {len(assign_tests)} اختبار/ات لـ {assign_emp}")

                            # Send email
                            if send_email_flag and assign_email:
                                ok, msg = send_test_email(assign_email, assign_emp, assign_tests,
                                    str(assign_deadline), st.session_state.get('user_name','المدير'))
                                if ok:
                                    st.success(f"📧 تم إرسال الإشعار إلى {assign_email}")
                                else:
                                    st.warning(f"⚠️ لم يتم الإرسال: {msg}")
                            elif send_email_flag and not assign_email:
                                st.warning("⚠️ لم يتم إدخال بريد إلكتروني")
                            st.rerun()
                        else:
                            st.error("يرجى إدخال اسم الموظف واختيار اختبار واحد على الأقل")

                with bc2:
                    if assign_email and assign_tests and assign_emp:
                        if st.button("📧 إرسال إشعار فقط (بدون تعيين)", key="send_only", use_container_width=True):
                            ok, msg = send_test_email(assign_email, assign_emp, assign_tests,
                                str(assign_deadline), st.session_state.get('user_name','المدير'))
                            if ok:
                                st.success(f"📧 تم إرسال الإشعار إلى {assign_email}")
                            else:
                                st.warning(f"⚠️ {msg}")

                # Show all assignments
                if st.session_state.test_assignments:
                    st.markdown("---")
                    st.markdown("### 📋 جميع الاختبارات المعيّنة")
                    adf = pd.DataFrame(st.session_state.test_assignments)
                    # Check completed
                    completed_names = {(r["الاسم"], r["type"]) for r in st.session_state.personality_results}
                    for i, row in adf.iterrows():
                        if (row["الموظف"], row["الاختبار"].replace(" (OCEAN)","")) in completed_names or \
                           (row["الموظف"], row["الاختبار"]) in completed_names:
                            adf.at[i, "الحالة"] = "✅ مكتمل"
                    st.dataframe(adf, use_container_width=True, hide_index=True)

                    k1,k2,k3 = st.columns(3)
                    total_a = len(adf)
                    done_a = len(adf[adf["الحالة"]=="✅ مكتمل"])
                    with k1: kpi("📋 إجمالي التعيينات", f"{total_a}")
                    with k2: kpi("✅ مكتمل", f"{done_a}")
                    with k3: kpi("⏳ قيد الانتظار", f"{total_a - done_a}")

                    # Admin-only delete
                    if st.session_state.get('user_role') == "مدير":
                        if st.button("🗑️ مسح التعيينات (مدير فقط)", key="assign_clr"):
                            db_delete_assignments()
                            st.session_state.test_assignments = []
                            st.rerun()
                    else:
                        st.caption("⚠️ حذف التعيينات متاح لمدير النظام فقط")

        # ===== PAGE: Big Five =====
        elif page == "🧠 Big Five (OCEAN)":
            hdr("🧠 Big Five - OCEAN Model", "نموذج العوامل الخمسة الكبرى للشخصية (25 سؤال)")
            is_emp = st.session_state.get('user_role') == "موظف"
            if is_emp:
                bf_name = st.session_state.get('user_name','')
                bf_dept = st.session_state.get('user_dept','')
                st.info(f"👤 مرحباً {bf_name} | القسم: {bf_dept}")
            else:
                b1, b2 = st.columns(2)
                with b1: bf_name = st.text_input("الاسم:", key="bf_name")
                with b2: bf_dept = st.text_input("القسم:", key="bf_dept")
            is_mand = any(a["الموظف"]==bf_name and "Big Five" in a["الاختبار"] for a in st.session_state.test_assignments)
            assigned_by = next((a["معيّن_بواسطة"] for a in st.session_state.test_assignments if a["الموظف"]==bf_name and "Big Five" in a["الاختبار"]), "")
            if is_mand: ibox(f"⚠️ هذا الاختبار إجباري - معيّن بواسطة: {assigned_by}", "warning")
            if bf_name:
                render_test("Big Five (OCEAN)", bf_name, bf_dept, is_mand, assigned_by)



        # ===== PAGE: Thomas PPA =====
        elif page == "📊 Thomas PPA":
            hdr("📊 Thomas PPA", "تحليل الملف الشخصي المهني (24 سؤال)")
            st.caption("مبني على نموذج DISC لتحليل السلوك المهني: الهيمنة، التأثير، الثبات، الامتثال")
            is_emp = st.session_state.get('user_role') == "موظف"
            if is_emp:
                tp_name = st.session_state.get('user_name','')
                tp_dept = st.session_state.get('user_dept','')
                st.info(f"👤 مرحباً {tp_name} | القسم: {tp_dept}")
            else:
                t1, t2 = st.columns(2)
                with t1: tp_name = st.text_input("الاسم:", key="tp_name")
                with t2: tp_dept = st.text_input("القسم:", key="tp_dept")
            is_mand = any(a["الموظف"]==tp_name and "Thomas" in a["الاختبار"] for a in st.session_state.test_assignments)
            assigned_by = next((a["معيّن_بواسطة"] for a in st.session_state.test_assignments if a["الموظف"]==tp_name and "Thomas" in a["الاختبار"]), "")
            if is_mand: ibox(f"⚠️ هذا الاختبار إجباري - معيّن بواسطة: {assigned_by}", "warning")
            if tp_name:
                render_test("Thomas PPA", tp_name, tp_dept, is_mand, assigned_by)



        # ===== PAGE: Hogan HPI =====
        elif page == "🔬 Hogan HPI":
            hdr("🔬 Hogan HPI", "مقياس هوجان للشخصية المهنية (28 سؤال)")
            st.caption("يقيس 7 مقاييس أساسية: التوازن النفسي، الطموح، الاجتماعية، الحساسية، الحصافة، الفضول، التعلم")
            is_emp = st.session_state.get('user_role') == "موظف"
            if is_emp:
                hg_name = st.session_state.get('user_name','')
                hg_dept = st.session_state.get('user_dept','')
                st.info(f"👤 مرحباً {hg_name} | القسم: {hg_dept}")
            else:
                h1, h2 = st.columns(2)
                with h1: hg_name = st.text_input("الاسم:", key="hg_name")
                with h2: hg_dept = st.text_input("القسم:", key="hg_dept")
            is_mand = any(a["الموظف"]==hg_name and "Hogan" in a["الاختبار"] for a in st.session_state.test_assignments)
            assigned_by = next((a["معيّن_بواسطة"] for a in st.session_state.test_assignments if a["الموظف"]==hg_name and "Hogan" in a["الاختبار"]), "")
            if is_mand: ibox(f"⚠️ هذا الاختبار إجباري - معيّن بواسطة: {assigned_by}", "warning")
            if hg_name:
                render_test("Hogan HPI", hg_name, hg_dept, is_mand, assigned_by)



        # ===== PAGE: MBTI =====
        elif page == "💡 MBTI":
            hdr("💡 MBTI", "مؤشر مايرز بريغز لأنماط الشخصية (32 سؤال)")
            st.caption("يحدد نمط الشخصية من 16 نمط عبر 4 أبعاد: الطاقة، المعلومات، القرارات، أسلوب الحياة")
            is_emp = st.session_state.get('user_role') == "موظف"
            if is_emp:
                mb_name = st.session_state.get('user_name','')
                mb_dept = st.session_state.get('user_dept','')
                st.info(f"👤 مرحباً {mb_name} | القسم: {mb_dept}")
            else:
                m1, m2 = st.columns(2)
                with m1: mb_name = st.text_input("الاسم:", key="mb_name")
                with m2: mb_dept = st.text_input("القسم:", key="mb_dept")
            is_mand = any(a["الموظف"]==mb_name and "MBTI" in a["الاختبار"] for a in st.session_state.test_assignments)
            assigned_by = next((a["معيّن_بواسطة"] for a in st.session_state.test_assignments if a["الموظف"]==mb_name and "MBTI" in a["الاختبار"]), "")
            if is_mand: ibox(f"⚠️ هذا الاختبار إجباري - معيّن بواسطة: {assigned_by}", "warning")
            if mb_name:
                render_test("MBTI", mb_name, mb_dept, is_mand, assigned_by)

        # ===== PAGE: DISC =====
        elif page == "💎 DISC":
            hdr("💎 اختبار DISC", "تقييم أنماط السلوك المهني الأربعة (24 سؤال)")
            st.caption("يقيس 4 أنماط سلوكية: الهيمنة (D)، التأثير (I)، الثبات (S)، الالتزام (C) مع تحديد النمط المركب")
            is_emp = st.session_state.get('user_role') == "موظف"
            if is_emp:
                disc_name = st.session_state.get('user_name','')
                disc_dept = st.session_state.get('user_dept','')
                st.info(f"👤 مرحباً {disc_name} | القسم: {disc_dept}")
            else:
                d1, d2 = st.columns(2)
                with d1: disc_name = st.text_input("الاسم:", key="disc_name")
                with d2: disc_dept = st.text_input("القسم:", key="disc_dept")
            is_mand = any(a["الموظف"]==disc_name and "DISC" in a["الاختبار"] for a in st.session_state.test_assignments)
            assigned_by = next((a["معيّن_بواسطة"] for a in st.session_state.test_assignments if a["الموظف"]==disc_name and "DISC" in a["الاختبار"]), "")
            if is_mand: ibox(f"⚠️ هذا الاختبار إجباري - معيّن بواسطة: {assigned_by}", "warning")
            if disc_name:
                render_test("DISC", disc_name, disc_dept, is_mand, assigned_by)



        # ===== PAGE: Reports =====
        elif page == "📈 تقارير الشخصية":
            hdr("📈 تقارير اختبارات الشخصية", "قاعدة بيانات دائمة - لا تُحذف إلا بموافقة مدير النظام")

            # Load from DB always
            db_results = db_load_results()
            total_db = db_count_results()

            if db_results:
                # DB Stats
                st.markdown("### 🗄️ إحصائيات قاعدة البيانات")
                dk1,dk2,dk3,dk4 = st.columns(4)
                with dk1: kpi("🗄️ إجمالي السجلات", f"{total_db}")
                types_count = {}
                for r in db_results:
                    types_count[r["type"]] = types_count.get(r["type"], 0) + 1
                with dk2: kpi("📊 أنواع الاختبارات", f"{len(types_count)}")
                emps = set(r["الاسم"] for r in db_results)
                with dk3: kpi("👥 الموظفين", f"{len(emps)}")
                mandatory_count = sum(1 for r in db_results if r.get("إجباري"))
                with dk4: kpi("📌 إجباري", f"{mandatory_count}")

                # Filter
                st.markdown("---")
                fc1, fc2 = st.columns(2)
                with fc1: test_filter = st.selectbox("تصفية حسب الاختبار:", ["الكل"] + ALL_TESTS, key="pt_f")
                with fc2: emp_filter = st.selectbox("تصفية حسب الموظف:", ["الكل"] + sorted(list(emps)), key="pt_emp")

                filtered = db_results
                if test_filter != "الكل":
                    filtered = [r for r in filtered if r["type"] == test_filter]
                if emp_filter != "الكل":
                    filtered = [r for r in filtered if r["الاسم"] == emp_filter]

                if filtered:
                    # Summary table
                    rows = []
                    for r in filtered:
                        row = {"#": r.get("id",""), "الاسم": r["الاسم"], "القسم": r["القسم"], "النوع": r["type"],
                            "التاريخ": r["التاريخ"], "إجباري": "نعم" if r.get("إجباري") else "اختياري",
                            "بواسطة": r.get("معيّن_بواسطة",""), "سُجّل": r.get("created_at","")}
                        if "mbti_type" in r and r["mbti_type"]: row["النمط"] = r["mbti_type"]
                        elif "dominant" in r and r["dominant"]: row["النمط السائد"] = r["dominant"]
                        rows.append(row)
                    st.dataframe(pd.DataFrame(rows), use_container_width=True, hide_index=True)

                    # Individual detailed report
                    st.markdown("### 📄 تقرير فردي مفصّل")
                    sel_options = [f"{r['الاسم']} ({r['type']}) - {r['التاريخ']}" for r in filtered]
                    sel = st.selectbox("اختر الموظف:", sel_options, key="pt_sel")
                    sel_idx = sel_options.index(sel)
                    sel_r = filtered[sel_idx]

                    st.markdown(f"**{sel_r['الاسم']}** | {sel_r['القسم']} | {sel_r['type']} | {sel_r['التاريخ']}")
                    if sel_r.get('إجباري'): st.caption(f"📌 اختبار إجباري - معيّن بواسطة: {sel_r.get('معيّن_بواسطة','')}")
                    if sel_r.get('created_at'): st.caption(f"🕐 تاريخ التسجيل: {sel_r.get('created_at','')}")

                    if sel_r["type"] == "Big Five":
                        for t, pct in sel_r["scores"].items():
                            info = BIG5_TRAITS.get(t, {"name":t,"en":t,"high":"","low":""})
                            level = "مرتفع" if pct >= 70 else ("متوسط" if pct >= 40 else "منخفض")
                            st.progress(pct/100, text=f"{info['name']} ({info['en']}): {pct}% - {level}")
                            st.caption(f"  {info['high'] if pct>=60 else info['low']}")
                    elif sel_r["type"] == "Thomas PPA":
                        for s, pct in sel_r["scores"].items():
                            info = THOMAS_SCALES.get(s, {"name":s,"en":s,"high":"","low":""})
                            st.progress(pct/100, text=f"{info['name']} ({info['en']}): {pct}%")
                            st.caption(f"  {info['high'] if pct>=60 else info['low']}")
                    elif sel_r["type"] == "Hogan HPI":
                        for s, pct in sel_r["scores"].items():
                            info = HOGAN_SCALES.get(s, {"name":s,"en":s,"impact":""})
                            st.progress(pct/100, text=f"{info['name']} ({info['en']}): {pct}%")
                            st.caption(f"  {info['impact']}")
                    elif sel_r["type"] == "MBTI":
                        t = sel_r.get("mbti_type","")
                        if t in MBTI_TYPES:
                            ti = MBTI_TYPES[t]
                            st.markdown(f"### 💡 النمط: {t} - {ti['name']}")
                            ibox(f"**الوصف:** {ti['desc']}")
                            ibox(f"**نقاط القوة:** {ti['strengths']}", "success")
                            ibox(f"**المسارات المهنية:** {ti['careers']}")
                    elif sel_r["type"] == "DISC":
                        dominant = sel_r.get("dominant","")
                        secondary = sel_r.get("secondary","")
                        if dominant in DISC_STYLES:
                            di = DISC_STYLES[dominant]
                            st.markdown(f"### 💎 النمط السائد: {di['name']} ({dominant}) | الثانوي: {secondary}")
                            for s, pct in sel_r["scores"].items():
                                info = DISC_STYLES.get(s, {"name":s,"en":s})
                                st.progress(pct/100, text=f"{info['name']} ({info['en']}): {pct}%")
                            ibox(f"**نقاط القوة:** {di['strengths']}", "success")
                            ibox(f"**التحديات:** {di['challenges']}", "warning")
                            ibox(f"**الوظائف المناسبة:** {di['careers']}")

                    # === Download individual report ===
                    st.markdown("### 📥 تحميل تقرير هذا الموظف")
                    dl1, dl2 = st.columns(2)
                    with dl1:
                        # PDF download
                        pdf_bytes = generate_employee_pdf(sel_r)
                        if pdf_bytes:
                            st.download_button("📄 تحميل PDF", data=pdf_bytes,
                                file_name=f"Report_{sel_r['الاسم']}_{sel_r['type']}_{sel_r['التاريخ']}.pdf",
                                mime="application/pdf", use_container_width=True)
                        else:
                            st.caption("⚠️ يرجى تثبيت fpdf2 لتصدير PDF")
                    with dl2:
                        # Excel download for individual
                        ind_ox = io.BytesIO()
                        with pd.ExcelWriter(ind_ox, engine='xlsxwriter') as w:
                            ind_row = {"الاسم": sel_r["الاسم"], "القسم": sel_r["القسم"], "النوع": sel_r["type"],
                                "التاريخ": sel_r["التاريخ"], "إجباري": "نعم" if sel_r.get("إجباري") else "لا"}
                            if sel_r.get("mbti_type"): ind_row["نمط MBTI"] = sel_r["mbti_type"]
                            if sel_r.get("dominant"): ind_row["النمط السائد"] = sel_r["dominant"]
                            for k, v in sel_r["scores"].items():
                                name = k
                                if k in BIG5_TRAITS: name = BIG5_TRAITS[k]["name"]
                                elif k in THOMAS_SCALES: name = THOMAS_SCALES[k]["name"]
                                elif k in HOGAN_SCALES: name = HOGAN_SCALES[k]["name"]
                                elif k in DISC_STYLES: name = DISC_STYLES[k]["name"]
                                ind_row[name] = f"{v}%"
                            pd.DataFrame([ind_row]).to_excel(w, sheet_name='النتيجة', index=False)
                            w.sheets['النتيجة'].right_to_left()
                        st.download_button("📊 تحميل Excel", data=ind_ox.getvalue(),
                            file_name=f"Report_{sel_r['الاسم']}_{sel_r['type']}_{sel_r['التاريخ']}.xlsx",
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", use_container_width=True)

                    # Comparison
                    same_type = [r for r in filtered if r["type"] == sel_r["type"]]
                    if len(same_type) >= 2 and sel_r["type"] in ["Big Five", "Hogan HPI", "DISC"]:
                        st.markdown("### 📊 مقارنة مع الآخرين")
                        fig = go.Figure()
                        for r in same_type:
                            if r["type"] == "Big Five":
                                names = [BIG5_TRAITS.get(t,{"name":t})["name"] for t in r["scores"]]
                            elif r["type"] == "Hogan HPI":
                                names = [HOGAN_SCALES.get(s,{"name":s})["name"] for s in r["scores"]]
                            elif r["type"] == "DISC":
                                names = [DISC_STYLES.get(s,{"name":s})["name"] for s in r["scores"]]
                            else:
                                names = list(r["scores"].keys())
                            vals = list(r["scores"].values()) + [list(r["scores"].values())[0]]
                            fig.add_trace(go.Scatterpolar(r=vals, theta=names+[names[0]], fill='toself', name=r["الاسم"]))
                        fig.update_layout(polar=dict(radialaxis=dict(range=[0,100])), title="مقارنة الملفات الشخصية", font=dict(family="Noto Sans Arabic"), height=500)
                        st.plotly_chart(fig, use_container_width=True)

                # Admin-only delete section
                st.markdown("---")
                if st.session_state.get('user_role') == "مدير":
                    st.markdown("### 🔒 إدارة قاعدة البيانات (مدير النظام فقط)")

                    del_tab1, del_tab2 = st.tabs(["🗑️ حذف سجل محدد", "⚠️ حذف الكل"])

                    with del_tab1:
                        del_id = st.number_input("رقم السجل المراد حذفه:", min_value=1, step=1, key="del_rid")
                        del_confirm = st.text_input("اكتب 'تأكيد الحذف' للمتابعة:", key="del_conf")
                        if st.button("🗑️ حذف السجل", key="del_one"):
                            if del_confirm == "تأكيد الحذف":
                                db_delete_result(int(del_id))
                                st.session_state.personality_results = db_load_results()
                                st.success(f"✅ تم حذف السجل #{del_id}")
                                st.rerun()
                            else:
                                st.error("يرجى كتابة 'تأكيد الحذف' للمتابعة")

                    with del_tab2:
                        st.error(f"⚠️ تحذير: سيتم حذف جميع السجلات ({total_db} سجل) نهائياً!")
                        del_all_confirm = st.text_input("اكتب 'حذف جميع البيانات' للمتابعة:", key="del_all_conf")
                        if st.button("⚠️ حذف جميع السجلات نهائياً", key="del_all"):
                            if del_all_confirm == "حذف جميع البيانات":
                                db_delete_all_results()
                                st.session_state.personality_results = []
                                st.success("✅ تم حذف جميع السجلات")
                                st.rerun()
                            else:
                                st.error("يرجى كتابة 'حذف جميع البيانات' للمتابعة")
                else:
                    ibox("🔒 بيانات الاختبارات محفوظة في قاعدة بيانات دائمة ولا يمكن حذفها إلا بموافقة مدير النظام.", "warning")

            else:
                ibox("لا توجد نتائج بعد. اذهب لأي اختبار وابدأ.", "warning")

        # ===== PAGE: Export =====
        elif page == "📥 تصدير الاختبارات":
            hdr("📥 تصدير نتائج اختبارات الشخصية", "تصدير Excel أو PDF لجميع النتائج من قاعدة البيانات")

            db_results = db_load_results()

            if db_results:
                st.success(f"📊 {len(db_results)} نتيجة محفوظة في قاعدة البيانات")

                export_type = st.radio("نوع التصدير:", ["📊 Excel شامل", "📄 PDF لجميع الموظفين"], horizontal=True, key="exp_type")

                if export_type == "📊 Excel شامل":
                    # Filter options
                    exp_filter = st.selectbox("تصفية:", ["الكل"] + ALL_TESTS, key="exp_f")
                    exp_data = db_results if exp_filter == "الكل" else [r for r in db_results if r["type"] == exp_filter]

                    ox = io.BytesIO()
                    with pd.ExcelWriter(ox, engine='xlsxwriter') as w:
                        # All results summary
                        rows = []
                        for r in exp_data:
                            row = {"رقم السجل": r.get("id",""), "الاسم": r["الاسم"], "القسم": r["القسم"], "النوع": r["type"],
                                "التاريخ": r["التاريخ"], "إجباري": "نعم" if r.get("إجباري") else "لا",
                                "معيّن بواسطة": r.get("معيّن_بواسطة",""), "تاريخ التسجيل": r.get("created_at","")}
                            if "mbti_type" in r and r.get("mbti_type"): row["نمط MBTI"] = r["mbti_type"]
                            if "dominant" in r and r.get("dominant"): row["النمط السائد"] = r["dominant"]
                            if "secondary" in r and r.get("secondary"): row["النمط الثانوي"] = r["secondary"]
                            for k, v in r.get("scores",{}).items():
                                name = k
                                if k in BIG5_TRAITS: name = BIG5_TRAITS[k]["name"]
                                elif k in THOMAS_SCALES: name = THOMAS_SCALES[k]["name"]
                                elif k in HOGAN_SCALES: name = HOGAN_SCALES[k]["name"]
                                elif k in DISC_STYLES: name = DISC_STYLES[k]["name"]
                                row[name] = f"{v}%"
                            rows.append(row)
                        pd.DataFrame(rows).to_excel(w, sheet_name='جميع النتائج', index=False)
                        w.sheets['جميع النتائج'].right_to_left()

                        # Per-test sheets
                        for test in ALL_TESTS:
                            test_data = [r for r in exp_data if r["type"] == test]
                            if test_data:
                                trows = []
                                for r in test_data:
                                    tr = {"الاسم": r["الاسم"], "القسم": r["القسم"], "التاريخ": r["التاريخ"]}
                                    for k, v in r.get("scores",{}).items():
                                        name = k
                                        if k in BIG5_TRAITS: name = BIG5_TRAITS[k]["name"]
                                        elif k in THOMAS_SCALES: name = THOMAS_SCALES[k]["name"]
                                        elif k in HOGAN_SCALES: name = HOGAN_SCALES[k]["name"]
                                        elif k in DISC_STYLES: name = DISC_STYLES[k]["name"]
                                        tr[name] = v
                                    trows.append(tr)
                                sname = test.replace(" (OCEAN)","")[:31]
                                pd.DataFrame(trows).to_excel(w, sheet_name=sname, index=False)
                                w.sheets[sname].right_to_left()

                        # Assignments sheet
                        db_assigns = db_load_assignments()
                        if db_assigns:
                            pd.DataFrame(db_assigns).to_excel(w, sheet_name='التعيينات', index=False)
                            w.sheets['التعيينات'].right_to_left()

                    st.download_button("📥 تحميل Excel الشامل", data=ox.getvalue(),
                        file_name=f"Personality_DB_{datetime.now().strftime('%Y%m%d_%H%M')}.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", type="primary", use_container_width=True)

                elif export_type == "📄 PDF لجميع الموظفين":
                    st.info("اختر موظف لتحميل تقريره بصيغة PDF")
                    pdf_sel = st.selectbox("الموظف:", [f"{r['الاسم']} ({r['type']}) - {r['التاريخ']}" for r in db_results], key="pdf_sel")
                    pdf_idx = [f"{r['الاسم']} ({r['type']}) - {r['التاريخ']}" for r in db_results].index(pdf_sel)
                    pdf_r = db_results[pdf_idx]

                    pdf_bytes = generate_employee_pdf(pdf_r)
                    if pdf_bytes:
                        st.download_button("📄 تحميل PDF", data=pdf_bytes,
                            file_name=f"Report_{pdf_r['الاسم']}_{pdf_r['type']}_{pdf_r['التاريخ']}.pdf",
                            mime="application/pdf", type="primary", use_container_width=True)
                    else:
                        ibox("⚠️ يرجى تثبيت مكتبة fpdf2 لتوليد ملفات PDF", "warning")

                    # Batch PDF - all employees
                    st.markdown("---")
                    st.markdown("### 📦 تصدير جميع التقارير")
                    st.caption(f"سيتم تصدير {len(db_results)} تقرير في ملف Excel واحد مع ورقة لكل اختبار")
            else:
                ibox("لا توجد بيانات للتصدير. أكمل اختباراً واحداً على الأقل.", "warning")


    # =========================================
    #       👥 USER MANAGEMENT
    # =========================================
    elif section == "👥 إدارة المستخدمين":
        user_management_page()


if __name__ == "__main__":
    main()

